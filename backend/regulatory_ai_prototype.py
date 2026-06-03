#!/usr/bin/env python3
"""SchoolFood AI v2.1 — Phase 2B: Database (Supabase), History, Feedback Backend"""

import base64
import json
from datetime import datetime, timezone, timedelta
from pathlib import Path

# Múi giờ Việt Nam UTC+7 — dùng cho mọi tính toán thời gian
VN_TZ = timezone(timedelta(hours=7))

def now_vn() -> datetime:
    """Trả về giờ hiện tại theo múi giờ Việt Nam (UTC+7)."""
    return datetime.now(VN_TZ)

import anthropic
import streamlit as st

# ── G1: Database Layer — Supabase ────────────────────────────────────────────
# Không dùng @cache_resource để tránh cache None từ trước khi có secrets
_sb_client = None
_sb_error  = ""

def _get_sb():
    """Kết nối Supabase — lazy init, không cache để đảm bảo luôn đọc secrets mới nhất."""
    global _sb_client, _sb_error

    # Đã kết nối rồi thì trả về luôn
    if _sb_client is not None:
        return _sb_client

    try:
        from supabase import create_client  # type: ignore

        # Thử đọc secrets theo nhiều cách
        url = ""
        key = ""
        try:
            url = st.secrets["SUPABASE_URL"]
            key = st.secrets["SUPABASE_ANON_KEY"]
        except (KeyError, FileNotFoundError):
            try:
                url = st.secrets.get("SUPABASE_URL", "")
                key = st.secrets.get("SUPABASE_ANON_KEY", "")
            except Exception:
                pass

        if not url or not key:
            _sb_error = "SUPABASE_URL hoặc SUPABASE_ANON_KEY chưa được thêm vào Secrets"
            return None

        if not url.startswith("https://"):
            _sb_error = f"SUPABASE_URL không hợp lệ: {url[:30]}..."
            return None

        _sb_client = create_client(str(url), str(key))
        _sb_error  = ""
        return _sb_client

    except ImportError:
        _sb_error = "Package 'supabase' chưa được cài (kiểm tra requirements.txt)"
    except Exception as e:
        _sb_error = f"Lỗi kết nối: {str(e)[:100]}"

    return None

def db_ok() -> bool:
    """Kiểm tra database có sẵn sàng không."""
    return _get_sb() is not None

def db_error_msg() -> str:
    """Trả về thông báo lỗi kết nối (nếu có)."""
    return _sb_error

def db_save_checklist(school: str, date_str: str, inspector: str, menu: str,
                      level: str, results: dict, notes: dict,
                      alert_level: str, pass_count: int, fail_count: int,
                      ai_narrative: str = "",
                      extra_results: dict | None = None,
                      check_type: str = "ban_giam_sat") -> str | None:
    """
    Lưu kết quả checklist vào Supabase.
    check_type: 'ban_giam_sat' | 'nha_cung_cap'
    Trả về session_id hoặc None nếu lỗi/không có DB.
    """
    sb = _get_sb()
    if not sb:
        return None
    try:
        sess = sb.table("checklist_sessions").insert({
            "school_name":    school or "Chưa nhập",
            "inspector_name": inspector or "",
            "check_date":     date_str,
            "menu_today":     menu or "",
            "school_level":   level,
            "check_type":     check_type,
            "alert_level":    alert_level,
            "total_items":    pass_count + fail_count,
            "pass_count":     pass_count,
            "fail_count":     fail_count,
            "ai_narrative":   ai_narrative[:2000] if ai_narrative else "",
        }).execute()
        if not sess.data:
            return None
        sid = sess.data[0]["id"]

        # Lưu từng điểm kiểm tra
        items = [
            {"session_id": sid, "item_code": k,
             "result": v.replace("✅ Đạt", "Đạt").replace("❌ Không Đạt", "Không Đạt"),
             "note": notes.get(k, ""),
             "is_critical": k in CRITICAL_ITEMS}
            for k, v in results.items() if v is not None
        ]
        if extra_results:
            items += [
                {"session_id": sid, "item_code": k,
                 "result": v.replace("✅ Đạt", "Đạt").replace("❌ Không Đạt", "Không Đạt"),
                 "note": "", "is_critical": False}
                for k, v in extra_results.items() if v is not None
            ]
        if items:
            sb.table("checklist_results").insert(items).execute()
        return sid
    except Exception as e:
        st.warning(f"⚠️ Không lưu được vào database: {e}", icon="💾")
        return None


def db_save_kiem_thuc(school: str, date_str: str, yte_name: str, menu: str,
                      all_results: dict, all_notes: dict, timestamps: dict,
                      pass_count: int, fail_count: int) -> str | None:
    """Lưu kết quả Kiểm thực 3 bước của Y Tế Học Đường."""
    sb = _get_sb()
    if not sb:
        return None
    try:
        sess = sb.table("checklist_sessions").insert({
            "school_name":    school or "Chưa nhập",
            "inspector_name": yte_name or "",
            "check_date":     date_str,
            "menu_today":     menu or "",
            "school_level":   "Y Tế Học Đường",
            "check_type":     "kiem_thuc_3_buoc",
            "alert_level":    "OK" if fail_count == 0 else "MAJOR",
            "total_items":    pass_count + fail_count,
            "pass_count":     pass_count,
            "fail_count":     fail_count,
        }).execute()
        if not sess.data:
            return None
        sid = sess.data[0]["id"]

        # Lưu từng điểm
        items = [
            {"session_id": sid, "item_code": k,
             "result": v.replace("✅ Đạt", "Đạt").replace("❌ Không Đạt", "Không Đạt"),
             "note": all_notes.get(k, ""),
             "is_critical": k in {"B3_05"}}
            for k, v in all_results.items() if v is not None
        ]
        if items:
            sb.table("checklist_results").insert(items).execute()

        # Lưu timestamp từng bước
        for step in KIEM_THUC:
            b = step["buoc"]
            ts = timestamps.get(b, "")
            step_results = {k: v for k, v in all_results.items()
                           if k.startswith(f"B{b}_")}
            sp = sum(1 for v in step_results.values() if v == "✅ Đạt")
            sf = sum(1 for v in step_results.values() if v == "❌ Không Đạt")
            # on_time: check if timestamp in window
            on_time = None
            if ts:
                try:
                    parts = step["time_window"].replace("–","-").split("-")
                    ws = sum(int(x)*m for x,m in zip(parts[0].strip().split(":"), [60,1]))
                    we = sum(int(x)*m for x,m in zip(parts[1].strip().split(":"), [60,1]))
                    tm = sum(int(x)*m for x,m in zip(ts.split(":")[:2], [60,1]))
                    on_time = ws <= tm <= we
                except Exception:
                    on_time = None
            sb.table("kiem_thuc_steps").insert({
                "session_id":  sid,
                "step_no":     b,
                "time_window": step["time_window"],
                "confirmed_at": ts,
                "on_time":     on_time,
                "pass_count":  sp,
                "fail_count":  sf,
            }).execute()
        return sid
    except Exception as e:
        st.warning(f"⚠️ Không lưu được vào database: {e}", icon="💾")
        return None


def db_save_feedback(school: str, category: str, content: str) -> bool:
    """Lưu feedback Phụ Huynh vào Supabase."""
    sb = _get_sb()
    if not sb:
        return False
    try:
        sb.table("parent_feedback").insert({
            "school_name": school or "Không rõ",
            "category":    category,
            "content":     content[:2000],
            "status":      "pending",
        }).execute()
        return True
    except Exception:
        return False


@st.cache_data(ttl=60, show_spinner=False)
def db_get_sessions(school: str = "", limit: int = 30) -> list:
    """Lấy lịch sử phiên kiểm tra — cache 60s để giảm lag khi chuyển tab."""
    sb = _get_sb()
    if not sb:
        return []
    try:
        q = sb.table("checklist_sessions").select("*").order("created_at", desc=True).limit(limit)
        if school:
            q = q.eq("school_name", school)
        return q.execute().data or []
    except Exception:
        return []


@st.cache_data(ttl=120)
def db_get_schools() -> list[str]:
    """Lấy danh sách tên trường distinct từ DB — cache 2 phút."""
    sb = _get_sb()
    if not sb:
        return []
    try:
        rows = sb.table("checklist_sessions").select("school_name").execute().data or []
        seen, result = set(), []
        for r in rows:
            n = r.get("school_name", "").strip()
            if n and n not in seen:
                seen.add(n)
                result.append(n)
        return sorted(result)
    except Exception:
        return []


# ── G3: Authentication & User Management ─────────────────────────────────────

ROLE_VN = {
    "phu_huynh":      "Phụ Huynh",
    "ban_giam_sat":   "Ban Giám Sát (Đại Diện PHHS)",
    "y_te_hoc_duong": "Y Tế Học Đường",
    "ban_giam_hieu":  "Ban Giám Hiệu",
}
ROLE_KEY = {v: k for k, v in ROLE_VN.items()}
ROLE_CLR_MAP = {
    "Phụ Huynh":                    "#2563EB",
    "Ban Giám Sát (Đại Diện PHHS)": "#7C3AED",
    "Y Tế Học Đường":               "#0D9488",
    "Ban Giám Hiệu":                "#B45309",
}


def _auth_sb():
    """Tạo fresh Supabase client cho auth — tránh chia sẻ auth state toàn cục."""
    try:
        from supabase import create_client
        url = st.secrets.get("SUPABASE_URL", "") if hasattr(st, "secrets") else ""
        key = st.secrets.get("SUPABASE_ANON_KEY", "") if hasattr(st, "secrets") else ""
        if url and key:
            return create_client(str(url), str(key))
    except Exception:
        pass
    return None


def db_auth_login(email: str, password: str) -> dict:
    """Đăng nhập email/password → {id, email, access_token} hoặc raise Exception."""
    client = _auth_sb()
    if not client:
        raise Exception("Database chưa kết nối — không thể đăng nhập")
    resp = client.auth.sign_in_with_password({"email": email.strip().lower(), "password": password})
    if not resp.user or not resp.session:
        raise Exception("Email hoặc mật khẩu không đúng")
    return {"id": str(resp.user.id), "email": resp.user.email,
            "access_token": resp.session.access_token}


def db_auth_signup(email: str, password: str) -> str:
    """Tạo tài khoản Supabase Auth → user_id hoặc raise."""
    client = _auth_sb()
    if not client:
        raise Exception("Database chưa kết nối")
    resp = client.auth.sign_up({"email": email.strip().lower(), "password": password})
    if not resp.user:
        raise Exception("Không tạo được tài khoản — email có thể đã tồn tại")
    return str(resp.user.id)


def db_auth_reset_password(email: str) -> bool:
    """Gửi email đặt lại mật khẩu."""
    client = _auth_sb()
    if not client:
        return False
    try:
        client.auth.reset_password_email(email.strip().lower())
        return True
    except Exception:
        return False


def db_get_profile(user_id: str) -> dict | None:
    """Lấy profile từ user_profiles table."""
    sb = _get_sb()
    if not sb:
        return None
    try:
        r = sb.table("user_profiles").select("*").eq("id", user_id).execute()
        return r.data[0] if r.data else None
    except Exception:
        return None


def db_save_profile(user_id: str, email: str, full_name: str,
                    role: str, school_name: str,
                    default_level: str = "Tiểu Học (6–11 tuổi)") -> bool:
    """Tạo hoặc cập nhật user profile. default_level dùng cho Y Tế đa cấp."""
    sb = _get_sb()
    if not sb:
        return False
    try:
        sb.table("user_profiles").upsert({
            "id": user_id, "email": email.strip().lower(),
            "full_name": full_name.strip(), "role": role,
            "school_name": school_name.strip() if school_name else "",
            "default_level": default_level,
            "is_active": True,
        }).execute()
        return True
    except Exception:
        return False


def db_get_all_profiles(school: str = "") -> list:
    """Lấy tất cả user profiles (BGH admin)."""
    sb = _get_sb()
    if not sb:
        return []
    try:
        q = sb.table("user_profiles").select("*").order("created_at", desc=False)
        if school:
            q = q.eq("school_name", school)
        return q.execute().data or []
    except Exception:
        return []


def db_toggle_profile(user_id: str, is_active: bool) -> bool:
    """Bật/tắt tài khoản."""
    sb = _get_sb()
    if not sb:
        return False
    try:
        sb.table("user_profiles").update({"is_active": is_active}).eq("id", user_id).execute()
        return True
    except Exception:
        return False


def db_get_feedback(school: str = "", status: str = "pending") -> list:
    """Lấy feedback Phụ Huynh."""
    sb = _get_sb()
    if not sb:
        return []
    try:
        q = sb.table("parent_feedback").select("*").eq("status", status) \
            .order("created_at", desc=True).limit(50)
        if school:
            q = q.eq("school_name", school)
        return q.execute().data or []
    except Exception:
        return []


def db_update_feedback_status(feedback_id: str, new_status: str):
    """Ban Giám Hiệu đánh dấu feedback đã xử lý."""
    sb = _get_sb()
    if not sb:
        return
    try:
        sb.table("parent_feedback").update({
            "status": new_status,
            "reviewed_at": now_vn().isoformat(),
        }).eq("id", feedback_id).execute()
    except Exception:
        pass


def db_get_all_feedbacks(school: str = "", limit: int = 100) -> list:
    """Lấy TẤT CẢ feedback không filter status — dùng cho tab_history và traceback."""
    sb = _get_sb()
    if not sb: return []
    try:
        q = (sb.table("parent_feedback").select("*")
             .order("created_at", desc=True).limit(limit))
        if school:
            q = q.eq("school_name", school)
        return q.execute().data or []
    except Exception:
        return []


def db_add_evidence(feedback_id: str, evidence_text: str, by_name: str) -> bool:
    """BGS/Y Tế thêm minh chứng, diễn giải cho complaint — status chuyển sang reviewed."""
    sb = _get_sb()
    if not sb: return False
    try:
        sb.table("parent_feedback").update({
            "evidence_text": evidence_text,
            "evidence_by":   by_name,
            "status":        "reviewed",
            "reviewed_at":   now_vn().isoformat(),
        }).eq("id", feedback_id).execute()
        return True
    except Exception:
        return False


def db_resolve_complaint(feedback_id: str, response_text: str, by_name: str) -> bool:
    """BGH đóng complaint, ghi phản hồi chính thức."""
    sb = _get_sb()
    if not sb: return False
    try:
        sb.table("parent_feedback").update({
            "status":        "resolved",
            "response_text": response_text,
            "response_by":   by_name,
            "reviewed_at":   now_vn().isoformat(),
        }).eq("id", feedback_id).execute()
        return True
    except Exception:
        return False

# ── Task#5: NCC Registry — track giấy phép & chứng nhận ATTP ─────────────────

def db_get_ncc_registry(school: str = "") -> list:
    """Lấy danh sách NCC đã đăng ký, kèm ngày hết hạn chứng nhận."""
    sb = _get_sb()
    if not sb: return []
    try:
        q = sb.table("ncc_registry").select("*").eq("is_active", True).order("ncc_name")
        if school: q = q.eq("school_name", school)
        return q.execute().data or []
    except Exception: return []


def db_save_ncc_registry(school: str, ncc_name: str, license_no: str,
                          license_expiry: str, attp_expiry: str,
                          phone: str = "", notes: str = "") -> bool:
    """Lưu hoặc cập nhật thông tin NCC trong registry."""
    sb = _get_sb()
    if not sb: return False
    try:
        # Upsert theo school + ncc_name
        existing = sb.table("ncc_registry").select("id")\
            .eq("school_name", school).eq("ncc_name", ncc_name).execute().data
        data = {
            "school_name": school, "ncc_name": ncc_name,
            "license_no": license_no or "",
            "license_expiry": license_expiry or None,
            "attp_expiry": attp_expiry or None,
            "phone": phone or "", "notes": notes or "",
            "is_active": True,
            "updated_at": now_vn().isoformat(),
        }
        if existing:
            sb.table("ncc_registry").update(data).eq("id", existing[0]["id"]).execute()
        else:
            sb.table("ncc_registry").insert(data).execute()
        return True
    except Exception: return False


def db_change_password(access_token: str, new_password: str) -> tuple[bool, str]:
    """Đổi mật khẩu người dùng đang đăng nhập. Trả về (success, error_msg)."""
    try:
        from supabase import create_client
        url = st.secrets.get("SUPABASE_URL","") if hasattr(st,"secrets") else ""
        key = st.secrets.get("SUPABASE_ANON_KEY","") if hasattr(st,"secrets") else ""
        if not url or not key:
            return False, "Database chưa kết nối"
        _c = create_client(str(url), str(key))
        # Dùng access_token để authenticate session trước khi update
        _c.auth.set_session(access_token, access_token)
        _c.auth.update_user({"password": new_password})
        return True, ""
    except Exception as _e:
        return False, str(_e)[:120]


# Thư mục gốc của repo deploy (04_Development/)
ROOT        = Path(__file__).parent.parent
# Tìm PDF: ưu tiên thư mục deploy, fallback về thư mục gốc project
LEGAL_DIR   = ROOT / "legal_docs"
if not LEGAL_DIR.exists():
    LEGAL_DIR = ROOT.parent / "07_Legal_Regulations"
# Tìm prompt: ưu tiên thư mục deploy
PROMPT_FILE = ROOT / "prompts/regulatory_qa_v1.md"
if not PROMPT_FILE.exists():
    PROMPT_FILE = ROOT.parent / "03_Product/AI_Prompts/regulatory_qa_v1.md"
# Model mặc định — Sonnet 4.6 cho tất cả text tasks
# Vision (ảnh): dùng Sonnet 4.6 (cũng hỗ trợ vision, rẻ hơn Opus ~5x)
MODEL         = "claude-sonnet-4-6"
MODEL_VISION  = "claude-sonnet-4-6"   # Thay vì opus — cùng chất lượng vision, tiết kiệm hơn
MAX_TOK       = 1200                   # Giảm từ 1500 — đủ cho hầu hết câu trả lời

# ── Định nghĩa mục bắt buộc & ngưỡng điểm ────────────────────────────────────
CRITICAL_ITEMS = {"C03", "C07", "C09", "C10", "C11", "C18", "C20"}
# C03: hạn dùng, C07: nhiệt độ nhận, C09: nhiệt độ chia, C10: thời gian nấu
# C11: màu/mùi, C18: sổ kiểm thực, C20: mẫu lưu
SCORE_EXCELLENT  = 18   # ≥ 18/20 → Đạt chuẩn
SCORE_ACCEPTABLE = 15   # 15–17/20 → Cần cải thiện
# < 15 → Không đạt

# ── Tiêu chuẩn dinh dưỡng theo cấp học (QĐ 3958/QĐ-BYT 2025) ────────────────
NUTRITION = {
    "Tiểu Học (6–11 tuổi)": {
        "short": "Tiểu Học",
        "kcal": "500–650 kcal",
        "pct_day": "30–40%",
        "protein_pct": "13–20%",
        "fat_pct": "25–30%",
        "carb_pct": "55–65%",
        "meat_g": 50,
        "veg_g": 80,
        "veg_range": "80–120g",
        "note": "Bữa phụ buổi chiều thêm sữa hoặc sản phẩm sữa (5–10% nhu cầu ngày)",
    },
    "THCS (12–15 tuổi)": {
        "short": "THCS",
        "kcal": "650–800 kcal",
        "pct_day": "35–42%",
        "protein_pct": "13–20%",
        "fat_pct": "20–30%",
        "carb_pct": "55–65%",
        "meat_g": 70,
        "veg_g": 100,
        "veg_range": "100–150g",
        "note": "Cá/hải sản tối thiểu 2–3 lần/tuần, đậu hũ/đậu 2 lần/tuần",
    },
    "THPT (16–18 tuổi)": {
        "short": "THPT",
        "kcal": "750–900 kcal",
        "pct_day": "38–45%",
        "protein_pct": "13–20%",
        "fat_pct": "20–30%",
        "carb_pct": "55–65%",
        "meat_g": 80,
        "veg_g": 120,
        "veg_range": "100–150g",
        "note": "Nam và nữ có nhu cầu năng lượng khác nhau — thực đơn nên đa dạng",
    },
}

# ── Hệ thống cảnh báo theo cấp độ ────────────────────────────────────────────
ALERT_SYSTEM = {
    "CRITICAL": {
        "icon": "🔴", "label": "CRITICAL — Nguy hiểm tức thì",
        "color": "#DC2626", "bg": "#FEF2F2", "border": "#FCA5A5",
        "triggers": [
            "Bất kỳ mục bắt buộc (*) nào bị KHÔNG ĐẠT",
            "Phát hiện dấu hiệu ngộ độc tập thể (≥ 2 học sinh có triệu chứng)",
            "Thức ăn hỏng rõ ràng: mốc, mùi chua, màu lạ bất thường",
        ],
        "notify": [
            "🏫 Hiệu Trưởng — GỌI ĐIỆN NGAY",
            "🏥 Y Tế Học Đường — NGAY LẬP TỨC",
            "👥 Ban Giám Sát (Đại Diện PHHS) — TRONG 5 PHÚT",
            "📞 Cấp cứu 115 — Nếu có học sinh bị ảnh hưởng",
            "🏛️ Sở Y Tế địa phương — TRONG 24 GIỜ (bắt buộc theo luật)",
        ],
        "timeframe": "Trong 5 phút",
        "action": "⛔ Tạm dừng bữa ăn ngay · Giữ nguyên mẫu thức ăn (không vứt, không rửa) · Cách ly học sinh bị ảnh hưởng",
    },
    "MAJOR": {
        "icon": "🟠", "label": "MAJOR — Xử lý trong ngày",
        "color": "#D97706", "bg": "#FFFBEB", "border": "#FCD34D",
        "triggers": [
            "Tổng điểm dưới 15/20 (dù không có critical item nào fail)",
            "Từ 3 mục KHÔNG ĐẠT trong cùng một nhóm",
            "Nhà Cung Cấp thiếu hóa đơn hoặc giấy tờ ATTP",
        ],
        "notify": [
            "🏫 Hiệu Trưởng — Trong 2 giờ",
            "🏥 Y Tế Học Đường — Ngay",
            "🏢 Nhà Cung Cấp — Yêu cầu giải trình và khắc phục",
        ],
        "timeframe": "Trong 2–4 giờ",
        "action": "⚠️ Yêu cầu Nhà Cung Cấp khắc phục trước bữa ăn tiếp theo · Kiểm tra lại trong 24h · Ghi vào hồ sơ theo dõi",
    },
    "MINOR": {
        "icon": "🟡", "label": "MINOR — Cải thiện trong tuần",
        "color": "#CA8A04", "bg": "#FEFCE8", "border": "#FDE68A",
        "triggers": [
            "Tổng điểm đạt 15–17/20 (dưới mức xuất sắc)",
            "Có 1–2 mục không bắt buộc bị KHÔNG ĐẠT",
        ],
        "notify": [
            "🏥 Y Tế Học Đường — Lưu hồ sơ theo dõi",
            "🏢 Nhà Cung Cấp — Thông báo cải thiện",
        ],
        "timeframe": "Trong 24–48 giờ",
        "action": "📝 Ghi vào hồ sơ · Yêu cầu cải thiện trong lần kiểm tra tiếp theo",
    },
    "OK": {
        "icon": "✅", "label": "ĐẠT CHUẨN — Lưu hồ sơ",
        "color": "#16A34A", "bg": "#F0FDF4", "border": "#86EFAC",
        "triggers": [
            "Tất cả mục bắt buộc (*) đều ĐẠT",
            "Tổng điểm từ 18/20 trở lên",
        ],
        "notify": [
            "🏫 Hiệu Trưởng — Báo cáo tổng hợp cuối tháng",
            "👨‍👩‍👧 Phụ Huynh — Chia sẻ kết quả tốt qua ứng dụng (tuỳ chọn)",
        ],
        "timeframe": "Cuối tháng",
        "action": "✅ Lưu báo cáo vào hồ sơ · Duy trì tiêu chuẩn",
    },
}

# ── Lịch kiểm tra & tần suất ─────────────────────────────────────────────────
SCHEDULE = [
    {
        "role": "🏥 Y Tế Học Đường",
        "freq": "Kiểm thực 3 bước: MỖI NGÀY có bữa ăn · Kiểm tra NCC: Khi nhận hàng (tối thiểu 1 lần/tuần)",
        "when": "Kiểm thực: 10:00–10:45 (trước bữa trưa 30–45 phút) · Kiểm NCC: Ngay khi xe giao hàng đến",
        "what": "Kiểm thực 3 bước (sổ bắt buộc TTLT 13/2016) · Checklist 12 điểm NCC khi nhận hàng",
        "notice": "Không cần báo trước",
        "report": "Lưu sổ tại bếp · Báo Hiệu Trưởng ngay khi có vấn đề · Gửi kết quả NCC định kỳ tháng",
        "color": "#2563EB",
    },
    {
        "role": "👥 Ban Giám Sát (Đại Diện PHHS)",
        "freq": "Checklist bữa ăn: 2 lần/tuần tối thiểu · Kiểm tra NCC toàn diện: 1 lần/tháng",
        "when": "Bữa ăn: Thứ 2–3 (báo trước 1 ngày) + 1 lần đột xuất · NCC: Cuối tháng hoặc khi nhận hàng",
        "what": "Checklist 20 điểm + ảnh minh chứng · Đánh giá NCC 12 điểm toàn diện (giấy phép + giao hàng)",
        "notice": "Bữa ăn: 1 lần báo trước ≥ 24h, 1 lần đột xuất · NCC: Có thể báo trước 1 ngày",
        "report": "Gửi báo cáo Hiệu Trưởng trong 24h · Báo cáo NCC tháng ghi nhận xếp loại A/B/C",
        "color": "#7C3AED",
    },
    {
        "role": "🏫 Ban Giám Hiệu",
        "freq": "1 lần / tháng",
        "when": "Tuần cuối mỗi tháng — xem tổng hợp + spot-check",
        "what": "Duyệt báo cáo tháng + kiểm tra xác suất 5 mục ngẫu nhiên",
        "notice": "Không cần báo trước",
        "report": "Tổng hợp báo cáo học kỳ gửi Sở GD&ĐT",
        "color": "#0D9488",
    },
    {
        "role": "🏛️ Sở GD&ĐT / Sở Y Tế",
        "freq": "1–2 lần / học kỳ",
        "when": "Không cố định — kiểm tra đột xuất hoàn toàn",
        "what": "Kiểm tra toàn diện + lấy mẫu xét nghiệm tại chỗ",
        "notice": "Không báo trước — trường cần luôn sẵn sàng",
        "report": "Kết quả gửi Bộ GD&ĐT và Bộ Y Tế",
        "color": "#B45309",
    },
]

# ── Kiểm thực 3 bước — dành riêng cho Y Tế Học Đường (TTLT 13/2016 Điều 9) ──
KIEM_THUC = [
    {
        "buoc": 1, "icon": "🥩",
        "label": "Bước 1 — TRƯỚC chế biến",
        "time_window": "8:00 – 9:30",
        "law": "TTLT 13/2016 Điều 9 – Khoản a",
        "desc": "Kiểm tra toàn bộ nguyên liệu đầu vào trước khi đưa vào sơ chế và chế biến",
        "color": "#2563EB",
        "items": [
            ("B1_01", "Thịt/cá có tem kiểm dịch thú y còn hiệu lực",
             "Xem tem trên bao bì hoặc trên thân con vật",
             "Có tem, còn hiệu lực, đúng ngày", "Không có tem hoặc tem đã hết hạn",
             "QCVN 8-1:2011/BYT"),
            ("B1_02", "Rau củ có hóa đơn nguồn gốc từ vựa có đăng ký",
             "Xem hóa đơn mua hàng trong ngày — phải có tên vựa, địa chỉ",
             "Hóa đơn hợp lệ, đúng ngày, có địa chỉ vựa rõ ràng", "Không có hóa đơn hoặc hóa đơn mờ",
             "NĐ 15/2018 Điều 11"),
            ("B1_03", "Tất cả nguyên liệu đóng gói còn hạn sử dụng ≥ 3 ngày",
             "Kiểm tra từng bao bì — ghi nhận loại nào hết hạn sớm nhất",
             "Tất cả còn ≥ 3 ngày so với hôm nay", "Bất kỳ loại nào hết hạn hoặc không có ngày",
             "Luật ATTP 55/2010 Điều 10"),
            ("B1_04", "Nhiệt độ tủ lạnh bảo quản thực phẩm sống < 5°C",
             "Đọc nhiệt kế gắn trên tủ hoặc đo bằng nhiệt kế thực phẩm",
             "< 5°C", "≥ 5°C — vùng nguy hiểm vi khuẩn tăng nhanh",
             "QCVN 8-1:2011/BYT + WHO Five Keys"),
            ("B1_05", "Thực phẩm sống và chín bảo quản tách biệt hoàn toàn",
             "Quan sát các ngăn tủ lạnh và khu vực sơ chế",
             "Riêng biệt rõ ràng, có nhãn phân loại", "Để chung — nguy cơ nhiễm chéo vi khuẩn",
             "NĐ 15/2018 Điều 12"),
        ]
    },
    {
        "buoc": 2, "icon": "🍳",
        "label": "Bước 2 — TRONG chế biến",
        "time_window": "9:30 – 10:30",
        "law": "TTLT 13/2016 Điều 9 – Khoản b",
        "desc": "Kiểm tra quá trình nấu ăn, vệ sinh nhân viên và an toàn trong bếp",
        "color": "#7C3AED",
        "items": [
            ("B2_01", "Thức ăn được nấu chín kỹ, nhiệt độ lõi ≥ 70°C",
             "Dùng nhiệt kế thực phẩm đo tại phần dày nhất",
             "≥ 70°C — vi khuẩn (Salmonella, E.coli) bị tiêu diệt", "< 70°C — chưa đủ chín kỹ",
             "WHO Five Keys to Safer Food — Key 4"),
            ("B2_02", "Nhân viên đeo khẩu trang, găng tay và tạp dề sạch",
             "Quan sát toàn bộ nhân viên trong bếp trong quá trình nấu",
             "100% nhân viên đeo đầy đủ và đúng cách", "Thiếu bảo hộ — nguy cơ lây nhiễm từ người",
             "NĐ 15/2018 Điều 13"),
            ("B2_03", "Dao và thớt riêng biệt cho thực phẩm sống và chín",
             "Kiểm tra màu sắc hoặc ký hiệu phân loại thớt/dao",
             "Có phân loại rõ ràng (màu hoặc ký hiệu S/C)", "Dùng chung — nhiễm chéo vi khuẩn",
             "HACCP nguyên tắc 2 + NĐ 15/2018"),
            ("B2_04", "Dụng cụ nấu ăn sạch sẽ, không rỉ sét, không nứt vỡ",
             "Quan sát nồi, chảo, muỗng, vá, rây... trước khi dùng",
             "Sạch, nguyên vẹn, không rỉ sét", "Rỉ sét hoặc nứt vỡ — trú ẩn vi khuẩn",
             "NĐ 15/2018 Điều 13"),
            ("B2_05", "Khu vực bếp sạch sẽ, không có côn trùng",
             "Quan sát mặt bếp, sàn, góc tường, ống thoát nước",
             "Sạch, không có ruồi/gián/kiến xuất hiện", "Bẩn hoặc có côn trùng — vector lây E.coli",
             "NĐ 15/2018 Điều 11"),
        ]
    },
    {
        "buoc": 3, "icon": "🍱",
        "label": "Bước 3 — SAU chế biến",
        "time_window": "10:30 – 11:00",
        "law": "TTLT 13/2016 Điều 9 – Khoản c",
        "desc": "Kiểm tra thức ăn trước khi phục vụ và lấy mẫu lưu nghiệm bắt buộc",
        "color": "#0D9488",
        "items": [
            ("B3_01", "Nhiệt độ thức ăn khi chia đúng chuẩn an toàn",
             "Đo nhiệt kế trực tiếp vào thức ăn tại thời điểm bắt đầu chia",
             "Nóng ≥ 60°C | Lạnh ≤ 5°C — an toàn hoàn toàn", "5°C < T < 60°C — vùng NGUY HIỂM",
             "QCVN 8-1:2011/BYT + HACCP CCP"),
            ("B3_02", "Thời gian từ khi nấu xong đến khi phục vụ < 2 giờ",
             "Ghi lại giờ nấu xong (từ bếp trưởng) và giờ bắt đầu chia",
             "Dưới 2 giờ — ngưỡng an toàn tuyệt đối", "Trên 4 giờ ở nhiệt độ phòng — nguy hiểm",
             "WHO Five Keys — Key 2"),
            ("B3_03", "Màu sắc và mùi vị thức ăn bình thường, không có dấu hiệu hỏng",
             "Quan sát màu từng món, ngửi mùi trực tiếp trước khi chia",
             "Màu tự nhiên đặc trưng, mùi thơm ngon", "Màu lạ, mùi chua/hôi/đắng bất thường",
             "QCVN 8-1:2011/BYT"),
            ("B3_04", "Khẩu phần ăn đủ theo định mức đã đăng ký với nhà trường",
             "Ước lượng hoặc cân thực tế một suất so với định mức đã ký hợp đồng",
             "Đạt ≥ 90% định mức đăng ký", "Thiếu >20% — vi phạm hợp đồng cung cấp",
             "QĐ 3958/QĐ-BYT 2025"),
            ("B3_05", "Đã lấy mẫu lưu nghiệm 24h từng món — nhãn đầy đủ",
             "Kiểm tra tủ lạnh lưu mẫu: mỗi món ≥ 100g, nhãn ghi tên món + giờ lấy + ngày",
             "Có mẫu đủ lượng từng món, nhãn đầy đủ 3 thông tin", "Không có mẫu lưu hoặc thiếu nhãn",
             "TTLT 13/2016 Điều 9 — BẮT BUỘC PHÁP LUẬT"),
        ]
    },
]

# ── G4: Checklist Nhà Cung Cấp (12 điểm) ────────────────────────────────────
SUPPLIER_ITEMS = [
    {"code": "S01", "icon": "📄", "critical": True,
     "desc": "Giấy phép CSSX/KDDV thực phẩm còn hiệu lực",
     "hint": "Kiểm tra số giấy phép, ngày cấp, ngày hết hạn trên văn bản gốc",
     "pass_std": "Còn hiệu lực, đúng địa chỉ cơ sở", "fail_std": "Hết hạn hoặc không xuất trình được",
     "law": "Luật ATTP 55/2010 Điều 34"},
    {"code": "S02", "icon": "🏅", "critical": True,
     "desc": "Giấy chứng nhận cơ sở đủ điều kiện ATTP còn hiệu lực",
     "hint": "Chứng nhận 3 năm/lần cấp bởi Sở Y Tế hoặc Sở NN&PTNT",
     "pass_std": "Còn hiệu lực, không quá 3 năm kể từ ngày cấp", "fail_std": "Hết hạn hoặc không có",
     "law": "NĐ 15/2018 Điều 11"},
    {"code": "S03", "icon": "🚚", "critical": True,
     "desc": "Xe/thùng vận chuyển cách nhiệt, sạch sẽ, không côn trùng",
     "hint": "Quan sát trực tiếp xe hoặc thùng lạnh khi giao hàng",
     "pass_std": "Kín, sạch, không mùi hôi, không côn trùng", "fail_std": "Bẩn, nứt vỡ hoặc có côn trùng",
     "law": "NĐ 15/2018 Điều 18"},
    {"code": "S04", "icon": "🌡️", "critical": True,
     "desc": "Nhiệt độ vận chuyển thực phẩm chín ≥ 60°C hoặc lạnh < 8°C",
     "hint": "Dùng nhiệt kế đo tại điểm trung tâm thùng hàng khi nhận",
     "pass_std": "Nóng ≥ 60°C | Lạnh < 8°C", "fail_std": "8°C–60°C — vùng nguy hiểm vi khuẩn",
     "law": "QCVN 8-1:2011/BYT + WHO Five Keys"},
    {"code": "S05", "icon": "🧾", "critical": True,
     "desc": "Hóa đơn và chứng từ nguồn gốc thực phẩm cho lô hàng hôm nay",
     "hint": "Hóa đơn phải ghi đúng ngày, tên hàng, đơn vị cung cấp có địa chỉ rõ ràng",
     "pass_std": "Hóa đơn đủ hàng, đúng ngày, có địa chỉ nhà cung cấp",
     "fail_std": "Không có hóa đơn hoặc hóa đơn viết tay mờ",
     "law": "NĐ 15/2018 Điều 11"},
    {"code": "S06", "icon": "🏷️", "critical": False,
     "desc": "Nhãn mác thực phẩm đóng gói đủ: ngày sản xuất, hạn dùng, nơi sản xuất",
     "hint": "Kiểm tra mẫu ngẫu nhiên 3 gói/thùng hàng",
     "pass_std": "Đủ 3 thông tin bắt buộc, chữ rõ, không tẩy xóa",
     "fail_std": "Thiếu 1 trong 3 thông tin hoặc nhãn bị che/xóa",
     "law": "NĐ 43/2017/NĐ-CP về nhãn hàng hóa"},
    {"code": "S07", "icon": "📋", "critical": False,
     "desc": "Thực đơn giao khớp với đơn đặt hàng và cam kết dinh dưỡng",
     "hint": "So sánh phiếu giao hàng hôm nay với thực đơn tuần đã ký duyệt",
     "pass_std": "Khớp hoàn toàn hoặc sai lệch < 5% có báo trước",
     "fail_std": "Thiếu món hoặc thay thế không báo trước",
     "law": "QĐ 3958/QĐ-BYT 2025"},
    {"code": "S08", "icon": "⚖️", "critical": False,
     "desc": "Khẩu phần đủ định lượng: thịt/cá ≥ định mức ký hợp đồng",
     "hint": "Cân ngẫu nhiên 3 suất ăn, so với định mức đã ký",
     "pass_std": "Đạt ≥ 90% định mức", "fail_std": "Thiếu > 10% định mức",
     "law": "QĐ 3958/QĐ-BYT 2025 + hợp đồng cung cấp"},
    {"code": "S09", "icon": "👷", "critical": False,
     "desc": "Nhân viên giao hàng đeo khẩu trang, găng tay, đồng phục sạch",
     "hint": "Quan sát tất cả người tiếp xúc trực tiếp với thực phẩm",
     "pass_std": "100% nhân viên đeo đầy đủ BHLĐ", "fail_std": "Bất kỳ ai thiếu BHLĐ",
     "law": "NĐ 15/2018 Điều 13"},
    {"code": "S10", "icon": "📦", "critical": False,
     "desc": "Dụng cụ đựng thực phẩm kín, sạch, không nứt vỡ",
     "hint": "Kiểm tra khay, hộp, nồi đựng khi giao",
     "pass_std": "Kín, sạch, nguyên vẹn", "fail_std": "Nứt vỡ hoặc bẩn",
     "law": "NĐ 15/2018 Điều 13"},
    {"code": "S11", "icon": "🧫", "critical": True,
     "desc": "Mẫu lưu thực phẩm được giao đúng quy định (≥ 3 mẫu, ≥ 100g/mẫu, nhãn đủ 3 thông tin)",
     "hint": "Kiểm tra tủ lạnh mẫu lưu: tên món + giờ lấy + ngày",
     "pass_std": "Đủ số lượng, đủ trọng lượng, nhãn đầy đủ",
     "fail_std": "Thiếu mẫu, thiếu lượng hoặc thiếu nhãn",
     "law": "TTLT 13/2016 Điều 9 — BẮT BUỘC PHÁP LUẬT"},
    {"code": "S12", "icon": "⏰", "critical": False,
     "desc": "Thời gian giao hàng đúng lịch (trước bữa ăn tối thiểu 30 phút)",
     "hint": "Ghi lại giờ xe đến cổng trường vs giờ bắt đầu phục vụ",
     "pass_std": "Đến đúng giờ hoặc sớm hơn ≥ 30 phút",
     "fail_std": "Trễ > 15 phút so với lịch hẹn",
     "law": "Điều khoản hợp đồng cung cấp"},
]
SUPPLIER_CRITICAL = {"S01", "S02", "S03", "S04", "S05", "S11"}
SUPPLIER_SCORE_PASS = 10   # ≥ 10/12 + critical OK → Loại A
SUPPLIER_SCORE_WARN = 8    # 8–9/12 → Loại B

# ── CSS ────────────────────────────────────────────────────────────────────────
def inject_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

    /* Xoá hoàn toàn khoảng trắng trên cùng */
    #MainMenu, footer, header { display: none !important; }
    [data-testid="stHeader"]  { display: none !important; }
    [data-testid="stToolbar"] { display: none !important; }
    .main .block-container {
        padding-top: 0.6rem !important;
        padding-bottom: 1.5rem !important;
        max-width: 1200px !important;
    }

    .stApp { background-color: #F0F4F8; }

    /* Ẩn hoàn toàn nút đóng sidebar — người dùng không thể đóng sidebar trên desktop */
    @media (min-width: 769px) {
        [data-testid="stSidebarCollapseButton"],
        button[title="Close sidebar"],
        button[aria-label="Close sidebar"],
        [data-testid="stSidebar"] > div:first-child > div > button,
        section[data-testid="stSidebar"] > div > div > div > button { display: none !important; }
    }
    /* Nút MỞ LẠI (nếu vẫn lỡ đóng trên mobile) — đỏ nổi bật */
    [data-testid="collapsedControl"],
    button[title="Open sidebar"],
    button[aria-label="Open sidebar"] {
        background: #DC2626 !important; border-radius: 0 14px 14px 0 !important;
        display: flex !important; visibility: visible !important;
        align-items: center !important; justify-content: center !important;
        width: 44px !important; min-height: 100px !important;
        opacity: 1 !important; z-index: 999999 !important;
        position: fixed !important; left: 0 !important; top: calc(50% - 50px) !important;
        box-shadow: 4px 0 20px rgba(220,38,38,0.5) !important;
        cursor: pointer !important; border: none !important;
    }
    [data-testid="collapsedControl"] svg,
    button[title="Open sidebar"] svg { fill: white !important; color: white !important; }

    /* sf-header không còn dùng — đã thay bằng HTML inline trong main() */
    .sf-header { display: none; }

    .sf-card {
        background: white; border-radius: 12px; padding: 18px 22px;
        margin-bottom: 12px; box-shadow: 0 1px 4px rgba(0,0,0,0.07);
        border: 1px solid #E8ECF0;
    }
    .sf-card-title { font-size: 1rem; font-weight: 600; color: #1E293B; margin-bottom: 5px; }
    .sf-card-body  { font-size: 0.875rem; color: #475569; line-height: 1.65; }

    /* Checklist item row */
    .cl-row {
        display: flex; align-items: flex-start; gap: 10px;
        padding: 10px 0; border-bottom: 1px solid #F1F5F9;
    }
    .cl-code {
        font-size: 0.72rem; font-weight: 700; color: #94A3B8;
        min-width: 32px; padding-top: 3px; letter-spacing: 0.03em;
    }
    .cl-desc {
        font-size: 0.9rem; font-weight: 500; color: #1E293B; flex: 1; line-height: 1.5;
    }
    .cl-sub {
        font-size: 0.78rem; color: #64748B; font-weight: 400; margin-top: 2px;
    }
    .badge-critical {
        display: inline-block; background: #FEE2E2; color: #991B1B;
        font-size: 0.68rem; font-weight: 700; padding: 2px 7px;
        border-radius: 10px; margin-left: 6px; vertical-align: middle;
        letter-spacing: 0.04em; border: 1px solid #FECACA;
    }
    .badge-pass {
        background: #DCFCE7; color: #166534;
        padding: 3px 10px; border-radius: 20px;
        font-size: 0.78rem; font-weight: 600; border: 1px solid #BBF7D0;
    }
    .badge-fail {
        background: #FEE2E2; color: #991B1B;
        padding: 3px 10px; border-radius: 20px;
        font-size: 0.78rem; font-weight: 600; border: 1px solid #FECACA;
    }

    /* Group title */
    .group-title {
        font-size: 0.78rem; font-weight: 700; color: #475569;
        text-transform: uppercase; letter-spacing: 0.08em;
        padding: 14px 0 8px 0; border-bottom: 2px solid #E2E8F0;
        margin-bottom: 4px;
    }

    /* Metric */
    .metric-box {
        background: white; border-radius: 10px; padding: 14px 10px;
        text-align: center; border: 1px solid #E8ECF0;
        min-height: 110px;
        display: flex; flex-direction: column;
        justify-content: center; align-items: center; gap: 2px;
    }
    .metric-num  { font-size: 2rem; font-weight: 700; line-height: 1.1; }
    .metric-lbl  { font-size: 0.78rem; color: #64748B; line-height: 1.3; }
    .c-green  { color: #16A34A; } .c-red  { color: #DC2626; }
    .c-blue   { color: #2563EB; } .c-orange{ color: #D97706; }

    /* Alert card */
    .alert-critical { background:#FEF2F2; border:2px solid #FCA5A5; border-radius:12px; padding:16px 20px; }
    .alert-major    { background:#FFFBEB; border:2px solid #FCD34D; border-radius:12px; padding:16px 20px; }
    .alert-minor    { background:#FEFCE8; border:2px solid #FDE68A; border-radius:12px; padding:16px 20px; }
    .alert-ok       { background:#F0FDF4; border:2px solid #86EFAC; border-radius:12px; padding:16px 20px; }
    .alert-title  { font-size: 1rem; font-weight: 700; margin-bottom: 8px; }
    .alert-body   { font-size: 0.85rem; line-height: 1.7; }

    /* Nutrition banner */
    .nutrition-banner {
        background: #EFF6FF; border: 1px solid #BFDBFE; border-radius: 10px;
        padding: 12px 16px; margin-bottom: 14px;
    }
    .nutrition-label { font-size: 0.78rem; color: #1D4ED8; font-weight: 700;
        text-transform: uppercase; letter-spacing: 0.06em; }
    .nutrition-grid { display: flex; gap: 20px; flex-wrap: wrap; margin-top: 6px; }
    .nutrition-item { font-size: 0.82rem; color: #1E40AF; }
    .nutrition-val  { font-weight: 700; }

    /* Schedule card */
    .schedule-card {
        background: white; border-radius: 10px; padding: 16px 20px;
        margin-bottom: 10px; border: 1px solid #E2E8F0;
        border-left: 4px solid var(--sc-color, #64748B);
    }
    .schedule-role  { font-size: 0.95rem; font-weight: 700; color: #1E293B; margin-bottom: 6px; }
    .schedule-row   { font-size: 0.83rem; color: #475569; margin: 3px 0; }
    .schedule-key   { font-weight: 600; color: #334155; min-width: 100px; display: inline-block; }

    /* Emergency */
    .emergency-header { background:#DC2626; color:white; border-radius:10px;
        padding:14px 18px; margin-bottom:16px; font-weight:700; font-size:1rem; }

    /* Divider */
    .sf-div { border-top: 1px solid #E2E8F0; margin: 16px 0; }

    /* Section header */
    .sec-hdr { font-size:0.72rem; font-weight:800; color:#334155;
        text-transform:uppercase; letter-spacing:0.1em; margin:16px 0 8px; }

    .stButton > button { border-radius: 8px !important; font-family:'Inter',sans-serif !important; font-weight:500 !important; }
    .stTextInput > div > div > input { border-radius: 8px !important; }
    div[data-testid="stChatInput"] textarea { font-family:'Inter',sans-serif !important; }

    /* Input disabled — override màu xám mặc định, giữ chữ đọc được */
    input:disabled, input[disabled] {
        -webkit-text-fill-color: #16A34A !important;
        color: #16A34A !important;
        opacity: 1 !important;
        font-weight: 600 !important;
        background-color: #F0FDF4 !important;
        border-color: #86EFAC !important;
        cursor: default !important;
    }

    /* ── Ẩn nút đóng sidebar trên desktop (phòng người dùng vô tình đóng) ── */
    @media (min-width: 769px) {
        [data-testid="stSidebarCollapseButton"],
        button[title="Close sidebar"],
        button[aria-label="Close sidebar"],
        [data-testid="stSidebar"] > div:first-child > div > button {
            display: none !important;
        }
    }

    /* ── Nút MỞ LẠI sidebar khi lỡ đóng — đỏ nổi bật, không thể bỏ qua ── */
    [data-testid="collapsedControl"],
    button[title="Open sidebar"],
    button[aria-label="Open sidebar"] {
        background: #DC2626 !important;
        border-radius: 0 14px 14px 0 !important;
        display: flex !important;
        visibility: visible !important;
        align-items: center !important;
        justify-content: center !important;
        width: 44px !important;
        min-height: 100px !important;
        opacity: 1 !important;
        z-index: 999999 !important;
        position: fixed !important;
        left: 0 !important;
        top: calc(50% - 50px) !important;
        box-shadow: 4px 0 16px rgba(220,38,38,0.4) !important;
        cursor: pointer !important;
        border: none !important;
        transition: width 0.2s ease !important;
    }
    [data-testid="collapsedControl"]::after,
    button[title="Open sidebar"]::after,
    button[aria-label="Open sidebar"]::after {
        content: "☰" !important;
        color: white !important;
        font-size: 1.3rem !important;
        font-weight: bold !important;
    }
    [data-testid="collapsedControl"] svg,
    button[title="Open sidebar"] svg,
    button[aria-label="Open sidebar"] svg {
        fill: white !important;
        color: white !important;
        opacity: 1 !important;
    }
    [data-testid="collapsedControl"]:hover,
    button[title="Open sidebar"]:hover,
    button[aria-label="Open sidebar"]:hover {
        background: #B91C1C !important;
        width: 52px !important;
    }

    /* ── Tabs spacing — giãn cách giữa các tab ── */
    .stTabs [data-baseweb="tab-list"] {
        gap: 6px !important;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 8px 18px !important;
        border-radius: 8px 8px 0 0 !important;
        font-size: 0.875rem !important;
        font-weight: 500 !important;
        letter-spacing: 0.01em !important;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        font-weight: 700 !important;
    }

    /* ── Mobile responsive ── */
    @media (max-width: 768px) {
        /* ── FIX MOBILE KEYBOARD: iOS tự zoom khi font < 16px → mất khoảng cách ── */
        input[type="text"],
        input[type="search"],
        input[type="email"],
        input[type="password"],
        input[type="number"],
        textarea,
        select {
            font-size: 16px !important;          /* Ngăn iOS Safari auto-zoom */
            -webkit-text-size-adjust: 100% !important;
        }
        /* Streamlit specific inputs */
        .stTextInput input,
        .stTextArea textarea,
        .stSelectbox select,
        [data-testid="stTextInput"] input,
        [data-testid="stTextArea"] textarea {
            font-size: 16px !important;
            touch-action: manipulation !important; /* Ngăn double-tap delay */
        }

        /* Header mới (inline HTML) — thu nhỏ trên mobile */
        .main .block-container { padding-left: 0.8rem !important; padding-right: 0.8rem !important; }

        /* Card padding giảm */
        .sf-card { padding: 14px 14px !important; }

        /* Metric boxes nhỏ hơn, font nhỏ hơn */
        .metric-box { min-height: 80px !important; padding: 10px 8px !important; }
        .metric-num { font-size: 1.5rem !important; }
        .metric-lbl { font-size: 0.7rem !important; }

        /* Nút bấm full width trên mobile */
        .stButton > button { min-height: 44px !important; font-size: 0.85rem !important; }

        /* Columns stack dọc trên mobile */
        [data-testid="stHorizontalBlock"] {
            flex-wrap: wrap !important;
        }
        [data-testid="column"] {
            min-width: 100% !important;
            flex: 1 1 100% !important;
        }

        /* Segmented control mobile */
        [data-baseweb="button-group"] {
            flex-direction: row !important;
            gap: 3px !important;
        }
        [data-baseweb="button-group"] button {
            font-size: 0.72rem !important;
            padding: 5px 6px !important;
        }

        /* Tab labels ngắn hơn */
        .stTabs [data-baseweb="tab"] {
            font-size: 0.75rem !important;
            padding: 8px 8px !important;
        }

        /* Alert boxes padding */
        .alert-critical, .alert-major, .alert-minor, .alert-ok {
            padding: 12px 14px !important;
        }

        /* Group title font nhỏ hơn */
        .group-title { font-size: 0.72rem !important; }

        /* Schedule cards */
        .schedule-card { padding: 12px 14px !important; }
    }

    @media (max-width: 480px) {
        .metric-num { font-size: 1.2rem !important; }
    }

    /* ── G6: Mobile tab navigation — cuộn ngang thay vì xuống dòng ── */
    @media (max-width: 768px) {
        /* Tab bar cuộn ngang, không xuống dòng */
        .stTabs [data-baseweb="tab-list"] {
            overflow-x: auto !important;
            overflow-y: hidden !important;
            flex-wrap: nowrap !important;
            -webkit-overflow-scrolling: touch !important;
            scrollbar-width: none !important;
            padding-bottom: 2px !important;
        }
        .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar { display: none !important; }
        .stTabs [data-baseweb="tab"] {
            white-space: nowrap !important;
            flex-shrink: 0 !important;
            padding: 8px 12px !important;
            font-size: 0.78rem !important;
        }
        /* Safe area cho iPhone có notch */
        .main .block-container {
            padding-bottom: max(1.5rem, env(safe-area-inset-bottom)) !important;
        }
        /* Nút bấm tối thiểu 44px (WCAG AA touch target) */
        .stButton > button {
            min-height: 44px !important;
        }
        /* Giảm padding header trên mobile */
        [data-testid="stMarkdownContainer"] div[style*="border-radius:16px"] {
            padding: 18px 16px !important;
        }
    }
    @media (max-width: 480px) {
        .stTabs [data-baseweb="tab"] {
            font-size: 0.72rem !important;
            padding: 6px 10px !important;
        }
    }

    /* ── Checklist segmented control animations (CSS-only, no JS needed) ── */
    /* Áp dụng cho cả direct children và div-wrapped children (tuỳ BaseWeb version) */
    [data-baseweb="button-group"] button {
        border-radius: 8px !important;
        font-size: 0.78rem !important;
        font-weight: 600 !important;
        padding: 5px 12px !important;
        transition: transform 0.15s ease, box-shadow 0.2s ease, background 0.2s ease !important;
        letter-spacing: 0.02em !important;
        cursor: pointer !important;
    }
    [data-baseweb="button-group"] button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 3px 8px rgba(0,0,0,0.12) !important;
    }

    /* Nút 1 — "Chưa chấm": viền nháy, text luôn đọc được */
    [data-baseweb="button-group"] > button:nth-child(1):not([aria-pressed="true"]):not([aria-checked="true"]),
    [data-baseweb="button-group"] > div:nth-child(1) > button:not([aria-pressed="true"]):not([aria-checked="true"]) {
        background: #F8FAFC !important;
        color: #374151 !important;          /* Đủ tối để đọc rõ */
        border: 1.5px dashed #CBD5E1 !important;
        animation: pendingBorder 2.5s ease-in-out infinite !important;
    }
    @keyframes pendingBorder {
        0%, 100% { border-color: #CBD5E1; background: #F8FAFC; }
        50%       { border-color: #F59E0B; background: #FFFBEB; }
        /* Viền chuyển từ xám → vàng hổ phách — "cần chú ý" */
    }

    /* Nút 2 — "✅ Đạt": xanh lá + pop animation khi chọn */
    [data-baseweb="button-group"] > button:nth-child(2):not([aria-pressed="true"]):not([aria-checked="true"]),
    [data-baseweb="button-group"] > div:nth-child(2) > button:not([aria-pressed="true"]):not([aria-checked="true"]) {
        background: #F0FDF4 !important; color: #16A34A !important;
        border: 1.5px solid #BBF7D0 !important;
    }
    [data-baseweb="button-group"] > button:nth-child(2)[aria-pressed="true"],
    [data-baseweb="button-group"] > button:nth-child(2)[aria-checked="true"],
    [data-baseweb="button-group"] > div:nth-child(2) > button[aria-pressed="true"],
    [data-baseweb="button-group"] > div:nth-child(2) > button[aria-checked="true"] {
        background: #16A34A !important; color: #FFFFFF !important;
        border: 1.5px solid #16A34A !important;
        box-shadow: 0 0 0 3px rgba(22,163,74,0.25), 0 4px 14px rgba(22,163,74,0.35) !important;
        animation: passClick 0.38s cubic-bezier(0.34, 1.56, 0.64, 1) forwards !important;
    }
    @keyframes passClick {
        0%   { transform: scale(0.88); }
        55%  { transform: scale(1.10); }
        100% { transform: scale(1.00); }
    }

    /* Nút 3 — "❌ Không Đạt": đỏ + shake animation khi chọn */
    [data-baseweb="button-group"] > button:nth-child(3):not([aria-pressed="true"]):not([aria-checked="true"]),
    [data-baseweb="button-group"] > div:nth-child(3) > button:not([aria-pressed="true"]):not([aria-checked="true"]) {
        background: #FFF1F2 !important; color: #DC2626 !important;
        border: 1.5px solid #FECACA !important;
    }
    [data-baseweb="button-group"] > button:nth-child(3)[aria-pressed="true"],
    [data-baseweb="button-group"] > button:nth-child(3)[aria-checked="true"],
    [data-baseweb="button-group"] > div:nth-child(3) > button[aria-pressed="true"],
    [data-baseweb="button-group"] > div:nth-child(3) > button[aria-checked="true"] {
        background: #DC2626 !important; color: #FFFFFF !important;
        border: 1.5px solid #DC2626 !important;
        box-shadow: 0 0 0 3px rgba(220,38,38,0.25), 0 4px 14px rgba(220,38,38,0.35) !important;
        animation: failShake 0.48s ease-out forwards !important;
    }
    @keyframes failShake {
        0%,100% { transform: translateX(0) scale(1);    }
        18%     { transform: translateX(-5px) scale(0.97); }
        36%     { transform: translateX(5px)  scale(0.97); }
        54%     { transform: translateX(-4px); }
        72%     { transform: translateX(4px);  }
        90%     { transform: translateX(-2px); }
    }

    /* Reminder banner */
    .reminder-banner {
        background: linear-gradient(135deg,#FEF9C3,#FFFBEB);
        border: 2px solid #F59E0B; border-radius: 12px;
        padding: 14px 20px; margin-bottom: 16px;
    }
    .reminder-title { font-weight: 700; color: #92400E; font-size: 0.95rem; }
    .reminder-body  { font-size: 0.83rem; color: #78350F; margin-top: 5px; line-height: 1.6; }
    .reminder-countdown {
        display: inline-block; background: #F59E0B; color: white;
        border-radius: 20px; padding: 2px 12px; font-weight: 700;
        font-size: 0.8rem; margin-left: 8px;
    }

    /* Validation error box */
    .validation-box {
        background: #FFF7ED; border: 2px solid #FB923C;
        border-radius: 10px; padding: 14px 18px; margin-bottom: 12px;
    }
    .validation-title { font-weight: 700; color: #9A3412; font-size: 0.9rem; margin-bottom: 8px; }
    .validation-item  { font-size: 0.83rem; color: #7C2D12; margin: 4px 0; }

    /* Progress bar completion */
    .completion-bar-wrap {
        background: #E2E8F0; border-radius: 20px; height: 8px; margin: 8px 0;
        overflow: hidden;
    }
    .completion-bar-fill {
        height: 100%; border-radius: 20px;
        transition: width 0.4s ease;
    }
    </style>
    """, unsafe_allow_html=True)


# ── Checklist theo cấp học ────────────────────────────────────────────────────
def get_checklist(level_key: str) -> list:
    n = NUTRITION.get(level_key, NUTRITION["Tiểu Học (6–11 tuổi)"])
    mg, vg, lvl = n["meat_g"], n["veg_g"], n["short"]
    fail_m, fail_v = int(mg * 0.7), int(vg * 0.7)
    return [
        ("📦 Nguồn gốc nguyên liệu", [
            ("C01", False, "Thịt/cá có tem kiểm dịch thú y còn hiệu lực",
             "Xem nhãn mác trên bao bì từng lô hàng",
             "Có tem còn hiệu lực", "Không có tem, hoặc tem đã hết hạn"),
            ("C02", False, "Rau củ có hóa đơn nguồn gốc từ vựa có đăng ký",
             "Yêu cầu xem hóa đơn mua hàng trong ngày",
             "Có hóa đơn hợp lệ đúng ngày", "Mua chợ không rõ nguồn gốc, không hóa đơn"),
            ("C03", True,  "Nguyên liệu đóng gói còn hạn sử dụng ≥ 3 ngày",
             "Kiểm tra date trên bao bì từng loại nguyên liệu",
             "Còn ≥ 3 ngày so với hôm nay", "Hết hạn, hoặc không có ngày sản xuất/hạn dùng"),
            ("C04", False, "Có hóa đơn mua hàng của ngày hôm nay đầy đủ",
             "Yêu cầu xem toàn bộ hóa đơn mua hàng trong ngày",
             "Có hóa đơn hợp lệ, đúng ngày, đủ loại", "Thiếu hóa đơn hoặc không xuất trình được"),
        ]),
        ("🌡️ Bảo quản & vận chuyển", [
            ("C05", False, "Nhiệt độ tủ lạnh thực phẩm sống dưới 5°C",
             "Đọc đồng hồ hoặc đo nhiệt kế tủ trực tiếp",
             "Dưới 5°C", "Từ 5°C trở lên — vùng nguy hiểm, vi khuẩn tăng gấp đôi mỗi 20 phút"),
            ("C06", False, "Thực phẩm sống và chín để riêng, có nhãn phân biệt rõ ràng",
             "Quan sát kho lạnh và tủ lạnh, kiểm tra nhãn",
             "Có ngăn riêng biệt, nhãn ghi rõ loại thực phẩm", "Để chung hoặc không có nhãn phân biệt"),
            ("C07", True,  "Nhiệt độ thức ăn khi nhận tại trường đạt từ 60°C trở lên",
             "Đo nhiệt kế thực phẩm trực tiếp vào từng nồi/khay khi nhận",
             "Từ 60°C trở lên — an toàn", "Dưới 60°C — đã nguội, nguy cơ nhiễm khuẩn cao"),
            ("C08", False, "Thùng/nồi vận chuyển kín nắp, sạch, không có mùi lạ",
             "Quan sát trực tiếp từng thùng chứa khi bàn giao",
             "Kín nắp, bề mặt sạch, không mùi lạ", "Hở nắp, bẩn, hoặc có mùi khó chịu"),
        ]),
        ("🍽️ Thức ăn khi phục vụ", [
            ("C09", True,  "Nhiệt độ thức ăn khi chia đúng chuẩn an toàn",
             "Đo nhiệt kế thực phẩm trực tiếp tại thời điểm bắt đầu chia",
             "Nóng ≥ 60°C · Lạnh ≤ 5°C", "5°C < T < 60°C — VÙNG NGUY HIỂM tuyệt đối"),
            ("C10", True,  "Thời gian từ khi nấu xong đến khi phục vụ dưới 2 giờ",
             "Hỏi bếp trưởng và đối chiếu nhật ký giờ nấu",
             "Dưới 2 giờ — an toàn hoàn toàn", "Trên 4 giờ ở nhiệt độ phòng — nguy hiểm"),
            ("C11", True,  "Màu sắc và mùi vị thức ăn bình thường, không có dấu hiệu hỏng",
             "Quan sát màu, ngửi mùi trực tiếp từng món; hỏi ý kiến bếp trưởng",
             "Màu tự nhiên đặc trưng, mùi thơm ngon đúng món", "Màu lạ, mùi chua/hôi, có nấm mốc/nhớt"),
            ("C12", False, f"Khẩu phần thịt/cá đạt tiêu chuẩn cấp {lvl}: từ {mg}g/học sinh",
             f"Cân hoặc ước lượng suất ăn thực tế, so với định mức {mg}g đã đăng ký",
             f"Từ {mg}g/học sinh trở lên (cấp {lvl})", f"Dưới {fail_m}g — thiếu hụt dưới 70% định mức"),
            ("C13", False, f"Rau xanh đủ khẩu phần cấp {lvl}: từ {vg}g/học sinh",
             f"Ước lượng lượng rau trong suất ăn thực tế, tiêu chuẩn {n['veg_range']}",
             f"Từ {vg}g/học sinh (cấp {lvl})", f"Dưới {fail_v}g — thiếu rau nghiêm trọng"),
        ]),
        ("🧼 Vệ sinh dụng cụ & nhân viên", [
            ("C14", False, "Bát đũa muỗng sạch, khô ráo, không còn vết thức ăn cũ",
             "Kiểm tra ngẫu nhiên 5–10 bộ dụng cụ",
             "Sạch, khô, không mùi", "Còn thức ăn cũ, ẩm ướt, hoặc có mùi hôi"),
            ("C15", False, "Nhân viên đeo khẩu trang và găng tay đúng cách khi chia cơm",
             "Quan sát trực tiếp toàn bộ nhân viên tham gia chia",
             "Tất cả đều đeo đầy đủ và đúng cách", "Có người không đeo hoặc đeo sai, dùng lại găng cũ"),
            ("C16", False, "Nhân viên không ho/hắt hơi trực tiếp vào thức ăn",
             "Quan sát liên tục trong suốt 15 phút chia cơm",
             "Quay mặt đi hoặc che kín miệng khi cần", "Ho/hắt hơi thẳng vào thức ăn — không che chắn"),
            ("C17", False, "Khu vực chia cơm gọn sạch, không có côn trùng",
             "Quan sát mặt bàn, sàn, cửa sổ và các góc xung quanh khu vực chia",
             "Sạch sẽ, không có ruồi/gián/kiến", "Bẩn, hoặc có côn trùng xuất hiện"),
        ]),
        ("📋 Hồ sơ & giấy tờ", [
            ("C18", True,  "Sổ kiểm thực 3 bước điền đầy đủ hôm nay, có chữ ký xác nhận",
             "Yêu cầu xem sổ tại bếp — đây là tài liệu bắt buộc theo pháp luật",
             "Ghi đủ 3 bước (trước/trong/sau chế biến), có chữ ký Y Tế", "Chưa điền, thiếu bước, hoặc không có chữ ký"),
            ("C19", False, "Thực đơn thực tế phục vụ khớp với thực đơn đã đăng ký trước",
             "So sánh menu treo tại bếp với các món đang thực tế phục vụ",
             "Khớp hoàn toàn với đăng ký", "Thay đổi món mà không thông báo trước"),
            ("C20", True,  "Có mẫu lưu thức ăn 24h từng món, đủ nhãn ngày giờ",
             "Yêu cầu xem tủ lạnh lưu mẫu — mỗi món cần ≥ 100g, nhãn đầy đủ",
             "Có mẫu từng món, nhãn ghi: tên món + giờ lấy mẫu + ngày", "Không có mẫu lưu, thiếu nhãn, hoặc mẫu không đủ lượng"),
        ]),
    ]

TOTAL_ITEMS = 20


# ── Helpers ────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=3600, show_spinner=False)
def load_legal_pdfs() -> str:
    try:
        import pypdf
    except ImportError:
        return ""
    parts = []
    for p in sorted(LEGAL_DIR.glob("*.pdf")):
        try:
            r = pypdf.PdfReader(str(p))
            txt = "\n".join(pg.extract_text() or "" for pg in r.pages).strip()
            if txt:
                parts.append(f"=== {p.name} ===\n{txt}")
        except Exception:
            pass
    return "\n\n".join(parts)


@st.cache_data(ttl=3600, show_spinner=False)
def build_system_prompt(role: str, level: str, loc: str) -> str:
    if PROMPT_FILE.exists():
        raw = PROMPT_FILE.read_text(encoding="utf-8")
        s = raw.find("```\n") + 4; e = raw.find("\n```", s)
        base = raw[s:e] if e != -1 and s > 3 else raw
    else:
        base = "Bạn là chuyên gia ATVSTP trường học tại Việt Nam. Trả lời đơn giản, trích dẫn pháp luật."
    base = base.replace("{user_role}", role).replace("{school_level}", level).replace("{location}", loc)
    pdfs = load_legal_pdfs()
    if pdfs:
        base += "\n\n=== VĂN BẢN PHÁP LUẬT ĐÍNH KÈM ===\nƯu tiên trích dẫn từ các văn bản sau:\n\n" + pdfs
    return base


def ask_claude(client, system: str, history: list, user_input: str) -> str:
    try:
        r = client.messages.create(
            model=MODEL, max_tokens=MAX_TOK,
            system=[{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}],
            messages=history + [{"role": "user", "content": user_input}],
        )
        return r.content[0].text
    except anthropic.AuthenticationError:
        return "❌ API Key không hợp lệ. Kiểm tra lại ở thanh cài đặt phía trên."
    except anthropic.RateLimitError:
        return "⚠️ Vượt giới hạn API. Thử lại sau vài giây."
    except Exception as e:
        return f"❌ Lỗi: {e}"


def determine_alert(results: dict, cl: list) -> str:
    """Xác định mức cảnh báo dựa trên kết quả checklist."""
    pass_set  = {c for c, v in results.items() if v == "✅ Đạt"}
    fail_set  = {c for c, v in results.items() if v == "❌ Không Đạt"}
    critical_fails = CRITICAL_ITEMS & fail_set
    pass_count = len(pass_set)

    if critical_fails:
        return "CRITICAL"
    if pass_count < SCORE_ACCEPTABLE:
        return "MAJOR"
    # Kiểm tra ≥3 FAIL cùng nhóm
    for _, items in cl:
        codes = {code for code, *_ in items}
        if len(codes & fail_set) >= 3:
            return "MAJOR"
    if pass_count < SCORE_EXCELLENT:
        return "MINOR"
    return "OK"


# ── Hệ thống nhắc nhở kiểm tra theo vai trò ───────────────────────────────────
_REMINDER_TIMES = {
    "Y Tế Học Đường": {
        "hour": 10, "min": 0,
        "days": [0, 1, 2, 3, 4],          # Thứ 2–6
        "task": "Kiểm thực 3 bước bữa trưa (trước 10:45)",
        "tab":  "✅ Checklist kiểm tra",
    },
    "Ban Giám Sát (Đại Diện PHHS)": {
        "hour": 10, "min": 0,
        "days": [0, 2],                    # Thứ 2, Thứ 4 (định kỳ có báo trước)
        "task": "Checklist 20 điểm định kỳ (Thứ 2 & Thứ 4)",
        "tab":  "✅ Checklist kiểm tra",
    },
    "Ban Giám Hiệu": {
        "hour": 9, "min": 0,
        "days": [0, 1, 2, 3, 4],
        "task": "Xem tổng hợp báo cáo tuần & duyệt checklist",
        "tab":  "📅 Lịch & thông báo",
        "last_week_only": True,
    },
}


def get_inspection_reminder(role: str) -> dict | None:
    """Trả về thông tin nhắc nhở nếu còn ≤ 15 phút đến giờ kiểm tra."""
    import calendar
    now = now_vn()
    if now.weekday() >= 5:          # Thứ 7, Chủ nhật — không có bữa bán trú
        return None

    info = _REMINDER_TIMES.get(role)
    if not info or now.weekday() not in info["days"]:
        return None

    if info.get("last_week_only"):
        last_day = calendar.monthrange(now.year, now.month)[1]
        if now.day < last_day - 6:  # Không phải tuần cuối tháng
            return None

    inspect_mins  = info["hour"] * 60 + info["min"]
    current_mins  = now.hour * 60 + now.minute
    mins_left     = inspect_mins - current_mins

    if 0 <= mins_left <= 15:
        return {
            "task":      info["task"],
            "time":      f"{info['hour']:02d}:{info['min']:02d}",
            "mins_left": mins_left,
            "tab":       info["tab"],
        }
    return None


def show_reminder_banner(role: str):
    """Hiển thị banner nhắc nhở và toast nếu trong vòng 15 phút trước giờ kiểm tra."""
    r = get_inspection_reminder(role)
    if not r:
        return
    key = f"reminded_{role}_{now_vn().strftime('%Y%m%d%H%M')[:12]}"
    if not st.session_state.get(key):
        st.toast(f"⏰ Còn {r['mins_left']} phút đến giờ kiểm tra!", icon="🔔")
        st.session_state[key] = True
    st.markdown(f"""
    <div class="reminder-banner">
        <div class="reminder-title">
            ⏰ NHẮC KIỂM TRA <span class="reminder-countdown">Còn {r['mins_left']} phút</span>
        </div>
        <div class="reminder-body">
            <b>Nhiệm vụ:</b> {r['task']}<br>
            <b>Giờ thực hiện:</b> {r['time']} · <b>Chuyển đến tab:</b> {r['tab']}
        </div>
    </div>
    """, unsafe_allow_html=True)


# ── Validation checklist trước khi xuất báo cáo ────────────────────────────────
def validate_checklist(results: dict, photos: dict) -> list[str]:
    """Trả về danh sách lỗi. Rỗng = hợp lệ, được phép xuất báo cáo."""
    errors = []

    unanswered = sorted(c for c, v in results.items() if v is None)
    if unanswered:
        errors.append(
            f"Còn {len(unanswered)} mục chưa được đánh giá: "
            + ", ".join(unanswered)
        )

    has_fail = any(v == "❌ Không Đạt" for v in results.values())
    total_photos = sum(
        (len(v) if isinstance(v, list) else 1)
        for v in photos.values() if v
    )
    if has_fail and total_photos == 0:
        errors.append(
            "Có mục KHÔNG ĐẠT — bắt buộc cung cấp ít nhất 1 ảnh minh chứng "
            "(mở phần '📷 Ảnh minh chứng' bên dưới mỗi nhóm để chụp/tải ảnh)"
        )

    return errors


# ── Cơ sở pháp lý dùng cho tất cả AI analysis ───────────────────────────────
_LEGAL_BASIS = """
CĂN CỨ PHÁP LÝ VÀ TIÊU CHUẨN KỸ THUẬT ÁP DỤNG:
1. NĐ 15/2018/NĐ-CP — Quy định điều kiện đảm bảo ATTP, yêu cầu bếp ăn tập thể
2. TTLT 13/2016/TTLT-BYT-BGDĐT — Công tác y tế trường học, kiểm soát bữa ăn học đường
3. QĐ 3958/QĐ-BYT ngày 25/12/2025 — Hướng dẫn dinh dưỡng bữa ăn học đường
4. QCVN 8-1:2011/BYT — Giới hạn ô nhiễm vi sinh vật (E.coli, Salmonella, Staphylococcus)
5. QCVN 8-2:2011/BYT — Giới hạn ô nhiễm kim loại nặng (chì, cadimi, thuỷ ngân)
6. QCVN 8-3:2012/BYT — Giới hạn dư lượng thuốc bảo vệ thực vật trong thực phẩm
7. Luật ATTP số 55/2010/QH12 — Khung pháp lý tổng thể về an toàn thực phẩm
8. WHO Five Keys to Safer Food — 5 nguyên tắc vệ sinh thực phẩm của WHO
9. Codex Alimentarius CAC/RCP 1-1969 — Quy phạm thực hành vệ sinh tổng quát
10. Nguyên tắc HACCP (Hazard Analysis Critical Control Points) — Phân tích mối nguy
VÙNG NHIỆT ĐỘ NGUY HIỂM: 5°C – 60°C (vi khuẩn tăng gấp đôi mỗi 20 phút)
"""

_VISUAL_CRITERIA = """
TIÊU CHÍ ĐÁNH GIÁ TRỰC QUAN (Căn cứ WHO Food Safety Visual Inspection Guide + QCVN):
- Màu sắc thực phẩm: bất thường (xanh, đen, xám) → nguy cơ nhiễm khuẩn/mốc
- Bề mặt: nhớt, ẩm ướt bất thường → dấu hiệu vi khuẩn phát triển
- Mùi: chua, hôi, khác lạ → phân huỷ protein hoặc nhiễm khuẩn
- Nấm mốc: đốm trắng/xanh/đen → Aspergillus, Penicillium, Fusarium
- Côn trùng: ruồi, gián, kiến → vector lây truyền Salmonella, E.coli
- Nhiệt độ: <60°C (nóng) hoặc >5°C (lạnh) → vùng nguy hiểm theo HACCP
- Sổ kiểm thực: thiếu chữ ký, thiếu bước → vi phạm TTLT 13/2016 Điều 9
- Dụng cụ: rỉ sét, nứt vỡ → nơi trú ẩn vi khuẩn, không đảm bảo NĐ 15/2018
"""


# ── AI #2: Phân tích ảnh rủi ro ATTP (Claude Vision) ─────────────────────────
def analyze_photo_ai(photo_bytes: bytes, group_name: str, api_key: str) -> dict:
    """Gọi Claude Vision phân tích ảnh ATTP dựa trên chuẩn WHO + QCVN + NĐ 15/2018."""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        b64 = base64.standard_b64encode(photo_bytes).decode()
        resp = client.messages.create(
            model=MODEL_VISION,
            max_tokens=500,   # JSON output ngắn gọn, không cần nhiều
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image",
                     "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}},
                    {"type": "text", "text": f"""Bạn là chuyên gia kiểm tra ATTP trường học Việt Nam.
Nhóm kiểm tra: {group_name}

{_VISUAL_CRITERIA}

Phân tích ảnh theo các tiêu chí trên. Trả lời JSON (không thêm text ngoài JSON):
{{
  "risk_level": "OK" hoặc "WARNING" hoặc "CRITICAL",
  "issues": ["vấn đề cụ thể + căn cứ pháp lý/kỹ thuật nếu có"],
  "positives": ["điểm đạt chuẩn quan sát được"],
  "recommendation": "hành động khắc phục cụ thể + thời hạn (để trống nếu OK)",
  "legal_ref": "văn bản pháp lý áp dụng chính (ví dụ: QCVN 8-1:2011/BYT)",
  "confidence": 0.85
}}"""}
                ]
            }]
        )
        text = resp.content[0].text.strip()
        s, e = text.find("{"), text.rfind("}") + 1
        return json.loads(text[s:e]) if s != -1 and e > s else \
               {"risk_level": "OK", "issues": [], "positives": [],
                "recommendation": "", "legal_ref": "", "confidence": 0.5}
    except Exception as ex:
        return {"risk_level": "ERROR", "issues": [str(ex)], "positives": [],
                "recommendation": "", "legal_ref": "", "confidence": 0}


# ── AI #1: Checklist động theo thực đơn ──────────────────────────────────────
def generate_extra_checklist(menu: str, school_level: str, api_key: str) -> list:
    """Tạo 3-5 điểm kiểm tra bổ sung đặc thù theo thực đơn, có căn cứ QCVN/NĐ 15."""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=MODEL,
            max_tokens=800,   # JSON array 3-5 items, giảm từ 1200
            messages=[{"role": "user", "content": f"""Bạn là chuyên gia ATTP trường học Việt Nam.
Thực đơn hôm nay ({school_level}): {menu}

{_LEGAL_BASIS}

Dựa trên từng nguyên liệu trong thực đơn, tạo 3-5 điểm kiểm tra ATTP BỔ SUNG
(ngoài 20 điểm chuẩn), đặc thù cho nguyên liệu đó, có dẫn chiếu quy chuẩn.

Trả lời JSON array (không thêm text khác):
[
  {{
    "code": "E01",
    "ingredient": "Tên nguyên liệu cụ thể",
    "desc": "Mô tả điểm kiểm tra (ngắn gọn, dưới 15 từ)",
    "how": "Cách kiểm tra thực tế (dụng cụ cần thiết nếu có)",
    "pass": "Tiêu chí đạt (kèm giá trị cụ thể nếu có: nhiệt độ, màu sắc...)",
    "fail": "Tiêu chí không đạt (dấu hiệu cụ thể cần nhận biết)",
    "is_critical": false,
    "legal_ref": "Căn cứ: QCVN... / NĐ 15/2018 Điều... / WHO...",
    "why": "Giải thích ngắn tại sao nguyên liệu này cần kiểm tra điểm này"
  }}
]"""}]
        )
        text = resp.content[0].text.strip()
        s, e = text.find("["), text.rfind("]") + 1
        items = json.loads(text[s:e]) if s != -1 and e > s else []
        for i, item in enumerate(items):
            item["code"] = f"E{i+1:02d}"
        return items
    except Exception:
        return []


# ── Anti-Fraud: Câu hỏi ngẫu nhiên chỉ người có mặt mới trả lời được ─────────
def generate_anti_fraud_questions(menu: str, school: str, date_str: str,
                                   school_level: str, api_key: str) -> list:
    """
    Tạo 2-3 câu hỏi xác thực ngẫu nhiên dựa trên bối cảnh thực tế hôm nay.
    Chỉ người THỰC SỰ có mặt tại bếp mới trả lời được — chống gian lận.
    Câu hỏi thay đổi mỗi ngày, không thể đoán trước.
    """
    try:
        _now = now_vn()
        _day_vn = ["Thứ Hai","Thứ Ba","Thứ Tư","Thứ Năm","Thứ Sáu","Thứ Bảy","Chủ Nhật"][_now.weekday()]
        _time_ctx = f"{_now.strftime('%H:%M')} {_day_vn}"
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=MODEL, max_tokens=500,
            messages=[{"role": "user", "content": f"""Bạn là kiểm tra viên ATTP đang ở tại bếp trường.
Bối cảnh: {school} · {_day_vn} {date_str} {_time_ctx} · Cấp {school_level}
Thực đơn hôm nay: {menu}

Tạo ĐÚNG 2 câu hỏi xác thực ngẫu nhiên — câu hỏi phải:
- Chỉ người ĐANG CÓ MẶT tại bếp lúc này mới biết đáp án
- Dựa trên quan sát trực tiếp (màu sắc, mùi, nhiệt kế, số lượng thực tế hôm nay...)
- KHÔNG thể trả lời bằng cách nhìn vào form hay đoán mò
- Thay đổi mỗi ngày, không lặp lại

Ví dụ tốt: "Nhiệt kế trên tủ lạnh số 1 đang hiển thị bao nhiêu độ?"
Ví dụ xấu: "Nhiệt độ tủ lạnh phải là bao nhiêu?" (đây là câu hỏi kiến thức, không xác thực)

Trả về JSON array (không thêm text):
[
  {{"q": "Câu hỏi quan sát cụ thể?", "hint": "Gợi ý nơi cần quan sát"}},
  {{"q": "Câu hỏi quan sát cụ thể khác?", "hint": "Gợi ý nơi cần quan sát"}}
]"""}]
        )
        text = resp.content[0].text.strip()
        s, e = text.find("["), text.rfind("]") + 1
        return json.loads(text[s:e]) if s != -1 and e > s else []
    except Exception:
        return []


# ── AI #3: Báo cáo ngôn ngữ tự nhiên ────────────────────────────────────────
def generate_ai_narrative(results: dict, notes: dict, alert_level: str,
                           school: str, date_str: str, menu: str,
                           pass_count: int, total: int,
                           level_key: str, api_key: str) -> str:
    """Claude viết tóm tắt báo cáo kiểm tra bằng tiếng Việt tự nhiên."""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        fail_list = [f"{c} ({notes.get(c,'').strip() or 'không có ghi chú'})"
                     for c, v in results.items() if v == "❌ Không Đạt"]
        context = (
            f"Trường: {school or '(chưa nhập)'} | Cấp: {level_key} | Ngày: {date_str}\n"
            f"Thực đơn: {menu or '(chưa nhập)'}\n"
            f"Kết quả: {pass_count}/{total} đạt chuẩn | Cảnh báo: {alert_level}\n"
            f"Mục không đạt: {', '.join(fail_list) if fail_list else 'Không có'}"
        )
        resp = client.messages.create(
            model=MODEL,
            max_tokens=400,   # Đoạn văn 120-160 từ, giảm từ 500
            messages=[{"role": "user", "content": f"""Viết đoạn tóm tắt báo cáo ATTP (120–160 từ) bằng tiếng Việt:
{context}

Yêu cầu:
- Văn phong chuyên nghiệp, dễ hiểu với phụ huynh và ban giám hiệu
- Nêu điểm tốt trước, điểm cần cải thiện sau
- Có khuyến nghị cụ thể nếu có mục không đạt
- Kết thúc bằng đánh giá tổng thể 1 câu
- Không dùng markdown header hay bullet points, viết thành đoạn văn"""
            }]
        )
        return resp.content[0].text.strip()
    except Exception as ex:
        return f"(Không tạo được tóm tắt AI: {ex})"


# ── Tạo báo cáo Word chuẩn chính phủ (.docx) ────────────────────────────────
def _docx_set_font(run, bold=False, size_pt=12, color=None):
    """Helper: thiết lập font Times New Roman cho run."""
    from docx.shared import Pt, RGBColor
    run.font.name     = "Times New Roman"
    run.font.size     = Pt(size_pt)
    run.font.bold     = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _docx_para(doc, text, bold=False, size=12, align="left", space_before=0, space_after=6):
    """Thêm paragraph với Times New Roman, trả về paragraph."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    _MAP = {"left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = _MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    _docx_set_font(run, bold=bold, size_pt=size)
    return p


def _docx_table_header(table, headers: list, bg_color="1B3B6F"):
    """Tô màu header row của table."""
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared  import Pt, RGBColor
    row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(row.cells, headers)):
        cell.text = ""
        run = cell.paragraphs[0].add_run(hdr)
        _docx_set_font(run, bold=True, size_pt=11, color=(255, 255, 255))
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg_color)
        tcPr.append(shd)


def generate_word_report(school: str, date_str: str, insp: str, menu: str,
                          level_key: str, results: dict, notes: dict,
                          pass_count: int, fail_count: int, alert_key: str,
                          cl: list, ai_narrative: str,
                          photo_analysis: dict) -> bytes:
    """
    Tạo báo cáo kiểm tra ATTP định dạng Word (.docx) chuẩn văn bản hành chính Việt Nam.
    Font Times New Roman, trình bày như báo cáo gửi cấp Bộ/Chính phủ.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns  import qn
    from docx.oxml     import OxmlElement
    from io import BytesIO

    doc = Document()

    # ── Lề trang (chuẩn văn bản hành chính Việt Nam) ─────────────────────────
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(3.0)   # lề trái rộng cho đóng bìa
    sec.right_margin  = Cm(2.0)

    # ── Quốc hiệu & Tiêu ngữ ─────────────────────────────────────────────────
    _docx_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
               bold=True, size=13, align="center", space_after=2)
    _docx_para(doc, "Độc lập – Tự do – Hạnh phúc",
               bold=True, size=12, align="center", space_after=2)
    # Đường kẻ ngang
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_line.add_run("───────────────────────")
    _docx_set_font(r, size_pt=11)
    p_line.paragraph_format.space_after = Pt(4)

    # Số hiệu và ngày
    vn_days = ["Thứ Hai","Thứ Ba","Thứ Tư","Thứ Năm","Thứ Sáu","Thứ Bảy","Chủ Nhật"]
    try:
        dt_obj   = datetime.strptime(date_str, "%d/%m/%Y")
        day_name = vn_days[dt_obj.weekday()]
        date_full = f"TP. Hồ Chí Minh, {day_name} ngày {dt_obj.day} tháng {dt_obj.month} năm {dt_obj.year}"
    except Exception:
        date_full = f"TP. Hồ Chí Minh, ngày ... tháng ... năm 2026"
    _docx_para(doc, date_full, size=12, align="right", space_before=4, space_after=12)

    # ── Tiêu đề báo cáo ──────────────────────────────────────────────────────
    _docx_para(doc, "BÁO CÁO", bold=True, size=14, align="center", space_before=6, space_after=2)
    _docx_para(doc, "KIỂM TRA AN TOÀN THỰC PHẨM BỮA ĂN HỌC ĐƯỜNG",
               bold=True, size=14, align="center", space_after=12)

    a = ALERT_SYSTEM.get(alert_key, {})
    _docx_para(doc,
               f"Kính gửi: Ban Giám hiệu Trường {school or '...'} và Sở Giáo dục & Đào tạo",
               size=12, align="left", space_after=8)

    # ── I. THÔNG TIN CHUNG ────────────────────────────────────────────────────
    _docx_para(doc, "I. THÔNG TIN CHUNG", bold=True, size=13, space_before=6, space_after=4)

    info_rows = [
        ("1. Cơ sở giáo dục",       school or "(chưa nhập)"),
        ("2. Cấp học",               level_key),
        ("3. Ngày kiểm tra",         date_str),
        ("4. Người kiểm tra",        insp or "(chưa nhập)"),
        ("5. Thực đơn hôm nay",      menu or "(chưa nhập)"),
        ("6. Mức cảnh báo",          f"{a.get('icon','')} {a.get('label','')}"),
        ("7. Kết quả tổng quát",     f"{pass_count} ĐẠT / {pass_count+fail_count} điểm đã kiểm tra"),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(info_rows):
        r0 = t.rows[i].cells[0].paragraphs[0].add_run(k)
        r1 = t.rows[i].cells[1].paragraphs[0].add_run(v)
        _docx_set_font(r0, bold=True, size_pt=11)
        _docx_set_font(r1, size_pt=11)
        t.rows[i].cells[0].width = Cm(6)
        t.rows[i].cells[1].width = Cm(10)
    doc.add_paragraph()

    # ── II. CĂN CỨ PHÁP LÝ ───────────────────────────────────────────────────
    _docx_para(doc, "II. CĂN CỨ PHÁP LÝ", bold=True, size=13, space_before=6, space_after=4)
    legal_refs = [
        "Nghị định số 15/2018/NĐ-CP ngày 02/02/2018 của Chính phủ — Quy định chi tiết thi hành một số điều của Luật An toàn thực phẩm;",
        "Thông tư liên tịch số 13/2016/TTLT-BYT-BGDĐT ngày 12/5/2016 — Quy định về công tác y tế trường học;",
        "Quyết định số 3958/QĐ-BYT ngày 25/12/2025 của Bộ Y tế — Hướng dẫn dinh dưỡng bữa ăn học đường;",
        "QCVN 8-1:2011/BYT — Quy chuẩn kỹ thuật quốc gia về giới hạn ô nhiễm vi sinh vật trong thực phẩm;",
        "QCVN 8-2:2011/BYT — Quy chuẩn kỹ thuật quốc gia về giới hạn ô nhiễm kim loại nặng trong thực phẩm;",
        "Luật An toàn thực phẩm số 55/2010/QH12 ngày 17/6/2010 của Quốc hội.",
    ]
    for ref in legal_refs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(ref)
        _docx_set_font(r, size_pt=11)
        p.paragraph_format.space_after = Pt(3)

    # ── III. KẾT QUẢ KIỂM TRA ────────────────────────────────────────────────
    _docx_para(doc, "III. KẾT QUẢ KIỂM TRA CHI TIẾT",
               bold=True, size=13, space_before=8, space_after=4)

    # Bảng chi tiết
    item_map = {code: (desc, is_crit, grp)
                for grp, items in cl for code, is_crit, desc, *_ in items}
    headers = ["Mã", "Bắt buộc", "Nội dung kiểm tra", "Kết quả", "Ghi chú"]
    tbl = doc.add_table(rows=1 + len(results), cols=5)
    tbl.style = "Table Grid"
    _docx_table_header(tbl, headers)

    col_widths = [Cm(1.2), Cm(1.8), Cm(8.5), Cm(2.5), Cm(3.0)]
    for i, (code, result) in enumerate(results.items()):
        if code not in item_map:
            continue
        desc, is_crit, _ = item_map[code]
        row = tbl.rows[i + 1]
        cells_data = [
            code,
            "★ Bắt buộc" if is_crit else "",
            desc,
            result.replace("✅ Đạt", "ĐẠT").replace("❌ Không Đạt", "KHÔNG ĐẠT").replace("Chưa chấm", "—"),
            notes.get(code, ""),
        ]
        for j, (cell, val, w) in enumerate(zip(row.cells, cells_data, col_widths)):
            cell.width = w
            is_fail = result == "❌ Không Đạt" and j == 3
            r = cell.paragraphs[0].add_run(val)
            _docx_set_font(r, bold=(j == 3 and is_fail), size_pt=10,
                           color=(192, 0, 0) if is_fail else None)
    doc.add_paragraph()

    # ── IV. ĐÁNH GIÁ VÀ NHẬN XÉT (AI narrative) ─────────────────────────────
    _docx_para(doc, "IV. ĐÁNH GIÁ VÀ NHẬN XÉT",
               bold=True, size=13, space_before=8, space_after=4)
    if ai_narrative and ai_narrative.startswith("(Không"):
        _docx_para(doc, "(Chưa tạo tóm tắt AI — vui lòng kết nối API key)",
                   size=11, align="justify")
    else:
        _docx_para(doc, ai_narrative or "(Chưa có đánh giá AI)",
                   size=12, align="justify", space_after=6)

    # Kết quả phân tích ảnh AI (nếu có)
    if photo_analysis:
        _docx_para(doc, "Kết quả phân tích ảnh minh chứng (AI):",
                   bold=True, size=12, space_before=4, space_after=2)
        RISK_VN = {"OK": "ĐẠT CHUẨN", "WARNING": "CẦN CHÚ Ý",
                   "CRITICAL": "NGUY HIỂM", "ERROR": "Không xác định"}
        for idx, (g_idx, r) in enumerate(photo_analysis.items()):
            lvl = r.get("risk_level", "OK")
            _docx_para(doc,
                       f"  • Nhóm {g_idx+1}: {RISK_VN.get(lvl, lvl)} "
                       f"(tin cậy {int(r.get('confidence',0.8)*100)}%) "
                       f"— {r.get('legal_ref','')}",
                       size=11, space_after=2)
            for issue in r.get("issues", []):
                _docx_para(doc, f"      → Vấn đề: {issue}", size=10, space_after=1)

    # ── V. KIẾN NGHỊ ─────────────────────────────────────────────────────────
    _docx_para(doc, "V. KIẾN NGHỊ VÀ YÊU CẦU XỬ LÝ",
               bold=True, size=13, space_before=8, space_after=4)
    fail_items = {c: v for c, v in results.items() if v == "❌ Không Đạt"}
    critical_fails = CRITICAL_ITEMS & set(fail_items.keys())

    if critical_fails:
        _docx_para(doc,
                   "1. YÊU CẦU KHẨN: Các mục bắt buộc dưới đây vi phạm — "
                   "đề nghị xử lý ngay trước bữa ăn tiếp theo:",
                   bold=True, size=12, space_after=2)
        for code in sorted(critical_fails):
            desc = item_map.get(code, ("",))[0]
            _docx_para(doc, f"   • {code}: {desc}", size=11, space_after=2)

    non_crit_fails = set(fail_items.keys()) - critical_fails
    if non_crit_fails:
        n = 2 if critical_fails else 1
        _docx_para(doc,
                   f"{n}. Các mục cần cải thiện trong 24 giờ:",
                   bold=True, size=12, space_after=2)
        for code in sorted(non_crit_fails):
            desc = item_map.get(code, ("",))[0]
            note = notes.get(code, "")
            _docx_para(doc,
                       f"   • {code}: {desc}" + (f" ({note})" if note else ""),
                       size=11, space_after=2)

    if not fail_items:
        _docx_para(doc,
                   "Bữa ăn đạt toàn bộ tiêu chí kiểm tra. "
                   "Đề nghị duy trì và tiếp tục theo dõi định kỳ.",
                   size=12, space_after=4)

    # ── VI. KẾT LUẬN ─────────────────────────────────────────────────────────
    _docx_para(doc, "VI. KẾT LUẬN", bold=True, size=13, space_before=8, space_after=4)
    _docx_para(doc,
               f"Trên đây là báo cáo kết quả kiểm tra An toàn thực phẩm bữa ăn học đường "
               f"tại {school or 'cơ sở giáo dục'} ngày {date_str}. "
               f"Tổng điểm đạt: {pass_count}/{pass_count+fail_count} điểm. "
               f"Đánh giá tổng thể: {a.get('label','—')}. "
               f"Báo cáo này được lập theo đúng quy định của TTLT 13/2016/TTLT-BYT-BGDĐT "
               f"và NĐ 15/2018/NĐ-CP.",
               size=12, align="justify", space_after=8)
    _docx_para(doc,
               "Kính đề nghị Ban Giám hiệu xem xét, chỉ đạo xử lý các vấn đề nêu trên "
               "và thông báo kết quả khắc phục cho Ban Giám sát trong vòng 24 giờ.",
               size=12, align="justify", space_after=16)

    # ── Chữ ký ───────────────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    sig_table.style = None  # Bỏ border cho bảng chữ ký

    left_cells = ["Đại diện Ban Giám Sát (Đại Diện PHHS)", "", "", ""]
    right_cells = [f"TP. Hồ Chí Minh, {date_str}", "Người kiểm tra", "", f"{insp or '(ký và ghi rõ họ tên)'}"]
    for i, (lc, rc) in enumerate(zip(left_cells, right_cells)):
        for cell, txt, bold in [(sig_table.rows[i].cells[0], lc, i==0),
                                 (sig_table.rows[i].cells[1], rc, i in (0,1))]:
            r = cell.paragraphs[0].add_run(txt)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _docx_set_font(r, bold=bold, size_pt=12)

    # Ghi chú cuối trang
    doc.add_paragraph()
    _docx_para(doc,
               "─────────────────────────────────────────────────────",
               size=10, align="center", space_before=8, space_after=2)
    _docx_para(doc,
               f"Báo cáo được tạo tự động bởi SchoolFood AI v2.0 — {now_vn().strftime('%d/%m/%Y %H:%M')} (GMT+7)   |   "
               f"Đường dây nóng Cục ATTP: 1800 6838 (miễn phí)",
               size=9, align="center", space_after=0)

    # ── Xuất ra bytes ─────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_word_incident(incident_log: list, school: str = "") -> bytes:
    """Tạo biên bản sự cố ngộ độc định dạng Word chuẩn hành chính."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)

    _docx_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
               bold=True, size=13, align="center", space_after=2)
    _docx_para(doc, "Độc lập – Tự do – Hạnh phúc",
               bold=True, size=12, align="center", space_after=2)
    _docx_para(doc, "────────────────────────────────",
               size=11, align="center", space_after=4)
    _docx_para(doc, f"TP. Hồ Chí Minh, ngày {now_vn().strftime('%d tháng %m năm %Y')}",
               size=12, align="right", space_after=12)
    _docx_para(doc, "BIÊN BẢN SỰ CỐ AN TOÀN THỰC PHẨM",
               bold=True, size=14, align="center", space_after=2)
    _docx_para(doc, "Nghi ngờ ngộ độc thực phẩm tại cơ sở giáo dục",
               bold=True, size=13, align="center", space_after=12)
    _docx_para(doc, f"Cơ sở giáo dục: {school or '....................'}",
               size=12, space_after=4)
    _docx_para(doc, f"Thời gian lập biên bản: {now_vn().strftime('%H:%M ngày %d/%m/%Y')}",
               size=12, space_after=8)

    _docx_para(doc, "DIỄN BIẾN SỰ CỐ (TIMELINE):",
               bold=True, size=13, space_before=4, space_after=4)
    for entry in incident_log:
        _docx_para(doc, entry, size=12, align="justify", space_after=3)

    _docx_para(doc, "CĂN CỨ PHÁP LÝ ÁP DỤNG:",
               bold=True, size=13, space_before=8, space_after=4)
    for ref in [
        "TTLT 13/2016/TTLT-BYT-BGDĐT — Xử lý ngộ độc thực phẩm tập thể tại trường học",
        "NĐ 15/2018/NĐ-CP — Trách nhiệm báo cáo khi xảy ra ngộ độc thực phẩm",
        "Luật ATTP 55/2010/QH12 — Nghĩa vụ báo cáo cơ quan quản lý trong 24 giờ",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(ref); _docx_set_font(r, size_pt=11)
        p.paragraph_format.space_after = Pt(3)

    _docx_para(doc, "Đường dây nóng Cục ATTP: 1800 6838 (miễn phí)  |  Cấp cứu: 115",
               bold=True, size=12, align="center", space_before=8, space_after=16)

    for label in ["Đại diện Nhà trường", "Người lập biên bản", "Đại diện Ban Giám sát"]:
        p = doc.add_paragraph(f"{'':>40}{label}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.runs[0]; _docx_set_font(r, bold=True, size_pt=12)
        doc.add_paragraph()

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ── AI #7: Trợ lý ứng phó sự cố ngộ độc ─────────────────────────────────────
def incident_ai_response(client: anthropic.Anthropic,
                          history: list, user_msg: str) -> str:
    """AI dẫn dắt xử lý sự cố ngộ độc từng bước, ghi nhận timeline."""
    INCIDENT_SYSTEM = """Bạn là chuyên gia ATTP khẩn cấp đang hỗ trợ xử lý sự cố ngộ độc thực phẩm học đường.
Nhiệm vụ: Dẫn dắt người dùng từng bước qua quy trình xử lý, hỏi thông tin còn thiếu, ghi nhận timeline.

QUY TRÌNH CHUẨN (TTLT 13/2016):
1. Dừng bữa ăn ngay
2. Gọi 115 nếu triệu chứng nặng (khó thở, co giật)
3. Giữ nguyên mẫu thức ăn (không vứt, không rửa)
4. Báo Hiệu trưởng + Y tế học đường ngay
5. Ghi chép số học sinh, triệu chứng, thời gian
6. Báo Sở Y tế trong 24h (bắt buộc nếu ≥2 người)

CÁCH PHẢN HỒI:
- Luôn xác nhận bước đã làm
- Hỏi thông tin cụ thể còn thiếu (số học sinh, triệu chứng, thời gian)
- Đưa ra bước tiếp theo rõ ràng
- Ghi nhận thông tin vào "SỔ GHI SỰ CỐ" format: [HH:MM] - nội dung
- Nếu tình huống nghiêm trọng (khó thở, co giật) → ưu tiên gọi 115 NGAY"""

    messages = history + [{"role": "user", "content": user_msg}]
    try:
        resp = client.messages.create(
            model=MODEL,
            max_tokens=500,   # Giảm từ 600 — hướng dẫn từng bước ngắn gọn
            system=[{"type": "text", "text": INCIDENT_SYSTEM,
                     "cache_control": {"type": "ephemeral"}}],  # Cache system prompt
            messages=messages,
        )
        return resp.content[0].text
    except Exception as ex:
        return f"❌ Lỗi kết nối: {ex}"


# ── TAB 1: Hỏi đáp AI ─────────────────────────────────────────────────────────
def tab_chat(api_key, role, level, loc):
    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">💬 Hỏi đáp pháp luật ATTP</div>
        <div class="sf-card-body">Đặt câu hỏi bằng tiếng Việt thông thường — AI trả lời dựa trên văn bản pháp luật thực tế.</div>
    </div>""", unsafe_allow_html=True)

    QUICK = {
        "Phụ Huynh": [
            "Phụ Huynh có quyền vào bếp kiểm tra không?",
            "Con tôi đau bụng sau bữa trưa — tôi phải làm gì ngay?",
            "Làm sao biết suất ăn của con đủ dinh dưỡng không?",
            "Tôi có thể yêu cầu xem thực đơn trước không?",
        ],
        "Ban Giám Sát (Đại Diện PHHS)": [
            "Ban Giám Sát có quyền gì theo pháp luật?",
            "Phát hiện thực phẩm hết hạn — phải báo cáo ai?",
            "Nhà Cung Cấp cần có những giấy tờ gì?",
            "Tiêu chuẩn khẩu phần tiểu học theo quy định 2025?",
        ],
        "Y Tế Học Đường": [
            "Sổ kiểm thực 3 bước cần ghi những gì?",
            "Mẫu lưu thức ăn cần lưu bao lâu và cách nào?",
            "Nhiệt độ bảo quản đúng chuẩn cho từng loại thực phẩm?",
            "Quy trình xử lý khi nghi ngờ ngộ độc tập thể?",
        ],
        "Ban Giám Hiệu": [
            "Hiệu Trưởng chịu trách nhiệm pháp lý như thế nào?",
            "Tiêu chí lựa chọn Nhà Cung Cấp suất ăn hợp lệ?",
            "Mức xử phạt khi xảy ra ngộ độc thực phẩm tại trường?",
            "Báo cáo định kỳ ATTP gửi Sở GD&ĐT như thế nào?",
        ],
    }
    questions = QUICK.get(role, QUICK["Phụ Huynh"])
    st.markdown('<div class="sec-hdr">Câu hỏi phổ biến cho vai trò của bạn</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    for i, q in enumerate(questions):
        if (c1 if i % 2 == 0 else c2).button(q, key=f"qq{i}", use_container_width=True):
            st.session_state.preset_q = q

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)

    if "messages" not in st.session_state:
        st.session_state.messages = []
    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    preset = st.session_state.pop("preset_q", None)
    user_input = st.chat_input("Nhập câu hỏi...") or preset
    if user_input:
        if not api_key:
            st.warning("⚠️ Vui lòng nhập API Key ở thanh cài đặt phía trên.")
            st.stop()
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)
        with st.chat_message("assistant"):
            with st.spinner("Đang tra cứu..."):
                sys = build_system_prompt(role, level, loc)
                hist = [{"role": m["role"], "content": m["content"]}
                        for m in st.session_state.messages[:-1]]
                ans = ask_claude(anthropic.Anthropic(api_key=api_key), sys, hist, user_input)
                st.markdown(ans)
                st.session_state.messages.append({"role": "assistant", "content": ans})
    if st.session_state.get("messages"):
        if st.button("🗑️ Xóa lịch sử", use_container_width=True):
            st.session_state.messages = []
            st.rerun()


# ── TAB 2: Checklist ──────────────────────────────────────────────────────────
def tab_checklist(api_key: str = ""):
    ai_on = bool(api_key)   # AI features chỉ hoạt động khi có API key

    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">✅ Checklist kiểm tra ATTP bữa ăn học đường</div>
        <div class="sf-card-body">20 điểm chuẩn hoá · Mục <span style="background:#FEE2E2;color:#991B1B;padding:1px 6px;border-radius:8px;font-weight:700;font-size:0.75rem">BẮT BUỘC</span> phải đạt tuyệt đối
        """ + (" · <span style='color:#2563EB;font-weight:600'>🤖 AI đang hoạt động</span>" if ai_on else " · <i style='color:#94A3B8'>AI tắt — checklist vẫn dùng được đầy đủ</i>") + """
        </div>
    </div>""", unsafe_allow_html=True)

    # Cấp học — lấy từ profile (locked) hoặc cho chọn nếu demo/admin
    _cl_default_lvl = st.session_state.get("user_profile", {}).get("default_level") or "Tiểu Học (6–11 tuổi)"
    _cl_is_locked = bool(st.session_state.get("auth_user")) and not st.session_state.get("is_super", False) and not st.session_state.get("auth_user", {}).get("demo", False)
    if _cl_is_locked and _cl_default_lvl in list(NUTRITION.keys()):
        level_key = _cl_default_lvl
        st.caption(f"📚 Cấp học: **{_cl_default_lvl}** · _Theo tài khoản_")
    else:
        level_key = st.selectbox(
            "📚 Cấp học đang kiểm tra", list(NUTRITION.keys()), key="cl_level",
        )
    n = NUTRITION[level_key]

    # Banner dinh dưỡng — thu gọn để không che checklist chính
    with st.expander(f"📊 Tiêu chuẩn dinh dưỡng cấp {n['short']} (tham khảo cho C12, C13)"):
        st.markdown(f"""<div class="nutrition-banner">
            <div class="nutrition-label">📊 Tiêu Chuẩn Dinh Dưỡng Bữa Trưa — Cấp {n['short']} (QĐ 3958/QĐ-BYT 2025)</div>
            <div class="nutrition-grid">
                <div class="nutrition-item">⚡ Năng lượng: <span class="nutrition-val">{n['kcal']}</span> ({n['pct_day']} nhu cầu ngày)</div>
                <div class="nutrition-item">🥩 Thịt/cá tối thiểu: <span class="nutrition-val">{n['meat_g']}g/học sinh</span></div>
                <div class="nutrition-item">🥦 Rau xanh: <span class="nutrition-val">{n['veg_range']}/học sinh</span></div>
                <div class="nutrition-item">💪 Protein: <span class="nutrition-val">{n['protein_pct']}</span> tổng năng lượng</div>
            </div>
            <div style="font-size:0.75rem;color:#1D4ED8;margin-top:6px">💡 {n['note']}</div>
        </div>""", unsafe_allow_html=True)

    # Thông tin kiểm tra
    st.markdown('<div class="sec-hdr">Thông tin buổi kiểm tra</div>', unsafe_allow_html=True)
    _us = st.session_state.get("user_school", "")
    c1, c2, c3 = st.columns(3)
    school = c1.text_input("Tên trường", value=_us,
                            disabled=bool(_us),
                            placeholder="VD: TH Nguyễn Du, Q.1")
    date   = c2.date_input("Ngày kiểm tra", value=datetime.today(), format="DD/MM/YYYY")
    insp   = c3.text_input("Người kiểm tra",
                            value=st.session_state.get("user_profile", {}).get("full_name", ""),
                            placeholder="Họ và tên")
    menu   = st.text_input("Thực đơn hôm nay",
                            placeholder="VD: Cơm, thịt kho trứng, rau muống xào tỏi, canh chua cá",
                            key="shared_menu")

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)

    if "cl_r"       not in st.session_state: st.session_state.cl_r       = {}
    if "cl_n"       not in st.session_state: st.session_state.cl_n       = {}
    if "cl_photos"  not in st.session_state: st.session_state.cl_photos  = {}
    if "cl_extra"   not in st.session_state: st.session_state.cl_extra   = []
    if "photo_analysis" not in st.session_state: st.session_state.photo_analysis = {}

    # ── Giải thích cấu trúc checklist ────────────────────────────────────────
    st.markdown(f"""
    <div style="background:#F8FAFC;border:1px solid #E2E8F0;border-radius:10px;
                padding:12px 16px;margin-bottom:12px;font-size:0.85rem">
        <b>📋 Cấu trúc kiểm tra:</b>
        <span style="color:#1E293B">
          &nbsp;&nbsp;🔵 <b>20 câu chuẩn</b> — luôn có, không cần API
          &nbsp;&nbsp;|&nbsp;&nbsp;
          🤖 <b>Câu hỏi AI bổ sung</b> — tùy chọn, theo thực đơn hôm nay
          {"&nbsp;&nbsp;|&nbsp;&nbsp;<span style='color:#2563EB'>💳 Mỗi lần tạo ≈ $0.004 credit</span>" if ai_on else ""}
        </span>
    </div>
    """, unsafe_allow_html=True)

    # ── AI #1: Tạo checklist bổ sung theo thực đơn ───────────────────────────
    if ai_on and menu:
        col_btn, col_status = st.columns([0.5, 0.5])
        if col_btn.button("🤖 Tạo câu hỏi bổ sung theo thực đơn (~$0.004)",
                          use_container_width=True, key="gen_extra"):
            with st.spinner("AI đang phân tích thực đơn và tra cứu QCVN..."):
                extras = generate_extra_checklist(menu, level_key, api_key)
                st.session_state.cl_extra = extras
                st.session_state.photo_analysis = {}
        if st.session_state.cl_extra:
            col_status.markdown(
                f'<span style="color:#16A34A;font-size:0.85rem">✅ Đã tạo '
                f'<b>{len(st.session_state.cl_extra)}</b> câu hỏi — '
                f'<span style="color:#94A3B8">căn cứ QCVN + NĐ 15/2018</span></span>',
                unsafe_allow_html=True,
            )
    elif ai_on and not menu:
        st.caption("💡 Nhập thực đơn hôm nay bên trên để AI tạo câu hỏi kiểm tra riêng.")
    else:
        st.caption("🔌 Kết nối API key ở thanh cài đặt phía trên để dùng tính năng tạo câu hỏi theo thực đơn.")

    # ── Anti-Fraud: Câu hỏi xác thực ngẫu nhiên ──────────────────────────────
    if ai_on and menu and school:
        if "cl_af_questions" not in st.session_state:
            st.session_state.cl_af_questions = []
        if "cl_af_answers" not in st.session_state:
            st.session_state.cl_af_answers = {}
        if "cl_af_verified" not in st.session_state:
            st.session_state.cl_af_verified = False

        with st.expander("🔐 Xác thực hiện diện — Chống gian lận", expanded=False):
            st.markdown(
                '<div style="background:#FFF7ED;border:1px solid #FED7AA;border-radius:8px;'
                'padding:10px 14px;font-size:0.82rem;color:#92400E;margin-bottom:8px">'
                '⚠️ <b>Anti-Fraud:</b> AI tạo câu hỏi ngẫu nhiên dựa trên bối cảnh thực tế hôm nay. '
                'Chỉ người <b>đang có mặt tại bếp</b> mới trả lời được. '
                'Câu hỏi thay đổi mỗi lần — không thể đoán trước hoặc copy từ lần trước.'
                '</div>',
                unsafe_allow_html=True,
            )
            _af_col1, _af_col2 = st.columns([2, 1])
            if _af_col1.button("🎲 Tạo câu hỏi xác thực ngẫu nhiên", key="gen_af",
                                use_container_width=True):
                with st.spinner("AI đang tạo câu hỏi dựa trên bối cảnh hôm nay..."):
                    _af_qs = generate_anti_fraud_questions(
                        menu=menu, school=school,
                        date_str=str(date), school_level=level_key, api_key=api_key
                    )
                st.session_state.cl_af_questions = _af_qs
                st.session_state.cl_af_answers = {}
                st.session_state.cl_af_verified = False

            if st.session_state.cl_af_questions:
                st.markdown("**Trả lời dựa trên quan sát TRỰC TIẾP tại bếp:**")
                _all_answered = True
                for _qi, _q in enumerate(st.session_state.cl_af_questions):
                    st.markdown(
                        f'<div style="background:#F8FAFC;border-radius:6px;padding:8px 12px;'
                        f'margin:4px 0;font-size:0.85rem;font-weight:600;color:#1E293B">'
                        f'❓ {_q["q"]}</div>'
                        f'<div style="font-size:0.75rem;color:#64748B;padding:0 4px 4px">'
                        f'💡 {_q.get("hint","")}</div>',
                        unsafe_allow_html=True,
                    )
                    _ans = st.text_input(
                        f"Câu trả lời {_qi+1}",
                        key=f"af_ans_{_qi}",
                        placeholder="Nhập kết quả quan sát trực tiếp...",
                        label_visibility="collapsed",
                    )
                    st.session_state.cl_af_answers[_qi] = _ans
                    if not _ans.strip():
                        _all_answered = False

                if _all_answered:
                    if st.button("✅ Xác nhận đã kiểm tra trực tiếp", key="af_confirm",
                                 type="primary", use_container_width=True):
                        st.session_state.cl_af_verified = True
                        st.success("✅ Đã xác thực hiện diện — kết quả kiểm tra được ghi nhận.")

                if st.session_state.cl_af_verified:
                    st.markdown(
                        '<div style="background:#DCFCE7;border-radius:6px;padding:6px 12px;'
                        'font-size:0.8rem;color:#166534;font-weight:600">'
                        '🛡️ Đã xác thực — Câu trả lời được lưu vào báo cáo</div>',
                        unsafe_allow_html=True,
                    )

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    cl = get_checklist(level_key)

    # ── FIX: Pre-initialize TRƯỚC KHI render để tránh reset state khi rerun ──
    # Đồng bộ cl_r từ widget state — nguồn sự thật là session_state[seg_{code}]
    # Không pre-init "Chưa chấm" → segmented_control bắt đầu với None (chưa chọn)
    for _, grp_items in cl:
        for item_code, *_ in grp_items:
            st.session_state.cl_r[item_code] = st.session_state.get(f"seg_{item_code}")

    pass_count = fail_count = 0

    for g_idx, (group_name, items) in enumerate(cl):
        # Đếm trạng thái trong nhóm → hiển thị summary trên header
        g_codes   = [c for c, *_ in items]
        g_pass    = sum(1 for c in g_codes if st.session_state.cl_r.get(c) == "✅ Đạt")
        g_fail    = sum(1 for c in g_codes if st.session_state.cl_r.get(c) == "❌ Không Đạt")
        g_pending = len(g_codes) - g_pass - g_fail
        summary_parts = []
        if g_pass:    summary_parts.append(f'<span style="color:#16A34A;font-weight:700">✅ {g_pass} đạt</span>')
        if g_fail:    summary_parts.append(f'<span style="color:#DC2626;font-weight:700">❌ {g_fail} không đạt</span>')
        if g_pending: summary_parts.append(f'<span style="color:#D97706">○ {g_pending} chưa chấm</span>')
        summary_html = (
            '<span style="font-size:0.72rem;margin-left:10px;font-weight:500">'
            + ' &nbsp;·&nbsp; '.join(summary_parts) + '</span>'
        )
        st.markdown(
            f'<div class="group-title">{group_name}{summary_html}</div>',
            unsafe_allow_html=True,
        )

        for (code, is_critical, desc, how, pass_cond, fail_cond) in items:
            # Đọc TRỰC TIẾP từ widget key → không lag, chuyển màu ngay lần bấm đầu tiên
            cur_state = st.session_state.get(f"seg_{code}")
            if cur_state == "✅ Đạt":
                row_left = "#16A34A"; row_bg = "#F0FDF4"
                code_clr = "#166534"; code_icon = "✅"
                state_label = '<span style="font-size:0.7rem;font-weight:700;color:#16A34A;margin-left:6px">ĐẠT</span>'
            elif cur_state == "❌ Không Đạt":
                row_left = "#DC2626"; row_bg = "#FFF5F5"
                code_clr = "#991B1B"; code_icon = "❌"
                state_label = '<span style="font-size:0.7rem;font-weight:700;color:#DC2626;margin-left:6px">KHÔNG ĐẠT</span>'
            else:  # None — chưa chọn
                row_left = "#F59E0B"; row_bg = "#FFFBEB"
                code_clr = "#D97706"; code_icon = "○"
                state_label = '<span style="font-size:0.7rem;color:#D97706;margin-left:6px">chưa chấm</span>'

            crit_badge = (
                '<span style="background:#FEE2E2;color:#991B1B;font-size:0.65rem;'
                'font-weight:700;padding:1px 6px;border-radius:8px;margin-left:6px;'
                'border:1px solid #FECACA">BẮT BUỘC</span>'
            ) if is_critical else ""

            col_desc, col_ctrl = st.columns([0.62, 0.38])

            with col_desc:
                st.markdown(
                    f'<div style="background:{row_bg};border-left:3px solid {row_left};'
                    f'border-radius:0 8px 8px 0;padding:10px 14px;margin:3px 0;'
                    f'transition:background 0.4s ease,border-color 0.4s ease">'
                    f'<div style="margin-bottom:4px;display:flex;align-items:center;flex-wrap:wrap;gap:4px">'
                    f'<span style="font-size:0.72rem;font-weight:800;color:{code_clr}">'
                    f'{code_icon} {code}</span>'
                    f'{crit_badge}{state_label}'
                    f'</div>'
                    f'<div style="font-size:0.88rem;font-weight:500;color:#1E293B;line-height:1.55">'
                    f'{desc}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                with st.expander("🔍 Hướng dẫn kiểm tra"):
                    st.markdown(
                        f"**Thực hiện:** {how}  \n"
                        f"**✅ Đạt khi:** {pass_cond}  \n"
                        f"**❌ Không đạt khi:** {fail_cond}"
                    )

            with col_ctrl:
                st.segmented_control(
                    label=code,
                    options=["✅ Đạt", "❌ Không Đạt"],
                    key=f"seg_{code}",
                    label_visibility="collapsed",
                )
                result = st.session_state.get(f"seg_{code}")  # None nếu chưa chọn
                st.session_state.cl_r[code] = result

                # Ghi chú — mở khi Không Đạt, đóng khi Đạt/chưa chấm
                with st.expander("📝 Ghi chú", expanded=(result == "❌ Không Đạt")):
                    note = st.text_input(
                        label=f"ghi_chú_{code}", label_visibility="collapsed",
                        placeholder="Ghi chú (nếu có)...", key=f"note_{code}",
                    )
                    st.session_state.cl_n[code] = note

            if result == "✅ Đạt":
                pass_count += 1
            elif result == "❌ Không Đạt":
                fail_count += 1

        # ── Ảnh minh chứng + AI phân tích ────────────────────────────────────
        with st.expander(
            f"📷 Ảnh minh chứng — {group_name}"
            + (" · 🤖 AI phân tích ảnh (~$0.015/ảnh)" if ai_on else "")
        ):
            # Hướng dẫn chụp ảnh đạt chuẩn
            st.markdown("""
            <div style="background:#FFFBEB;border:1px solid #FCD34D;border-radius:8px;
                        padding:10px 14px;margin-bottom:10px;font-size:0.8rem;color:#78350F">
                <b>📐 Chụp ảnh để AI phân tích chính xác:</b><br>
                ✅ Đủ sáng, không ngược sáng &nbsp;|&nbsp;
                ✅ Cách 20–50cm, chụp thẳng &nbsp;|&nbsp;
                ✅ Ảnh nét, không bị mờ &nbsp;|&nbsp;
                ✅ Chụp rõ vùng cần kiểm tra<br>
                ❌ Tránh: tối, mờ nhòe, góc nghiêng >45°, che khuất vùng cần xem<br>
                <span style="color:#92400E"><b>Lưu ý:</b> Claude Vision phân tích dấu hiệu
                <b>nhìn thấy bằng mắt</b> — không thể phát hiện vi khuẩn vô hình
                hay đo nhiệt độ thực tế.</span>
            </div>
            """, unsafe_allow_html=True)

            photo_col1, photo_col2 = st.columns(2)
            with photo_col1:
                st.caption("📱 Chụp ảnh (camera điện thoại/webcam)")
                cam = st.camera_input("Chụp ảnh", key=f"cam_{g_idx}",
                                      label_visibility="collapsed")
                if cam:
                    st.session_state.cl_photos[f"cam_{g_idx}"] = cam
                    st.success("✅ Đã lưu ảnh chụp")
            with photo_col2:
                st.caption("💻 Tải ảnh từ thư viện máy (tối đa 3 ảnh/mục)")
                upl = st.file_uploader(
                    "Tải ảnh", type=["jpg", "jpeg", "png", "heic"],
                    key=f"upl_{g_idx}", label_visibility="collapsed",
                    accept_multiple_files=True,
                )
                if upl:
                    if len(upl) > 3:
                        st.warning("⚠️ Chỉ lưu 3 ảnh đầu tiên (giới hạn 3 ảnh/mục)")
                        upl = upl[:3]
                    st.session_state.cl_photos[f"upl_{g_idx}"] = upl
                    st.success(f"✅ Đã tải {len(upl)}/3 ảnh")

            # ── AI #2: Nút phân tích ảnh ──────────────────────────────────
            active_photo = (
                st.session_state.cl_photos.get(f"cam_{g_idx}") or
                (st.session_state.cl_photos.get(f"upl_{g_idx}") or [None])[0]
            )
            if ai_on and active_photo:
                if st.button(f"🔍 Phân tích ảnh với AI", key=f"analyze_{g_idx}",
                             use_container_width=True):
                    with st.spinner("AI đang phân tích ảnh..."):
                        photo_bytes = (active_photo.read()
                                       if hasattr(active_photo, "read")
                                       else active_photo.getvalue())
                        result = analyze_photo_ai(photo_bytes, group_name, api_key)
                        st.session_state.photo_analysis[g_idx] = result

                # Hiển thị kết quả phân tích
                if g_idx in st.session_state.photo_analysis:
                    r = st.session_state.photo_analysis[g_idx]
                    lvl  = r.get("risk_level", "OK")
                    clr  = {"OK": "#16A34A", "WARNING": "#D97706",
                            "CRITICAL": "#DC2626", "ERROR": "#64748B"}.get(lvl, "#64748B")
                    bg   = {"OK": "#F0FDF4", "WARNING": "#FFFBEB",
                            "CRITICAL": "#FEF2F2", "ERROR": "#F8FAFC"}.get(lvl, "#F8FAFC")
                    icon = {"OK": "✅", "WARNING": "⚠️",
                            "CRITICAL": "🚨", "ERROR": "❓"}.get(lvl, "❓")
                    conf = int(r.get("confidence", 0.8) * 100)

                    issues_html = "".join(
                        f"<li style='color:#DC2626'>{i}</li>" for i in r.get("issues", [])
                    ) or ""
                    pos_html = "".join(
                        f"<li style='color:#16A34A'>{p}</li>" for p in r.get("positives", [])
                    ) or ""

                    st.markdown(f"""
                    <div style="background:{bg};border-left:4px solid {clr};
                                border-radius:8px;padding:12px 16px;margin-top:8px">
                        <div style="font-weight:700;color:{clr};margin-bottom:6px">
                            {icon} Kết quả AI: <b>{lvl}</b>
                            <span style="font-weight:400;font-size:0.75rem;
                                         color:#64748B;margin-left:8px">
                                Độ tin cậy: {conf}%</span>
                        </div>
                        {"<ul style='margin:4px 0;padding-left:16px'>" + issues_html + "</ul>" if issues_html else ""}
                        {"<ul style='margin:4px 0;padding-left:16px'>" + pos_html + "</ul>" if pos_html else ""}
                        {"<div style='font-size:0.82rem;color:#475569;margin-top:6px'><b>Khuyến nghị:</b> " + r.get("recommendation","") + "</div>" if r.get("recommendation") else ""}
                    </div>""", unsafe_allow_html=True)
            elif not ai_on and active_photo:
                st.caption("💡 Kết nối AI để phân tích ảnh tự động.")

        st.markdown("<br>", unsafe_allow_html=True)

    # ── Kết quả & Cảnh báo ───────────────────────────────────────────────────
    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-hdr">Kết quả tổng hợp</div>', unsafe_allow_html=True)

    total_answered = pass_count + fail_count
    critical_fails = CRITICAL_ITEMS & {c for c, v in st.session_state.cl_r.items() if v == "❌ Không Đạt"}

    m1, m2, m3, m4, m5 = st.columns(5)
    m1.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">Đã kiểm tra</div>
        <div class="metric-num c-blue">{total_answered}</div>
        <div class="metric-lbl">/ {TOTAL_ITEMS} điểm</div>
    </div>""", unsafe_allow_html=True)
    m2.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">✅ Đạt</div>
        <div class="metric-num c-green">{pass_count}</div>
        <div class="metric-lbl">điểm</div>
    </div>""", unsafe_allow_html=True)
    m3.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">❌ Không Đạt</div>
        <div class="metric-num c-red">{fail_count}</div>
        <div class="metric-lbl">điểm</div>
    </div>""", unsafe_allow_html=True)
    crit_color = "c-red" if critical_fails else "c-green"
    crit_num   = str(len(critical_fails)) if critical_fails else "✓"
    crit_sub   = "vi phạm" if critical_fails else "Tất cả đạt"
    m4.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">Mục bắt buộc</div>
        <div class="metric-num {crit_color}">{crit_num}</div>
        <div class="metric-lbl">{crit_sub} / {len(CRITICAL_ITEMS)} mục</div>
    </div>""", unsafe_allow_html=True)
    pct = int(pass_count / total_answered * 100) if total_answered else 0
    pct_color = "c-green" if pct >= 90 else "c-orange" if pct >= 75 else "c-red"
    m5.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">Tỷ lệ đạt</div>
        <div class="metric-num {pct_color}">{pct}%</div>
        <div class="metric-lbl">trong số đã chấm</div>
    </div>""", unsafe_allow_html=True)

    # Alert level
    if total_answered >= 10:
        alert_key = determine_alert(st.session_state.cl_r, cl)
        a = ALERT_SYSTEM[alert_key]
        css_cls = f"alert-{alert_key.lower()}"
        notify_html = "".join(f"<li>{n}</li>" for n in a["notify"])
        st.markdown(f"""
        <div class="{css_cls}" style="margin-top:16px">
            <div class="alert-title" style="color:{a['color']}">{a['icon']} {a['label']}</div>
            <div class="alert-body">
                <b>Hành động:</b> {a['action']}<br><br>
                <b>Thông báo đến:</b>
                <ul style="margin:4px 0 0 16px;padding:0">{notify_html}</ul>
                <b>Thời hạn:</b> {a['timeframe']}
            </div>
        </div>
        """, unsafe_allow_html=True)

        if critical_fails:
            codes_str = ", ".join(sorted(critical_fails))
            st.error(f"🚨 Mục bắt buộc vi phạm: **{codes_str}** — Xem lại ngay và báo Ban Giám Hiệu!")

    # ── Thanh tiến độ hoàn thành ─────────────────────────────────────────────
    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    answered_count = sum(1 for v in st.session_state.cl_r.values() if v is not None)
    pct_done = int(answered_count / TOTAL_ITEMS * 100)
    bar_color = "#16A34A" if pct_done == 100 else "#2563EB" if pct_done >= 50 else "#F59E0B"
    remaining  = TOTAL_ITEMS - answered_count
    done_label = "✅ Đã hoàn thành toàn bộ!" if remaining == 0 else f"Còn {remaining} mục chưa đánh giá"
    st.markdown(f"""
    <div style="margin-bottom:12px">
        <div style="display:flex;justify-content:space-between;margin-bottom:4px">
            <span style="font-size:0.82rem;font-weight:600;color:#334155">
                Tiến độ đánh giá: {answered_count}/{TOTAL_ITEMS} mục
            </span>
            <span style="font-size:0.8rem;color:{'#16A34A' if remaining==0 else '#D97706'};font-weight:600">
                {done_label}
            </span>
        </div>
        <div class="completion-bar-wrap">
            <div class="completion-bar-fill"
                 style="width:{pct_done}%;background:{bar_color}"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Validation trước khi xuất báo cáo ────────────────────────────────────
    errors = validate_checklist(st.session_state.cl_r, st.session_state.cl_photos)
    if errors:
        err_html = "".join(f'<div class="validation-item">⚠️ {e}</div>' for e in errors)
        st.markdown(f"""
        <div class="validation-box">
            <div class="validation-title">⛔ Chưa thể xuất báo cáo — Cần hoàn thành các mục sau:</div>
            {err_html}
        </div>
        """, unsafe_allow_html=True)

    can_submit = len(errors) == 0

    # ── Khóa giờ submit (chống đánh trước/sau giờ quy định) ──────────────────
    _now_vn   = now_vn()
    _hour_now = _now_vn.hour + _now_vn.minute / 60
    _is_super = st.session_state.get("is_super", False)
    # BGS: 7:00–13:30 | bỏ qua nếu is_super (admin test)
    _in_window = _is_super or (7.0 <= _hour_now <= 13.5)
    if not _in_window and can_submit:
        st.warning(
            f"⏰ Ngoài giờ kiểm tra hợp lệ (7:00–13:30). Hiện tại: "
            f"{_now_vn.strftime('%H:%M')} — "
            "Kết quả nộp ngoài giờ không được tính vào hồ sơ chính thức."
        )
        can_submit = False

    # ── Hiển thị checklist bổ sung từ AI (nếu có) ────────────────────────────
    if st.session_state.cl_extra:
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="sec-hdr">🤖 Câu hỏi bổ sung AI ({len(st.session_state.cl_extra)} câu) — theo thực đơn hôm nay</div>',
            unsafe_allow_html=True,
        )
        if "cl_extra_r" not in st.session_state:
            st.session_state.cl_extra_r = {}
        for item in st.session_state.cl_extra:
            code = item.get("code", "?")
            desc = item.get("desc", "")
            ingr = item.get("ingredient", "")
            why  = item.get("why", "")
            seg_key = f"seg_extra_{code}"
            # Đọc trực tiếp từ widget key để màu cập nhật ngay
            _ex_cur = st.session_state.get(seg_key)
            if _ex_cur == "✅ Đạt":
                _ex_left, _ex_bg, _ex_icon = "#16A34A", "#F0FDF4", "✅"
                _ex_lbl = '<span style="font-size:0.7rem;font-weight:700;color:#16A34A;margin-left:6px">ĐẠT</span>'
            elif _ex_cur == "❌ Không Đạt":
                _ex_left, _ex_bg, _ex_icon = "#DC2626", "#FFF5F5", "❌"
                _ex_lbl = '<span style="font-size:0.7rem;font-weight:700;color:#DC2626;margin-left:6px">KHÔNG ĐẠT</span>'
            else:
                _ex_left, _ex_bg, _ex_icon = "#2563EB", "#EFF6FF", "🤖"
                _ex_lbl = '<span style="font-size:0.7rem;color:#2563EB;margin-left:6px">chưa chấm</span>'
            col_d, col_c = st.columns([0.65, 0.35])
            with col_d:
                st.markdown(
                    f'<div style="background:{_ex_bg};border-left:3px solid {_ex_left};'
                    f'border-radius:0 8px 8px 0;padding:8px 14px;margin:3px 0">'
                    f'<div style="display:flex;align-items:center;gap:6px;margin-bottom:4px">'
                    f'<span style="font-size:0.7rem;font-weight:800;color:{_ex_left}">{_ex_icon} {code}</span>'
                    f'<span style="font-size:0.72rem;color:#1D4ED8;background:#DBEAFE;'
                    f'padding:1px 6px;border-radius:8px">{ingr}</span>'
                    f'{_ex_lbl}</div>'
                    f'<div style="font-size:0.88rem;font-weight:500;color:#1E293B">{desc}</div>'
                    + (f'<div style="font-size:0.75rem;color:#64748B;margin-top:2px">{why}</div>' if why else '')
                    + '</div>', unsafe_allow_html=True,
                )
                with st.expander("Hướng dẫn"):
                    st.markdown(
                        f"**Kiểm tra:** {item.get('how','')}  \n"
                        f"**✅ Đạt:** {item.get('pass','')}  \n"
                        f"**❌ Không đạt:** {item.get('fail','')}"
                    )
            with col_c:
                st.segmented_control(
                    code, ["✅ Đạt", "❌ Không Đạt"],
                    key=seg_key, label_visibility="collapsed",
                )
                st.session_state.cl_extra_r[code] = st.session_state.get(seg_key)
        if st.button("🗑️ Xoá câu hỏi AI", use_container_width=True):
            st.session_state.cl_extra = []
            st.rerun()

    # ── Nút tạo báo cáo ──────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)

    # Guard chống duplicate save: mỗi tổ hợp (trường + ngày + người kt) chỉ lưu 1 lần
    _save_guard_key = f"cl_saved_{school}_{date}_{insp}"
    _already_saved  = st.session_state.get(_save_guard_key, False)

    if st.button(
        "📄 Tải báo cáo" if can_submit else "⛔ Hoàn thành đủ 20 mục để xuất báo cáo",
        type="primary" if can_submit else "secondary",
        disabled=not can_submit,
        use_container_width=True,
        key="cl_submit_btn",
    ):
        alert_key = determine_alert(st.session_state.cl_r, cl)
        date_vn   = date.strftime("%d/%m/%Y")
        date_iso  = date.strftime("%Y-%m-%d")
        report    = _build_report(school, date_vn, insp, menu, level_key,
                                  st.session_state.cl_r, st.session_state.cl_n,
                                  pass_count, fail_count, alert_key, cl)
        photo_count = sum(
            (len(v) if isinstance(v, list) else 1)
            for v in st.session_state.cl_photos.values() if v
        )

        # ── G1: Auto-save vào Supabase ────────────────────────────────────────
        # ── AI #3: Tóm tắt ngôn ngữ tự nhiên ────────────────────────────────
        narrative = ""
        if ai_on:
            with st.spinner("🤖 AI đang viết tóm tắt báo cáo..."):
                narrative = generate_ai_narrative(
                    st.session_state.cl_r, st.session_state.cl_n,
                    alert_key, school, date_vn, menu,
                    pass_count, pass_count + fail_count, level_key, api_key,
                )
                st.markdown(f"""
                <div style="background:#F8FAFC;border:1px solid #E2E8F0;
                            border-radius:10px;padding:16px 20px;margin-bottom:12px">
                    <div style="font-size:0.75rem;font-weight:700;color:#2563EB;
                                margin-bottom:8px">🤖 TÓM TẮT BÁO CÁO — AI tạo tự động</div>
                    <div style="font-size:0.9rem;color:#1E293B;line-height:1.7">{narrative}</div>
                </div>""", unsafe_allow_html=True)
                report = f"TÓM TẮT (AI):\n{narrative}\n\n{'='*64}\n\n" + report

        if photo_count:
            st.info(f"📷 Kèm {photo_count} ảnh + "
                    f"{len(st.session_state.photo_analysis)} kết quả phân tích AI")

        # ── G1: Auto-save vào Supabase (guard chống duplicate) ───────────────
        if not _already_saved:
            extra_r    = st.session_state.get("cl_extra_r", {})
            session_id = db_save_checklist(
                school, date_iso, insp, menu, level_key,
                st.session_state.cl_r, st.session_state.cl_n,
                alert_key, pass_count, fail_count,
                ai_narrative=narrative,
                extra_results=extra_r or None,
            )
            if session_id:
                st.session_state[_save_guard_key] = True   # Đánh dấu đã lưu
                # DB saved silently
            elif db_ok():
                st.warning("⚠️ Lưu database thất bại — báo cáo vẫn tải được bình thường")
        else:
            st.info("💾 Báo cáo này đã được lưu trước đó.")

        # ── Tải báo cáo Word (.docx) ─────────────────────────────────────────
        with st.spinner("⚙️ Đang tạo file Word..."):
            docx_bytes = generate_word_report(
                school, date_vn, insp, menu, level_key,
                st.session_state.cl_r, st.session_state.cl_n,
                pass_count, fail_count, alert_key, cl,
                narrative,
                st.session_state.photo_analysis,
            )
        fname_docx = f"BaoCao_ATTP_{(school or 'Truong').replace(' ','_')}_{date.strftime('%d-%m-%Y')}.docx"
        st.download_button(
            "⬇️ Tải báo cáo (.docx)",
            data=docx_bytes, file_name=fname_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, type="primary",
        )
        # Vẫn giữ txt để xem nhanh
        with st.expander("Xem trước nội dung (text)"):
            st.text(report)


def _build_report(school, date, insp, menu, level_key, results, notes,
                  pass_count, fail_count, alert_key, cl) -> str:
    a = ALERT_SYSTEM.get(alert_key, {})
    lines = [
        "=" * 64,
        "   BÁO CÁO KIỂM TRA AN TOÀN THỰC PHẨM BỮA ĂN HỌC ĐƯỜNG",
        "=" * 64,
        f"   Trường          : {school or '(chưa nhập)'}",
        f"   Cấp học         : {level_key}",
        f"   Ngày kiểm tra   : {date}",
        f"   Người kiểm tra  : {insp or '(chưa nhập)'}",
        f"   Thực đơn        : {menu or '(chưa nhập)'}",
        f"   Thời gian tạo   : {now_vn().strftime('%d/%m/%Y %H:%M')}",
        "",
        f"   KẾT QUẢ   : {pass_count} ĐẠT / {pass_count + fail_count} điểm đã kiểm tra",
        f"   MỨC CẢNH BÁO: {a.get('icon','')} {a.get('label','—')}",
        f"   HÀNH ĐỘNG : {a.get('action','—')}",
        "",
        "-" * 64,
        "   CHI TIẾT TỪNG ĐIỂM KIỂM TRA",
        "-" * 64,
    ]
    item_map = {
        code: (desc, is_crit, group)
        for group, items in cl
        for (code, is_crit, desc, *_) in items
    }
    cur_grp = None
    for code, result in results.items():
        if code not in item_map:
            continue
        desc, is_crit, grp = item_map[code]
        if grp != cur_grp:
            cur_grp = grp
            lines += ["", f"   [ {grp} ]"]
        crit = " [BẮT BUỘC]" if is_crit else ""
        status = result.replace("✅ Đạt", "ĐẠT").replace("❌ Không Đạt", "KHÔNG ĐẠT").replace("Chưa chấm", "—")
        note = notes.get(code, "")
        note_str = f"  → {note}" if note else ""
        lines.append(f"   {code}{crit:<12} [{status:<12}] {desc}{note_str}")

    notify_str = "\n".join(f"     • {n}" for n in a.get("notify", []))

    # Anti-fraud verification section
    _af_qs = cl.get("af_questions", []) if isinstance(cl, dict) else []
    _af_verified = cl.get("af_verified", False) if isinstance(cl, dict) else False
    _af_section = ""
    if _af_qs or _af_verified:
        _af_section = "\n   XÁC THỰC HIỆN DIỆN (Anti-Fraud):\n"
        for i, q in enumerate(_af_qs):
            _ans = cl.get(f"af_ans_{i}", "—") if isinstance(cl, dict) else "—"
            _af_section += f"     Q{i+1}: {q.get('q','')}\n     A{i+1}: {_ans}\n"
        _af_section += f"     Trạng thái: {'✓ Đã xác thực' if _af_verified else '⚠ Chưa xác thực'}\n"

    lines += [
        "", "-" * 64,
        "   THÔNG BÁO ĐẾN:", notify_str,
        "", "-" * 64,
        "   Căn cứ pháp lý:",
        "     • NĐ 15/2018/NĐ-CP — Điều kiện ATTP cơ sở kinh doanh thực phẩm",
        "     • TTLT 13/2016/TTLT-BYT-BGDĐT — Y tế trường học",
        "     • QĐ 3958/QĐ-BYT 2025 — Dinh dưỡng bữa ăn học đường",
        "     • Đường dây nóng Cục ATTP: 1800 6838 (miễn phí)",
        "-" * 64,
        "   Ký tên người kiểm tra: _____________________________",
        "   SchoolFood AI v1.1",
        "=" * 64,
    ]
    return "\n".join(lines)


# ── TAB 3: Lịch & Thông báo ───────────────────────────────────────────────────
def tab_schedule():
    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">📅 Tần Suất Kiểm Tra & Hệ Thống Thông Báo</div>
        <div class="sf-card-body">Quy định rõ ai kiểm tra gì, khi nào, báo cáo ai — để không có khoảng trống trách nhiệm.</div>
    </div>""", unsafe_allow_html=True)

    st.markdown('<div class="sec-hdr">Lịch kiểm tra theo vai trò</div>', unsafe_allow_html=True)

    for s in SCHEDULE:
        st.markdown(f"""
        <div class="schedule-card" style="--sc-color:{s['color']}">
            <div class="schedule-role">{s['role']}</div>
            <div class="schedule-row"><span class="schedule-key">🔁 Tần suất:</span> {s['freq']}</div>
            <div class="schedule-row"><span class="schedule-key">⏰ Thời điểm:</span> {s['when']}</div>
            <div class="schedule-row"><span class="schedule-key">📋 Nội dung:</span> {s['what']}</div>
            <div class="schedule-row"><span class="schedule-key">📢 Báo trước:</span> {s['notice']}</div>
            <div class="schedule-row"><span class="schedule-key">📤 Báo cáo:</span> {s['report']}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-hdr">Hệ thống cảnh báo theo cấp độ</div>', unsafe_allow_html=True)

    for key, a in ALERT_SYSTEM.items():
        triggers_html = "".join(f"<li>{t}</li>" for t in a["triggers"])
        notify_html   = "".join(f"<li>{n}</li>" for n in a["notify"])
        css_cls = f"alert-{key.lower()}"
        st.markdown(f"""
        <div class="{css_cls}" style="margin-bottom:12px">
            <div class="alert-title" style="color:{a['color']}">{a['icon']} {a['label']}</div>
            <div class="alert-body">
                <b>Kích hoạt khi:</b>
                <ul style="margin:4px 0 8px 16px;padding:0">{triggers_html}</ul>
                <b>Thông báo đến:</b>
                <ul style="margin:4px 0 8px 16px;padding:0">{notify_html}</ul>
                <b>Thời hạn xử lý:</b> {a['timeframe']}<br>
                <b>Hành động:</b> {a['action']}
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-hdr">Tiêu chí đánh giá tổng thể</div>', unsafe_allow_html=True)

    st.markdown("""<div class="sf-card">
    <table style="width:100%;border-collapse:collapse;font-size:0.875rem">
    <tr style="background:#F8FAFC;font-weight:600;color:#475569">
        <td style="padding:10px 14px;border-bottom:2px solid #E2E8F0">Điều kiện</td>
        <td style="padding:10px 14px;border-bottom:2px solid #E2E8F0">Kết quả</td>
        <td style="padding:10px 14px;border-bottom:2px solid #E2E8F0">Mức cảnh báo</td>
    </tr>
    <tr>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9">Có bất kỳ mục BẮT BUỘC nào không đạt</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9;color:#DC2626;font-weight:600">Tạm dừng bữa ăn</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9"><span style="background:#FEE2E2;color:#991B1B;padding:2px 8px;border-radius:8px;font-size:0.78rem;font-weight:700">🔴 CRITICAL</span></td>
    </tr>
    <tr style="background:#FAFAFA">
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9">Tổng điểm đạt dưới 15/20</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9;color:#D97706;font-weight:600">Khắc phục trong ngày</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9"><span style="background:#FEF9C3;color:#854D0E;padding:2px 8px;border-radius:8px;font-size:0.78rem;font-weight:700">🟠 MAJOR</span></td>
    </tr>
    <tr>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9">Tổng điểm đạt 15–17/20</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9;color:#CA8A04;font-weight:600">Cải thiện trong tuần</td>
        <td style="padding:10px 14px;border-bottom:1px solid #F1F5F9"><span style="background:#FEFCE8;color:#854D0E;padding:2px 8px;border-radius:8px;font-size:0.78rem;font-weight:700">🟡 MINOR</span></td>
    </tr>
    <tr style="background:#FAFAFA">
        <td style="padding:10px 14px">Tất cả mục BẮT BUỘC đạt + Tổng ≥ 18/20</td>
        <td style="padding:10px 14px;color:#16A34A;font-weight:600">Lưu hồ sơ bình thường</td>
        <td style="padding:10px 14px"><span style="background:#DCFCE7;color:#166534;padding:2px 8px;border-radius:8px;font-size:0.78rem;font-weight:700">✅ ĐẠT CHUẨN</span></td>
    </tr>
    </table>
    <div style="font-size:0.78rem;color:#94A3B8;margin-top:12px">
        * 7 mục BẮT BUỘC: C03 (hạn dùng) · C07 (nhiệt độ nhận) · C09 (nhiệt độ chia) · C10 (thời gian nấu) · C11 (màu/mùi) · C18 (sổ kiểm thực) · C20 (mẫu lưu)
    </div>
    </div>""", unsafe_allow_html=True)


# ── TAB 4: Khẩn cấp ──────────────────────────────────────────────────────────
def tab_emergency(api_key: str = ""):
    st.markdown('<div class="emergency-header">🚨 XỬ LÝ KHẨN CẤP KHI NGHI NGỜ NGỘ ĐỘC THỰC PHẨM</div>',
                unsafe_allow_html=True)

    # ── AI #7: Chế độ ứng phó sự cố trực tiếp ───────────────────────────────
    if api_key:
        st.markdown("""<div class="sf-card" style="border:2px solid #DC2626">
            <div class="sf-card-title" style="color:#DC2626">
                🤖 Chế độ ứng phó sự cố — AI hỗ trợ trực tiếp
            </div>
            <div class="sf-card-body">
                AI sẽ hỏi từng bước, ghi nhận timeline và tạo biên bản sự cố tự động.
            </div>
        </div>""", unsafe_allow_html=True)

        if "incident_active" not in st.session_state:
            st.session_state.incident_active = False
        if "incident_history" not in st.session_state:
            st.session_state.incident_history = []
        if "incident_log" not in st.session_state:
            st.session_state.incident_log = []

        col_start, col_end = st.columns(2)
        if col_start.button("🚨 Bắt đầu xử lý sự cố", type="primary",
                            use_container_width=True,
                            disabled=st.session_state.incident_active):
            st.session_state.incident_active = True
            st.session_state.incident_history = []
            st.session_state.incident_log = [
                f"[{now_vn().strftime('%H:%M')}] — Bắt đầu ghi nhận sự cố"
            ]
            st.rerun()

        if col_end.button("⏹ Kết thúc & Tạo biên bản", use_container_width=True,
                          disabled=not st.session_state.incident_active):
            st.session_state.incident_active = False
            st.rerun()

        if st.session_state.incident_active:
            st.error("🔴 **ĐANG XỬ LÝ SỰ CỐ** — Trả lời AI từng bước")

            for msg in st.session_state.incident_history:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

            # Khởi động với câu hỏi đầu tiên
            if not st.session_state.incident_history:
                client  = anthropic.Anthropic(api_key=api_key)
                opening = incident_ai_response(client, [],
                    "Tôi nghi ngờ có học sinh bị ngộ độc thực phẩm sau bữa ăn. Tôi cần làm gì?")
                st.session_state.incident_history.append(
                    {"role": "assistant", "content": opening})
                st.rerun()

            user_msg = st.chat_input("Mô tả tình huống...", key="incident_input")
            if user_msg:
                st.session_state.incident_history.append(
                    {"role": "user", "content": user_msg})
                st.session_state.incident_log.append(
                    f"[{now_vn().strftime('%H:%M')}] Người dùng: {user_msg}")
                client = anthropic.Anthropic(api_key=api_key)
                with st.spinner("AI đang phân tích..."):
                    ai_reply = incident_ai_response(
                        client,
                        [{"role": m["role"], "content": m["content"]}
                         for m in st.session_state.incident_history[:-1]],
                        user_msg,
                    )
                st.session_state.incident_history.append(
                    {"role": "assistant", "content": ai_reply})
                st.session_state.incident_log.append(
                    f"[{now_vn().strftime('%H:%M')}] AI: {ai_reply[:100]}...")
                st.rerun()

        elif st.session_state.incident_log:
            # Hiển thị biên bản sau khi kết thúc
            log_text = "\n".join(st.session_state.incident_log)
            st.download_button(
                "📄 Tải biên bản sự cố (.txt)",
                data=f"BIÊN BẢN SỰ CỐ ATTP\n{'='*40}\n{log_text}",
                file_name=f"SuCo_ATTP_{now_vn().strftime('%d-%m-%Y_%H%M')}.txt",
                mime="text/plain", use_container_width=True,
            )
            with st.expander("Xem biên bản"):
                st.text(log_text)

        # Khi kết thúc sự cố → lưu ngày để Phase 2 tự điền
        if not st.session_state.incident_active and st.session_state.incident_log:
            st.session_state["incident_trace_date"] = now_vn().strftime("%Y-%m-%d")
            st.markdown(
                '<div style="background:#FEF2F2;border:1.5px solid #FCA5A5;border-radius:8px;'
                'padding:10px 16px;margin:8px 0;font-size:0.85rem;color:#991B1B">'
                '🔍 <b>Sự cố đã ghi nhận.</b> Tiếp theo: Dùng công cụ '
                '"Điều tra & Báo cáo" bên dưới để truy vết nguyên nhân và tạo báo cáo gửi cấp trên.'
                '</div>', unsafe_allow_html=True,
            )
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
        st.markdown("**Hoặc dùng hướng dẫn tĩnh bên dưới:**")
    steps = [
        ("🛑", "Bước 1 — DỪNG BỮA ĂN NGAY",
         "Yêu cầu tất cả học sinh ngừng ăn. Không để thêm bất kỳ ai ăn thêm."),
        ("📞", "Bước 2 — GỌI CẤP CỨU 115 nếu cần",
         "Gọi **115** ngay nếu có học sinh: co giật, khó thở, mất ý thức, nôn ra máu."),
        ("🥣", "Bước 3 — GIỮ NGUYÊN MẪU THỨC ĂN",
         "**Không vứt, không rửa, không đổ** bất kỳ thức ăn nào. Đây là bằng chứng xét nghiệm nguyên nhân."),
        ("🔔", "Bước 4 — BÁO NGAY HIỆU TRƯỞNG & Y TẾ HỌC ĐƯỜNG",
         "Gọi điện trực tiếp (không nhắn tin). Cung cấp: số học sinh bị, triệu chứng, giờ ăn, tên món."),
        ("📝", "Bước 5 — GHI CHÉP ĐẦY ĐỦ",
         "Ghi ngay: số học sinh ảnh hưởng · triệu chứng cụ thể · thời gian bắt đầu · các món đã ăn · diễn biến theo thời gian."),
        ("🏥", "Bước 6 — BÁO SỞ Y TẾ trong 24 giờ",
         "Từ 2 người bị trở lên: bắt buộc báo Sở Y Tế địa phương trong 24h theo quy định TTLT 13/2016."),
    ]
    for icon, title, body in steps:
        st.markdown(f"""<div class="sf-card">
            <div class="sf-card-title">{icon} {title}</div>
            <div class="sf-card-body">{body}</div>
        </div>""", unsafe_allow_html=True)

    # ── Task#6: Truy vết sự cố ngộ độc — BGH + Y Tế ────────────────────────
    _role_em = st.session_state.get("user_role","")
    _can_trace = _role_em in ("Ban Giám Hiệu","Y Tế Học Đường") or st.session_state.get("is_super")
    if _can_trace and db_ok():
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="background:linear-gradient(135deg,#1E3A5F,#7F1D1D);'
            'border-radius:12px;padding:14px 20px;margin-bottom:12px">'
            '<div style="color:white;font-size:1rem;font-weight:700">'
            '🔍 Giai đoạn 2 — Điều tra & Báo cáo</div>'
            '<div style="color:#FECACA;font-size:0.8rem">'
            'Truy vết nguyên nhân: bữa ăn, NCC, kết quả kiểm tra ngày đó · '
            'Claude AI soạn báo cáo gửi cấp trên</div></div>',
            unsafe_allow_html=True,
        )
        # Auto pre-fill từ ngày sự cố Phase 1
        _tr_default = now_vn().date()
        if st.session_state.get("incident_trace_date"):
            try:
                import datetime as _dtem
                _tr_default = _dtem.date.fromisoformat(st.session_state["incident_trace_date"])
            except Exception: pass
        _tr_c1, _tr_c2 = st.columns([1, 2])
        _tr_date = _tr_c1.date_input("📅 Ngày sự cố", value=_tr_default,
                                      key="tr_date_em", format="DD/MM/YYYY")
        _tr_school = _tr_c2.text_input("🏫 Trường",
                                        value=st.session_state.get("user_school",""),
                                        key="tr_school_em", placeholder="Tên trường...")

        if st.button("🔍 Truy vết ngay", key="tr_search_em", type="primary"):
            _tr_ds = str(_tr_date)
            try:
                _tr_ses = (_get_sb().table("checklist_sessions").select("*")
                           .eq("check_date", _tr_ds)
                           .eq("school_name", _tr_school).execute().data or [])
            except Exception:
                _tr_ses = []

            if _tr_ses:
                st.markdown(
                    f'<div style="background:#FFF5F5;border:1px solid #FCA5A5;'
                    f'border-radius:10px;padding:14px 18px;margin:8px 0">'
                    f'<div style="font-weight:700;color:#991B1B;margin-bottom:8px">'
                    f'📋 Kết quả truy vết ngày {_tr_date.strftime("%d/%m/%Y")} — {_tr_school}</div>',
                    unsafe_allow_html=True,
                )
                _tr_summary = []
                for _s in _tr_ses:
                    _ct = _s.get("check_type","")
                    _lbl = {"ban_giam_sat":"BGS Checklist 20 điểm",
                            "kiem_thuc_3_buoc":"Y Tế Kiểm thực 3 bước",
                            "nha_cung_cap":"NCC Giao hàng"}.get(_ct, _ct)
                    _al = _s.get("alert_level","OK")
                    _alc = "#DC2626" if _al=="CRITICAL" else "#D97706" if _al=="MAJOR" else "#16A34A"
                    _menu = (_s.get("menu_today","") or "Không có thông tin")[:80]
                    _insp = _s.get("inspector_name","—")
                    st.markdown(
                        f'<div style="background:white;border-radius:6px;padding:8px 12px;'
                        f'margin:4px 0;border-left:3px solid {_alc}">'
                        f'<b>{_lbl}</b> · Người KT: {_insp}<br>'
                        f'Thực đơn: {_menu}<br>'
                        f'Kết quả: <span style="color:{_alc};font-weight:700">{_al}</span> '
                        f'({_s.get("pass_count",0)}/{_s.get("total_items",0)} điểm đạt)'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                    _tr_summary.append(
                        f"- {_lbl}: {_al}, {_s.get('pass_count',0)}/{_s.get('total_items',0)} đạt, "
                        f"thực đơn: {_menu[:50]}"
                    )
                st.markdown('</div>', unsafe_allow_html=True)

                # AI Report
                if api_key:
                    if st.button("🤖 AI tạo báo cáo sự cố gửi cấp trên", key="tr_ai_em",
                                  type="primary", use_container_width=True):
                        _tr_prompt = (
                            f"Bạn là chuyên gia ATTP trường học Việt Nam. "
                            f"Có nghi ngờ ngộ độc thực phẩm tại trường {_tr_school} ngày {_tr_date.strftime('%d/%m/%Y')}.\n\n"
                            f"Dữ liệu kiểm tra hệ thống:\n" + "\n".join(_tr_summary) + "\n\n"
                            f"Hãy viết báo cáo gửi cấp trên theo mẫu văn bản hành chính Việt Nam, gồm:\n"
                            f"1. Tiêu đề + quốc hiệu\n"
                            f"2. Tóm tắt sự cố (ngày, địa điểm, số người nghi ngờ)\n"
                            f"3. Phân tích kết quả kiểm tra trước sự cố (có đáng ngờ không)\n"
                            f"4. Nguyên nhân có thể (theo HACCP)\n"
                            f"5. Biện pháp đã/cần thực hiện ngay\n"
                            f"6. Kiến nghị (tạm dừng, xét nghiệm, báo Sở Y Tế)\n"
                            f"7. Chữ ký placeholder\n"
                            f"Văn phong chính thức, súc tích, dẫn điều khoản pháp luật liên quan."
                        )
                        with st.spinner("🤖 Claude đang soạn báo cáo..."):
                            try:
                                _tr_client = anthropic.Anthropic(api_key=api_key)
                                _tr_resp = _tr_client.messages.create(
                                    model=MODEL, max_tokens=1200,
                                    messages=[{"role":"user","content":_tr_prompt}]
                                )
                                _tr_report = _tr_resp.content[0].text if _tr_resp.content else ""
                                st.session_state["tr_ai_report"] = _tr_report
                            except Exception as _tre:
                                st.error(f"Lỗi AI: {_tre}")

                if st.session_state.get("tr_ai_report"):
                    st.markdown(
                        '<div style="background:white;border:1px solid #E2E8F0;'
                        'border-radius:10px;padding:16px 18px;margin-top:8px">'
                        '<div style="font-weight:700;color:#1B3B6F;margin-bottom:10px">'
                        '📄 Báo cáo sự cố — Claude AI soạn thảo</div>',
                        unsafe_allow_html=True,
                    )
                    st.text_area("Nội dung báo cáo (copy để gửi/in)",
                                 value=st.session_state["tr_ai_report"],
                                 height=350, key="tr_report_display",
                                 label_visibility="collapsed")
                    st.caption("💡 Copy toàn bộ nội dung, dán vào Word → chỉnh font Times New Roman 13 → in/gửi.")
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info(f"Không tìm thấy dữ liệu kiểm tra ngày {_tr_date.strftime('%d/%m/%Y')} tại {_tr_school or '(chọn trường)'}.")



    st.markdown('<div class="sec-hdr">Số điện thoại quan trọng</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    c1.markdown('<div class="metric-box"><div class="metric-lbl">Cấp Cứu</div><div class="metric-num c-red">115</div><div class="metric-lbl">Miễn phí · 24/7</div></div>', unsafe_allow_html=True)
    c2.markdown('<div class="metric-box"><div class="metric-lbl">Cục ATTP</div><div class="metric-num c-blue" style="font-size:1.4rem">1800 6838</div><div class="metric-lbl">Miễn phí · Giờ hành chính</div></div>', unsafe_allow_html=True)
    c3.markdown('<div class="metric-box"><div class="metric-lbl">Cảnh Sát</div><div class="metric-num c-orange">113</div><div class="metric-lbl">Khi có hành vi cố ý</div></div>', unsafe_allow_html=True)

# ── TAB 5: Về ứng dụng ───────────────────────────────────────────────────────
# ── Tab riêng cho Phụ Huynh (chỉ xem, không thực hiện checklist) ─────────────
def tab_parent_view(api_key: str = ""):
    """View dành cho Phụ Huynh — kết quả từ DB, traffic light, feedback tracking."""
    import pandas as _pd_pv

    school = st.session_state.get("user_school", "")
    today_str = now_vn().strftime("%Y-%m-%d")

    # ── Section 1: Kết quả an toàn hôm nay (traffic light) ───────────────────
    st.markdown('<div class="sec-hdr">🛡️ Bữa ăn hôm nay an toàn không?</div>',
                unsafe_allow_html=True)

    # Lấy kết quả từ DB (Y Tế + BGS cùng ngày)
    _sessions_today: list = []
    if db_ok():
        try:
            _q = (_get_sb().table("checklist_sessions")
                  .select("check_type,pass_count,total_items,alert_level,menu_today,check_date,school_name")
                  .eq("check_date", today_str)
                  .order("created_at", desc=True)
                  .limit(30))
            if school:
                _q = _q.eq("school_name", school)
            _sessions_today = _q.execute().data or []
        except Exception:
            pass

    _yte_sessions  = [s for s in _sessions_today if s.get("check_type") == "kiem_thuc_3_buoc"]
    _bgs_sessions  = [s for s in _sessions_today if s.get("check_type") == "ban_giam_sat"]

    def _avg_pct(sessions):
        if not sessions: return None
        pcts = [s["pass_count"] / max(s["total_items"], 1) * 100 for s in sessions
                if s.get("total_items", 0) > 0]
        return sum(pcts) / len(pcts) if pcts else None

    _yte_pct = _avg_pct(_yte_sessions)
    _bgs_pct = _avg_pct(_bgs_sessions)

    # Tính tổng hợp: Y Tế 40%, BGS 60% (BGS độc lập hơn)
    if _yte_pct is not None and _bgs_pct is not None:
        _combined_pct = _bgs_pct * 0.60 + _yte_pct * 0.40
        pass  # sources shown via _sources_html pills below
    elif _yte_pct is not None:
        _combined_pct = _yte_pct
        pass
    elif _bgs_pct is not None:
        _combined_pct = _bgs_pct
        pass
    else:
        _combined_pct = None

    # ── Helper: auto-emoji cho từng món ăn ───────────────────────────────────
    def _dish_emoji(dish: str) -> str:
        d = dish.lower()
        if any(k in d for k in ["cơm", "cháo"]): return "🍚"
        if any(k in d for k in ["canh", "súp", "lẩu"]): return "🍲"
        if any(k in d for k in ["phở", "bún", "hủ tiếu", "mì", "miến"]): return "🍜"
        if any(k in d for k in ["gà"]): return "🍗"
        if any(k in d for k in ["tôm"]): return "🦐"
        if any(k in d for k in ["cá", "cua", "mực", "bạch tuộc"]): return "🐟"
        if any(k in d for k in ["heo", "thịt", "sườn", "chả", "xúc xích"]): return "🥩"
        if any(k in d for k in ["bò"]): return "🥩"
        if any(k in d for k in ["trứng"]): return "🥚"
        if any(k in d for k in ["rau", "xào", "luộc", "salad"]): return "🥬"
        if any(k in d for k in ["đậu", "hũ", "tofu"]): return "🫘"
        if any(k in d for k in ["bánh", "sandwich"]): return "🥐"
        if any(k in d for k in ["trái cây", "tráng miệng"]): return "🍎"
        if any(k in d for k in ["sữa", "yaourt"]): return "🥛"
        return "🍽️"

    if _combined_pct is not None:
        _pct_r = round(_combined_pct)
        if _pct_r >= 90:
            _grad   = "linear-gradient(135deg, #16A34A 0%, #15803D 100%)"
            _icon   = "🛡️"
            _lbl    = "BỮA ĂN AN TOÀN"
            _sub    = "Bé có thể ăn ngon và an toàn hôm nay!"
            _pill_c = "rgba(255,255,255,0.25)"
        elif _pct_r >= 75:
            _grad   = "linear-gradient(135deg, #D97706 0%, #B45309 100%)"
            _icon   = "⚠️"
            _lbl    = "CẦN THEO DÕI"
            _sub    = "Có một số điểm cần cải thiện — nhà trường đang xử lý"
            _pill_c = "rgba(255,255,255,0.25)"
        else:
            _grad   = "linear-gradient(135deg, #DC2626 0%, #991B1B 100%)"
            _icon   = "🚨"
            _lbl    = "CÓ VẤN ĐỀ"
            _sub    = "Nhà trường đang khắc phục — Ban Giám Hiệu đã được thông báo"
            _pill_c = "rgba(255,255,255,0.2)"

        _sources_html = ""
        if _yte_pct is not None:
            _sources_html += (f'<span style="background:rgba(255,255,255,0.18);border-radius:12px;'
                              f'padding:3px 12px;font-size:0.75rem;color:white;margin:3px">'
                              f'🏥 Y Tế Học Đường {_yte_pct:.0f}%</span>')
        if _bgs_pct is not None:
            _sources_html += (f'<span style="background:rgba(255,255,255,0.18);border-radius:12px;'
                              f'padding:3px 12px;font-size:0.75rem;color:white;margin:3px">'
                              f'👥 Ban Giám Sát {_bgs_pct:.0f}%</span>')

        st.markdown(
            f'<div style="background:{_grad};border-radius:20px;padding:32px 24px;'
            f'text-align:center;position:relative;overflow:hidden;margin-bottom:14px;'
            f'box-shadow:0 8px 32px rgba(0,0,0,0.15)">'
            # decorative circles
            f'<div style="position:absolute;top:-40px;right:-40px;width:180px;height:180px;'
            f'border-radius:50%;background:rgba(255,255,255,0.06)"></div>'
            f'<div style="position:absolute;bottom:-50px;left:-30px;width:140px;height:140px;'
            f'border-radius:50%;background:rgba(255,255,255,0.05)"></div>'
            # icon
            f'<div style="font-size:4rem;line-height:1;margin-bottom:12px">{_icon}</div>'
            # label
            f'<div style="color:white;font-size:2rem;font-weight:900;letter-spacing:0.05em;'
            f'text-shadow:0 2px 8px rgba(0,0,0,0.2)">{_lbl}</div>'
            # sub
            f'<div style="color:rgba(255,255,255,0.88);font-size:1rem;margin-top:8px">{_sub}</div>'
            # score pill
            f'<div style="display:inline-block;background:{_pill_c};border-radius:24px;'
            f'padding:8px 24px;margin-top:16px;border:1px solid rgba(255,255,255,0.3)">'
            f'<span style="color:white;font-size:1.4rem;font-weight:800">{_pct_r}%</span>'
            f'<span style="color:rgba(255,255,255,0.75);font-size:0.85rem;margin-left:6px">'
            f'đạt chuẩn ATTP</span></div>'
            # sources
            f'<div style="margin-top:16px;display:flex;justify-content:center;flex-wrap:wrap;gap:4px">'
            f'{_sources_html}</div>'
            # timestamp
            f'<div style="color:rgba(255,255,255,0.55);font-size:0.75rem;margin-top:12px">'
            f'Cập nhật: {now_vn().strftime("%H:%M · %d/%m/%Y")}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        # Chưa có kết quả — design hấp dẫn
        _now_h = now_vn().hour
        _next_check = "9:30–11:00" if _now_h < 9 else ("Đang trong giờ kiểm tra" if _now_h < 11 else "Đã qua giờ Y Tế — chờ BGS kiểm tra")
        st.markdown(
            '<div style="background:linear-gradient(135deg,#F8FAFC,#EFF6FF);'
            'border:2px dashed #BFDBFE;border-radius:20px;padding:32px 24px;text-align:center;'
            'margin-bottom:14px">'
            '<div style="font-size:3.5rem;margin-bottom:12px">⏳</div>'
            '<div style="font-size:1.2rem;font-weight:800;color:#1E293B">'
            'Kết quả chưa có hôm nay</div>'
            '<div style="font-size:0.88rem;color:#64748B;margin-top:8px;line-height:1.8">'
            f'🏥 Y Tế Học Đường kiểm tra: <b>9:30–11:00</b><br>'
            f'👥 Ban Giám Sát kiểm tra <b>2 lần/tuần</b>'
            '</div>'
            f'<div style="display:inline-block;background:#DBEAFE;border-radius:12px;'
            f'padding:6px 18px;margin-top:14px;font-size:0.82rem;color:#1D4ED8;font-weight:600">'
            f'⏱️ {_next_check}</div>'
            '</div>',
            unsafe_allow_html=True,
        )

    # ── Section 2: Thực đơn hôm nay ──────────────────────────────────────────
    st.markdown('<div class="sec-hdr">🍱 Thực đơn hôm nay</div>', unsafe_allow_html=True)

    # Lấy thực đơn từ session mới nhất trong ngày
    _menu_today = ""
    for _s in _sessions_today:
        _m = (_s.get("menu_today") or "").strip()
        if _m and not _m.startswith("NCC:") and not _m.startswith("Nhà CC:"):
            _menu_today = _m
            break
    if not _menu_today:
        _menu_today = st.session_state.get("shared_menu", "").strip()

    if _menu_today:
        # Parse từng món + auto emoji
        import re as _re
        _dishes = [d.strip() for d in _re.split(r"[,،،\n]+", _menu_today) if d.strip()]
        _dish_html = ""
        for _dish in _dishes:
            _em = _dish_emoji(_dish)
            _dish_html += (
                f'<div style="display:flex;align-items:center;gap:10px;'
                f'padding:8px 12px;border-radius:8px;background:rgba(22,163,74,0.06);'
                f'margin:4px 0">'
                f'<span style="font-size:1.4rem">{_em}</span>'
                f'<span style="font-size:0.92rem;color:#1E293B;font-weight:500">{_dish}</span>'
                f'</div>'
            )
        st.markdown(
            f'<div style="background:white;border:1px solid #E2E8F0;border-radius:14px;'
            f'padding:16px 18px;box-shadow:0 1px 4px rgba(0,0,0,0.05)">'
            f'<div style="font-size:0.75rem;color:#16A34A;font-weight:700;margin-bottom:8px;'
            f'display:flex;align-items:center;gap:6px">'
            f'<span style="background:#DCFCE7;border-radius:8px;padding:2px 10px">✅ Thực đơn đã được cập nhật</span>'
            f'</div>'
            f'{_dish_html}'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div class="sf-card" style="border-left:4px solid #94A3B8">'
            '<div style="font-size:0.85rem;color:#64748B;line-height:1.7">'
            '📋 Thực đơn hôm nay chưa được cập nhật trên hệ thống.<br>'
            '👉 Xem thực đơn tại bảng thông báo trước phòng bếp '
            'hoặc hỏi Y tế học đường, giáo viên chủ nhiệm.'
            '</div></div>',
            unsafe_allow_html=True,
        )

    # ── Section 3: Lịch sử 7 ngày ────────────────────────────────────────────
    st.markdown('<div class="sec-hdr">📅 Lịch sử 7 ngày gần nhất</div>',
                unsafe_allow_html=True)
    if db_ok():
        try:
            _hq = (_get_sb().table("checklist_sessions")
                   .select("check_date,check_type,pass_count,total_items,alert_level")
                   .in_("check_type", ["ban_giam_sat", "kiem_thuc_3_buoc"])
                   .order("check_date", desc=True).limit(60))
            if school:
                _hq = _hq.eq("school_name", school)
            _hist = _hq.execute().data or []
        except Exception:
            _hist = []

        if _hist:
            _by_date: dict = {}
            for h in _hist:
                _d = h.get("check_date", "")
                _ti = h.get("total_items") or 0
                if not _d or _ti == 0: continue
                _by_date.setdefault(_d, []).append(h["pass_count"] / _ti * 100)
            _dates = sorted(_by_date.keys())[-7:]
            _DAY_VN = ["T2","T3","T4","T5","T6","T7","CN"]
            if _dates:
                _cols = st.columns(len(_dates))
                for i, _d in enumerate(_dates):
                    _avg = round(sum(_by_date[_d]) / len(_by_date[_d]))
                    _bg  = "#DCFCE7" if _avg >= 90 else "#FEF9C3" if _avg >= 75 else "#FEE2E2"
                    _tc  = "#16A34A" if _avg >= 90 else "#D97706" if _avg >= 75 else "#DC2626"
                    _dd  = f"{_d[8:10]}/{_d[5:7]}" if len(_d) >= 10 else _d
                    try:
                        import datetime as _dmod
                        _dow = _DAY_VN[_dmod.date.fromisoformat(_d).weekday()]
                    except Exception:
                        _dow = ""
                    _cols[i].markdown(
                        f'<div style="text-align:center;padding:6px 2px">'
                        f'<div style="font-size:0.65rem;color:#94A3B8;font-weight:600;'
                        f'text-transform:uppercase;margin-bottom:4px">{_dow}</div>'
                        f'<div style="background:{_bg};border-radius:10px;padding:10px 6px;'
                        f'border:2px solid {_tc}20">'
                        f'<div style="font-size:1.5rem">{"🟢" if _avg>=90 else "🟡" if _avg>=75 else "🔴"}</div>'
                        f'<div style="font-size:0.85rem;font-weight:800;color:{_tc};margin-top:4px">{_avg}%</div>'
                        f'<div style="font-size:0.65rem;color:#64748B;margin-top:2px">{_dd}</div>'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("Chưa có lịch sử kiểm tra.")
        else:
            st.caption("Chưa có lịch sử kiểm tra.")
    else:
        st.caption("Kết nối database để xem lịch sử.")

    # ── Section 4: Gửi phản hồi ───────────────────────────────────────────────
    st.markdown(
        '<div style="background:linear-gradient(135deg,#EFF6FF,#F5F3FF);'
        'border-radius:14px;padding:20px 22px;margin:14px 0 10px 0;'
        'border:1px solid #DBEAFE">'
        '<div style="font-size:1rem;font-weight:700;color:#1E293B;margin-bottom:4px">'
        '📤 Gửi phản hồi về bữa ăn</div>'
        '<div style="font-size:0.82rem;color:#64748B">'
        'Ý kiến của bạn giúp cải thiện bữa ăn cho các bé — Ban Giám Hiệu sẽ xem xét trong 1–2 ngày</div>'
        '</div>',
        unsafe_allow_html=True,
    )
    loai = st.selectbox(
        "Loại phản hồi",
        ["— Chọn loại phản hồi —",
         "🍽️ Chất lượng thức ăn (khẩu phần, hương vị)",
         "🧹 Vệ sinh bếp ăn (nghi ngờ không sạch)",
         "🚨 Nghi ngờ ngộ độc / dấu hiệu bất thường",
         "🥗 Thiếu dinh dưỡng theo chuẩn",
         "📋 Thực đơn không khớp thông báo",
         "💬 Góp ý khác"],
        label_visibility="collapsed",
    )
    noi_dung = st.text_area(
        "Mô tả cụ thể", height=100,
        placeholder="VD: Hôm nay 03/06, con kể thức ăn có mùi lạ ở bữa trưa...",
        label_visibility="collapsed",
    )
    if st.button("📤 Gửi phản hồi", type="primary", use_container_width=True):
        if loai.startswith("—") or not noi_dung.strip():
            st.warning("⚠️ Vui lòng chọn loại và điền nội dung.")
        elif "Ngộ độc" in loai:
            st.error("🚨 Nghi ngờ ngộ độc: gọi **115** ngay và xem tab **🚨 Khẩn cấp**.")
        else:
            _school_fb = school or st.session_state.get("kt_school", "") or "Chưa nhập"
            saved = db_save_feedback(_school_fb, loai, noi_dung)
            if saved:
                st.session_state["ph_last_feedback"] = {"loai": loai, "noi_dung": noi_dung,
                                                          "time": now_vn().strftime("%H:%M %d/%m")}
                st.success("✅ Đã ghi nhận — Ban Giám Hiệu xem xét trong 1–2 ngày làm việc.")
            else:
                st.warning("Chưa kết nối database — phản hồi chưa được lưu.")

    # ── Section 5: Theo dõi phản hồi trong trường ────────────────────────────
    st.markdown('<div class="sec-hdr">📬 Theo dõi phản hồi — Phụ Huynh trong trường</div>',
                unsafe_allow_html=True)

    if db_ok():
        _fb_school = school or st.session_state.get("kt_school", "")
        try:
            _fq = (_get_sb().table("parent_feedback").select("*")
                   .order("created_at", desc=True).limit(30))
            if _fb_school:
                _fq = _fq.eq("school_name", _fb_school)
            _fbs = _fq.execute().data or []
        except Exception:
            _fbs = []

        if _fbs:
            _ph_open   = [f for f in _fbs if f.get("status") != "resolved"]
            _ph_closed = [f for f in _fbs if f.get("status") == "resolved"]

            # Phản hồi đang chờ/đang xem
            if _ph_open:
                st.markdown(
                    f'<div style="font-size:0.82rem;font-weight:600;color:#D97706;margin:6px 0 4px">'
                    f'⏳ Đang chờ xử lý ({len(_ph_open)})</div>',
                    unsafe_allow_html=True,
                )
                for _fb in _ph_open:
                    _st = _fb.get("status","pending")
                    _st_lbl = "💬 Đang xem xét" if _st=="reviewed" else "⏳ Chờ xử lý"
                    _st_clr = "#2563EB" if _st=="reviewed" else "#D97706"
                    _st_bg  = "#EFF6FF" if _st=="reviewed" else "#FFFBEB"
                    _ev = _fb.get("evidence_text","") or ""
                    _dt_ph = f"{(_fb.get('created_at') or '')[:10].replace('-','/')[8:10]}/{(_fb.get('created_at') or '')[:10][5:7]}/{(_fb.get('created_at') or '')[:10][:4]}"
                    _cat_ph = _fb.get("category","")
                    st.markdown(
                        f'<div style="background:{_st_bg};border:1.5px solid {_st_clr}40;'
                        f'border-left:4px solid {_st_clr};border-radius:10px;'
                        f'padding:12px 16px;margin:6px 0">'
                        f'<div style="display:flex;justify-content:space-between;margin-bottom:5px">'
                        f'<span style="font-size:0.75rem;color:#64748B">📅 {_dt_ph} · {_cat_ph}</span>'
                        f'<span style="background:white;color:{_st_clr};font-size:0.7rem;'
                        f'font-weight:700;padding:2px 10px;border-radius:10px;'
                        f'border:1px solid {_st_clr}40">{_st_lbl}</span>'
                        f'</div>'
                        f'<div style="font-size:0.9rem;color:#1E293B;font-weight:500">'
                        f'{_fb.get("content","")}</div>'
                        + (f'<div style="font-size:0.78rem;color:#2563EB;margin-top:6px;'
                           f'background:white;border-radius:6px;padding:5px 10px">'
                           f'📋 Nhà trường đang xem xét: {_ev}</div>' if _ev else '')
                        + '</div>',
                        unsafe_allow_html=True,
                    )

            # Phản hồi đã xử lý — kiểu khác biệt hoàn toàn
            if _ph_closed:
                st.markdown(
                    f'<div style="background:#1E293B;border-radius:8px;padding:8px 14px;'
                    f'margin:14px 0 6px;font-size:0.82rem;color:#94A3B8">'
                    f'✅ Đã xử lý xong — {len(_ph_closed)} phản hồi</div>',
                    unsafe_allow_html=True,
                )
                for _fb in _ph_closed[:5]:
                    _rep = _fb.get("response_text","") or ""
                    _rby = _fb.get("response_by","") or ""
                    _raw_dt = (_fb.get("created_at") or "")[:10]
                    _dt_c  = f"{_raw_dt[8:10]}/{_raw_dt[5:7]}/{_raw_dt[:4]}" if len(_raw_dt)>=10 else _raw_dt
                    _raw_rdt = (_fb.get("reviewed_at") or "")[:10]
                    _rdt_c = f"{_raw_rdt[8:10]}/{_raw_rdt[5:7]}/{_raw_rdt[:4]}" if len(_raw_rdt)>=10 else _raw_rdt
                    st.markdown(
                        f'<div style="background:#F8FAFC;border:1px solid #E2E8F0;'
                        f'border-radius:8px;padding:10px 14px;margin:4px 0">'
                        f'<div style="display:flex;justify-content:space-between;margin-bottom:4px">'
                        f'<span style="font-size:0.73rem;color:#94A3B8">📅 {_dt_c} · {_fb.get("category","")}</span>'
                        f'<span style="font-size:0.7rem;color:#16A34A;font-weight:600">'
                        f'✅ Đóng ngày {_rdt_c}</span>'
                        f'</div>'
                        f'<div style="font-size:0.82rem;color:#64748B;font-style:italic">'
                        f'"{_fb.get("content","")}"</div>'
                        + (f'<div style="font-size:0.75rem;color:#166534;margin-top:5px;'
                           f'background:#F0FDF4;border-radius:5px;padding:4px 8px">'
                           f'💬 BGH ({_rby}): {_rep}</div>' if _rep else '')
                        + '</div>',
                        unsafe_allow_html=True,
                    )
        else:
            st.info("Chưa có phản hồi nào. Hãy là người đầu tiên gửi ý kiến!")
    else:
        st.caption("Kết nối database để xem phản hồi.")

    # ── Section 6: Quyền pháp lý + liên kết AI ──────────────────────────────
    st.markdown(
        '<div style="background:linear-gradient(135deg,#EFF6FF,#F5F3FF);'
        'border-radius:12px;padding:16px 20px;margin:14px 0 8px 0;'
        'border:1px solid #DBEAFE">'
        '<div style="font-size:0.95rem;font-weight:700;color:#1E293B;margin-bottom:8px">'
        '⚖️ Quyền của Phụ Huynh theo pháp luật</div>'
        '<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">'
        + "".join([
            f'<div style="background:white;border-radius:8px;padding:10px 14px">'
            f'<b style="font-size:0.85rem;color:#1E293B">{t}</b>'
            f'<div style="font-size:0.78rem;color:#64748B;margin-top:3px">{d}</div>'
            f'</div>'
            for t, d in [
                ("📋 Xem thực đơn hàng ngày", "Nhà trường phải công khai thực đơn — bảng trước phòng bếp hoặc app này."),
                ("👥 Yêu cầu giám sát độc lập", "Đề nghị Ban Đại Diện PHHS kiểm tra đột xuất bếp ăn bất kỳ lúc nào."),
                ("📤 Phản ánh chất lượng", "Gửi phản hồi qua form bên trên hoặc liên hệ trực tiếp Hiệu Trưởng."),
                ("📊 Tiếp cận kết quả kiểm tra", "Báo cáo ATTP là tài liệu công khai — yêu cầu Ban Giám Hiệu cung cấp."),
            ]
        ])
        + '</div>'
        + '<div style="margin-top:12px;background:rgba(124,58,237,0.08);border-radius:8px;'
          'padding:10px 14px;font-size:0.82rem;color:#5B21B6;display:flex;align-items:center;gap:8px">'
          '💬 <b>Có thắc mắc về quyền phụ huynh, ATTP, hay cách xử lý khi con bị ảnh hưởng?</b>'
          '&nbsp; → Dùng tab <b>💬 Hỏi đáp AI</b> để đặt câu hỏi bất kỳ lúc nào — '
          'AI trả lời dựa trên pháp luật Việt Nam.'
          '</div>'
        + '</div>',
        unsafe_allow_html=True,
    )

    # ── Hỏi đáp AI ────────────────────────────────────────────────────────────
    if api_key:
        st.markdown('<div class="sec-hdr">💬 Hỏi AI về bữa ăn &amp; quyền Phụ Huynh</div>',
                    unsafe_allow_html=True)
        q = st.text_input("Câu hỏi", placeholder="VD: Con đau bụng sau bữa trưa, tôi cần làm gì?")
        if q:
            from anthropic import Anthropic
            with st.spinner("AI đang trả lời..."):
                _sys = build_system_prompt("Phụ Huynh", "tiểu học", "Việt Nam")
                st.info(ask_claude(Anthropic(api_key=api_key), _sys, [], q))


# ── Kiểm thực 3 bước — Tab dành riêng Y Tế Học Đường ────────────────────────
def tab_kiem_thuc(api_key: str = "", level: str = "Tiểu Học (6–11 tuổi)"):
    """Kiểm thực 3 bước theo TTLT 13/2016 — dành riêng cho Y Tế Học Đường."""
    ai_on = bool(api_key)
    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">🏥 Kiểm thực 3 bước — Y Tế Học Đường</div>
        <div class="sf-card-body">
            Căn cứ: <b>TTLT 13/2016/TTLT-BYT-BGDĐT Điều 9</b> — Thực hiện hàng ngày, song song
            với sổ kiểm thực giấy bắt buộc. Mỗi bước ghi nhận <b>timestamp tự động</b>.
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Chọn cấp học + Tiêu chuẩn dinh dưỡng (collapsible) ──────────────────
    _kt_levels  = ["Tiểu Học (6–11 tuổi)", "THCS (12–15 tuổi)", "THPT (16–18 tuổi)"]
    _kt_default = st.session_state.get("user_profile", {}).get("default_level") or level
    _kt_locked  = bool(st.session_state.get("auth_user")) and not st.session_state.get("is_super", False) and not st.session_state.get("auth_user", {}).get("demo", False)
    if _kt_locked and _kt_default in _kt_levels:
        _kt_level = _kt_default
        st.caption(f"📚 Cấp học: **{_kt_default}** · _Theo tài khoản_")
    else:
        _kt_lvl_idx = _kt_levels.index(_kt_default) if _kt_default in _kt_levels else 0
        _kt_level = st.selectbox(
            "📚 Cấp học đang kiểm thực", _kt_levels, index=_kt_lvl_idx,
            key="kt_level_selector",
            help="Cấp học ảnh hưởng tiêu chuẩn dinh dưỡng",
        )
    level = _kt_level

    n_yte = NUTRITION.get(level, NUTRITION[list(NUTRITION.keys())[0]])
    with st.expander(f"📊 Tiêu chuẩn dinh dưỡng cấp {n_yte['short']} (tham khảo B3_04)"):
        st.markdown(
            f'<div class="nutrition-banner">'
            f'<div class="nutrition-label">📊 Tiêu Chuẩn Dinh Dưỡng Bữa Trưa — Cấp {n_yte["short"]} (QĐ 3958/QĐ-BYT 2025)</div>'
            f'<div class="nutrition-grid">'
            f'<div class="nutrition-item">⚡ Năng lượng: <span class="nutrition-val">{n_yte["kcal"]}</span> ({n_yte["pct_day"]} nhu cầu ngày)</div>'
            f'<div class="nutrition-item">🥩 Thịt/cá tối thiểu: <span class="nutrition-val">{n_yte["meat_g"]}g/học sinh</span></div>'
            f'<div class="nutrition-item">🥦 Rau xanh: <span class="nutrition-val">{n_yte["veg_range"]}/học sinh</span></div>'
            f'<div class="nutrition-item">💪 Protein: <span class="nutrition-val">{n_yte["protein_pct"]}</span> tổng năng lượng</div>'
            f'</div>'
            f'<div style="font-size:0.75rem;color:#1D4ED8;margin-top:6px">💡 {n_yte["note"]}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # Thông tin buổi kiểm tra
    st.markdown('<div class="sec-hdr">Thông tin ca kiểm thực</div>', unsafe_allow_html=True)
    _us_kt = st.session_state.get("user_school", "")
    kc1, kc2, kc3, kc4 = st.columns(4)
    # Tên trường — disabled khi đã có từ profile (đồng nhất style với Ngày/Y Tế/Thực đơn)
    kt_school = kc1.text_input("Tên trường", value=_us_kt,
                                disabled=bool(_us_kt),
                                placeholder="TH Nguyễn Du, Q.1", key="kt_school")
    kt_date   = kc2.date_input("Ngày", value=datetime.today(), format="DD/MM/YYYY",
                                key="kt_date")
    kt_name   = kc3.text_input("Y Tế Học Đường",
                                value=st.session_state.get("user_profile", {}).get("full_name", ""),
                                placeholder="Họ và tên", key="kt_name")
    kt_menu   = kc4.text_input("Thực đơn hôm nay",
                                placeholder="Cơm, thịt kho, rau...",
                                key="kt_menu_yte")

    # ── AI: Tạo câu hỏi bổ sung theo thực đơn ───────────────────────────────
    if "kt_extra" not in st.session_state: st.session_state.kt_extra = []
    if "kt_photo_analysis" not in st.session_state: st.session_state.kt_photo_analysis = {}

    st.markdown(
        f'<div style="background:#F8FAFC;border:1px solid #E2E8F0;border-radius:10px;'
        f'padding:10px 16px;margin-bottom:10px;font-size:0.82rem">'
        f'📋 <b>15 câu kiểm thực chuẩn</b> (luôn có) &nbsp;|&nbsp; '
        f'{"🤖 <b>AI bổ sung</b> theo thực đơn (~$0.004) &nbsp;|&nbsp; 📷 <b>AI phân tích ảnh</b> (~$0.015/ảnh)" if ai_on else "<i style=color:#94A3B8>Kết nối API key để dùng AI phân tích ảnh và tạo câu hỏi theo thực đơn</i>"}'
        f'</div>',
        unsafe_allow_html=True,
    )

    if ai_on and kt_menu:
        cb1, cb2 = st.columns([0.5, 0.5])
        if cb1.button("🤖 Tạo câu hỏi bổ sung theo thực đơn (~$0.004)",
                      use_container_width=True, key="kt_gen_extra"):
            with st.spinner("AI đang phân tích thực đơn..."):
                extras = generate_extra_checklist(kt_menu, "Y Tế Học Đường", api_key)
                st.session_state.kt_extra = extras
        if st.session_state.kt_extra:
            cb2.markdown(
                f'<span style="color:#16A34A;font-size:0.85rem;line-height:3">'
                f'✅ Đã tạo {len(st.session_state.kt_extra)} câu bổ sung</span>',
                unsafe_allow_html=True,
            )
    elif ai_on and not kt_menu:
        st.caption("💡 Nhập thực đơn để AI tạo câu hỏi kiểm tra riêng theo nguyên liệu hôm nay.")

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)

    # Khởi tạo session state cho 3 bước
    for b in [1, 2, 3]:
        if f"kt_b{b}_r" not in st.session_state: st.session_state[f"kt_b{b}_r"] = {}
        if f"kt_b{b}_n" not in st.session_state: st.session_state[f"kt_b{b}_n"] = {}
        if f"kt_b{b}_done" not in st.session_state: st.session_state[f"kt_b{b}_done"] = None
        if f"kt_b{b}_photos" not in st.session_state: st.session_state[f"kt_b{b}_photos"] = {}
        # Không pre-init kt_seg_* → segmented_control bắt đầu None (chưa chọn)

    # ── Render từng bước ──────────────────────────────────────────────────────
    for step in KIEM_THUC:
        b      = step["buoc"]
        clr    = step["color"]
        done_t = st.session_state.get(f"kt_b{b}_done")

        # Header bước
        status_badge = (
            f'<span style="background:#DCFCE7;color:#166534;font-size:0.75rem;'
            f'font-weight:700;padding:2px 10px;border-radius:12px;margin-left:8px">'
            f'✅ Hoàn thành lúc {done_t}</span>' if done_t else
            f'<span style="background:#FEF9C3;color:#92400E;font-size:0.75rem;'
            f'font-weight:700;padding:2px 10px;border-radius:12px;margin-left:8px">'
            f'⏳ Chưa hoàn thành</span>'
        )
        st.markdown(
            f'<div style="display:flex;align-items:center;padding:12px 0 4px 0;'
            f'border-bottom:2px solid {clr}">'
            f'<span style="font-size:1.5rem;margin-right:10px">{step["icon"]}</span>'
            f'<div><div style="font-size:1rem;font-weight:700;color:{clr}">'
            f'{step["label"]}</div>'
            f'<div style="font-size:0.78rem;color:#64748B">'
            f'⏰ {step["time_window"]} &nbsp;·&nbsp; {step["law"]}</div></div>'
            f'{status_badge}</div>',
            unsafe_allow_html=True,
        )
        st.caption(step["desc"])

        # 5 câu hỏi của bước này
        for (code, desc, how, pass_cond, fail_cond, legal) in step["items"]:
            # Đọc trực tiếp từ widget key → màu đổi ngay lần bấm đầu tiên
            cur = st.session_state.get(f"kt_seg_{code}")
            if cur == "✅ Đạt":
                row_bg, row_left = "#F0FDF4", "#16A34A"
                code_clr, code_icon, lbl = "#166534", "✅", "ĐẠT"
            elif cur == "❌ Không Đạt":
                row_bg, row_left = "#FFF5F5", "#DC2626"
                code_clr, code_icon, lbl = "#991B1B", "❌", "KHÔNG ĐẠT"
            else:  # None — chưa chọn
                row_bg, row_left = "#FFFBEB", "#F59E0B"
                code_clr, code_icon, lbl = "#D97706", "○", "chưa kiểm tra"

            col_d, col_c = st.columns([0.62, 0.38])
            with col_d:
                is_critical_item = code == "B3_05"  # Mẫu lưu là bắt buộc pháp lý
                crit_badge = (
                    '<span style="background:#FEE2E2;color:#991B1B;font-size:0.65rem;'
                    'font-weight:700;padding:1px 6px;border-radius:8px;margin-left:4px;'
                    'border:1px solid #FECACA">BẮT BUỘC PHÁP LUẬT</span>'
                    if is_critical_item else ""
                )
                st.markdown(
                    f'<div style="background:{row_bg};border-left:3px solid {row_left};'
                    f'border-radius:0 8px 8px 0;padding:10px 14px;margin:3px 0">'
                    f'<span style="font-size:0.7rem;font-weight:800;color:{code_clr}">'
                    f'{code_icon} {code}</span>'
                    f'<span style="font-size:0.68rem;color:#64748B;margin-left:6px">{legal}</span>'
                    f'{crit_badge}'
                    f'<div style="font-size:0.875rem;font-weight:500;color:#1E293B;margin-top:3px">'
                    f'{desc}</div></div>',
                    unsafe_allow_html=True,
                )
                with st.expander("🔍 Hướng dẫn"):
                    st.markdown(
                        f"**Thực hiện:** {how}  \n"
                        f"**✅ Đạt khi:** {pass_cond}  \n"
                        f"**❌ Không đạt khi:** {fail_cond}"
                    )
            with col_c:
                st.segmented_control(
                    code, ["✅ Đạt", "❌ Không Đạt"],
                    key=f"kt_seg_{code}", label_visibility="collapsed",
                )
                st.session_state[f"kt_b{b}_r"][code] = \
                    st.session_state.get(f"kt_seg_{code}")
                note = st.text_input(
                    f"ghi chú {code}", label_visibility="collapsed",
                    placeholder="Ghi chú...", key=f"kt_note_{code}",
                )
                st.session_state[f"kt_b{b}_n"][code] = note

        # ── Ảnh minh chứng + AI Vision ───────────────────────────────────────
        exp_label = (
            f"📷 Ảnh minh chứng Bước {b} (tối đa 4 ảnh/bước)"
            + (" · 🤖 AI phân tích (~$0.015/ảnh)" if ai_on else "")
        )
        with st.expander(exp_label):
            st.caption("📐 Ảnh nét, đủ sáng, chụp thẳng cách 20–50cm. Camera: 1 ảnh · Tải lên: tối đa 3 ảnh")
            pc1, pc2 = st.columns(2)
            with pc1:
                st.caption("📱 Chụp ảnh (1 ảnh)")
                cam = st.camera_input("Chụp ảnh", key=f"kt_cam_{b}",
                                      label_visibility="collapsed")
                if cam:
                    st.session_state[f"kt_b{b}_photos"]["cam"] = cam
                    st.success("✅ Đã lưu ảnh chụp")
            with pc2:
                st.caption("💻 Tải ảnh từ máy (tối đa 3 ảnh)")
                upl = st.file_uploader("Tải ảnh", type=["jpg","jpeg","png","heic"],
                                       key=f"kt_upl_{b}", label_visibility="collapsed",
                                       accept_multiple_files=True)
                if upl:
                    if len(upl) > 3:
                        st.warning("⚠️ Chỉ lưu 3 ảnh đầu tiên")
                        upl = upl[:3]
                    st.session_state[f"kt_b{b}_photos"]["upl"] = upl
                    st.success(f"✅ Đã tải {len(upl)}/3 ảnh")

            # ── AI Vision phân tích ảnh ───────────────────────────────────
            active_photo = (
                st.session_state[f"kt_b{b}_photos"].get("cam") or
                (st.session_state[f"kt_b{b}_photos"].get("upl") or [None])[0]
            )
            if ai_on and active_photo:
                if st.button(f"🔍 Phân tích ảnh Bước {b} với AI",
                             key=f"kt_analyze_{b}", use_container_width=True):
                    with st.spinner("AI đang phân tích ảnh..."):
                        photo_bytes = (active_photo.read()
                                       if hasattr(active_photo, "read")
                                       else active_photo.getvalue())
                        result = analyze_photo_ai(photo_bytes, step["label"], api_key)
                        st.session_state.kt_photo_analysis[b] = result

                if b in st.session_state.kt_photo_analysis:
                    r = st.session_state.kt_photo_analysis[b]
                    lvl  = r.get("risk_level", "OK")
                    clr  = {"OK":"#16A34A","WARNING":"#D97706","CRITICAL":"#DC2626"}.get(lvl,"#64748B")
                    bg   = {"OK":"#F0FDF4","WARNING":"#FFFBEB","CRITICAL":"#FEF2F2"}.get(lvl,"#F8FAFC")
                    icon = {"OK":"✅","WARNING":"⚠️","CRITICAL":"🚨"}.get(lvl,"❓")
                    conf = int(r.get("confidence",0.8)*100)
                    issues_html = "".join(f"<li style='color:#DC2626'>{i}</li>" for i in r.get("issues",[]))
                    pos_html    = "".join(f"<li style='color:#16A34A'>{p}</li>" for p in r.get("positives",[]))
                    st.markdown(
                        f'<div style="background:{bg};border-left:4px solid {clr};'
                        f'border-radius:8px;padding:10px 14px;margin-top:8px">'
                        f'<div style="font-weight:700;color:{clr};margin-bottom:4px">'
                        f'{icon} {lvl} — Độ tin cậy: {conf}%</div>'
                        f'{"<ul style=margin:2px 0;padding-left:14px>"+issues_html+"</ul>" if issues_html else ""}'
                        f'{"<ul style=margin:2px 0;padding-left:14px>"+pos_html+"</ul>" if pos_html else ""}'
                        f'{"<div style=font-size:0.8rem;color:#475569;margin-top:4px><b>Khuyến nghị:</b> "+r.get("recommendation","")+"</div>" if r.get("recommendation") else ""}'
                        f'</div>', unsafe_allow_html=True
                    )
            elif not ai_on and active_photo:
                st.caption("💡 Kết nối API key để AI phân tích ảnh tự động.")

        # Nút xác nhận hoàn thành bước
        b_results = st.session_state.get(f"kt_b{b}_r", {})
        b_answered = sum(1 for v in b_results.values() if v is not None)
        b_total    = len(step["items"])
        b_fail     = sum(1 for v in b_results.values() if v == "❌ Không Đạt")

        col_btn, col_prog = st.columns([0.4, 0.6])
        with col_btn:
            # Kiểm thực Y Tế: chỉ xác nhận 8:00–11:30 (ngoài giờ → warning)
            _kt_now   = now_vn()
            _kt_hour  = _kt_now.hour + _kt_now.minute / 60
            _kt_ok_time = st.session_state.get("is_super", False) or (8.0 <= _kt_hour <= 11.5)
            can_confirm = (b_answered == b_total) and _kt_ok_time
            if b_answered == b_total and not _kt_ok_time:
                st.warning(f"⏰ Ngoài giờ kiểm thực hợp lệ (8:00–11:30) · Hiện: {_kt_now.strftime('%H:%M')}")

            if done_t:
                # ── Đã xác nhận — hiển thị trạng thái hoàn thành rõ ràng ──
                st.markdown(
                    f'<div style="background:#DCFCE7;border:2px solid #16A34A;'
                    f'border-radius:10px;padding:12px 16px;text-align:center">'
                    f'<div style="font-weight:700;color:#166534;font-size:0.95rem">'
                    f'✅ Bước {b} đã hoàn thành</div>'
                    f'<div style="color:#16A34A;font-size:0.85rem;margin-top:4px">'
                    f'🕐 Xác nhận lúc <b>{done_t}</b> '
                    f'({["Thứ 2","Thứ 3","Thứ 4","Thứ 5","Thứ 6","Thứ 7","CN"][now_vn().weekday()]})'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )
                # Nút bỏ xác nhận (nếu cần chỉnh lại)
                if st.button(f"↩ Bỏ xác nhận Bước {b}", key=f"kt_undo_{b}",
                             use_container_width=True):
                    st.session_state[f"kt_b{b}_done"] = None
                    st.rerun()

            elif can_confirm:
                # ── Đủ điều kiện xác nhận ──
                if st.button(
                    f"☑ Xác nhận hoàn thành Bước {b}",
                    key=f"kt_confirm_{b}",
                    type="primary",
                    use_container_width=True,
                ):
                    ts = now_vn().strftime("%H:%M:%S")  # Lưu giờ:phút:giây
                    st.session_state[f"kt_b{b}_done"] = ts
                    st.toast(f"✅ Bước {b} xác nhận lúc {ts}", icon="🕐")
                    st.rerun()
            else:
                # ── Chưa đủ điều kiện ──
                st.markdown(
                    f'<div style="background:#F8FAFC;border:1px dashed #CBD5E1;'
                    f'border-radius:10px;padding:12px 16px;text-align:center;color:#94A3B8;'
                    f'font-size:0.85rem">'
                    f'⏳ Còn <b>{b_total - b_answered}</b> câu chưa kiểm tra</div>',
                    unsafe_allow_html=True,
                )

        with col_prog:
            pct = int(b_answered / b_total * 100) if b_total else 0
            fail_note = f" — ⚠️ {b_fail} không đạt" if b_fail else ""
            bar_color = "#16A34A" if pct == 100 else "#2563EB"
            st.markdown(
                f'<div style="margin-top:8px">'
                f'<div style="font-size:0.78rem;color:#64748B;margin-bottom:3px">'
                f'{b_answered}/{b_total} câu đã kiểm tra{fail_note}</div>'
                f'<div style="background:#E2E8F0;border-radius:20px;height:8px">'
                f'<div style="width:{pct}%;background:{bar_color};'
                f'border-radius:20px;height:100%;transition:width 0.4s ease"></div></div>'
                f'{"<div style=font-size:0.75rem;color:#16A34A;margin-top:4px>✅ Tất cả câu đã kiểm tra</div>" if pct==100 else ""}'
                f'</div>',
                unsafe_allow_html=True,
            )

        st.markdown("<br>", unsafe_allow_html=True)

    # ── Câu hỏi AI bổ sung (nếu có) ──────────────────────────────────────────
    if st.session_state.kt_extra:
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="sec-hdr">🤖 Câu hỏi bổ sung AI ({len(st.session_state.kt_extra)} câu) — theo thực đơn hôm nay</div>',
            unsafe_allow_html=True,
        )
        if "kt_extra_r" not in st.session_state: st.session_state.kt_extra_r = {}
        for item in st.session_state.kt_extra:
            code = item.get("code","?"); desc = item.get("desc","")
            ingr = item.get("ingredient",""); legal_ref = item.get("legal_ref","")
            col_d, col_c = st.columns([0.65, 0.35])
            with col_d:
                st.markdown(
                    f'<div style="background:#EFF6FF;border-left:3px solid #2563EB;'
                    f'border-radius:0 8px 8px 0;padding:8px 14px;margin:3px 0">'
                    f'<span style="font-size:0.7rem;font-weight:800;color:#2563EB">🤖 {code}</span>'
                    f'<span style="font-size:0.72rem;color:#1D4ED8;margin-left:6px;'
                    f'background:#DBEAFE;padding:1px 6px;border-radius:8px">{ingr}</span>'
                    f'<div style="font-size:0.86rem;font-weight:500;color:#1E293B;margin-top:3px">{desc}</div>'
                    f'{"<div style=font-size:0.72rem;color:#64748B;margin-top:2px>"+legal_ref+"</div>" if legal_ref else ""}'
                    f'</div>', unsafe_allow_html=True,
                )
                if item.get("how") or item.get("pass"):
                    with st.expander("Hướng dẫn"):
                        st.markdown(
                            f"**Kiểm tra:** {item.get('how','')}  \n"
                            f"**✅ Đạt:** {item.get('pass','')}  \n"
                            f"**❌ Không đạt:** {item.get('fail','')}"
                        )
            with col_c:
                seg_key = f"kt_seg_extra_{code}"
                st.segmented_control(
                    code, ["✅ Đạt", "❌ Không Đạt"],
                    key=seg_key, label_visibility="collapsed",
                )
                st.session_state.kt_extra_r[code] = st.session_state.get(seg_key)
        if st.button("🗑️ Xoá câu hỏi AI", use_container_width=True, key="kt_del_extra"):
            st.session_state.kt_extra = []; st.rerun()

    # ── Tổng kết 3 bước ───────────────────────────────────────────────────────
    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-hdr">Tổng kết ca kiểm thực</div>', unsafe_allow_html=True)

    all_results = {}
    for b in [1, 2, 3]:
        all_results.update(st.session_state.get(f"kt_b{b}_r", {}))
    total_items  = sum(len(s["items"]) for s in KIEM_THUC)
    total_done   = sum(1 for v in all_results.values() if v is not None)
    total_pass   = sum(1 for v in all_results.values() if v == "✅ Đạt")
    total_fail   = sum(1 for v in all_results.values() if v == "❌ Không Đạt")
    steps_done   = sum(1 for b in [1, 2, 3]
                       if st.session_state.get(f"kt_b{b}_done"))

    m1, m2, m3, m4 = st.columns(4)
    m1.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">Bước hoàn thành</div>
        <div class="metric-num {'c-green' if steps_done==3 else 'c-orange'}">{steps_done}/3</div>
    </div>""", unsafe_allow_html=True)
    m2.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">✅ Đạt chuẩn</div>
        <div class="metric-num c-green">{total_pass}</div>
        <div class="metric-lbl">/ {total_items} điểm</div>
    </div>""", unsafe_allow_html=True)
    m3.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">❌ Không đạt</div>
        <div class="metric-num c-red">{total_fail}</div>
        <div class="metric-lbl">điểm</div>
    </div>""", unsafe_allow_html=True)
    m4.markdown(f"""<div class="metric-box">
        <div class="metric-lbl">Mẫu lưu B3_05</div>
        <div class="metric-num {'c-green' if all_results.get('B3_05')=='✅ Đạt' else 'c-red'}">
            {'✅' if all_results.get('B3_05')=='✅ Đạt' else '❌'}
        </div>
        <div class="metric-lbl">{'Đã lưu' if all_results.get('B3_05')=='✅ Đạt' else 'Chưa lưu!'}</div>
    </div>""", unsafe_allow_html=True)

    # ── Nút xuất sổ kiểm thực ────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    all_notes = {}
    for b in [1, 2, 3]:
        all_notes.update(st.session_state.get(f"kt_b{b}_n", {}))

    # Timestamps từng bước để đưa vào báo cáo
    timestamps = {
        b: st.session_state.get(f"kt_b{b}_done") or ""
        for b in [1, 2, 3]
    }

    # Điều kiện xuất: PHẢI hoàn thành ĐỦ cả 3 bước VÀ đủ 15 câu
    all_steps_confirmed = all(timestamps[b] for b in [1, 2, 3])
    all_items_answered  = (total_done == total_items)
    can_export = all_steps_confirmed and all_items_answered

    # Hiển thị checklist điều kiện xuất
    cond_rows = [
        (all_items_answered, f"Đã kiểm tra đủ 15/15 câu hỏi ({total_done}/{total_items})"),
        (bool(timestamps[1]),
         f"Bước 1 đã xác nhận{' lúc ' + timestamps[1] if timestamps[1] else ' — chưa xác nhận'}"),
        (bool(timestamps[2]),
         f"Bước 2 đã xác nhận{' lúc ' + timestamps[2] if timestamps[2] else ' — chưa xác nhận'}"),
        (bool(timestamps[3]),
         f"Bước 3 đã xác nhận{' lúc ' + timestamps[3] if timestamps[3] else ' — chưa xác nhận'}"),
    ]
    cond_html = "".join(
        f'<div style="font-size:0.82rem;padding:3px 0;color:{"#16A34A" if ok else "#DC2626"}">'
        f'{"✅" if ok else "❌"} {label}</div>'
        for ok, label in cond_rows
    )
    st.markdown(
        f'<div style="background:#F8FAFC;border:1px solid #E2E8F0;border-radius:10px;'
        f'padding:14px 18px;margin-bottom:12px">'
        f'<div style="font-size:0.8rem;font-weight:700;color:#334155;margin-bottom:6px">'
        f'Điều kiện để xuất sổ kiểm thực:</div>'
        f'{cond_html}</div>',
        unsafe_allow_html=True,
    )

    if can_export:
        # G1: Auto-save kiểm thực — guard chống duplicate
        date_vn_kt    = kt_date.strftime("%d/%m/%Y")
        date_iso_kt   = kt_date.strftime("%Y-%m-%d")
        _kt_guard_key = f"kt_saved_{kt_school}_{date_vn_kt}_{kt_name}"
        if not st.session_state.get(_kt_guard_key, False):
            sid_kt = db_save_kiem_thuc(
                kt_school, date_iso_kt, kt_name, kt_menu,
                all_results, all_notes, timestamps,
                total_pass, total_fail,
            )
            if sid_kt:
                st.session_state[_kt_guard_key] = True
                # DB saved silently
        else:
            st.info("💾 Sổ kiểm thực này đã được lưu trước đó.")

        with st.spinner("Đang tạo sổ kiểm thực..."):
            docx_bytes = generate_so_kiem_thuc_docx(
                kt_school, date_vn_kt,
                kt_name, kt_menu, all_results, all_notes, timestamps,
            )
        fname = f"SoKiemThuc_{(kt_school or 'Truong').replace(' ','_')}_{kt_date.strftime('%d-%m-%Y')}.docx"
        st.download_button(
            "📋 Tải báo cáo (.docx)",
            data=docx_bytes, file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, type="primary",
        )
    else:
        missing = sum(1 for ok, _ in cond_rows if not ok)
        st.button(
            f"⛔ Chưa đủ điều kiện xuất — còn {missing} điều kiện chưa đáp ứng",
            disabled=True, use_container_width=True,
        )


def generate_so_kiem_thuc_docx(school: str, date_str: str, yte_name: str,
                                menu: str, results: dict, notes: dict,
                                timestamps: dict | None = None) -> bytes:
    """Tạo Sổ Kiểm Thực 3 Bước chuẩn TTLT 13/2016 — Times New Roman.
    timestamps: {1: "HH:MM:SS", 2: "HH:MM:SS", 3: "HH:MM:SS"} để xác minh timeline.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)

    # Quốc hiệu
    _docx_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
               bold=True, size=13, align="center", space_after=2)
    _docx_para(doc, "Độc lập – Tự do – Hạnh phúc",
               bold=True, size=12, align="center", space_after=2)
    _docx_para(doc, "────────────────────────────",
               size=10, align="center", space_after=4)
    _docx_para(doc, f"TP. Hồ Chí Minh, ngày {date_str}",
               size=12, align="right", space_after=10)

    _docx_para(doc, "SỔ KIỂM THỰC 3 BƯỚC",
               bold=True, size=15, align="center", space_after=2)
    _docx_para(doc, "BỮA ĂN HỌC ĐƯỜNG",
               bold=True, size=13, align="center", space_after=2)
    _docx_para(doc,
               "(Căn cứ: TTLT số 13/2016/TTLT-BYT-BGDĐT ngày 12/5/2016 – Điều 9)",
               size=11, align="center", space_after=10)

    # Thông tin chung
    info_rows = [
        ("Cơ sở giáo dục", school or "..."),
        ("Ngày kiểm thực", date_str),
        ("Nhân viên y tế thực hiện", yte_name or "..."),
        ("Thực đơn bữa ăn", menu or "..."),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(info_rows):
        r0 = t.rows[i].cells[0].paragraphs[0].add_run(k)
        r1 = t.rows[i].cells[1].paragraphs[0].add_run(v)
        _docx_set_font(r0, bold=True, size_pt=11)
        _docx_set_font(r1, size_pt=11)
        t.rows[i].cells[0].width = Cm(6)
    doc.add_paragraph()

    # Chi tiết từng bước — bao gồm timestamp xác nhận để xác minh timeline
    ts_map = timestamps or {}
    for step in KIEM_THUC:
        b_num = step["buoc"]
        ts    = ts_map.get(b_num, "")
        # Xác minh timeline — kiểm tra ĐẦY ĐỦ khung giờ (start <= ts <= end)
        def _parse_hhmm(t: str) -> int:
            """Chuyển 'HH:MM' hoặc 'HH:MM:SS' thành tổng phút."""
            parts = t.strip().split(":")
            return int(parts[0]) * 60 + int(parts[1]) if len(parts) >= 2 else -1

        # Phân tích "8:00 – 9:30" → start=480 phút, end=570 phút
        w_parts = step["time_window"].replace("–", "-").replace("—", "-").split("-")
        w_start = _parse_hhmm(w_parts[0]) if len(w_parts) >= 1 else 0
        w_end   = _parse_hhmm(w_parts[1]) if len(w_parts) >= 2 else 1439

        ts_mins  = _parse_hhmm(ts) if ts else -1
        on_time  = (w_start <= ts_mins <= w_end) if ts_mins >= 0 else None

        ts_label = (
            f" ✅ Xác nhận lúc {ts} — đúng khung giờ ({step['time_window']})"
            if on_time is True else
            f" ⚠️ NGOÀI KHUNG GIỜ: Xác nhận lúc {ts} (yêu cầu: {step['time_window']})"
            if on_time is False else
            " ❌ Chưa xác nhận"
        )

        _docx_para(doc,
                   f"{step['icon']}  {step['label']}  ({step['time_window']})  —  {step['law']}",
                   bold=True, size=13, space_before=8, space_after=2)

        # Dòng thời gian — highlight đỏ nếu ngoài khung giờ
        p_ts = doc.add_paragraph()
        p_ts.paragraph_format.space_after = Pt(4)
        r_ts = p_ts.add_run(f"⏱ Thời gian thực hiện:{ts_label}")
        _docx_set_font(r_ts, bold=(on_time is False), size_pt=11,
                       color=(192, 0, 0) if on_time is False else
                       (22, 163, 74) if on_time is True else (100, 116, 139))

        _docx_para(doc, step["desc"], size=11, space_after=4)

        tbl = doc.add_table(rows=1 + len(step["items"]), cols=4)
        tbl.style = "Table Grid"
        _docx_table_header(tbl, ["Mã", "Nội dung kiểm tra", "Kết quả", "Ghi chú"])
        widths = [Cm(1.5), Cm(8.5), Cm(2.5), Cm(4.5)]
        for i, (code, desc, *_) in enumerate(step["items"]):
            row  = tbl.rows[i + 1]
            res  = results.get(code, "Chưa kiểm tra")
            note = notes.get(code, "")
            status = res.replace("✅ Đạt", "ĐẠT").replace("❌ Không Đạt", "KHÔNG ĐẠT").replace("Chưa kiểm tra", "—")
            cells_data = [code, desc, status, note]
            for j, (cell, val, w) in enumerate(zip(row.cells, cells_data, widths)):
                cell.width = w
                r = cell.paragraphs[0].add_run(val)
                is_fail = status == "KHÔNG ĐẠT" and j == 2
                _docx_set_font(r, bold=is_fail, size_pt=10,
                               color=(192, 0, 0) if is_fail else None)
        doc.add_paragraph()

    # Chữ ký
    _docx_para(doc, f"Ngày kiểm thực: {date_str}   |   Giờ hoàn thành: ___:___",
               size=11, space_before=10, space_after=10)
    sig = doc.add_table(rows=3, cols=2)
    for i, (l, r_) in enumerate([
        ("Xác nhận Hiệu Trưởng", "Nhân viên Y Tế Học Đường"),
        ("(ký, đóng dấu)", "(ký và ghi rõ họ tên)"),
        ("", yte_name or ""),
    ]):
        for cell, txt in [(sig.rows[i].cells[0], l), (sig.rows[i].cells[1], r_)]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell.paragraphs[0].add_run(txt)
            _docx_set_font(r, bold=(i == 0), size_pt=12)

    _docx_para(doc, "─" * 50, size=9, align="center", space_before=14, space_after=2)
    _docx_para(doc,
               "Sổ kiểm thực này được tạo bằng SchoolFood AI v2.0 — "
               f"song song với sổ giấy bắt buộc theo TTLT 13/2016 Điều 9",
               size=9, align="center")

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ── Nội dung sổ tay — dùng chung cho tab_guide() và generate_user_manual_docx()
MANUAL_CONTENT = {
    "title": "SỔ TAY HƯỚNG DẪN SỬ DỤNG SCHOOLFOOD AI",
    "subtitle": "Nền tảng giám sát An toàn Thực phẩm bữa ăn học đường",
    "version": "v2.0 — Tháng 6/2026",
    "sections": [
        {
            "id": "intro",
            "icon": "🍱",
            "title": "1. SchoolFood AI là gì?",
            "content": (
                "SchoolFood AI là nền tảng công nghệ giúp các bên liên quan tại trường học "
                "giám sát An toàn Thực phẩm (ATTP) bữa ăn bán trú một cách hệ thống, có căn "
                "cứ pháp luật và dễ thực hiện.\n\n"
                "App KHÔNG thay thế xét nghiệm vi sinh học trong phòng thí nghiệm. "
                "App hỗ trợ kiểm tra cảm quan, ghi nhận hồ sơ và cảnh báo sớm các "
                "dấu hiệu rủi ro nhìn thấy bằng mắt thường."
            ),
            "subsections": [
                ("Vấn đề đang giải quyết",
                 "Mỗi ngày hơn 1 triệu học sinh Hà Nội và hàng triệu học sinh cả nước ăn "
                 "bán trú tại trường. Phụ huynh lo lắng nhưng không có công cụ giám sát; "
                 "Ban Giám sát muốn kiểm tra nhưng không biết kiểm tra gì đúng chuẩn; "
                 "Y tế học đường ghi sổ giấy rồi cất vào tủ. SchoolFood AI lấp đầy khoảng trống này."),
                ("Ai nên dùng app này",
                 "• Phụ Huynh — xem thực đơn, kết quả kiểm tra, phản hồi\n"
                 "• Ban Giám Sát (Đại Diện PHHS) — thực hiện checklist 20 điểm và tạo báo cáo\n"
                 "• Y Tế Học Đường — số hoá kiểm thực 3 bước hàng ngày\n"
                 "• Ban Giám Hiệu — xem tổng quan và duyệt báo cáo"),
            ]
        },
        {
            "id": "roles",
            "icon": "👥",
            "title": "2. Hướng dẫn theo vai trò — Tab & Chức năng",
            "subsections": [
                ("👨‍👩‍👧 Phụ Huynh — 4 Tab",
                 "Đăng nhập bằng tài khoản được Ban Giám Hiệu cấp.\n\n"
                 "Tab 💬 Hỏi đáp AI: Đặt câu hỏi tiếng Việt về ATTP, dinh dưỡng, quyền phụ huynh\n"
                 "Tab 🍱 Góc Phụ Huynh:\n"
                 "  • Xem kết quả bữa ăn hôm nay (🟢 An toàn / 🟡 Cần theo dõi / 🔴 Có vấn đề)\n"
                 "  • Xem thực đơn hôm nay\n"
                 "  • Lịch sử 7 ngày gần nhất\n"
                 "  • Gửi phản hồi/complaint về bữa ăn — theo dõi trạng thái ⏳→💬→✅\n"
                 "Tab 🚨 Khẩn cấp: Quy trình 6 bước khi nghi ngờ ngộ độc · Số điện thoại 115, 1800 6838\n"
                 "Tab 📖 Hướng dẫn: Sổ tay đầy đủ + quyền pháp lý\n\n"
                 "⚠️ Phụ huynh KHÔNG trực tiếp vào bếp kiểm tra — chỉ Ban Đại Diện PHHS chính thức mới có quyền."),

                ("👥 Ban Giám Sát (Đại Diện PHHS) — 7 Tab",
                 "Đăng nhập bằng tài khoản được Ban Giám Hiệu cấp.\n\n"
                 "Tab 💬 Hỏi đáp AI: Tư vấn pháp luật, câu hỏi nghiệp vụ kiểm tra\n"
                 "Tab ✅ Checklist kiểm tra (NHIỆM VỤ CHÍNH · 2 lần/tuần tối thiểu):\n"
                 "  • 1 lần báo trước ≥ 24h · 1 lần đột xuất KHÔNG báo trước\n"
                 "  • 20 điểm chuẩn hóa · 7 mục bắt buộc BẮT BUỘC có ảnh minh chứng khi Không Đạt\n"
                 "  • Xuất báo cáo Word gửi Ban Giám Hiệu trong 24 giờ\n"
                 "Tab 🏭 Nhà Cung Cấp (KIỂM TRA TOÀN DIỆN · 1 lần/tháng):\n"
                 "  • 12 điểm theo NĐ 15/2018 · S01–S12 · Xếp loại A/B/C\n"
                 "  • Kiểm tra giao hàng khi có mặt: S03–S12 (10 điểm)\n"
                 "Tab 📊 Lịch sử: Dashboard kết quả + theo dõi complaint (chỉ đọc)\n"
                 "Tab 📅 Lịch & Chuẩn mực: Lịch kiểm tra + hệ thống cảnh báo 4 cấp\n"
                 "Tab 🚨 Khẩn cấp: Quy trình sự cố\n"
                 "Tab 📖 Hướng dẫn: Sổ tay đầy đủ"),

                ("🏥 Y Tế Học Đường — 6 Tab",
                 "Đăng nhập bằng tài khoản được Ban Giám Hiệu cấp.\n"
                 "Chọn cấp học đúng với lớp bạn phụ trách (Tiểu Học / THCS / THPT) — ảnh hưởng tiêu chuẩn dinh dưỡng.\n\n"
                 "Tab 💬 Hỏi đáp AI: Tư vấn kỹ thuật ATTP, câu hỏi nghiệp vụ\n"
                 "Tab 🏥 Kiểm thực 3 bước (NHIỆM VỤ CHÍNH · HÀNG NGÀY 9:30–11:00):\n"
                 "  • Bước 1 (8:00–9:30 | Trước chế biến): tem kiểm dịch, hóa đơn, hạn dùng, nhiệt độ tủ lạnh\n"
                 "  • Bước 2 (9:30–10:30 | Trong chế biến): nhiệt độ nấu ≥70°C, vệ sinh dụng cụ, nhân viên\n"
                 "  • Bước 3 (10:30–11:00 | Sau chế biến): nhiệt độ chia, màu mùi, khẩu phần, mẫu lưu 24h\n"
                 "  → Song song ghi sổ kiểm thực giấy (bắt buộc theo TTLT 13/2016)\n"
                 "Tab 🏭 Nhà Cung Cấp (KIỂM TRA GIAO HÀNG · Mỗi lần NCC giao):\n"
                 "  • 10 điểm S03–S12: xe vận chuyển, nhiệt độ, hóa đơn, nhãn mác, mẫu lưu\n"
                 "  • Cung cấp minh chứng cho complaint của Phụ Huynh (Y Tế là người trực tiếp tại bếp)\n"
                 "Tab 📊 Lịch sử: Xem kết quả + cung cấp minh chứng complaint\n"
                 "Tab 🚨 Khẩn cấp: Quy trình sự cố\n"
                 "Tab 📖 Hướng dẫn: Sổ tay đầy đủ"),

                ("🏫 Ban Giám Hiệu — 6 Tab",
                 "Tài khoản đầu tiên do Admin hệ thống (Quản trị viên) cấp.\n"
                 "Ban Giám Hiệu tự tạo tài khoản cho Y Tế, BGS, Phụ Huynh trong trường.\n\n"
                 "Tab 💬 Hỏi đáp AI: Tư vấn pháp luật, quản lý chất lượng bữa ăn\n"
                 "Tab 📊 Lịch sử & Phản hồi (DASHBOARD CHÍNH):\n"
                 "  • Tổng quan kết quả kiểm tra bữa ăn + nhà cung cấp\n"
                 "  • Phát hiện bất thường tự động (điểm quá đều, BGS/Y Tế chênh lệch)\n"
                 "  • Quản lý complaint: xem minh chứng từ Y Tế → đóng task + phản hồi chính thức\n"
                 "  • Luồng: ⏳ Chờ xử lý → 💬 Y Tế thêm minh chứng → ✅ BGH đóng + phản hồi\n"
                 "Tab 📅 Lịch & Chuẩn mực: Lịch + hệ thống cảnh báo 4 cấp\n"
                 "Tab 🚨 Khẩn cấp: Quy trình sự cố\n"
                 "Tab 📖 Hướng dẫn: Sổ tay đầy đủ\n"
                 "Tab 👤 Quản lý tài khoản:\n"
                 "  • Tạo tài khoản cho Y Tế (có chọn cấp học mặc định), BGS, Phụ Huynh\n"
                 "  • Tối đa 2 tài khoản BGH · Chỉ tạo cho trường mình quản lý"),
            ]
        },
        {
            "id": "complaint",
            "icon": "📬",
            "title": "2b. Hệ thống Phản hồi & Complaint",
            "content": "Luồng xử lý complaint 3 bước — từ Phụ Huynh đến Ban Giám Hiệu:",
            "subsections": [
                ("⏳ Bước 1 — Phụ Huynh gửi phản hồi",
                 "Phụ Huynh vào Tab 🍱 Góc Phụ Huynh → Gửi phản hồi → Chọn loại + Mô tả cụ thể\n"
                 "Trạng thái: ⏳ Chờ xử lý\n"
                 "Phụ Huynh có thể theo dõi trạng thái complaint đã gửi trong cùng tab."),
                ("💬 Bước 2 — Y Tế Học Đường thêm minh chứng",
                 "Y Tế vào Tab 📊 Lịch sử → Thấy complaint → Bấm '📝 Thêm minh chứng'\n"
                 "Nhập diễn giải + tải ảnh/file minh chứng (jpg, png, pdf)\n"
                 "Trạng thái tự động chuyển: ⏳ → 💬 Đang xem xét\n"
                 "Ban Giám Hiệu sẽ thấy minh chứng này khi xem xét."),
                ("✅ Bước 3 — Ban Giám Hiệu đóng task",
                 "BGH vào Tab 📊 Lịch sử → Xem minh chứng từ Y Tế → Bấm '🔒 Đóng task'\n"
                 "Nhập phản hồi chính thức → Lưu\n"
                 "Trạng thái: ✅ Đóng DD/MM/YYYY\n"
                 "Phụ Huynh thấy phản hồi chính thức của BGH trong Tab Góc Phụ Huynh.\n\n"
                 "⚠️ Complaint quá 2 ngày chưa xử lý → Được highlight đỏ trong tab Lịch sử của BGH"),
                ("📊 Visualize complaint",
                 "Ban Giám Hiệu thấy 2 biểu đồ trong tab Lịch sử:\n"
                 "• Bar chart: số complaint theo ngày (30 ngày gần nhất)\n"
                 "• Donut chart: phân bố theo loại (Chất lượng / Vệ sinh / Nghi ngờ ngộ độc / ...)"),
            ]
        },
        {
            "id": "accounts",
            "icon": "🔐",
            "title": "2c. Quản lý tài khoản & Trường đa cấp",
            "subsections": [
                ("Quy trình tạo tài khoản",
                 "Bước 1 — Admin hệ thống tạo tài khoản BGH đầu tiên cho trường\n"
                 "Bước 2 — BGH đăng nhập → Tab 👤 Quản lý tài khoản → Tạo tài khoản cho:\n"
                 "  • Y Tế Học Đường: Chọn cấp học mặc định (TH / THCS / THPT)\n"
                 "  • Ban Giám Sát (Đại Diện PHHS)\n"
                 "  • Phụ Huynh\n"
                 "Bước 3 — Người dùng nhận mật khẩu tạm từ BGH → Đăng nhập → Đổi mật khẩu qua 'Quên mật khẩu'"),
                ("Trường đa cấp (TH + THCS hoặc TH + THCS + THPT)",
                 "Phương án được khuyến nghị:\n\n"
                 "CÁCH 1 — Tên trường có kèm cấp:\n"
                 "  • 'Trường ABC (TH)' → 1 bộ tài khoản riêng\n"
                 "  • 'Trường ABC (THCS)' → 1 bộ tài khoản riêng\n"
                 "  → Dữ liệu tách biệt rõ ràng, dễ quản lý\n\n"
                 "CÁCH 2 — Cùng tên trường, Y Tế phân theo cấp mặc định:\n"
                 "  • Y Tế TH: tài khoản 'default_level = Tiểu Học'\n"
                 "  • Y Tế THCS: tài khoản 'default_level = THCS'\n"
                 "  → Dữ liệu chung, level tự điền theo tài khoản\n"
                 "  → BGH thấy tổng hợp cả 2 cấp trong tab Lịch sử\n\n"
                 "BGS và Phụ Huynh chọn level khi làm checklist (per session) nên dùng được cho cả 2 cách."),
                ("Bảo mật và phân quyền dữ liệu",
                 "• Phụ Huynh / BGS / Y Tế: chỉ thấy dữ liệu trường mình (locked)\n"
                 "• BGH: thấy tất cả dữ liệu trường mình\n"
                 "• Admin hệ thống: thấy tất cả trường, có role switcher để test\n"
                 "• BGH không tạo được tài khoản cho trường khác\n"
                 "• Tối đa 2 tài khoản BGH mỗi trường"),
            ]
        },
        {
            "id": "checklist",
            "icon": "✅",
            "title": "3. Hướng dẫn Checklist 20 điểm",
            "content": (
                "Checklist 20 điểm được xây dựng theo NĐ 15/2018/NĐ-CP và TTLT 13/2016/TTLT-BYT-BGDĐT. "
                "Chia làm 5 nhóm, mỗi nhóm kiểm tra một khía cạnh của chuỗi thực phẩm."
            ),
            "subsections": [
                ("5 nhóm kiểm tra",
                 "📦 Nhóm 1 — Nguồn gốc nguyên liệu (C01–C04): tem kiểm dịch, hóa đơn, hạn dùng\n"
                 "🌡️ Nhóm 2 — Bảo quản & vận chuyển (C05–C08): nhiệt độ tủ lạnh, tách biệt sống/chín\n"
                 "🍽️ Nhóm 3 — Thức ăn khi phục vụ (C09–C13): nhiệt độ chia, thời gian nấu, khẩu phần\n"
                 "🧼 Nhóm 4 — Vệ sinh dụng cụ & nhân viên (C14–C17): dụng cụ sạch, bảo hộ\n"
                 "📋 Nhóm 5 — Hồ sơ & giấy tờ (C18–C20): sổ kiểm thực, thực đơn, mẫu lưu"),
                ("7 mục BẮT BUỘC — phải đạt tuyệt đối",
                 "C03 — Hạn sử dụng nguyên liệu ≥ 3 ngày\n"
                 "C07 — Nhiệt độ thức ăn khi nhận ≥ 60°C\n"
                 "C09 — Nhiệt độ thức ăn khi chia: nóng ≥ 60°C, lạnh ≤ 5°C\n"
                 "C10 — Thời gian nấu đến phục vụ < 2 giờ\n"
                 "C11 — Màu sắc và mùi thức ăn bình thường\n"
                 "C18 — Sổ kiểm thực 3 bước điền đầy đủ, có chữ ký\n"
                 "C20 — Có mẫu lưu thức ăn 24h từng món\n\n"
                 "⚠️ Nếu bất kỳ mục nào fail → mức cảnh báo tự động là CRITICAL bất kể điểm tổng."),
                ("Ngưỡng điểm và đánh giá tổng thể",
                 "18–20 điểm ĐẠT → ✅ Đạt chuẩn ATTP — lưu hồ sơ bình thường\n"
                 "15–17 điểm ĐẠT → ⚠️ Cần cải thiện — yêu cầu khắc phục trong 24h\n"
                 "< 15 điểm ĐẠT → ❌ Không đạt — báo ngay Ban Giám Hiệu"),
                ("Tại sao 3 checklist có điểm giống nhau? — Không phải trùng lặp",
                 "Nhiều tiêu chí xuất hiện ở cả Checklist BGS (C*), Kiểm thực Y Tế (B*) và Đánh giá NCC (S*).\n"
                 "Đây là thiết kế CÓ CHỦ Ý theo nguyên tắc HACCP — kiểm soát nhiều điểm trên cùng 1 chuỗi:\n\n"
                 "Ví dụ — Nhiệt độ thực phẩm:\n"
                 "• S04 (NCC): Nhiệt độ khi giao hàng ≥ 60°C hay < 8°C → Nhà CC có giữ đúng chuẩn không?\n"
                 "• B1_04 (Y Tế Bước 1): Tủ lạnh < 5°C khi nhận hàng → Sau giao có bảo quản đúng không?\n"
                 "• C07 (BGS): Nhiệt độ thức ăn khi nhận ≥ 60°C → Tại thời điểm kiểm tra BGS có đúng không?\n"
                 "• B3_01 (Y Tế Bước 3): Nhiệt độ chia ≥ 60°C → Đến lúc phục vụ học sinh còn đúng không?\n\n"
                 "Ví dụ — Mẫu lưu thức ăn:\n"
                 "• S11 (NCC): NCC có giao đủ mẫu lưu không? → Trách nhiệm nhà cung cấp\n"
                 "• C20 / B3_05 (BGS/Y Tế): Mẫu lưu đã được lấy và bảo quản đúng không? → Trách nhiệm nhà trường\n\n"
                 "→ Mỗi vai trò kiểm tra cùng 1 tiêu chí nhưng ở THỜI ĐIỂM và GÓC ĐỘ TRÁCH NHIỆM khác nhau.\n"
                 "Đây là cơ chế 'Defense in Depth' — nếu 1 tầng bỏ sót, tầng kế tiếp vẫn phát hiện được."),
                ("Câu hỏi AI bổ sung theo thực đơn",
                 "Ngoài 20 câu chuẩn, khi nhập thực đơn và bấm 🤖 Tạo câu hỏi bổ sung, AI sẽ tạo "
                 "thêm 3–5 điểm kiểm tra đặc thù cho từng nguyên liệu dựa trên QCVN 8-1/8-2/8-3.\n\n"
                 "Ví dụ: thực đơn có cá → AI thêm check nhiệt độ bảo quản cá tươi theo QCVN 8-1.\n"
                 "Tính năng này CẦN credit API (~$0.004/lần). Không bấm → vẫn đánh giá đủ 20 câu."),
                ("Hướng dẫn chụp ảnh minh chứng",
                 "Khi nào cần ảnh: bắt buộc nếu có mục KHÔNG ĐẠT; khuyến khích cho tất cả các lần kiểm tra.\n\n"
                 "Chụp đạt chuẩn:\n"
                 "✅ Đủ sáng, không bóng đổ che khuất\n"
                 "✅ Cách 20–50cm, chụp thẳng góc\n"
                 "✅ Ảnh nét, không bị mờ\n"
                 "✅ Thấy rõ vùng cần kiểm tra\n"
                 "❌ Tránh: tối, mờ, nghiêng >45°, ảnh quá xa\n\n"
                 "AI phân tích ảnh (~$0.015/ảnh) phát hiện: màu bất thường, mốc, côn trùng, "
                 "sổ kiểm thực thiếu ký, thực phẩm để sai vị trí.\n"
                 "AI KHÔNG thể: phát hiện vi khuẩn vô hình, đo nhiệt độ thực, đảm bảo 100% chính xác."),
            ]
        },
        {
            "id": "alert",
            "icon": "🔔",
            "title": "4. Hệ thống cảnh báo 4 cấp",
            "content": (
                "Bảng cảnh báo đầy đủ (kích hoạt, thời hạn, người nhận thông báo) xem trong "
                "tab **📅 Lịch & Chuẩn mực**. Tóm tắt nhanh để tra cứu:"
            ),
            "subsections": [
                ("🔴 CRITICAL — trong 5 phút",
                 "Bất kỳ mục BẮT BUỘC nào KHÔNG ĐẠT → DỪNG bữa ăn · Giữ mẫu · Gọi Hiệu Trưởng + 115"),
                ("🟠 MAJOR — trong 2–4 giờ",
                 "Tổng điểm < 15/20 → Yêu cầu nhà cung cấp khắc phục trước bữa ăn tiếp theo"),
                ("🟡 MINOR — trong 24 giờ",
                 "Tổng điểm 15–17/20 → Ghi hồ sơ, thông báo cải thiện"),
                ("✅ ĐẠT CHUẨN",
                 "Tất cả BẮT BUỘC đạt + Tổng ≥ 18/20 → Lưu hồ sơ bình thường"),
            ]
        },
        {
            "id": "ai",
            "icon": "🤖",
            "title": "5. Tính năng AI — hướng dẫn và giới hạn",
            "subsections": [
                ("Cần credit API không?",
                 "CẦN credit (~$5 = dùng 1–2 tháng pilot):\n"
                 "• Hỏi đáp pháp luật: ~$0.003/câu\n"
                 "• Tạo checklist theo thực đơn: ~$0.004/lần\n"
                 "• Phân tích ảnh (Vision): ~$0.015/ảnh\n"
                 "• Báo cáo ngôn ngữ tự nhiên: ~$0.004/lần\n\n"
                 "KHÔNG cần credit (luôn dùng được):\n"
                 "• Checklist 20 câu · Xuất báo cáo Word · Lịch nhắc nhở · Hướng dẫn khẩn cấp"),
                ("AI phân tích ảnh hoạt động thế nào?",
                 "Ảnh được gửi đến Claude Opus (model AI đa phương thức của Anthropic). "
                 "AI đánh giá dựa trên tiêu chí cảm quan từ:\n"
                 "• WHO Food Safety Visual Inspection Guide\n"
                 "• QCVN 8-1:2011/BYT (dấu hiệu nhiễm vi sinh)\n"
                 "• NĐ 15/2018/NĐ-CP (điều kiện vệ sinh bếp ăn)\n\n"
                 "Kết quả trả về: mức rủi ro (OK/WARNING/CRITICAL), vấn đề phát hiện, "
                 "khuyến nghị khắc phục, căn cứ pháp lý áp dụng, độ tin cậy (%)."),
                ("Giới hạn quan trọng của AI",
                 "AI KHÔNG THỂ thay thế:\n"
                 "❌ Xét nghiệm vi sinh học (Salmonella, E.coli) — cần phòng thí nghiệm\n"
                 "❌ Đo nhiệt độ thực tế của thức ăn — cần nhiệt kế vật lý\n"
                 "❌ Ngửi mùi thực phẩm — cần người kiểm tra trực tiếp\n"
                 "❌ Đảm bảo 100% không có mối nguy — không có hệ thống nào làm được\n\n"
                 "AI phù hợp nhất cho: tạo hồ sơ, ghi nhận bằng chứng, hỗ trợ ra quyết định."),
            ]
        },
        {
            "id": "emergency",
            "icon": "🚨",
            "title": "6. Xử lý khẩn cấp ngộ độc thực phẩm",
            "content": "Quy trình 6 bước theo TTLT 13/2016/TTLT-BYT-BGDĐT:",
            "subsections": [
                ("Bước 1 — DỪNG BỮA ĂN NGAY",
                 "Yêu cầu toàn bộ học sinh ngừng ăn. Không để thêm bất kỳ ai ăn thêm."),
                ("Bước 2 — GỌI 115 nếu có triệu chứng nặng",
                 "Gọi 115 ngay khi học sinh có: co giật, khó thở, mất ý thức, nôn ra máu. "
                 "Không chờ đợi — mỗi phút là quan trọng."),
                ("Bước 3 — GIỮ NGUYÊN MẪU THỨC ĂN",
                 "KHÔNG vứt, KHÔNG rửa, KHÔNG đổ bất kỳ thức ăn nào. "
                 "Đây là bằng chứng duy nhất để xét nghiệm xác định nguyên nhân."),
                ("Bước 4 — BÁO NGAY Hiệu Trưởng + Y Tế học đường",
                 "Gọi điện trực tiếp, không nhắn tin. "
                 "Cung cấp: số học sinh bị, triệu chứng, giờ ăn, các món đã ăn."),
                ("Bước 5 — GHI CHÉP ĐẦY ĐỦ",
                 "Ghi ngay vào app (tab 🚨 Khẩn cấp → Bắt đầu xử lý sự cố) hoặc giấy: "
                 "số học sinh, triệu chứng, thời gian phát sinh, diễn biến."),
                ("Bước 6 — BÁO SỞ Y TẾ trong 24 giờ",
                 "Từ 2 người bị trở lên: bắt buộc báo Sở Y Tế địa phương trong 24h.\n"
                 "Đường dây nóng Cục ATTP: 1800 6838 (miễn phí, giờ hành chính)"),
            ]
        },
        {
            "id": "faq",
            "icon": "❓",
            "title": "7. Câu hỏi thường gặp (FAQ)",
            "faq": [
                ("Checklist 20 câu dựa trên luật nào?",
                 "NĐ 15/2018/NĐ-CP (điều kiện bếp ăn tập thể), TTLT 13/2016/TTLT-BYT-BGDĐT "
                 "(y tế trường học), QĐ 3958/QĐ-BYT 2025 (dinh dưỡng học đường), "
                 "QCVN 8-1/8-2/8-3 (giới hạn ô nhiễm vi sinh, kim loại nặng, thuốc trừ sâu)."),
                ("Nhà cung cấp từ chối cho xem sổ kiểm thực, phải làm sao?",
                 "Đây là vi phạm pháp luật. Sổ kiểm thực 3 bước là tài liệu bắt buộc theo TTLT 13/2016 "
                 "và phải được cung cấp khi Ban Giám Sát yêu cầu. "
                 "Ghi nhận vào báo cáo (C18 = KHÔNG ĐẠT) và báo Hiệu Trưởng + Sở Y Tế ngay."),
                ("Không có API key có dùng được app không?",
                 "Có. Checklist 20 câu, xuất báo cáo Word, lịch nhắc nhở, hướng dẫn khẩn cấp "
                 "hoàn toàn dùng được mà không cần API key hay credit. "
                 "Chỉ các tính năng AI (hỏi đáp pháp luật, phân tích ảnh, câu hỏi theo thực đơn) mới cần."),
                ("20 câu chuẩn hay câu hỏi AI bổ sung quan trọng hơn?",
                 "20 câu chuẩn là NỀN TẢNG — bắt buộc phải hoàn thành. "
                 "Câu hỏi AI bổ sung là THÊM VÀO — giúp phát hiện rủi ro đặc thù theo thực đơn hôm nay. "
                 "Không bấm tạo câu hỏi AI → vẫn đánh giá đầy đủ theo 20 câu."),
                ("AI phân tích ảnh có chính xác 100% không?",
                 "Không. AI phân tích dựa trên dấu hiệu nhìn thấy bằng mắt với độ tin cậy thường 75–90%. "
                 "Kết quả là gợi ý để người kiểm tra quyết định — không tự động thay thế phán xét của con người. "
                 "Với ảnh tốt (đủ sáng, nét, rõ), độ chính xác cao hơn."),
                ("Báo cáo Word gửi cho ai và lưu ở đâu?",
                 "Gửi cho: Hiệu Trưởng (trong 24h sau kiểm tra), Ban Giám Hiệu (để duyệt ký), "
                 "Sở GD&ĐT (khi được yêu cầu theo Điều 9 TTLT 13/2016). "
                 "Lưu: thư mục hồ sơ ATTP của trường (lưu tối thiểu 2 năm theo quy định)."),
                ("Phụ huynh có được vào bếp kiểm tra không?",
                 "Phụ huynh đơn lẻ KHÔNG có quyền tự vào bếp. "
                 "Ban Đại Diện Cha Mẹ Học Sinh (BĐDCMHS) được bầu chính thức mới có quyền "
                 "giám sát định kỳ theo quy chế dân chủ và TTLT 13/2016."),
                ("Khoảng cách giữa hai lần kiểm tra là bao nhiêu?",
                 "Ban Giám Sát (Đại Diện PHHS): tối thiểu 2 lần/tuần (1 báo trước ≥24h, 1 đột xuất). "
                 "Y Tế Học Đường: mỗi ngày có bữa ăn, lúc 10:00. "
                 "Ban Giám Hiệu: 1 lần/tháng (tuần cuối tháng). "
                 "Sở GD&ĐT/Sở Y Tế: 1–2 lần/học kỳ (đột xuất)."),
                ("Vùng nhiệt độ nguy hiểm là gì?",
                 "5°C – 60°C là vùng nhiệt độ vi khuẩn phát triển nhanh nhất "
                 "(tăng gấp đôi mỗi 20 phút ở nhiệt độ phòng). "
                 "Thức ăn nóng phải ≥60°C; thức ăn lạnh phải ≤5°C; "
                 "thức ăn đã nấu không được để ở nhiệt độ phòng quá 2 giờ."),
                ("Mẫu lưu thức ăn (lưu nghiệm) là gì và cần lưu bao lâu?",
                 "Mỗi món ăn trong bữa cần lấy ≥100g làm mẫu, cho vào túi kín, ghi nhãn "
                 "(tên món + giờ lấy + ngày), bảo quản trong tủ lạnh ≤5°C. "
                 "Lưu tối thiểu 24 giờ sau bữa ăn. Dùng để xét nghiệm khi nghi ngờ ngộ độc."),
            ]
        },
        {
            "id": "legal",
            "icon": "⚖️",
            "title": "8. Căn cứ pháp lý tóm tắt",
            "legal_refs": [
                ("Luật ATTP số 55/2010/QH12",
                 "Khung pháp lý tổng thể về an toàn thực phẩm tại Việt Nam. "
                 "Điều quan trọng: Điều 53 — nghĩa vụ báo cáo ngộ độc thực phẩm."),
                ("Nghị định 15/2018/NĐ-CP",
                 "Điều kiện đảm bảo ATTP cho cơ sở sản xuất kinh doanh thực phẩm, bao gồm bếp ăn tập thể. "
                 "Điều quan trọng: Điều 11–15 về điều kiện bếp ăn tập thể và nhà cung cấp suất ăn."),
                ("TTLT 13/2016/TTLT-BYT-BGDĐT",
                 "Công tác y tế trường học, bao gồm kiểm soát ATTP bữa ăn học đường. "
                 "Điều quan trọng: Điều 9 về kiểm thực 3 bước và lưu mẫu thức ăn."),
                ("Quyết định 3958/QĐ-BYT ngày 25/12/2025",
                 "Hướng dẫn dinh dưỡng bữa ăn học đường (văn bản mới nhất). "
                 "Quy định khẩu phần năng lượng, protein, rau theo từng cấp học."),
                ("QCVN 8-1:2011/BYT",
                 "Giới hạn ô nhiễm vi sinh vật (Salmonella, E.coli, Staphylococcus aureus, v.v.) "
                 "trong từng loại thực phẩm."),
                ("QCVN 8-2:2011/BYT",
                 "Giới hạn ô nhiễm kim loại nặng (chì, cadimi, thuỷ ngân, asen) trong thực phẩm."),
                ("Nghị định 115/2018/NĐ-CP",
                 "Mức xử phạt vi phạm hành chính về ATTP. "
                 "Phạt từ 10–100 triệu đồng; gây tử vong → xử lý hình sự."),
            ]
        },
        {
            "id": "glossary",
            "icon": "📚",
            "title": "9. Bảng thuật ngữ",
            "terms": [
                ("ATTP", "An toàn thực phẩm"),
                ("ATVSTP", "An toàn vệ sinh thực phẩm (cách gọi cũ, nay dùng ATTP)"),
                ("PHHS", "Phụ huynh học sinh"),
                ("Ban Đại Diện PHHS (BĐDCMHS)", "Ban Đại Diện Cha Mẹ Học Sinh — tổ chức đại diện chính thức của phụ huynh"),
                ("Kiểm thực 3 bước", "Quy trình kiểm tra thực phẩm trước/trong/sau chế biến theo luật"),
                ("Lưu mẫu / Lưu nghiệm", "Giữ lại mẫu thức ăn 24h để xét nghiệm nếu cần"),
                ("Vùng nhiệt độ nguy hiểm", "5°C – 60°C: vi khuẩn tăng gấp đôi mỗi 20 phút"),
                ("HACCP", "Phân tích mối nguy và kiểm soát điểm tới hạn — tiêu chuẩn quốc tế"),
                ("Codex Alimentarius", "Bộ tiêu chuẩn thực phẩm quốc tế của FAO/WHO"),
                ("Claude Vision", "Mô hình AI đa phương thức của Anthropic — phân tích hình ảnh"),
                ("CRITICAL / MAJOR / MINOR", "Ba cấp độ cảnh báo trong app (tới hạn / nghiêm trọng / nhỏ)"),
                ("API key", "Mật khẩu truy cập dịch vụ AI — lấy tại console.anthropic.com"),
                ("Credit", "Tín dụng thanh toán API AI — $5 dùng được khoảng 1–2 tháng pilot"),
            ]
        },
    ]
}


# ── Render nội dung sổ tay trong tab ─────────────────────────────────────────
def tab_guide():
    """Tab hướng dẫn sử dụng đầy đủ — sổ tay điện tử tích hợp trong app."""
    mc = MANUAL_CONTENT
    st.markdown(f"""<div class="sf-card">
        <div class="sf-card-title">📖 {mc['title']}</div>
        <div class="sf-card-body">{mc['subtitle']} · {mc['version']}</div>
    </div>""", unsafe_allow_html=True)

    # Nút tải sổ tay Word
    if st.button("⬇️ Tải Sổ Tay PDF/Word (.docx) để in và đào tạo",
                 type="primary", use_container_width=True):
        with st.spinner("Đang tạo file Word..."):
            docx_bytes = generate_manual_docx()
        st.download_button(
            "📥 Tải ngay — Sổ Tay SchoolFood AI.docx",
            data=docx_bytes,
            file_name="So_Tay_SchoolFood_AI_v2.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)

    # Render từng section
    for section in mc["sections"]:
        with st.expander(f"{section['icon']} {section['title']}", expanded=False):
            if "content" in section:
                st.markdown(section["content"])

            if "subsections" in section:
                for sub_title, sub_content in section["subsections"]:
                    st.markdown(
                        f'<div style="background:#F8FAFC;border-left:3px solid #2563EB;'
                        f'border-radius:0 8px 8px 0;padding:10px 14px;margin:8px 0">'
                        f'<b style="color:#1E293B;font-size:0.9rem">{sub_title}</b>'
                        f'</div>', unsafe_allow_html=True
                    )
                    # Render nội dung với line breaks
                    for line in sub_content.split("\n"):
                        if line.strip():
                            st.markdown(
                                f'<div style="font-size:0.875rem;color:#334155;'
                                f'padding:2px 14px;line-height:1.65">{line}</div>',
                                unsafe_allow_html=True
                            )
                    st.markdown("")

            # FAQ
            if "faq" in section:
                for q, a in section["faq"]:
                    with st.expander(f"❓ {q}"):
                        st.markdown(
                            f'<div style="font-size:0.875rem;color:#334155;line-height:1.7">{a}</div>',
                            unsafe_allow_html=True
                        )

            # Căn cứ pháp lý
            if "legal_refs" in section:
                for name, desc in section["legal_refs"]:
                    st.markdown(
                        f'<div class="sf-card" style="padding:12px 16px;margin:6px 0">'
                        f'<span class="role-tag">{name}</span>'
                        f'<div style="font-size:0.83rem;color:#475569;margin-top:6px">{desc}</div>'
                        f'</div>', unsafe_allow_html=True
                    )

            # Thuật ngữ
            if "terms" in section:
                for term, definition in section["terms"]:
                    st.markdown(
                        f'<div style="display:flex;gap:12px;padding:5px 0;'
                        f'border-bottom:1px solid #F1F5F9;font-size:0.875rem">'
                        f'<span style="min-width:180px;font-weight:700;color:#1E293B">{term}</span>'
                        f'<span style="color:#475569">{definition}</span>'
                        f'</div>', unsafe_allow_html=True
                    )

    # ── Tài liệu pháp luật đã tải vào hệ thống ───────────────────────────────
    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-hdr">📋 Tài liệu pháp luật tích hợp trong app</div>',
                unsafe_allow_html=True)
    pdfs = sorted(LEGAL_DIR.glob("*.pdf"))
    if pdfs:
        for p in pdfs:
            st.success(f"✅ {p.name} ({p.stat().st_size // 1024} KB)")
    else:
        st.info("Chưa có file PDF trong thư mục 07_Legal_Regulations/ — AI vẫn tư vấn dựa trên kiến thức đào tạo.")
    try:
        import pypdf  # noqa: F401
        st.caption("✅ pypdf đã cài — AI đọc được nội dung các văn bản pháp luật thực tế")
    except ImportError:
        st.caption("pypdf chưa cài — AI dùng kiến thức đào tạo, không đọc file PDF trực tiếp")

    # ── Về ứng dụng ───────────────────────────────────────────────────────────
    st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="sf-card" style="border-left:3px solid #94A3B8">'
        '<div class="sf-card-title">ℹ️ Về SchoolFood AI v2.0</div>'
        '<div class="sf-card-body">'
        'Nền tảng giám sát An toàn Thực phẩm bữa ăn học đường — giúp mỗi bên thực hiện '
        'đúng vai trò, đúng thời điểm, có bằng chứng rõ ràng.<br>'
        'Xây dựng theo: Luật ATTP 55/2010 · NĐ 15/2018 · TTLT 13/2016 · QĐ 3958/QĐ-BYT 2025 · '
        'Hỗ trợ AI: Claude Sonnet 4.6 (Anthropic)<br>'
        '<b>Cập nhật:</b> 06/2026 &nbsp;·&nbsp; '
        '<b>Đường dây nóng Cục ATTP:</b> 1800 6838 (miễn phí) &nbsp;·&nbsp; '
        '<b>Cấp cứu:</b> 115'
        '</div></div>',
        unsafe_allow_html=True,
    )

    # ── Đổi mật khẩu — đặt trong Hướng dẫn, không làm rối header ─────────────
    _auth_user_guide = st.session_state.get("auth_user")
    _is_demo_guide   = (_auth_user_guide or {}).get("demo", False) if _auth_user_guide else True
    if _auth_user_guide and not _is_demo_guide:
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="background:#F8FAFC;border:1px solid #E2E8F0;border-radius:12px;'
            'padding:16px 20px;margin-top:8px">'
            '<div style="font-size:0.9rem;font-weight:700;color:#1E293B;margin-bottom:4px">'
            '🔑 Đổi mật khẩu</div>'
            '<div style="font-size:0.78rem;color:#64748B;margin-bottom:12px">'
            'Sau khi đổi, dùng mật khẩu mới cho lần đăng nhập tiếp theo.</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        _gpw_c1, _gpw_c2 = st.columns(2)
        _gpw1 = _gpw_c1.text_input("Mật khẩu mới (≥ 6 ký tự)", type="password",
                                     placeholder="Mật khẩu mới...", key="gpw_new")
        _gpw2 = _gpw_c2.text_input("Xác nhận mật khẩu mới", type="password",
                                     placeholder="Nhập lại...", key="gpw_confirm")
        if st.button("💾 Cập nhật mật khẩu", key="gpw_btn", type="primary",
                     use_container_width=False):
            if not _gpw1 or not _gpw2:
                st.warning("Vui lòng nhập cả 2 trường.")
            elif len(_gpw1) < 6:
                st.warning("Mật khẩu cần ≥ 6 ký tự.")
            elif _gpw1 != _gpw2:
                st.error("Mật khẩu xác nhận không khớp.")
            else:
                _tok_g = (_auth_user_guide or {}).get("access_token", "")
                _ok_g, _err_g = db_change_password(_tok_g, _gpw1)
                if _ok_g:
                    st.success("✅ Đã đổi mật khẩu! Dùng mật khẩu mới lần sau đăng nhập.")
                else:
                    st.error(f"❌ {_err_g}")


# ── Tạo sổ tay Word chuyên nghiệp ─────────────────────────────────────────────
def generate_manual_docx() -> bytes:
    """Tạo file Word sổ tay hướng dẫn chuẩn hành chính, font Times New Roman."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    mc  = MANUAL_CONTENT
    doc = Document()

    # Lề trang
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)

    # Bìa
    _docx_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
               bold=True, size=13, align="center", space_after=2)
    _docx_para(doc, "Độc lập – Tự do – Hạnh phúc",
               bold=True, size=12, align="center", space_after=2)
    _docx_para(doc, "────────────────────────────────",
               size=11, align="center", space_after=20)
    _docx_para(doc, mc["title"], bold=True, size=16, align="center",
               space_before=10, space_after=6)
    _docx_para(doc, mc["subtitle"], bold=True, size=13, align="center", space_after=4)
    _docx_para(doc, mc["version"], size=12, align="center", space_after=30)
    _docx_para(doc, "Dành cho: Phụ Huynh · Ban Giám Sát (Đại Diện PHHS) · Y Tế Học Đường · Ban Giám Hiệu",
               size=12, align="center", space_after=4)
    _docx_para(doc, f"Cập nhật: {now_vn().strftime('%d/%m/%Y')}",
               size=11, align="center", space_after=4)
    _docx_para(doc, "Đường dây nóng Cục ATTP: 1800 6838 (miễn phí) | Cấp cứu: 115",
               bold=True, size=12, align="center", space_after=0)

    doc.add_page_break()

    # Mục lục
    _docx_para(doc, "MỤC LỤC", bold=True, size=14, align="center",
               space_before=0, space_after=10)
    for section in mc["sections"]:
        _docx_para(doc, f"  {section['title']}", size=12, space_after=4)
    doc.add_page_break()

    # Nội dung từng section
    for section in mc["sections"]:
        _docx_para(doc, f"{section['icon']}  {section['title']}",
                   bold=True, size=14, space_before=12, space_after=6)

        if "content" in section:
            for para in section["content"].split("\n\n"):
                if para.strip():
                    _docx_para(doc, para.strip(), size=12, align="justify",
                               space_after=4)

        if "subsections" in section:
            for sub_title, sub_content in section["subsections"]:
                _docx_para(doc, sub_title, bold=True, size=12,
                           space_before=6, space_after=3)
                for line in sub_content.split("\n"):
                    if line.strip():
                        _docx_para(doc, line, size=11, align="justify",
                                   space_after=2)

        if "faq" in section:
            for q, a in section["faq"]:
                _docx_para(doc, f"Hỏi: {q}", bold=True, size=12,
                           space_before=6, space_after=2)
                _docx_para(doc, f"Đáp: {a}", size=11, align="justify",
                           space_after=4)

        if "legal_refs" in section:
            for name, desc in section["legal_refs"]:
                _docx_para(doc, name, bold=True, size=12, space_before=4, space_after=2)
                _docx_para(doc, desc, size=11, align="justify", space_after=4)

        if "terms" in section:
            tbl = doc.add_table(rows=1 + len(section["terms"]), cols=2)
            tbl.style = "Table Grid"
            _docx_table_header(tbl, ["Thuật ngữ / Ký hiệu", "Giải thích"])
            for i, (term, definition) in enumerate(section["terms"]):
                r0 = tbl.rows[i+1].cells[0].paragraphs[0].add_run(term)
                r1 = tbl.rows[i+1].cells[1].paragraphs[0].add_run(definition)
                _docx_set_font(r0, bold=True, size_pt=11)
                _docx_set_font(r1, size_pt=11)

        doc.add_paragraph()

    # Trang cuối
    doc.add_page_break()
    _docx_para(doc, "─" * 50, size=10, align="center", space_before=20, space_after=6)
    _docx_para(doc, f"Sổ tay được tạo tự động bởi SchoolFood AI v2.0",
               size=10, align="center", space_after=2)
    _docx_para(doc, f"Ngày tạo: {now_vn().strftime('%d/%m/%Y %H:%M')} (GMT+7)",
               size=10, align="center", space_after=2)
    _docx_para(doc, "Thông tin tham khảo — vui lòng xác nhận với cơ quan có thẩm quyền khi cần thiết",
               size=9, align="center", space_after=0)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


def tab_history(role: str = "", school_filter: str = ""):
    """Live dashboard chuyên nghiệp — biểu đồ cột, tròn, đường, phân bố."""
    import pandas as pd
    import plotly.express as px
    import plotly.graph_objects as go
    from io import BytesIO

    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">📊 Live Dashboard — Lịch sử kiểm tra ATTP</div>
        <div class="sf-card-body">
            Dữ liệu realtime từ database · Biểu đồ tương tác · Xuất Excel chuẩn
        </div>
    </div>""", unsafe_allow_html=True)

    # Thử kết nối và hiện lỗi cụ thể
    _get_sb()  # Trigger kết nối để lấy error message
    if not db_ok():
        err = db_error_msg()
        # Kiểm tra xem secrets có tồn tại không
        has_url = bool(st.secrets.get("SUPABASE_URL", "") if hasattr(st, "secrets") else "")
        has_key = bool(st.secrets.get("SUPABASE_ANON_KEY", "") if hasattr(st, "secrets") else "")
        diag = []
        if not has_url: diag.append("❌ SUPABASE_URL chưa có trong Secrets")
        if not has_key: diag.append("❌ SUPABASE_ANON_KEY chưa có trong Secrets")
        if has_url and has_key: diag.append("✅ Secrets có đủ 2 key — nhưng kết nối thất bại")
        diag_html = "<br>".join(diag)

        st.markdown(
            f'<div style="background:#FFF7ED;border:2px solid #FB923C;border-radius:12px;'
            f'padding:20px 24px">'
            f'<div style="font-size:1.1rem;font-weight:700;color:#9A3412;margin-bottom:8px">'
            f'📦 Database chưa được kết nối</div>'
            f'<div style="font-size:0.88rem;color:#7C2D12;line-height:1.9">'
            f'{diag_html}<br>'
            f'{"<b>Lỗi chi tiết:</b> " + err if err else ""}'
            f'</div>'
            f'<div style="margin-top:12px;font-size:0.82rem;color:#92400E">'
            f'Thêm vào Streamlit Secrets (Settings → Secrets):<br>'
            f'<code style="background:#FED7AA;padding:2px 8px;border-radius:4px;display:block;margin:4px 0">'
            f'SUPABASE_URL = "https://vmvensiremofatkylcuw.supabase.co"</code>'
            f'<code style="background:#FED7AA;padding:2px 8px;border-radius:4px;display:block;margin:4px 0">'
            f'SUPABASE_ANON_KEY = "eyJhbGci..."</code>'
            f'</div></div>',
            unsafe_allow_html=True,
        )
        return

    # ── Bộ lọc ────────────────────────────────────────────────────────────────
    # Nếu đã đăng nhập và không phải BGH → khóa theo trường của mình
    _lock     = bool(school_filter) and not st.session_state.get("is_super", False) and bool(st.session_state.get("auth_user"))

    if _lock:
        # Locked: hiện tên trường + view_mode cùng 1 hàng compact
        _lk1, _lk2 = st.columns([3, 1])
        _lk1.markdown(
            f'<div style="background:#F0FDF4;border:1px solid #86EFAC;border-radius:6px;'
            f'padding:4px 12px;font-size:0.82rem;color:#166534;'
            f'display:flex;align-items:center;height:38px;margin-top:2px">'
            f'🏫 <b>{school_filter}</b></div>',
            unsafe_allow_html=True,
        )
        view_mode = _lk2.selectbox("Hiển thị", ["Tất cả", "🍱 Bữa ăn", "🏭 Nhà Cung Cấp"],
                                    label_visibility="collapsed")
        _school_sel = school_filter
    else:
        # Admin/BGH: selectbox chọn trường + view_mode
        _fl1, _fl2 = st.columns([3, 1])
        _schools_db  = db_get_schools()
        _school_opts = ["Tất cả trường"] + _schools_db
        _default_idx = (_school_opts.index(school_filter)
                        if school_filter and school_filter in _school_opts else 0)
        _school_sel  = _fl1.selectbox(
            "Lọc theo tên trường", options=_school_opts, index=_default_idx,
            label_visibility="collapsed",
            help="Gõ tên trường để tìm nhanh",
        )
        view_mode = _fl2.selectbox("Hiển thị", ["Tất cả", "🍱 Bữa ăn", "🏭 Nhà Cung Cấp"],
                                    label_visibility="collapsed")

    school_input = "" if _school_sel == "Tất cả trường" else _school_sel
    sessions = db_get_sessions(school=school_input, limit=200)

    # ── 🚨 ALERT BANNER — hiện ngay đầu trang cho BGH/Admin ──────────────────
    # Điều kiện: BGH hoặc Admin (is_super), có DB, có hoặc không có sessions
    # → luôn kiểm tra tần suất, kể cả khi chưa có dữ liệu
    if (role == "Ban Giám Hiệu" or st.session_state.get("is_super")) and db_ok():
        import pandas as _pd_alert
        _now_alert = _pd_alert.Timestamp(now_vn().date())
        _quick_alerts = []

        # Task#7: NCC monthly check
        _this_m = now_vn().strftime("%Y-%m")
        _ncc_m = [s for s in sessions if s.get("check_type")=="nha_cung_cap"
                  and (s.get("check_date","") or "").startswith(_this_m)
                  and len(s.get("check_date","")) >= 7]
        if not _ncc_m:
            _quick_alerts.append(("⏰", f"Tháng {now_vn().month:02d} chưa có đánh giá NCC toàn diện (12 điểm)"))

        # BGS tần suất (lấy từ sessions)
        _bgs_dates = sorted([s["check_date"] for s in sessions
                             if s.get("check_type") == "ban_giam_sat" and s.get("check_date")],
                            reverse=True)
        if _bgs_dates:
            _last_bgs_dt = _pd_alert.Timestamp(_bgs_dates[0])
            _bgs_gap = (_now_alert - _last_bgs_dt).days
            _bgs_7d = sum(1 for d in _bgs_dates
                          if _pd_alert.Timestamp(d) >= _now_alert - _pd_alert.Timedelta(days=7))
            if _bgs_gap > 7:
                _quick_alerts.append(("🚨", f"BGS chưa kiểm tra {_bgs_gap} ngày — vi phạm NĐ 15/2018"))
            elif _bgs_7d < 2:
                _quick_alerts.append(("⏰", f"BGS chỉ kiểm tra {_bgs_7d} lần/tuần (yêu cầu ≥ 2)"))
        else:
            _quick_alerts.append(("🚨", "BGS chưa có lần kiểm tra nào — cần báo cáo ngay"))

        # Y Tế tần suất
        _yte_dates = sorted([s["check_date"] for s in sessions
                             if s.get("check_type") == "kiem_thuc_3_buoc" and s.get("check_date")],
                            reverse=True)
        if _yte_dates:
            _last_yte_dt = _pd_alert.Timestamp(_yte_dates[0])
            _yte_gap = (_now_alert - _last_yte_dt).days
            if _yte_gap > 3:
                _quick_alerts.append(("🚨", f"Y Tế chưa kiểm thực {_yte_gap} ngày — vi phạm TTLT 13/2016 Điều 9"))
        else:
            _quick_alerts.append(("🚨", "Y Tế chưa có kiểm thực nào — vi phạm TTLT 13/2016"))

        # Task#5: NCC cert expiry alert (≤ 30 ngày)
        try:
            _ncc_certs = db_get_ncc_registry(school=school_input)
            _today_chk = now_vn().date()
            import datetime as _dt_chk
            for _nc in _ncc_certs:
                for _f, _lbl in [("license_expiry","GP"), ("attp_expiry","ATTP")]:
                    _e = _nc.get(_f)
                    if _e:
                        _e_dt = _dt_chk.date.fromisoformat(_e)
                        _dl = (_e_dt - _today_chk).days
                        if _dl <= 0:
                            _quick_alerts.append(("🚨", f"NCC {_nc['ncc_name']}: Chứng nhận {_lbl} đã HẾT HẠN ({_e_dt.strftime('%d/%m/%Y')})"))
                        elif _dl <= 30:
                            _quick_alerts.append(("⏰", f"NCC {_nc['ncc_name']}: Chứng nhận {_lbl} hết hạn sau {_dl} ngày"))
        except Exception:
            pass

        # Complaints quá hạn (> 2 ngày)
        try:
            _pending_fb_alert = db_get_feedback(school=school_input, status="pending")
            _overdue_fb = 0
            for _f in _pending_fb_alert:
                try:
                    _f_dt = __import__("datetime").datetime.fromisoformat(
                        (_f.get("created_at",""))[:19]).replace(
                        tzinfo=__import__("datetime").timezone.utc)
                    if (now_vn() - _f_dt).days >= 2:
                        _overdue_fb += 1
                except Exception:
                    pass
            if _overdue_fb > 0:
                _quick_alerts.append(("⏰", f"{_overdue_fb} phản hồi PH quá 2 ngày chưa xử lý"))
        except Exception:
            pass

        # Chỉ hiện banner khi có THỰC SỰ có alert — không hiện "0 cảnh báo"
        if _quick_alerts:
            _crit_cnt = sum(1 for lvl, _ in _quick_alerts if lvl == "🚨")
            _warn_cnt = sum(1 for lvl, _ in _quick_alerts if lvl == "⏰")
            _top_bg   = "#FEF2F2" if _crit_cnt > 0 else "#FFFBEB"
            _top_bd   = "#FCA5A5" if _crit_cnt > 0 else "#FCD34D"
            _top_tc   = "#991B1B" if _crit_cnt > 0 else "#78350F"
            # Tiêu đề: chỉ nêu số > 0
            _title_parts = []
            if _crit_cnt > 0: _title_parts.append(f"{_crit_cnt} vấn đề nghiêm trọng")
            if _warn_cnt > 0: _title_parts.append(f"{_warn_cnt} cảnh báo")
            _title_str = " · ".join(_title_parts) + " — cần xử lý"
            _alert_html = "".join(
                f'<div style="margin:3px 0;font-size:0.85rem">{lvl} {msg}</div>'
                for lvl, msg in _quick_alerts
            )
            st.markdown(
                f'<div style="background:{_top_bg};border:2px solid {_top_bd};'
                f'border-radius:12px;padding:14px 18px;margin-bottom:12px">'
                f'<div style="font-weight:700;color:{_top_tc};margin-bottom:6px;font-size:0.92rem">'
                f'{"🚨" if _crit_cnt > 0 else "⚠️"} {_title_str}</div>'
                f'{_alert_html}'
                f'</div>',
                unsafe_allow_html=True,
            )

    # ── Feedback Phụ Huynh — luôn hiện TRƯỚC early return ────────────────────
    # ── Hệ thống Complaint — expandable, hiện TRƯỚC early return ─────────────
    # BGH: mở mặc định (họ cần xử lý) · BGS/Y Tế: đóng mặc định
    if role in ("Ban Giám Hiệu", "Ban Giám Sát (Đại Diện PHHS)", "Y Tế Học Đường") and db_ok():
        import plotly.graph_objects as _go_fb
        import re as _re_fb

        _all_fb    = db_get_all_feedbacks(school=school_input, limit=200)
        _is_bgh_fb = (role == "Ban Giám Hiệu")
        # Chỉ Y Tế Học Đường được thêm minh chứng
        _user_name = st.session_state.get("user_profile", {}).get("full_name", role)

        def _fmt_date_fb(iso_str: str) -> str:
            """ISO date → DD/MM/YYYY chuẩn Việt Nam."""
            d = (iso_str or "")[:10]
            return f"{d[8:10]}/{d[5:7]}/{d[:4]}" if len(d) >= 10 else d

        # ── Header nổi bật (giống section Bữa Ăn/NCC) ─────────────────────────
        st.markdown(
            '<div style="background:linear-gradient(135deg,#4C1D95 0%,#7C3AED 100%);'
            'border-radius:12px;padding:14px 22px;margin:20px 0 14px 0">'
            '<div style="color:white;font-size:1.05rem;font-weight:700;margin-bottom:2px">'
            '📬 PHẢN HỒI PHỤ HUYNH — QUẢN LÝ & THEO DÕI</div>'
            '<div style="color:#DDD6FE;font-size:0.8rem">'
            + ("BGH: Đọc minh chứng từ BGS/Y Tế → Đóng task & ghi phản hồi chính thức"
               if _is_bgh_fb else
               "Thêm minh chứng/diễn giải để hỗ trợ Ban Giám Hiệu xử lý phản hồi")
            + ' &nbsp;·&nbsp; Luồng: ⏳ Chờ → 💬 Đang xem (BGS thêm minh chứng) → ✅ Đóng (BGH)'
            + '</div></div>',
            unsafe_allow_html=True,
        )

        if _all_fb:
            # ── Metrics tổng quan ─────────────────────────────────────────────
            _total_fb    = len(_all_fb)
            _pending_fb  = sum(1 for f in _all_fb if f.get("status") == "pending")
            _reviewed_fb = sum(1 for f in _all_fb if f.get("status") == "reviewed")
            _resolved_fb = sum(1 for f in _all_fb if f.get("status") == "resolved")
            _overdue_fb  = sum(
                1 for f in _all_fb
                if f.get("status") == "pending"
                and f.get("created_at","")
                and (now_vn() - __import__("datetime").datetime.fromisoformat(
                    f["created_at"][:19]).replace(
                    tzinfo=__import__("datetime").timezone.utc)).days >= 2
            )
            _fm1,_fm2,_fm3,_fm4 = st.columns(4)
            _fm1.markdown(f'<div class="metric-box"><div class="metric-lbl">Tổng phản hồi</div>'
                          f'<div class="metric-num c-blue">{_total_fb}</div></div>',
                          unsafe_allow_html=True)
            _fm2.markdown(f'<div class="metric-box"><div class="metric-lbl">⏳ Chờ xử lý</div>'
                          f'<div class="metric-num c-orange">{_pending_fb}</div></div>',
                          unsafe_allow_html=True)
            _fm3.markdown(f'<div class="metric-box"><div class="metric-lbl">💬 Đang xem</div>'
                          f'<div class="metric-num c-blue">{_reviewed_fb}</div></div>',
                          unsafe_allow_html=True)
            _fm4.markdown(f'<div class="metric-box"><div class="metric-lbl">✅ Đã xử lý</div>'
                          f'<div class="metric-num c-green">{_resolved_fb}</div></div>',
                          unsafe_allow_html=True)
            if _overdue_fb:
                st.markdown(
                    f'<div style="background:#FEE2E2;border:1px solid #FCA5A5;border-radius:8px;'
                    f'padding:8px 16px;margin:8px 0;font-size:0.85rem;color:#991B1B;font-weight:600">'
                    f'🚨 {_overdue_fb} phản hồi quá 2 ngày chưa xử lý — cần ưu tiên giải quyết!</div>',
                    unsafe_allow_html=True,
                )

            # ── Chi tiết: Biểu đồ + Danh sách — collapsible ─────────────────
            st.markdown("<br>", unsafe_allow_html=True)  # khoảng cách như section Bữa Ăn / NCC
            _detail_label = ("📊 Chi tiết biểu đồ & danh sách xử lý"
                             + (f" ({_pending_fb} chờ, {_reviewed_fb} đang xem)"
                                if _pending_fb + _reviewed_fb > 0 else ""))
            with st.expander(_detail_label, expanded=_is_bgh_fb):
              _fc1, _fc2 = st.columns(2)
              with _fc1:
                # Bar chart: số complaint theo ngày (30 ngày gần nhất)
                try:
                    _dates_fb = {}
                    for f in _all_fb:
                        _d = (f.get("created_at","") or "")[:10]
                        if _d:
                            _dates_fb[_d] = _dates_fb.get(_d, 0) + 1
                    _sorted_dates = sorted(_dates_fb.keys())[-30:]
                    _date_labels = [f"{d[8:10]}/{d[5:7]}" for d in _sorted_dates]
                    _fig_bar_fb = _go_fb.Figure(_go_fb.Bar(
                        x=_date_labels, y=[_dates_fb[d] for d in _sorted_dates],
                        marker_color="#7C3AED", opacity=0.8,
                        hovertemplate="%{x}: %{y} phản hồi<extra></extra>",
                    ))
                    _fig_bar_fb.update_layout(
                        plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                        font=dict(size=11), margin=dict(l=10,r=10,t=30,b=10), height=200,
                        title=dict(text="📅 Phản hồi theo ngày (30 ngày gần nhất)",
                                   font=dict(size=12, color="#1B3B6F")),
                        xaxis=dict(showgrid=False, tickangle=-30),
                        yaxis=dict(showgrid=True, gridcolor="#E2E8F0", dtick=1),
                    )
                    st.plotly_chart(_fig_bar_fb, use_container_width=True)
                except Exception:
                    pass

              with _fc2:
                # Donut: by category
                try:
                    # Keywords mapping để legend ngắn gọn rõ ràng
                    _CAT_KW = {
                        "Chất lượng": "Chất lượng thức ăn",
                        "Vệ sinh":    "Vệ sinh bếp ăn",
                        "Nghi ngờ":   "Nghi ngờ ngộ độc",
                        "Thiếu":      "Thiếu dinh dưỡng",
                        "Thực đơn":   "Thực đơn không khớp",
                        "Góp ý":      "Góp ý khác",
                        "Khác":       "Khác",
                    }
                    _cats_fb = {}
                    for f in _all_fb:
                        _raw_c = (f.get("category","") or "Khác")
                        # Map sang keyword cố định
                        _c = "Khác"
                        for _kw, _lbl in _CAT_KW.items():
                            if _kw.lower() in _raw_c.lower():
                                _c = _lbl; break
                        _cats_fb[_c] = _cats_fb.get(_c, 0) + 1
                    _fig_pie_fb = _go_fb.Figure(_go_fb.Pie(
                        labels=list(_cats_fb.keys()), values=list(_cats_fb.values()),
                        hole=0.45, textfont_size=11,
                        marker_colors=["#7C3AED","#2563EB","#DC2626","#D97706","#0D9488","#64748B"],
                        hovertemplate="%{label}: %{value} (%{percent})<extra></extra>",
                    ))
                    _fig_pie_fb.update_layout(
                        plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                        font=dict(size=11), margin=dict(l=10,r=10,t=30,b=10), height=200,
                        title=dict(text="🏷️ Phản hồi theo loại",
                                   font=dict(size=12, color="#1B3B6F")),
                        showlegend=True,
                        legend=dict(font=dict(size=10), orientation="v", x=1, y=0.5),
                    )
                    st.plotly_chart(_fig_pie_fb, use_container_width=True)
                except Exception:
                    pass

              # ── Danh sách complaint chưa xử lý ────────────────────────────────
              _open_fb = [f for f in _all_fb if f.get("status") != "resolved"]
              if _open_fb:
                st.markdown(
                    f'<div class="sec-hdr" style="color:#DC2626">⏳ Chờ xử lý '
                    f'({len(_open_fb)} phản hồi)</div>',
                    unsafe_allow_html=True,
                )
                if _is_bgh_fb:
                    st.caption("BGH: Đọc minh chứng từ BGS/Y Tế và đóng task khi đã xử lý xong.")
                else:
                    st.caption("Cung cấp minh chứng/diễn giải để hỗ trợ Ban Giám Hiệu xử lý.")

                for _fb in _open_fb:
                    _fid   = _fb.get("id","")
                    _fdt   = (_fb.get("created_at","") or "")[:10]
                    _fcat  = _fb.get("category","")
                    _fcnt  = _fb.get("content","")
                    _fst   = _fb.get("status","pending")
                    _fevtx = _fb.get("evidence_text","") or ""
                    _fevby = _fb.get("evidence_by","") or ""
                    _fsch  = _fb.get("school_name","")

                    # Kiểm tra quá hạn (> 2 ngày)
                    _is_overdue = False
                    try:
                        _created = __import__("datetime").datetime.fromisoformat(
                            (_fb.get("created_at","") or "")[:19]).replace(
                            tzinfo=__import__("datetime").timezone.utc)
                        _is_overdue = (now_vn() - _created).days >= 2
                    except Exception:
                        pass

                    _fborder = "#DC2626" if _is_overdue else ("#2563EB" if _fst=="reviewed" else "#F59E0B")
                    _fbg     = "#FFF5F5" if _is_overdue else ("#EFF6FF" if _fst=="reviewed" else "#FFFBEB")
                    _fst_lbl = "💬 Đang xem xét" if _fst=="reviewed" else "⏳ Chờ xử lý"
                    _fst_clr = "#2563EB" if _fst=="reviewed" else "#D97706"

                    _overdue_tag = (
                        '<span style="background:#FEE2E2;color:#DC2626;font-size:0.68rem;'
                        'font-weight:700;padding:2px 8px;border-radius:8px;margin-left:6px">'
                        '🚨 QUÁ HẠN</span>' if _is_overdue else ""
                    )
                    st.markdown(
                        f'<div style="background:{_fbg};border:1.5px solid {_fborder};'
                        f'border-radius:10px;padding:12px 16px;margin:6px 0">'
                        f'<div style="display:flex;justify-content:space-between;align-items:center;'
                        f'flex-wrap:wrap;gap:6px;margin-bottom:6px">'
                        f'<span style="font-size:0.75rem;color:#64748B">'
                        f'📅 {_fdt} · 🏫 {_fsch} · {_fcat}</span>'
                        f'<span style="display:flex;align-items:center">'
                        f'<span style="background:{_fst_clr}20;color:{_fst_clr};font-size:0.72rem;'
                        f'font-weight:700;padding:2px 10px;border-radius:10px">{_fst_lbl}</span>'
                        f'{_overdue_tag}</span>'
                        f'</div>'
                        f'<div style="font-size:0.9rem;color:#1E293B;margin-bottom:6px">{_fcnt}</div>'
                        + (f'<div style="background:white;border-radius:6px;padding:8px 10px;'
                           f'font-size:0.8rem;color:#1D4ED8;margin-top:4px">'
                           f'📋 <b>Minh chứng từ {_fevby}:</b> {_fevtx}</div>'
                           if _fevtx else '')
                        + '</div>',
                        unsafe_allow_html=True,
                    )

                    # Form action theo vai trò
                    # Chỉ Y Tế Học Đường mới thêm minh chứng (họ có mặt trực tiếp)
                    if role == "Y Tế Học Đường":
                        _ev_lbl = "📝 Cập nhật minh chứng" if _fevtx else "📝 Thêm minh chứng"
                        with st.expander(_ev_lbl):
                            if _fevtx:
                                st.markdown(
                                    f'<div style="background:#EFF6FF;border-radius:6px;'
                                    f'padding:8px 12px;font-size:0.82rem;color:#1D4ED8;margin-bottom:8px">'
                                    f'📋 <b>Minh chứng đã gửi ({_fevby}):</b> {_fevtx}</div>',
                                    unsafe_allow_html=True,
                                )
                            _evtxt_new = st.text_area(
                                "Mô tả hành động đã thực hiện / kiểm tra",
                                key=f"ev_{_fid}", height=90,
                                placeholder="VD: Đã kiểm tra lô hàng cá ngày 03/06 lúc 10:30, nhiệt độ 62°C đạt chuẩn. Mùi vị bình thường. Đã báo cáo Ban Giám Hiệu...",
                            )
                            _ev_file = st.file_uploader(
                                "📎 Tải lên minh chứng (ảnh / PDF ≤ 5MB)",
                                type=["jpg","jpeg","png","pdf"],
                                key=f"ev_file_{_fid}", accept_multiple_files=False,
                            )
                            _ev_file_note = ""
                            if _ev_file:
                                if _ev_file.type.startswith("image"):
                                    st.image(_ev_file, width=220, caption="Ảnh đính kèm")
                                else:
                                    st.caption(f"📄 File đính kèm: {_ev_file.name}")
                                _ev_file_note = f" [File: {_ev_file.name}]"
                            if st.button("📤 Gửi minh chứng", key=f"ev_btn_{_fid}",
                                         type="primary", use_container_width=True):
                                _final_ev = (_evtxt_new.strip() or _fevtx or "") + _ev_file_note
                                if _final_ev.strip():
                                    if db_add_evidence(_fid, _final_ev, _user_name):
                                        st.success("✅ Đã gửi — Ban Giám Hiệu sẽ xem xét.")
                                        st.rerun()
                                else:
                                    st.warning("Vui lòng nhập mô tả hoặc tải file.")

                    elif _is_bgh_fb:
                        with st.expander("🔒 Đóng task & ghi phản hồi chính thức"):
                            if _fevtx:
                                st.markdown(
                                    f'<div style="background:#EFF6FF;border-radius:6px;'
                                    f'padding:8px 12px;font-size:0.82rem;color:#1D4ED8;margin-bottom:8px">'
                                    f'📋 <b>Minh chứng từ {_fevby}:</b> {_fevtx}</div>',
                                    unsafe_allow_html=True,
                                )
                            _reptxt = st.text_area(
                                "Phản hồi chính thức của Ban Giám Hiệu",
                                key=f"rep_{_fid}", height=80,
                                placeholder="VD: Nhà trường đã xem xét và xác nhận. Sẽ yêu cầu nhà cung cấp đổi lô hàng và kiểm tra thêm. Cảm ơn phụ huynh đã phản ánh.",
                            )
                            if st.button("🔒 Đóng task & lưu phản hồi", key=f"rep_btn_{_fid}",
                                         type="primary", use_container_width=True):
                                _rtext = _reptxt.strip() or "Ban Giám Hiệu đã xem xét và xử lý."
                                if db_resolve_complaint(_fid, _rtext, _user_name):
                                    st.success("✅ Đã đóng task!")
                                    st.rerun()

              else:
                st.success("✅ Tất cả phản hồi đã được xử lý!")

              # ── Phần đã đóng — phân biệt rõ với phần mở ─────────────────────
              _closed_fb = [f for f in _all_fb if f.get("status") == "resolved"][:10]
              if _closed_fb:
                st.markdown(
                    '<div style="background:linear-gradient(135deg,#1E293B,#334155);'
                    'border-radius:10px;padding:10px 18px;margin:16px 0 8px 0">'
                    f'<span style="color:white;font-size:0.9rem;font-weight:700">'
                    f'✅ {len(_closed_fb)} phản hồi đã xử lý gần nhất</span>'
                    '</div>',
                    unsafe_allow_html=True,
                )
                for _cf in _closed_fb:
                    _cdt  = _fmt_date_fb(_cf.get("created_at",""))
                    # Keyword ngắn cho category
                    _ccat_raw = _cf.get("category","")
                    _ccat = "Khác"
                    for _kw2, _lbl2 in _CAT_KW.items():
                        if _kw2.lower() in _ccat_raw.lower():
                            _ccat = _lbl2; break
                    _ccnt = (_cf.get("content","") or "").strip()
                    _crep = (_cf.get("response_text","") or "").strip()
                    _crby = (_cf.get("response_by","") or "").strip()
                    _crev = _fmt_date_fb(_cf.get("reviewed_at",""))
                    _cev  = (_cf.get("evidence_text","") or "").strip()
                    _ceby = (_cf.get("evidence_by","") or "").strip()
                    st.markdown(
                        f'<div style="background:#F8FAFC;border:1px solid #E2E8F0;'
                        f'border-radius:8px;padding:12px 16px;margin:5px 0">'
                        f'<div style="display:flex;justify-content:space-between;'
                        f'align-items:center;margin-bottom:6px">'
                        f'<span style="font-size:0.75rem;color:#64748B;font-weight:600">'
                        f'📅 {_cdt} &nbsp;·&nbsp; {_ccat}</span>'
                        f'<span style="background:#DCFCE7;color:#166534;font-size:0.7rem;'
                        f'font-weight:700;padding:2px 10px;border-radius:8px">'
                        f'✅ Đóng {_crev}</span>'
                        f'</div>'
                        f'<div style="font-size:0.88rem;color:#334155;margin-bottom:6px">'
                        f'"{_ccnt}"</div>'
                        + (f'<div style="background:#EFF6FF;border-radius:6px;'
                           f'padding:5px 10px;margin-bottom:4px;font-size:0.78rem;color:#1D4ED8">'
                           f'📋 <b>Minh chứng Y Tế ({_ceby}):</b> {_cev}</div>'
                           if _cev and _ceby else '')
                        + (f'<div style="background:#F0FDF4;border-radius:6px;'
                           f'padding:5px 10px;font-size:0.78rem;color:#166534">'
                           f'💬 <b>BGH phản hồi ({_crby}):</b> {_crep}</div>'
                           if _crep else
                           '<div style="font-size:0.75rem;color:#94A3B8;font-style:italic">'
                           'BGH chưa ghi phản hồi văn bản</div>')
                        + '</div>',
                        unsafe_allow_html=True,
                    )
        else:
            st.info("Chưa có phản hồi nào từ Phụ Huynh.")
        st.markdown('<div class="sf-div"></div>', unsafe_allow_html=True)

    if not sessions:
        st.info("Chưa có dữ liệu lịch sử. Thực hiện kiểm tra và tạo báo cáo lần đầu.")
        return

    # ── Chuẩn bị dataframe chung ──────────────────────────────────────────────
    ALERT_VN = {"OK": "Đạt chuẩn", "MINOR": "Cần cải thiện",
                "MAJOR": "Không đạt", "CRITICAL": "Nguy hiểm"}
    TYPE_VN  = {"ban_giam_sat": "Ban Giám Sát", "kiem_thuc_3_buoc": "Y Tế (3 bước)",
                "nha_cung_cap": "Nhà cung cấp"}
    MEAL_TYPES = {"ban_giam_sat", "kiem_thuc_3_buoc"}

    rows_meal, rows_ncc = [], []
    for s in sessions:
        ctype = s.get("check_type", "")
        pct   = s.get("pass_count", 0) / max(s.get("total_items", 20), 1) * 100
        base  = {
            "Ngày":           s.get("check_date", ""),
            "Trường":         s.get("school_name", ""),
            "Người kiểm tra": s.get("inspector_name", ""),
            "Tỷ lệ đạt (%)":  round(pct, 1),
            "Điểm đạt":       s.get("pass_count", 0),
            "Điểm không đạt": s.get("fail_count", 0),
            "Tổng điểm":      s.get("total_items", 20),
            "Cấp cảnh báo":   s.get("alert_level", ""),
            "Đánh giá":       ALERT_VN.get(s.get("alert_level", ""), s.get("alert_level", "")),
        }
        if ctype in MEAL_TYPES:
            rows_meal.append({**base, "Loại": TYPE_VN.get(ctype, ctype)})
        elif ctype == "nha_cung_cap":
            # Trích tên NCC từ menu_today ("NCC: Tên NCC")
            raw_menu = s.get("menu_today", "")
            ncc_name = raw_menu.replace("NCC:", "").strip() if raw_menu.startswith("NCC:") else raw_menu
            rating = "A" if s.get("pass_count", 0) >= 10 else \
                     "B" if s.get("pass_count", 0) >= 8 else "C"
            rows_ncc.append({**base, "Nhà Cung Cấp": ncc_name or "—", "Xếp loại": f"Loại {rating}"})

    df_meal = pd.DataFrame(rows_meal)
    df_ncc  = pd.DataFrame(rows_ncc)

    show_meal = view_mode in ("Tất cả", "🍱 Bữa ăn") and not df_meal.empty
    show_ncc  = view_mode in ("Tất cả", "🏭 Nhà Cung Cấp") and not df_ncc.empty

    if not show_meal and not show_ncc:
        st.info("Chưa có dữ liệu cho loại hiển thị đã chọn.")
        return

    # Thiết lập chung cho biểu đồ
    _CHART_LAYOUT = dict(
        plot_bgcolor="white", paper_bgcolor="#F8FAFC",
        font=dict(family="Inter, sans-serif", size=12, color="#334155"),
        title_font=dict(size=14, color="#1B3B6F", family="Inter"),
        margin=dict(l=16, r=16, t=40, b=16),
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1: BỮA ĂN
    # ══════════════════════════════════════════════════════════════════════════
    if show_meal:
        st.markdown(
            '<div style="background:linear-gradient(135deg,#0F2651 0%,#1D4ED8 100%);'
            'border-radius:12px;padding:14px 22px;margin:16px 0 14px 0">'
            '<div style="color:white;font-size:1.05rem;font-weight:700;margin-bottom:2px">'
            '🍱 KẾT QUẢ KIỂM TRA BỮA ĂN</div>'
            '<div style="color:#BFDBFE;font-size:0.8rem">'
            'Checklist 20 điểm (Ban Giám Sát) &nbsp;·&nbsp; Kiểm thực 3 bước (Y Tế Học Đường)'
            '</div></div>',
            unsafe_allow_html=True,
        )

        df = df_meal
        total   = len(df)
        avg_pct = df["Tỷ lệ đạt (%)"].mean()
        crit_ct = (df["Cấp cảnh báo"] == "CRITICAL").sum()
        ok_ct   = (df["Cấp cảnh báo"] == "OK").sum()

        n_split    = max(1, len(df) // 3)
        recent_avg = df.head(n_split)["Tỷ lệ đạt (%)"].mean()
        older_avg  = df.tail(n_split)["Tỷ lệ đạt (%)"].mean()
        delta      = recent_avg - older_avg
        if abs(delta) < 1:
            trend_icon, trend_text, trend_bg, trend_tc = (
                ("✅", "Đang tốt",    "#DCFCE7", "#16A34A") if avg_pct >= 90
                else ("⚠️", "Cần theo dõi", "#FEF9C3", "#CA8A04"))
        elif delta > 0:
            trend_icon, trend_text, trend_bg, trend_tc = "📈","Đang cải thiện","#DBEAFE","#2563EB"
        else:
            trend_icon, trend_text, trend_bg, trend_tc = "📉","Đang giảm","#FEE2E2","#DC2626"

        k1, k2, k3, k4, k5 = st.columns(5)
        for _col, _val, _lbl, _clr in [
            (k1, str(total),        "Tổng lần kiểm tra",  "c-blue"),
            (k2, f"{avg_pct:.0f}%", "Trung bình đạt",     "c-green" if avg_pct >= 90 else "c-orange"),
            (k3, str(crit_ct),      "Mức CRITICAL",        "c-red" if crit_ct > 0 else "c-green"),
            (k4, str(ok_ct),        "Đạt chuẩn (OK)",     "c-green"),
        ]:
            _col.markdown(
                f'<div class="metric-box" style="text-align:center">'
                f'<div class="metric-lbl">{_lbl}</div>'
                f'<div class="metric-num {_clr}" style="font-size:2rem">{_val}</div>'
                f'</div>', unsafe_allow_html=True,
            )
        k5.markdown(
            f'<div class="metric-box" style="text-align:center;background:{trend_bg};border:1px solid {trend_tc}">'
            f'<div class="metric-lbl">Xu hướng</div>'
            f'<div class="metric-num" style="font-size:1.5rem;color:{trend_tc};line-height:1.2">'
            f'{trend_icon}<br><span style="font-size:0.75rem;font-weight:700">{trend_text}</span></div>'
            f'</div>', unsafe_allow_html=True,
        )
        # So sánh tháng này vs tháng trước (thay benchmark 1 dòng)
        try:
            _now_cmp = now_vn()
            _cur_m   = _now_cmp.strftime("%Y-%m")
            import datetime as _dt_cmp
            _prev_dt = (_now_cmp.replace(day=1) - _dt_cmp.timedelta(days=1))
            _prev_m  = _prev_dt.strftime("%Y-%m")

            def _month_avg(m_str):
                _rs = [s for s in sessions
                       if (s.get("check_date","") or "").startswith(m_str)
                       and s.get("check_type") in ("ban_giam_sat","kiem_thuc_3_buoc")]
                if not _rs: return None
                return sum(r["pass_count"]/max(r["total_items"],1)*100 for r in _rs)/len(_rs)

            _cur_avg  = _month_avg(_cur_m)
            _prev_avg = _month_avg(_prev_m)

            if _cur_avg is not None and _prev_avg is not None:
                _delta = _cur_avg - _prev_avg
                _cmp_ic  = "📈" if _delta > 0 else "📉" if _delta < -1 else "➡️"
                _cmp_c   = "#16A34A" if _delta > 0 else "#DC2626" if _delta < -1 else "#64748B"
                _cmp_txt = f"+{_delta:.1f}%" if _delta > 0 else f"{_delta:.1f}%"
                st.markdown(
                    f'<div style="background:white;border:1px solid #E2E8F0;border-radius:8px;'
                    f'padding:8px 14px;margin:6px 0;display:flex;align-items:center;gap:12px">'
                    f'<span style="font-size:1.5rem">{_cmp_ic}</span>'
                    f'<div style="flex:1">'
                    f'<b style="font-size:0.85rem;color:#1E293B">'
                    f'Tháng {_now_cmp.month:02d}: {_cur_avg:.0f}%</b>'
                    f'<span style="font-size:0.78rem;color:#64748B;margin-left:8px">'
                    f'vs tháng {_prev_dt.month:02d}: {_prev_avg:.0f}%</span>'
                    f'</div>'
                    f'<span style="font-size:0.85rem;font-weight:700;color:{_cmp_c}">{_cmp_txt}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            elif _cur_avg is not None:
                st.markdown(
                    f'<div style="background:white;border:1px solid #E2E8F0;border-radius:8px;'
                    f'padding:8px 14px;margin:6px 0">'
                    f'📅 <b>Tháng {_now_cmp.month:02d}:</b> TB {_cur_avg:.0f}% '
                    f'<span style="color:#64748B;font-size:0.78rem">'
                    f'(Chưa có dữ liệu tháng trước để so sánh)</span></div>',
                    unsafe_allow_html=True,
                )
        except Exception:
            pass

        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("📊 Biểu đồ & Bảng chi tiết — Bữa ăn", expanded=False):

            # Hàng 1: Xu hướng + Phân bố cảnh báo
            ch1, ch2 = st.columns([3, 2])
            with ch1:
                df_agg = (df.groupby("Ngày")["Tỷ lệ đạt (%)"].mean()
                          .reset_index().sort_values("Ngày").tail(30))
                try:
                    df_agg["Ngày_fmt"] = pd.to_datetime(df_agg["Ngày"]).dt.strftime("%d/%m")
                except Exception:
                    df_agg["Ngày_fmt"] = df_agg["Ngày"]
                fig_line = go.Figure()
                fig_line.add_trace(go.Scatter(
                    x=df_agg["Ngày_fmt"], y=df_agg["Tỷ lệ đạt (%)"].round(1),
                    mode="lines+markers+text",
                    line=dict(color="#2563EB", width=3, shape="spline"),
                    marker=dict(size=10, color="#2563EB", line=dict(width=2.5, color="white")),
                    text=[f"{v:.0f}%" for v in df_agg["Tỷ lệ đạt (%)"]],
                    textposition="top center",
                    textfont=dict(size=12, color="#1E293B", family="Inter"),
                    hovertemplate="Ngày %{x}<br>Tỷ lệ đạt: <b>%{y:.1f}%</b><extra></extra>",
                ))
                fig_line.add_hline(y=90, line_dash="dot", line_color="#DC2626", line_width=2,
                                   annotation_text=" Ngưỡng 90%",
                                   annotation_font=dict(color="#DC2626", size=11, family="Inter"),
                                   annotation_position="bottom right")
                fig_line.add_hrect(y0=90, y1=110, fillcolor="#DCFCE7", opacity=0.15,
                                   layer="below", line_width=0)
                fig_line.update_layout(
                    **_CHART_LAYOUT, height=300,
                    title="📈 Xu hướng tỷ lệ đạt bữa ăn theo ngày",
                    xaxis=dict(title="Ngày kiểm tra", showgrid=False,
                               tickangle=0 if len(df_agg) <= 10 else -30),
                    yaxis=dict(title="Tỷ lệ đạt (%)", range=[0, 110],
                               ticksuffix="%", showgrid=True, gridcolor="#E2E8F0", dtick=20),
                    showlegend=False,
                )
                st.plotly_chart(fig_line, use_container_width=True)

            with ch2:
                alert_counts = df["Đánh giá"].value_counts().reset_index()
                alert_counts.columns = ["Mức", "Số lần"]
                fig_pie = go.Figure(go.Pie(
                    labels=alert_counts["Mức"], values=alert_counts["Số lần"],
                    hole=0.45,
                    marker_colors=[{"Đạt chuẩn": "#16A34A", "Cần cải thiện": "#F59E0B",
                                    "Không đạt": "#D97706", "Nguy hiểm": "#DC2626"
                                    }.get(m, "#64748B") for m in alert_counts["Mức"]],
                    textfont_size=13, textinfo="percent+label",
                    hovertemplate="%{label}<br>%{value} lần (%{percent})<extra></extra>",
                ))
                fig_pie.update_layout(
                    **_CHART_LAYOUT, height=300,
                    title="🥧 Phân bố mức cảnh báo bữa ăn",
                    showlegend=False,
                    annotations=[dict(text=f"<b>{total}</b><br>lần",
                                      x=0.5, y=0.5, showarrow=False,
                                      font_size=15, font_color="#1E293B")],
                )
                st.plotly_chart(fig_pie, use_container_width=True)

            # Hàng 2: Theo tuần + Theo loại kiểm tra
            ch3, ch4 = st.columns([3, 2])
            with ch3:
                try:
                    df_wk = df.copy()
                    _wk_dt = pd.to_datetime(df_wk["Ngày"], errors="coerce")
                    _wk_periods = _wk_dt.dt.to_period("W")
                    df_wk["Tuần"] = _wk_periods.apply(
                        lambda p: f"{p.start_time.strftime('%d/%m')}-{p.end_time.strftime('%d/%m')}"
                        if pd.notna(p) else "?")
                    _wk_order  = _wk_periods.dropna().sort_values().unique()
                    _wk_labels = [f"{p.start_time.strftime('%d/%m')}-{p.end_time.strftime('%d/%m')}"
                                  for p in _wk_order]
                    wk_cnt = df_wk.groupby(["Tuần", "Đánh giá"]).size().reset_index(name="Số lần")
                    _CMAP = {"Đạt chuẩn": "#16A34A", "Cần cải thiện": "#F59E0B",
                             "Không đạt": "#F97316", "Nguy hiểm": "#DC2626"}
                    fig_bar = px.bar(wk_cnt, x="Tuần", y="Số lần", color="Đánh giá",
                                     color_discrete_map=_CMAP, barmode="stack", text="Số lần",
                                     category_orders={"Tuần": _wk_labels})
                    fig_bar.update_traces(
                        textposition="inside",
                        insidetextanchor="middle",   # canh giữa trong mỗi segment
                        textfont_size=12,
                        textfont_color="white",
                    )
                    fig_bar.update_layout(
                        **{**_CHART_LAYOUT, "margin": dict(l=16, r=130, t=40, b=16)},
                        height=300, title="📊 Số lần kiểm tra bữa ăn theo tuần",
                        xaxis=dict(title="Tuần", tickangle=-20, showgrid=False),
                        yaxis=dict(title="Số lần", showgrid=True, gridcolor="#E2E8F0"),
                        legend=dict(orientation="v", x=1.02, y=0.5, xanchor="left",
                                    yanchor="middle", title_text="Đánh giá",
                                    bgcolor="rgba(255,255,255,0.8)",
                                    bordercolor="#E2E8F0", borderwidth=1),
                    )
                    st.plotly_chart(fig_bar, use_container_width=True)
                except Exception as _e:
                    st.info(f"Cần thêm dữ liệu. ({_e})")

            with ch4:
                type_stats = (df.groupby("Loại")
                              .agg(avg_pct=("Tỷ lệ đạt (%)", "mean"),
                                   count=("Tỷ lệ đạt (%)", "count"))
                              .reset_index().sort_values("avg_pct", ascending=True))
                fig_type = go.Figure(go.Bar(
                    x=type_stats["avg_pct"].round(1), y=type_stats["Loại"],
                    orientation="h",
                    marker_color=["#16A34A" if v >= 90 else "#F59E0B" if v >= 80 else "#F97316"
                                  for v in type_stats["avg_pct"]],
                    # Text đặt TRONG bar, canh giữa — tránh bị cắt khi bar gần 100%
                    text=[f"{v:.0f}%  ({c} lần)"
                          for v, c in zip(type_stats["avg_pct"], type_stats["count"])],
                    textposition="inside",
                    insidetextanchor="middle",
                    textfont=dict(size=12, color="white"),
                    hovertemplate="<b>%{y}</b><br>TB đạt: <b>%{x:.1f}%</b><extra></extra>",
                ))
                fig_type.add_vline(x=90, line_dash="dot", line_color="#DC2626", line_width=1.5,
                                   annotation_text=" Chuẩn 90%",
                                   annotation_font=dict(size=11, color="#DC2626"),
                                   annotation_position="top")
                fig_type.update_layout(
                    **{**_CHART_LAYOUT, "margin": dict(l=10, r=20, t=40, b=16)},
                    height=300, title="🔍 Tỷ lệ đạt TB theo loại kiểm tra",
                    xaxis=dict(title="Tỷ lệ đạt (%)", range=[0, 105],
                               ticksuffix="%", showgrid=True, gridcolor="#E2E8F0"),
                    yaxis=dict(title="", showgrid=False, automargin=True),
                    showlegend=False,
                )
                st.plotly_chart(fig_type, use_container_width=True)

            # Top 10 fail bữa ăn — lọc theo session IDs của trường đã chọn
            st.markdown('<div class="sec-hdr">🔴 Top 10 điểm không đạt nhiều nhất (bữa ăn)</div>',
                        unsafe_allow_html=True)
            try:
                sb = _get_sb()
                if sb:
                    # Lấy session IDs của bữa ăn đã được lọc theo trường
                    _meal_ids = [s["id"] for s in sessions
                                 if s.get("check_type") in ("ban_giam_sat", "kiem_thuc_3_buoc")
                                 and s.get("id")]
                    if not _meal_ids:
                        st.info("Chưa có dữ liệu điểm không đạt.")
                    else:
                        # Chia nhỏ thành batch ≤ 100 IDs (giới hạn Supabase .in_())
                        _all_items: list = []
                        for _i in range(0, len(_meal_ids), 100):
                            _batch = _meal_ids[_i:_i+100]
                            _r = sb.table("checklist_results")\
                                .select("item_code,result")\
                                .in_("session_id", _batch)\
                                .in_("result", ["Không Đạt", "❌ Không Đạt"])\
                                .execute()
                            _all_items += (_r.data or [])

                        if not _all_items:
                            st.info("Chưa có điểm không đạt cho trường / bộ lọc đã chọn.")
                        else:
                            _KW = {
                                "C01":"Tem kiểm dịch thịt/cá","C02":"Hóa đơn nguồn gốc rau củ",
                                "C03":"Hạn sử dụng nguyên liệu","C04":"Hóa đơn mua hàng ngày",
                                "C05":"Nhiệt độ tủ lạnh < 5°C","C06":"Tách biệt thực phẩm sống/chín",
                                "C07":"Nhiệt độ nhận hàng ≥ 60°C","C08":"Thùng vận chuyển kín, sạch",
                                "C09":"Nhiệt độ chia ăn đúng chuẩn","C10":"Thời gian nấu → phục vụ < 2h",
                                "C11":"Màu sắc & mùi vị thức ăn","C12":"Khẩu phần thịt/cá đủ định mức",
                                "C13":"Khẩu phần rau xanh đủ định mức","C14":"Dụng cụ ăn sạch, khô ráo",
                                "C15":"Đeo khẩu trang & găng tay","C16":"Không ho/hắt hơi vào thức ăn",
                                "C17":"Khu vực chia cơm sạch","C18":"Sổ kiểm thực 3 bước đủ chữ ký",
                                "C19":"Thực đơn khớp đăng ký","C20":"Mẫu lưu thức ăn 24h đủ nhãn",
                                "B1_01":"Tem kiểm dịch (B1)","B1_02":"Hóa đơn rau củ (B1)",
                                "B1_03":"Hạn sử dụng (B1)","B1_04":"Nhiệt độ tủ lạnh (B1)",
                                "B1_05":"Tách biệt sống/chín (B1)","B2_01":"Nấu chín ≥ 70°C (B2)",
                                "B2_02":"Bảo hộ lao động (B2)","B2_03":"Dao thớt riêng sống/chín (B2)",
                                "B2_04":"Dụng cụ nấu sạch (B2)","B2_05":"Bếp sạch (B2)",
                                "B3_01":"Nhiệt độ chia đúng chuẩn (B3)","B3_02":"Thời gian nấu→chia (B3)",
                                "B3_03":"Màu sắc & mùi vị (B3)","B3_04":"Khẩu phần đủ định mức (B3)",
                                "B3_05":"Mẫu lưu 24h đủ nhãn (B3)",
                            }
                            df_it = pd.DataFrame(_all_items)
                            df_it = df_it[df_it["item_code"].str.startswith(("C", "B"))]
                            df_it["Tên điểm"] = df_it["item_code"].map(lambda c: _KW.get(c, c))
                            top_fail = (df_it.groupby("Tên điểm").size()
                                        .reset_index(name="Số lần")
                                        .sort_values("Số lần").tail(10))
                            n = len(top_fail)
                            fig_hbar = go.Figure(go.Bar(
                                x=top_fail["Số lần"], y=top_fail["Tên điểm"], orientation="h",
                                marker_color=[f"rgba(220,38,38,{0.25+0.75*i/max(n-1,1)})"
                                              for i in range(n)],
                                text=top_fail["Số lần"], textposition="outside",
                                textfont=dict(size=12),
                                hovertemplate="<b>%{y}</b><br>Không đạt: %{x} lần<extra></extra>",
                            ))
                            fig_hbar.update_layout(
                                plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                                font=dict(family="Inter, sans-serif", size=12, color="#334155"),
                                # Không set title key → Plotly không render chữ nào
                                margin=dict(l=10, r=60, t=4, b=16),
                                height=max(280, n * 50),
                                xaxis=dict(title="Số lần không đạt", showgrid=True,
                                           gridcolor="#E2E8F0", dtick=1),
                                yaxis=dict(showgrid=False, tickfont=dict(size=11), automargin=True),
                            )
                            st.plotly_chart(fig_hbar, use_container_width=True)
            except Exception as _e:
                st.warning(f"Không thể tải top fail: {_e}")

            # Bảng chi tiết bữa ăn
            st.markdown('<div class="sec-hdr">📋 Bảng chi tiết — Bữa ăn</div>', unsafe_allow_html=True)
            _dm = df_meal.drop(columns=["Cấp cảnh báo"], errors="ignore").copy()
            _dm["Ngày"] = pd.to_datetime(_dm["Ngày"], errors="coerce").dt.strftime("%d/%m/%Y").fillna(_dm["Ngày"])
            st.dataframe(_dm, use_container_width=True, hide_index=True)

        # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2: NHÀ CUNG CẤP
    # ══════════════════════════════════════════════════════════════════════════
    # NCC section — hiện banner tím dù có hay không có data
    st.markdown(
        '<div style="background:linear-gradient(135deg,#3B0764 0%,#7C3AED 100%);'
        'border-radius:12px;padding:14px 22px;margin:24px 0 14px 0">'
        '<div style="color:white;font-size:1.05rem;font-weight:700;margin-bottom:2px">'
        '🏭 KẾT QUẢ ĐÁNH GIÁ NHÀ CUNG CẤP</div>'
        '<div style="color:#DDD6FE;font-size:0.8rem">'
        'Checklist 12 điểm giao hàng · Xếp loại A (≥10/12) · B (8–9/12) · C (<8)'
        '</div></div>',
        unsafe_allow_html=True,
    )

    if not show_ncc:
        # Empty state — hiện thông tin hướng dẫn thay vì trống
        _ncc_empty_bg = "#F5F3FF" if view_mode in ("Tất cả","🏭 Nhà Cung Cấp") else "transparent"
        st.markdown(
            f'<div style="background:{_ncc_empty_bg};border:1px dashed #DDD6FE;'
            f'border-radius:12px;padding:24px;text-align:center;margin:4px 0">'
            f'<div style="font-size:2rem;margin-bottom:8px">🏭</div>'
            f'<div style="font-size:0.95rem;font-weight:600;color:#5B21B6;margin-bottom:6px">'
            f'Chưa có kết quả đánh giá Nhà Cung Cấp</div>'
            f'<div style="font-size:0.82rem;color:#64748B;line-height:1.7">'
            f'Y Tế Học Đường kiểm tra giao hàng hàng ngày (S03–S12 · 10 điểm)<br>'
            f'Ban Giám Sát kiểm tra toàn diện 1 lần/tháng (S01–S12 · 12 điểm)<br>'
            f'<b>→ Vào tab 🏭 Nhà Cung Cấp để bắt đầu đánh giá lần đầu</b>'
            f'</div></div>',
            unsafe_allow_html=True,
        )
    elif show_ncc:

        ncc_total = len(df_ncc)
        ncc_a = sum(1 for v in df_ncc["Xếp loại"] if v == "Loại A")
        ncc_b = sum(1 for v in df_ncc["Xếp loại"] if v == "Loại B")
        ncc_c = sum(1 for v in df_ncc["Xếp loại"] if v == "Loại C")

        n1, n2, n3, n4 = st.columns(4)
        for _nc, _nv, _nl, _ncl in [
            (n1, str(ncc_total), "Tổng lần kiểm tra NCC",  "c-blue"),
            (n2, str(ncc_a),     "✅ Loại A — Đạt chuẩn",   "c-green"),
            (n3, str(ncc_b),     "🟡 Loại B — Cần cải thiện","c-orange"),
            (n4, str(ncc_c),     "🔴 Loại C — Không đạt",   "c-red" if ncc_c > 0 else "c-green"),
        ]:
            _nc.markdown(
                f'<div class="metric-box" style="text-align:center">'
                f'<div class="metric-lbl">{_nl}</div>'
                f'<div class="metric-num {_ncl}" style="font-size:2rem">{_nv}</div>'
                f'</div>', unsafe_allow_html=True,
            )
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("📊 Biểu đồ & Bảng chi tiết — Nhà Cung Cấp", expanded=False):

            # Hàng: Hiệu suất NCC theo tên + Xu hướng theo thời gian
            nc1, nc2 = st.columns([3, 2])
            with nc1:
                if "Nhà Cung Cấp" in df_ncc.columns and df_ncc["Nhà Cung Cấp"].nunique() > 0:
                    ncc_perf = (df_ncc.groupby("Nhà Cung Cấp")
                                .agg(avg_pct=("Tỷ lệ đạt (%)", "mean"),
                                     count=("Tỷ lệ đạt (%)", "count"))
                                .reset_index().sort_values("avg_pct", ascending=True))
                    fig_ncc = go.Figure(go.Bar(
                        x=ncc_perf["avg_pct"].round(1), y=ncc_perf["Nhà Cung Cấp"],
                        orientation="h",
                        marker_color=["#16A34A" if v >= 83 else "#F59E0B" if v >= 67 else "#DC2626"
                                      for v in ncc_perf["avg_pct"]],
                        text=[f"{v:.0f}%  ({c} lần)"
                              for v, c in zip(ncc_perf["avg_pct"], ncc_perf["count"])],
                        textposition="inside", insidetextanchor="middle",
                        textfont=dict(size=12, color="white"),
                        hovertemplate="<b>%{y}</b><br>TB đạt: <b>%{x:.1f}%</b><extra></extra>",
                    ))
                    fig_ncc.add_vline(x=83, line_dash="dot", line_color="#16A34A", line_width=1.5,
                                      annotation_text=" Loại A (83%)",
                                      annotation_font=dict(size=10, color="#16A34A"),
                                      annotation_position="top")
                    fig_ncc.add_vline(x=67, line_dash="dot", line_color="#F59E0B", line_width=1.5,
                                      annotation_text=" Loại B (67%)",
                                      annotation_font=dict(size=10, color="#F59E0B"),
                                      annotation_position="bottom")
                    fig_ncc.update_layout(
                        **{**_CHART_LAYOUT, "margin": dict(l=10, r=20, t=40, b=16)},
                        height=max(280, len(ncc_perf) * 52),
                        title="🏭 Hiệu suất ATTP theo nhà cung cấp",
                        xaxis=dict(title="Tỷ lệ đạt TB (%)", range=[0, 105],
                                   ticksuffix="%", showgrid=True, gridcolor="#E2E8F0"),
                        yaxis=dict(showgrid=False, automargin=True, tickfont=dict(size=11)),
                        showlegend=False,
                    )
                    st.plotly_chart(fig_ncc, use_container_width=True)

            with nc2:
                # Xu hướng xếp loại NCC theo thời gian
                ncc_trend = (df_ncc.groupby(["Ngày", "Xếp loại"])
                             .size().reset_index(name="Số lần")
                             .sort_values("Ngày"))
                try:
                    ncc_trend["Ngày_fmt"] = pd.to_datetime(ncc_trend["Ngày"]).dt.strftime("%d/%m")
                except Exception:
                    ncc_trend["Ngày_fmt"] = ncc_trend["Ngày"]
                fig_ncc_trend = px.bar(
                    ncc_trend, x="Ngày_fmt", y="Số lần", color="Xếp loại",
                    color_discrete_map={"Loại A": "#16A34A", "Loại B": "#F59E0B", "Loại C": "#DC2626"},
                    barmode="stack", text="Số lần",
                    category_orders={"Xếp loại": ["Loại A", "Loại B", "Loại C"]},
                )
                fig_ncc_trend.update_traces(
                    textposition="inside", insidetextanchor="middle",
                    textfont_size=11, textfont_color="white",
                )
                fig_ncc_trend.update_layout(
                    **{**_CHART_LAYOUT, "margin": dict(l=16, r=110, t=40, b=16)},
                    height=max(280, len(ncc_perf) * 52) if "ncc_perf" in dir() else 280,
                    title="📅 Xếp loại NCC theo ngày kiểm tra",
                    xaxis=dict(title="Ngày", showgrid=False, tickangle=-20),
                    yaxis=dict(title="Số lần", showgrid=True, gridcolor="#E2E8F0"),
                    legend=dict(orientation="v", x=1.02, y=0.5, xanchor="left",
                                yanchor="middle", title_text="Xếp loại",
                                bgcolor="rgba(255,255,255,0.8)",
                                bordercolor="#E2E8F0", borderwidth=1),
                )
                st.plotly_chart(fig_ncc_trend, use_container_width=True)

            # Top 10 lỗi NCC nhiều nhất
            st.markdown('<div class="sec-hdr">🟣 Top điểm không đạt nhiều nhất (Nhà Cung Cấp)</div>',
                        unsafe_allow_html=True)
            try:
                sb = _get_sb()
                if sb:
                    _ncc_ids = [s["id"] for s in sessions
                                if s.get("check_type") == "nha_cung_cap" and s.get("id")]
                    if _ncc_ids:
                        _ncc_items: list = []
                        for _i in range(0, len(_ncc_ids), 100):
                            _batch = _ncc_ids[_i:_i+100]
                            _r = sb.table("checklist_results")\
                                .select("item_code,result")\
                                .in_("session_id", _batch)\
                                .in_("result", ["Không Đạt", "❌ Không Đạt"])\
                                .execute()
                            _ncc_items += (_r.data or [])
                        if _ncc_items:
                            # Map S01-S12 từ SUPPLIER_ITEMS constant
                            _NCC_KW = {it["code"]: f"{it['icon']} {it['desc']}"
                                       for it in SUPPLIER_ITEMS}
                            _NCC_KW.update({  # backward compat nếu desc dài
                                "S01": "📄 Giấy phép CSSX/KDDV",
                                "S02": "🏅 Chứng nhận ATTP cơ sở",
                                "S03": "🚚 Xe/thùng vận chuyển",
                                "S04": "🌡️ Nhiệt độ vận chuyển",
                                "S05": "🧾 Hóa đơn nguồn gốc hôm nay",
                                "S06": "🏷️ Nhãn mác thực phẩm",
                                "S07": "📋 Thực đơn khớp đặt hàng",
                                "S08": "⚖️ Khẩu phần đủ định mức",
                                "S09": "👷 BHLĐ nhân viên giao hàng",
                                "S10": "📦 Dụng cụ đựng thực phẩm",
                                "S11": "🧫 Mẫu lưu thực phẩm 24h",
                                "S12": "⏰ Thời gian giao hàng đúng lịch",
                            })
                            df_ncc_it = pd.DataFrame(_ncc_items)
                            df_ncc_it = df_ncc_it[df_ncc_it["item_code"].str.startswith("S")]
                            df_ncc_it["Tên điểm"] = df_ncc_it["item_code"].map(
                                lambda c: _NCC_KW.get(c, c))
                            top_ncc = (df_ncc_it.groupby("Tên điểm").size()
                                       .reset_index(name="Số lần")
                                       .sort_values("Số lần").tail(12))
                            n_n = len(top_ncc)
                            fig_ncc_fail = go.Figure(go.Bar(
                                x=top_ncc["Số lần"], y=top_ncc["Tên điểm"], orientation="h",
                                marker_color=[f"rgba(124,58,237,{0.25+0.75*i/max(n_n-1,1)})"
                                              for i in range(n_n)],
                                text=top_ncc["Số lần"], textposition="outside",
                                textfont=dict(size=12),
                                hovertemplate="<b>%{y}</b><br>Không đạt: %{x} lần<extra></extra>",
                            ))
                            fig_ncc_fail.update_layout(
                                plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                                font=dict(family="Inter, sans-serif", size=12, color="#334155"),
                                margin=dict(l=10, r=60, t=4, b=16),
                                height=max(280, n_n * 50),
                                xaxis=dict(title="Số lần không đạt", showgrid=True,
                                           gridcolor="#E2E8F0", dtick=1),
                                yaxis=dict(showgrid=False, tickfont=dict(size=11), automargin=True),
                            )
                            st.plotly_chart(fig_ncc_fail, use_container_width=True)
                        else:
                            st.info("Chưa có điểm không đạt nào từ nhà cung cấp.")
                    else:
                        st.info("Chưa có dữ liệu kiểm tra nhà cung cấp.")
            except Exception as _ne:
                st.warning(f"Không thể tải top fail NCC: {_ne}")

            # Bảng chi tiết NCC
            st.markdown('<div class="sec-hdr">📋 Bảng chi tiết — Nhà Cung Cấp</div>',
                        unsafe_allow_html=True)
            _dn = df_ncc.drop(columns=["Cấp cảnh báo", "Đánh giá"], errors="ignore").copy()
            _dn["Ngày"] = pd.to_datetime(_dn["Ngày"], errors="coerce").dt.strftime("%d/%m/%Y").fillna(_dn["Ngày"])
            # Sắp xếp cột hợp lý hơn
            _dn_cols = ["Ngày", "Trường", "Nhà Cung Cấp", "Người kiểm tra",
                        "Xếp loại", "Tỷ lệ đạt (%)", "Điểm đạt", "Điểm không đạt", "Tổng điểm"]
            _dn = _dn[[c for c in _dn_cols if c in _dn.columns]]
            st.dataframe(_dn, use_container_width=True, hide_index=True)


    # ── Cross-validation & Anomaly Detection — CHỈ BGH ────────────────────
    # BGS và Y Tế không thấy phần này
    if role == "Ban Giám Hiệu" or st.session_state.get("is_super"):
        _anomalies = []

        # 1-3. Kiểm tra chất lượng dữ liệu (cần có df_meal)
        if show_meal and not df_meal.empty:
            _streak_high = (df_meal["Tỷ lệ đạt (%)"] >= 95).sum()
            if _streak_high >= 10 and len(df_meal) >= 10:
                _anomalies.append(
                    f"📊 <b>Điểm quá đồng đều</b>: {_streak_high}/{len(df_meal)} lần đạt ≥ 95% "
                    "— kết quả thực tế hiếm khi hoàn hảo liên tục. Đề xuất kiểm tra đột xuất."
                )

        # 2. BGS và Y Tế cùng ngày chênh lệch > 20%
        if show_meal and not df_meal.empty and "Loại" in df_meal.columns:
            _bgs_df = df_meal[df_meal["Loại"] == "Ban Giám Sát"]
            _yte_df = df_meal[df_meal["Loại"] == "Y Tế (3 bước)"]
            if not _bgs_df.empty and not _yte_df.empty:
                _bgs_avg = _bgs_df["Tỷ lệ đạt (%)"].mean()
                _yte_avg = _yte_df["Tỷ lệ đạt (%)"].mean()
                if abs(_bgs_avg - _yte_avg) > 20:
                    _anomalies.append(
                        f"⚠️ <b>Chênh lệch BGS vs Y Tế</b>: BGS trung bình {_bgs_avg:.0f}% — "
                        f"Y Tế {_yte_avg:.0f}% (chênh {abs(_bgs_avg-_yte_avg):.0f}%). "
                        "Cần đối chiếu để đảm bảo tính trung thực."
                    )

        # 3. Feedback tăng đột biến trong khi điểm vẫn cao
        if db_ok() and show_meal and not df_meal.empty:
            try:
                _sch = school_input if school_input else ""
                _fb_count = len(db_get_feedback(school=_sch, status="pending"))
                _avg_score = df_meal["Tỷ lệ đạt (%)"].mean()
                if _fb_count >= 3 and _avg_score >= 90:
                    _anomalies.append(
                        f"📣 <b>Mâu thuẫn phản hồi</b>: {_fb_count} phản hồi phụ huynh chưa xử lý "
                        f"trong khi điểm kiểm tra trung bình {_avg_score:.0f}%. "
                        "Điều tra nguyên nhân bất cân xứng."
                    )
            except Exception:
                pass

        # 4. BGS + Y Tế tần suất — dùng sessions trực tiếp (không cần df_meal)
        # → Fix: luôn chạy kể cả khi chưa có dữ liệu trong view hiện tại
        try:
            _now_dt2 = pd.Timestamp(now_vn().date())
            # Lấy dates từ raw sessions (đã fetch ở trên, không cần thêm DB call)
            _s_bgs_dates = sorted([s["check_date"] for s in sessions
                                   if s.get("check_type") == "ban_giam_sat" and s.get("check_date")],
                                  reverse=True)
            _s_yte_dates = sorted([s["check_date"] for s in sessions
                                   if s.get("check_type") == "kiem_thuc_3_buoc" and s.get("check_date")],
                                  reverse=True)

            # BGS tần suất
            if _s_bgs_dates:
                _s_bgs_last = pd.Timestamp(_s_bgs_dates[0])
                _s_bgs_gap  = (_now_dt2 - _s_bgs_last).days
                _s_bgs_7d   = sum(1 for d in _s_bgs_dates
                                  if pd.Timestamp(d) >= _now_dt2 - pd.Timedelta(days=7))
                if _s_bgs_gap > 7:
                    _anomalies.append(f"🚨 <b>[BGH → BGS] Không kiểm tra</b>: "
                                      f"Ban Giám Sát chưa kiểm tra <b>{_s_bgs_gap} ngày</b> — "
                                      f"vi phạm NĐ 15/2018 (≥ 2 lần/tuần). Liên hệ ngay.")
                elif _s_bgs_7d < 2:
                    _anomalies.append(f"⏰ <b>[BGH → BGS] Tần suất thấp</b>: "
                                      f"7 ngày qua chỉ {_s_bgs_7d} lần (yêu cầu ≥ 2 lần/tuần).")
            else:
                _anomalies.append("🚨 <b>[BGH → BGS] Chưa có kiểm tra nào</b>: "
                                  "Ban Giám Sát chưa thực hiện. Liên hệ ngay để lên lịch.")

            # Y Tế tần suất
            if _s_yte_dates:
                _s_yte_last = pd.Timestamp(_s_yte_dates[0])
                _s_yte_gap  = (_now_dt2 - _s_yte_last).days
                if _s_yte_gap > 3:
                    _anomalies.append(f"🚨 <b>[BGH → Y Tế] Không kiểm thực</b>: "
                                      f"Y Tế chưa kiểm thực <b>{_s_yte_gap} ngày</b> — "
                                      f"vi phạm TTLT 13/2016 Điều 9 (hàng ngày). Cần xử lý ngay.")
            else:
                _anomalies.append("🚨 <b>[BGH → Y Tế] Chưa có kiểm thực</b>: "
                                  "Y Tế chưa thực hiện kiểm thực 3 bước. Vi phạm TTLT 13/2016.")
        except Exception:
            pass

        if _anomalies:
            st.markdown('<div class="sec-hdr" style="color:#991B1B;font-weight:700">🔍 PHÁT HIỆN BẤT THƯỜNG — CHI TIẾT</div>',
                        unsafe_allow_html=True)
            for _a in _anomalies:
                _is_crit_a = _a.startswith("🚨")
                _a_bg  = "#FEF2F2" if _is_crit_a else "#FFFBEB"
                _a_bd  = "#FCA5A5" if _is_crit_a else "#FCD34D"
                _a_clr = "#991B1B" if _is_crit_a else "#78350F"
                st.markdown(
                    f'<div style="background:{_a_bg};border:1.5px solid {_a_bd};'
                    f'border-radius:8px;padding:10px 16px;margin:5px 0;'
                    f'font-size:0.85rem;color:{_a_clr}">'
                    f'{_a}</div>',
                    unsafe_allow_html=True,
                )

            # ── Claude AI phân tích tổng hợp bất thường ─────────────────
            _api_key_hist = st.session_state.get("api_key_stored", "")
            import os as _os_hist
            _api_key_hist = (
                (st.secrets.get("ANTHROPIC_API_KEY","") if hasattr(st,"secrets") else "")
                or _os_hist.environ.get("ANTHROPIC_API_KEY","")
            )
            if _api_key_hist:
                st.markdown(
                    '<div style="background:#F5F3FF;border:1px solid #DDD6FE;'
                    'border-radius:8px;padding:10px 14px;margin-top:8px;'
                    'font-size:0.8rem;color:#5B21B6">'
                    '🤖 <b>AI Cross-Validation</b>: Claude phân tích tổng hợp tất cả bất thường, '
                    'đánh giá mức độ rủi ro thực sự và đề xuất hành động ưu tiên.</div>',
                    unsafe_allow_html=True,
                )
                if st.button("🤖 Claude phân tích bất thường tổng hợp",
                             key="ai_anomaly_analysis", use_container_width=False):
                    _anom_text = "\n".join(f"- {a}" for a in _anomalies)
                    _stats_txt = ""
                    if show_meal and not df_meal.empty:
                        _stats_txt = (
                            f"Tổng {len(df_meal)} lần kiểm tra · "
                            f"Trung bình {df_meal['Tỷ lệ đạt (%)'].mean():.0f}% · "
                            f"CRITICAL: {(df_meal['Cấp cảnh báo']=='CRITICAL').sum()} lần"
                        )
                    _ai_prompt = (
                        f"Bạn là chuyên gia ATTP trường học Việt Nam. Dưới đây là các bất thường "
                        f"phát hiện tự động trong hệ thống giám sát bữa ăn học đường:\n\n"
                        f"{_anom_text}\n\n"
                        f"Thống kê tổng hợp: {_stats_txt}\n\n"
                        f"Hãy phân tích:\n"
                        f"1. Đánh giá mức độ rủi ro thực sự (thấp/trung/cao/nguy hiểm)\n"
                        f"2. Mối liên hệ giữa các bất thường (có phải cùng 1 nguyên nhân không)\n"
                        f"3. 3 hành động ưu tiên nhất cần thực hiện ngay\n"
                        f"4. Dấu hiệu nào cần giám sát thêm\n"
                        f"Trả lời ngắn gọn, súc tích bằng tiếng Việt (~200 từ)."
                    )
                    try:
                        with st.spinner("🤖 Claude đang phân tích..."):
                            _ai_client = anthropic.Anthropic(api_key=_api_key_hist)
                            _ai_resp = _ai_client.messages.create(
                                model=MODEL, max_tokens=600,
                                messages=[{"role": "user", "content": _ai_prompt}]
                            )
                        _ai_analysis = _ai_resp.content[0].text if _ai_resp.content else ""
                        st.markdown(
                            f'<div style="background:#F5F3FF;border:1px solid #7C3AED;'
                            f'border-radius:10px;padding:14px 16px;margin-top:8px">'
                            f'<div style="font-size:0.85rem;font-weight:700;color:#6D28D9;'
                            f'margin-bottom:8px">🤖 Claude AI — Phân tích Cross-Validation</div>'
                            f'<div style="font-size:0.85rem;color:#1E293B;line-height:1.7">'
                            f'{_ai_analysis}</div></div>',
                            unsafe_allow_html=True,
                        )
                    except Exception as _ae:
                        st.error(f"Lỗi AI: {_ae}")
    
    st.markdown('<div class="sec-hdr">⬇️ Xuất báo cáo Excel</div>', unsafe_allow_html=True)

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font as XFont, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        wb.remove(wb.active)  # xoá sheet mặc định, tự tạo 2 sheet riêng

        # Styles dùng chung
        THIN    = Side(style="thin", color="CBD5E1")
        BORDER  = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
        HDR_FILL_MEAL = PatternFill("solid", fgColor="1B3B6F")   # xanh navy — bữa ăn
        HDR_FILL_NCC  = PatternFill("solid", fgColor="4C1D95")   # tím — NCC
        HDR_FONT  = XFont(name="Times New Roman", size=13, bold=True, color="FFFFFF")
        HDR_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)

        ALERT_BG = {
            "Nguy hiểm":     "FEE2E2",
            "Không đạt":     "FEF9C3",
            "Cần cải thiện": "FEFCE8",
        }
        RATING_BG = {"Loại A": "DCFCE7", "Loại B": "FEF9C3", "Loại C": "FEE2E2"}

        def _write_sheet(ws, df_ws, title_txt, subtitle_txt, hdr_fill, rating_col=None):
            """Ghi dữ liệu vào worksheet với header chuyên nghiệp."""
            if df_ws.empty:
                ws.cell(row=1, column=1, value="Chưa có dữ liệu.")
                return
            ncols = len(df_ws.columns)
            col_last = get_column_letter(ncols)

            # Quốc hiệu
            ws.merge_cells(f"A1:{col_last}1")
            _c = ws["A1"]
            _c.value = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM — Độc lập – Tự do – Hạnh phúc"
            _c.font = XFont(name="Times New Roman", size=11, bold=True)
            _c.alignment = Alignment(horizontal="center")
            ws.row_dimensions[1].height = 18

            # Tiêu đề
            ws.merge_cells(f"A2:{col_last}2")
            _c = ws["A2"]
            _c.value = title_txt
            _c.font = XFont(name="Times New Roman", size=14, bold=True, color="1B3B6F")
            _c.alignment = Alignment(horizontal="center")
            ws.row_dimensions[2].height = 26

            # Subtitle
            ws.merge_cells(f"A3:{col_last}3")
            _c = ws["A3"]
            _c.value = subtitle_txt
            _c.font = XFont(name="Times New Roman", size=11, italic=True, color="475569")
            _c.alignment = Alignment(horizontal="center")
            ws.row_dimensions[3].height = 15

            # Header cột
            for ci, cname in enumerate(df_ws.columns, 1):
                _c = ws.cell(row=4, column=ci, value=cname)
                _c.font = HDR_FONT; _c.fill = hdr_fill
                _c.alignment = HDR_ALIGN; _c.border = BORDER
            ws.row_dimensions[4].height = 22

            # Dữ liệu
            for ri, row_data in enumerate(df_ws.itertuples(index=False), 5):
                row_vals = list(row_data)
                # Chọn màu nền theo loại sheet
                if rating_col is not None and rating_col < len(row_vals):
                    bg = RATING_BG.get(str(row_vals[rating_col]), "EFF6FF" if ri % 2 == 0 else "FFFFFF")
                else:
                    last_val = str(row_vals[-1]) if row_vals else ""
                    bg = ALERT_BG.get(last_val, "EFF6FF" if ri % 2 == 0 else "FFFFFF")
                fill = PatternFill("solid", fgColor=bg)
                for ci, val in enumerate(row_vals, 1):
                    _c = ws.cell(row=ri, column=ci, value=val)
                    _c.font = XFont(name="Times New Roman", size=13)
                    _c.fill = fill; _c.border = BORDER
                    _c.alignment = Alignment(vertical="center",
                                             horizontal="left" if ci <= 3 else "center")
                ws.row_dimensions[ri].height = 18

            # Auto-width
            for ci, cname in enumerate(df_ws.columns, 1):
                mx = max(len(str(cname)),
                         max((len(str(ws.cell(row=r, column=ci).value or ""))
                              for r in range(4, len(df_ws) + 5)), default=0))
                ws.column_dimensions[get_column_letter(ci)].width = min(mx + 3, 42)
            ws.freeze_panes = "A5"

        # ── Sheet 1: Bữa ăn ──────────────────────────────────────────────────
        ws_meal = wb.create_sheet("🍱 Bữa Ăn")
        if show_meal and not df_meal.empty:
            _dm_xl = df_meal.drop(columns=["Cấp cảnh báo"], errors="ignore").copy()
            _dm_xl["Ngày"] = (pd.to_datetime(_dm_xl["Ngày"], errors="coerce")
                              .dt.strftime("%d/%m/%Y").fillna(_dm_xl["Ngày"]))
            meal_avg = _dm_xl["Tỷ lệ đạt (%)"].mean() if "Tỷ lệ đạt (%)" in _dm_xl.columns else 0
            _write_sheet(
                ws_meal, _dm_xl,
                title_txt="KẾT QUẢ KIỂM TRA AN TOÀN THỰC PHẨM BỮA ĂN HỌC ĐƯỜNG",
                subtitle_txt=(f"Xuất ngày: {now_vn().strftime('%d/%m/%Y %H:%M')} | "
                              f"Tổng: {len(_dm_xl)} lần | Trung bình đạt: {meal_avg:.0f}%"),
                hdr_fill=HDR_FILL_MEAL,
            )
        else:
            ws_meal.cell(row=1, column=1, value="Không có dữ liệu bữa ăn trong bộ lọc hiện tại.")

        # ── Sheet 2: Nhà Cung Cấp ────────────────────────────────────────────
        ws_ncc = wb.create_sheet("🏭 Nhà Cung Cấp")
        if show_ncc and not df_ncc.empty:
            _dn_xl = df_ncc.drop(columns=["Cấp cảnh báo", "Đánh giá"], errors="ignore").copy()
            _dn_xl["Ngày"] = (pd.to_datetime(_dn_xl["Ngày"], errors="coerce")
                              .dt.strftime("%d/%m/%Y").fillna(_dn_xl["Ngày"]))
            # Sắp xếp cột hợp lý
            _ncc_cols = ["Ngày", "Trường", "Nhà Cung Cấp", "Người kiểm tra",
                         "Xếp loại", "Tỷ lệ đạt (%)", "Điểm đạt", "Điểm không đạt", "Tổng điểm"]
            _dn_xl = _dn_xl[[col for col in _ncc_cols if col in _dn_xl.columns]]
            # Tìm vị trí cột "Xếp loại" để tô màu A/B/C
            _rating_idx = list(_dn_xl.columns).index("Xếp loại") if "Xếp loại" in _dn_xl.columns else None
            ncc_avg = _dn_xl["Tỷ lệ đạt (%)"].mean() if "Tỷ lệ đạt (%)" in _dn_xl.columns else 0
            ncc_a = (_dn_xl["Xếp loại"] == "Loại A").sum() if "Xếp loại" in _dn_xl.columns else 0
            ncc_b = (_dn_xl["Xếp loại"] == "Loại B").sum() if "Xếp loại" in _dn_xl.columns else 0
            ncc_c = (_dn_xl["Xếp loại"] == "Loại C").sum() if "Xếp loại" in _dn_xl.columns else 0
            _write_sheet(
                ws_ncc, _dn_xl,
                title_txt="KẾT QUẢ ĐÁNH GIÁ NHÀ CUNG CẤP SUẤT ĂN HỌC ĐƯỜNG",
                subtitle_txt=(f"Xuất ngày: {now_vn().strftime('%d/%m/%Y %H:%M')} | "
                              f"Tổng: {len(_dn_xl)} lần | TB đạt: {ncc_avg:.0f}% | "
                              f"Loại A: {ncc_a} · Loại B: {ncc_b} · Loại C: {ncc_c}"),
                hdr_fill=HDR_FILL_NCC,
                rating_col=_rating_idx,
            )
        else:
            ws_ncc.cell(row=1, column=1, value="Không có dữ liệu nhà cung cấp trong bộ lọc hiện tại.")

        buf = BytesIO(); wb.save(buf); buf.seek(0)
        st.download_button(
            "⬇️ Tải báo cáo (.xlsx)",
            data=buf.getvalue(),
            file_name=f"BaoCao_ATTP_{now_vn().strftime('%d-%m-%Y')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True, type="primary",
        )
    except Exception as e:
        st.error(f"Lỗi xuất Excel: {e}")

    # ── Báo cáo tháng tổng hợp — BGH gửi Sở GD&ĐT ───────────────────────────
    if role in ("Ban Giám Hiệu",) or st.session_state.get("is_super"):
        st.markdown('<div class="sec-hdr">📋 Báo cáo tháng tổng hợp — Gửi Sở GD&ĐT</div>',
                    unsafe_allow_html=True)
        _mc1, _mc2, _mc3 = st.columns([1.5, 1.5, 2])
        _sel_month = _mc1.selectbox(
            "Tháng", list(range(1, 13)),
            index=now_vn().month - 1,
            format_func=lambda x: f"Tháng {x:02d}",
            key="rpt_month", label_visibility="collapsed"
        )
        _sel_year  = _mc2.number_input("Năm", value=now_vn().year, min_value=2024,
                                        max_value=2030, step=1, key="rpt_year",
                                        label_visibility="collapsed")
        if _mc3.button("📋 Tạo báo cáo tháng", use_container_width=True,
                        key="gen_monthly_rpt"):
            try:
                from docx import Document as _Doc
                from docx.shared import Pt as _Pt, Cm as _Cm, RGBColor as _RGB
                from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
                from io import BytesIO as _BIO

                # Lọc sessions trong tháng
                _month_sessions = [s for s in sessions if (
                    s.get("check_date","").startswith(f"{_sel_year}-{_sel_month:02d}")
                )]
                _meal_ses = [s for s in _month_sessions if s.get("check_type") in {"ban_giam_sat","kiem_thuc_3_buoc"}]
                _ncc_ses  = [s for s in _month_sessions if s.get("check_type") == "nha_cung_cap"]
                _bgs_ses  = [s for s in _month_sessions if s.get("check_type") == "ban_giam_sat"]
                _yte_ses  = [s for s in _month_sessions if s.get("check_type") == "kiem_thuc_3_buoc"]
                _meal_avg = round(sum(s["pass_count"]/max(s["total_items"],1)*100 for s in _meal_ses)/max(len(_meal_ses),1),1) if _meal_ses else 0
                _crit_ct  = sum(1 for s in _meal_ses if s.get("alert_level")=="CRITICAL")
                # Complaints trong tháng
                try:
                    _month_fb = _get_sb().table("parent_feedback").select("status,category,created_at")\
                        .gte("created_at", f"{_sel_year}-{_sel_month:02d}-01")\
                        .lt("created_at", f"{_sel_year}-{_sel_month+1 if _sel_month<12 else 1:02d}-01")\
                        .execute().data or []
                except Exception:
                    _month_fb = []
                _fb_total   = len(_month_fb)
                _fb_resolved= sum(1 for f in _month_fb if f.get("status")=="resolved")

                doc = _Doc()
                for _sec in doc.sections:
                    _sec.top_margin = _sec.bottom_margin = _Cm(2.5)
                    _sec.left_margin = _Cm(3.0); _sec.right_margin = _Cm(2.0)

                def _rp(p, txt, bold=False, sz=13, align=None, color=None):
                    p.alignment = getattr(_ALIGN, align or "LEFT", _ALIGN.LEFT)
                    r = p.add_run(txt); r.bold = bold
                    r.font.name = "Times New Roman"; r.font.size = _Pt(sz)
                    if color: r.font.color.rgb = color
                    return r

                _rp(doc.add_paragraph(), "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, sz=13, align="CENTER")
                _rp(doc.add_paragraph(), "Độc lập – Tự do – Hạnh phúc", bold=True, sz=13, align="CENTER")
                doc.add_paragraph("")
                _rp(doc.add_paragraph(),
                    f"BÁO CÁO THÁNG {_sel_month:02d}/{_sel_year}\nAN TOÀN THỰC PHẨM BỮA ĂN HỌC ĐƯỜNG",
                    bold=True, sz=15, align="CENTER")
                _rp(doc.add_paragraph(), f"Trường: {school_input or 'Tất cả trường'}", sz=13)
                doc.add_paragraph("")

                # Bảng tổng hợp
                _rp(doc.add_paragraph(), "I. TỔNG HỢP KIỂM TRA BỮA ĂN", bold=True, sz=13)
                tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
                # Header row
                _hdr = tbl.rows[0].cells
                for _ci, _htxt in enumerate(["Chỉ số", "Số liệu"]):
                    _hdr[_ci].paragraphs[0].alignment = getattr(_ALIGN, "CENTER", 0)
                    _hrr = _hdr[_ci].paragraphs[0].add_run(_htxt)
                    _hrr.bold = True; _hrr.font.name = "Times New Roman"; _hrr.font.size = _Pt(12)
                    from docx.oxml.ns import qn as _qn
                    from docx.oxml import OxmlElement as _OE
                    _tc = _hdr[_ci]._tc
                    _tcp = _tc.get_or_add_tcPr()
                    _shd = _OE("w:shd")
                    _shd.set(_qn("w:fill"), "1B3B6F"); _shd.set(_qn("w:color"), "auto")
                    _shd.set(_qn("w:val"), "clear"); _tcp.append(_shd)
                    _hrr.font.color.rgb = _RGB(0xFF, 0xFF, 0xFF)
                for _r, _v in [
                    ("Số buổi kiểm tra (Ban Giám Sát)", str(len(_bgs_ses))),
                    ("Số lần kiểm thực 3 bước (Y Tế)", str(len(_yte_ses))),
                    ("Tỷ lệ đạt trung bình", f"{_meal_avg}%"),
                    ("Số lần CRITICAL", str(_crit_ct)),
                    ("Số phản hồi Phụ Huynh", str(_fb_total)),
                    ("Đã xử lý", str(_fb_resolved)),
                    ("Số lần đánh giá NCC", str(len(_ncc_ses))),
                ]:
                    _row = tbl.add_row().cells
                    for ci, val in enumerate([_r, _v]):
                        _pp = _row[ci].paragraphs[0]
                        _rr = _pp.add_run(val)
                        _rr.font.name = "Times New Roman"; _rr.font.size = _Pt(12)
                        if ci == 0: _rr.bold = True

                # Kết luận
                doc.add_paragraph("")
                _rp(doc.add_paragraph(), "II. KẾT LUẬN VÀ KIẾN NGHỊ", bold=True, sz=13)
                _concl_txt = (
                    f"Trong tháng {_sel_month:02d}/{_sel_year}, nhà trường đã thực hiện "
                    f"{len(_bgs_ses)} lần kiểm tra (Ban Giám Sát) và {len(_yte_ses)} lần kiểm thực "
                    f"3 bước (Y Tế Học Đường). Tỷ lệ đạt trung bình: {_meal_avg}%."
                )
                if _crit_ct > 0:
                    _concl_txt += f" Ghi nhận {_crit_ct} lần cảnh báo mức CRITICAL — đã xử lý theo quy trình."
                _concl_txt += f"\n\nPhản hồi Phụ Huynh: {_fb_total} phản hồi, {_fb_resolved} đã xử lý."
                _rp(doc.add_paragraph(), _concl_txt, sz=13)

                doc.add_paragraph("")
                _sign_dt = f"......., ngày {now_vn().strftime('%d')} tháng {now_vn().strftime('%m')} năm {now_vn().strftime('%Y')}"
                _rp(doc.add_paragraph(), _sign_dt, sz=13, align="RIGHT")
                _sig_tbl = doc.add_table(rows=1, cols=2)
                _s1, _s2 = _sig_tbl.rows[0].cells
                for _cell, _txt in [(_s1, "NGƯỜI LẬP BÁO CÁO\n(Ký, ghi rõ họ tên)"),
                                     (_s2, "HIỆU TRƯỞNG\n(Ký tên, đóng dấu)")]:
                    _cell.paragraphs[0].alignment = _ALIGN.CENTER
                    _rr = _cell.paragraphs[0].add_run(_txt)
                    _rr.bold = True; _rr.font.name = "Times New Roman"; _rr.font.size = _Pt(12)

                _mbuf = _BIO(); doc.save(_mbuf); _mbuf.seek(0)
                st.download_button(
                    f"⬇️ Tải báo cáo tháng {_sel_month:02d}/{_sel_year} (.docx)",
                    data=_mbuf.getvalue(),
                    file_name=f"BaoCao_Thang{_sel_month:02d}_{_sel_year}_{(school_input or 'TatCaTruong').replace(' ','_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary",
                )
                st.caption(f"✅ Báo cáo: {len(_meal_ses)} lần kiểm tra · {_meal_avg}% đạt · {_fb_total} phản hồi PH")
            except Exception as _me:
                st.error(f"Lỗi tạo báo cáo: {_me}")


    # Feedback section đã được chuyển lên trước early return ở đầu hàm
    # Feedback section đã được chuyển lên trước early return ở đầu hàm

def tab_supplier(api_key: str = "", role: str = ""):
    """G4: Checklist kiểm tra nhà cung cấp suất ăn 12 điểm."""
    st.markdown(
        '<div class="sf-card">'
        '<div class="sf-card-title">🏭 Đánh Giá Nhà Cung Cấp Suất Ăn</div>'
        '<div class="sf-card-body">'
        'Checklist 12 điểm theo Luật ATTP 55/2010 · NĐ 15/2018 · 6 mục bắt buộc (*) · '
        'Khi chấm <b>Không Đạt</b>: bắt buộc có ghi chú <i>hoặc</i> ảnh minh chứng · '
        'AI Vision phân tích hình ảnh từng mục'
        '</div></div>',
        unsafe_allow_html=True,
    )

    # ── Chọn chế độ kiểm tra ──────────────────────────────────────────────────
    check_mode = st.radio(
        "Chế độ kiểm tra",
        ["🚚 Kiểm tra khi nhận hàng (S03–S12 · 10 mục)", "📋 Kiểm tra toàn diện (S01–S12 · 12 mục)"],
        horizontal=True, key="sup_mode",
    )
    _delivery_mode = check_mode.startswith("🚚")

    _is_yte = (role == "Y Tế Học Đường")

    if _delivery_mode:
        _active_items    = [it for it in SUPPLIER_ITEMS if it["code"] not in ("S01", "S02")]
        _active_critical = {c for c in SUPPLIER_CRITICAL if c not in ("S01", "S02")}
        if _is_yte:
            _bg, _bd, _tc = "#EFF6FF", "#BFDBFE", "#1E40AF"
            _role_icon = "🏥 Y Tế Học Đường"
            _info = (
                '📅 <b>Tần suất:</b> Mỗi ngày có bữa ăn, khi NCC giao hàng vào bếp<br>'
                '📋 <b>Nội dung:</b> Kiểm tra 10 mục S03–S12 ngay tại điểm giao<br>'
                '⚡ <b>Căn cứ:</b> TTLT 13/2016 Điều 9 khoản a — kiểm tra nguyên liệu đầu vào'
            )
        else:
            _bg, _bd, _tc = "#F0FDF4", "#BBF7D0", "#166534"
            _role_icon = "👥 Ban Giám Sát (Đại Diện PHHS)"
            _info = (
                '📅 <b>Tần suất:</b> Khi có mặt lúc NCC giao hàng trong lịch kiểm tra thường kỳ<br>'
                '📋 <b>Nội dung:</b> Kiểm tra 10 mục S03–S12, ảnh minh chứng khi có mục Không Đạt<br>'
                '💡 <b>Lưu ý:</b> S01–S02 bỏ qua — được kiểm tra riêng trong đợt toàn diện hàng tháng'
            )
    else:
        _active_items    = SUPPLIER_ITEMS
        _active_critical = SUPPLIER_CRITICAL
        if _is_yte:
            _bg, _bd, _tc = "#F5F3FF", "#DDD6FE", "#5B21B6"
            _role_icon = "🏥 Y Tế Học Đường"
            _info = (
                '📅 <b>Tần suất:</b> Khi phát hiện vi phạm liên tục hoặc theo yêu cầu Ban Giám Hiệu<br>'
                '📋 <b>Nội dung:</b> Đủ 12 mục — bổ sung kiểm tra giấy phép (S01) và chứng nhận ATTP (S02)<br>'
                '📤 <b>Báo cáo:</b> Gửi Ban Giám Hiệu trong 24 giờ sau kiểm tra'
            )
        else:
            _bg, _bd, _tc = "#FFF7ED", "#FED7AA", "#9A3412"
            _role_icon = "👥 Ban Giám Sát (Đại Diện PHHS)"
            _info = (
                '📅 <b>Tần suất:</b> <b>1 lần/tháng</b> — cuối tháng hoặc trước khi gia hạn hợp đồng<br>'
                '📋 <b>Nội dung:</b> Đủ 12 mục — đặc biệt xác minh S01 (giấy phép) và S02 (chứng nhận ATTP)<br>'
                '📤 <b>Báo cáo:</b> Gửi Ban Giám Hiệu · Lưu hồ sơ phục vụ xét duyệt hợp đồng'
            )

    st.markdown(
        f'<div style="background:{_bg};border:1px solid {_bd};border-radius:9px;'
        f'padding:11px 16px;margin-bottom:12px">'
        f'<div style="font-size:0.83rem;font-weight:700;color:{_tc};margin-bottom:5px">'
        f'{_role_icon}</div>'
        f'<div style="font-size:0.78rem;color:{_tc};line-height:1.75">{_info}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Ngưỡng điểm tỷ lệ theo số mục đang kiểm tra
    _n_active = len(_active_items)
    _score_pass = round(_n_active * SUPPLIER_SCORE_PASS / len(SUPPLIER_ITEMS))  # tỷ lệ với 12 mục
    _score_warn = round(_n_active * SUPPLIER_SCORE_WARN / len(SUPPLIER_ITEMS))

    # ── Thông tin chung ────────────────────────────────────────────────────────
    c1, c2 = st.columns(2)
    sup_school    = c1.text_input("🏫 Tên trường", placeholder="VD: TH Nguyễn Du",
                                   key="sup_school", max_chars=100)
    sup_name      = c2.text_input("🏭 Tên nhà cung cấp", placeholder="VD: Công ty TNHH Bếp Xanh",
                                   key="sup_name", max_chars=100)
    c3, c4 = st.columns(2)
    sup_inspector = c3.text_input("👤 Người kiểm tra", placeholder="Họ và tên",
                                   key="sup_inspector", max_chars=80)
    sup_date      = c4.date_input("📅 Ngày kiểm tra", value=now_vn().date(),
                                   key="sup_date", format="DD/MM/YYYY")
    sup_contract  = st.text_input("📃 Số hợp đồng cung cấp (nếu có)", placeholder="VD: HĐ-2025-001",
                                   key="sup_contract", max_chars=50)

    # Session state
    for _sk in ("sup_r", "sup_notes", "sup_imgs", "sup_vision"):
        if _sk not in st.session_state:
            st.session_state[_sk] = {}

    # ── Checklist (số mục phụ thuộc chế độ) ──────────────────────────────────
    _hdr_mode = "🚚 Kiểm tra giao hàng" if _delivery_mode else "📋 Kiểm tra toàn diện"
    st.markdown(
        f'<div class="sec-hdr">{_hdr_mode} — {_n_active} mục cần hoàn thành</div>',
        unsafe_allow_html=True,
    )
    st.caption(
        f"• Phải chấm đủ cả {_n_active} mục (Đạt hoặc Không Đạt)  "
        "• Khi Không Đạt: bắt buộc điền Ghi chú ≥ 10 ký tự HOẶC tải ảnh minh chứng  "
        "• Tối đa 1 ảnh/mục · Ảnh được phân tích tự động bằng Claude Vision"
    )

    pass_count = fail_count = 0

    for item in _active_items:
        code     = item["code"]
        is_crit  = item["critical"]

        # Đọc TRỰC TIẾP từ widget key → màu đổi ngay lần bấm đầu tiên (không lag)
        cur = st.session_state.get(f"sup_seg_{code}")
        is_fail   = (cur == "❌ Không Đạt")
        has_img   = code in st.session_state.sup_imgs
        has_note  = len(st.session_state.sup_notes.get(code, "").strip()) >= 10
        need_evid = is_fail and not has_img and not has_note

        # Màu theo trạng thái — giống hệt checklist tab
        if cur == "✅ Đạt":
            row_left = "#16A34A"; row_bg = "#F0FDF4"
            code_icon = "✅"
            state_lbl = '<span style="font-size:0.7rem;font-weight:700;color:#16A34A;margin-left:6px">ĐẠT</span>'
        elif cur == "❌ Không Đạt":
            row_left = "#F97316" if need_evid else "#DC2626"
            row_bg   = "#FFF7ED" if need_evid else "#FFF5F5"
            code_icon = "❌"
            state_lbl = '<span style="font-size:0.7rem;font-weight:700;color:#DC2626;margin-left:6px">KHÔNG ĐẠT</span>'
        else:  # None — chưa chọn
            row_left = "#F59E0B"; row_bg = "#FFFBEB"
            code_icon = "○"
            state_lbl = '<span style="font-size:0.7rem;color:#D97706;margin-left:6px">chưa chấm</span>'

        crit_badge = (
            '<span style="background:#FEE2E2;color:#991B1B;font-size:0.65rem;font-weight:700;'
            'padding:1px 6px;border-radius:8px;margin-left:6px;border:1px solid #FECACA">BẮT BUỘC</span>'
        ) if is_crit else ""

        col_desc, col_ctrl = st.columns([0.60, 0.40])

        with col_desc:
            st.markdown(
                f'<div style="background:{row_bg};border-left:3px solid {row_left};'
                f'border-radius:0 8px 8px 0;padding:10px 14px;margin:4px 0;'
                f'transition:background 0.4s ease,border-color 0.4s ease">'
                f'<div style="margin-bottom:4px;display:flex;align-items:center;flex-wrap:wrap;gap:4px">'
                f'<span style="font-size:0.75rem;font-weight:800;color:{row_left}">'
                f'{code_icon} {code}</span>'
                f'{crit_badge}{state_lbl}'
                + (f'<span style="font-size:0.68rem;color:#EA580C;font-weight:600;margin-left:8px">'
                   f'⚠️ Cần bổ sung ghi chú/ảnh</span>' if need_evid else '')
                + f'</div>'
                f'<div style="font-size:0.87rem;font-weight:500;color:#1E293B;line-height:1.55">'
                f'{item["icon"]} {item["desc"]}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            with st.expander("🔍 Hướng dẫn & tiêu chuẩn"):
                st.markdown(
                    f"**💡 Thực hiện:** {item['hint']}  \n"
                    f"**✅ Đạt khi:** {item['pass_std']}  \n"
                    f"**❌ Không đạt khi:** {item['fail_std']}  \n"
                    f"**📖 Căn cứ pháp lý:** {item['law']}"
                )

        with col_ctrl:
            st.segmented_control(
                label=code,
                options=["✅ Đạt", "❌ Không Đạt"],
                key=f"sup_seg_{code}",
                label_visibility="collapsed",
            )
            result = st.session_state.get(f"sup_seg_{code}")
            st.session_state.sup_r[code] = result

        if result == "✅ Đạt":
            pass_count += 1
        elif result == "❌ Không Đạt":
            fail_count += 1

        # Ghi chú + Ảnh — trong expander, mở tự động khi Không Đạt
        _note_req = (result == "❌ Không Đạt")
        _ev_open  = _note_req  # mở khi không đạt để nhắc nhập bằng chứng
        _ev_label = ("📝 Ghi chú & Ảnh minh chứng — BẮT BUỘC"
                     if _note_req else "📝 Ghi chú & Ảnh (tuỳ chọn)")
        with st.expander(_ev_label, expanded=_ev_open):
            note = st.text_area(
                "Ghi chú", key=f"sup_note_{code}", max_chars=300, height=68,
                label_visibility="collapsed",
                placeholder=(
                    "Mô tả lỗi cụ thể, VD: Xe không có thùng cách nhiệt, nhiệt kế đo 35°C..."
                    if _note_req else "Ghi chú thêm nếu cần..."
                ),
            )
            st.session_state.sup_notes[code] = note
            uploaded = st.file_uploader(
                "📷 Ảnh minh chứng (1 ảnh · jpg/png ≤ 5 MB)",
                type=["jpg", "jpeg", "png"],
                key=f"sup_img_{code}", accept_multiple_files=False,
            )
            if uploaded is not None:
                raw = uploaded.read()
                if len(raw) <= 5 * 1024 * 1024:
                    st.session_state.sup_imgs[code] = {
                        "bytes": raw,
                        "type": uploaded.type or "image/jpeg",
                        "name": uploaded.name,
                    }
                else:
                    st.warning(f"[{code}] Ảnh vượt 5 MB.")

        # Hiện ảnh đã tải + Vision (ngoài expander để luôn thấy)
        if code in st.session_state.sup_imgs:
            _img_data = st.session_state.sup_imgs[code]
            _i1, _i2 = st.columns([1, 3])
            _i1.image(_img_data["bytes"], width=160, caption=f"Ảnh [{code}]")

            if api_key:
                if _i2.button(
                    f"🔍 Phân tích ảnh [{code}] với Claude Vision",
                    key=f"sup_vis_btn_{code}",
                ):
                    import base64 as _b64mod
                    _b64str = _b64mod.b64encode(_img_data["bytes"]).decode()
                    _mtype  = _img_data.get("type", "image/jpeg")
                    _vp = (
                        f"Đây là hình ảnh minh chứng kiểm tra mục [{code}]: \"{item['desc']}\" "
                        f"trong biên bản kiểm tra nhà cung cấp suất ăn học đường Việt Nam. "
                        f"Tiêu chuẩn ĐẠT: {item['pass_std']}. "
                        f"Tiêu chuẩn KHÔNG ĐẠT: {item['fail_std']}. "
                        f"Người kiểm tra chấm: {result}. "
                        f"Hãy quan sát hình ảnh và cho biết trong 2–3 câu ngắn gọn: "
                        f"hình ảnh thể hiện điều gì, có phù hợp với kết quả đã chấm không, "
                        f"và những gì cần ghi nhận. Viết bằng tiếng Việt."
                    )
                    try:
                        _client = anthropic.Anthropic(api_key=api_key)
                        with _i2:
                            with st.spinner("Claude Vision đang phân tích hình ảnh..."):
                                _vmsg = _client.messages.create(
                                    model=MODEL_VISION, max_tokens=350,
                                    messages=[{"role": "user", "content": [
                                        {"type": "image", "source": {
                                            "type": "base64",
                                            "media_type": _mtype,
                                            "data": _b64str,
                                        }},
                                        {"type": "text", "text": _vp},
                                    ]}],
                                )
                        _vtext = _vmsg.content[0].text if _vmsg.content else ""
                        st.session_state.sup_vision[code] = _vtext
                    except Exception as _ve:
                        _i2.error(f"Lỗi Vision [{code}]: {_ve}")

            if code in st.session_state.sup_vision:
                _i2.markdown(
                    f'<div style="background:#EEF2FF;border-left:3px solid #6366F1;'
                    f'border-radius:6px;padding:8px 12px;font-size:0.8rem;color:#3730A3;margin-top:6px">'
                    f'🤖 <b>Vision [{code}]:</b> {st.session_state.sup_vision[code]}</div>',
                    unsafe_allow_html=True,
                )

        st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)

    # ── KPI realtime ──────────────────────────────────────────────────────────
    _checked = pass_count + fail_count  # số mục đã chọn (không tính None)
    st.markdown("<br>", unsafe_allow_html=True)
    m1, m2, m3, m4 = st.columns(4)
    pct = round(pass_count / max(_checked, 1) * 100) if _checked else 0
    crit_fails = [c for c, v in st.session_state.sup_r.items()
                  if v == "❌ Không Đạt" and c in _active_critical]
    if _checked < _n_active:
        alert_key, rating = "OK", "—"   # chưa chấm xong → chưa xếp loại
    elif crit_fails:
        alert_key, rating = "CRITICAL", "C"
    elif pass_count < _score_warn:
        alert_key, rating = "MAJOR", "C"
    elif pass_count < _score_pass:
        alert_key, rating = "MINOR", "B"
    else:
        alert_key, rating = "OK", "A"

    rating_color = {"A": "#16A34A", "B": "#F59E0B", "C": "#DC2626", "—": "#64748B"}[rating]

    m1.markdown(f'<div class="metric-box"><div class="metric-lbl">Đã kiểm tra</div>'
                f'<div class="metric-num c-blue">{_checked}</div>'
                f'<div class="metric-lbl">/ {_n_active} mục</div></div>',
                unsafe_allow_html=True)
    m2.markdown(f'<div class="metric-box"><div class="metric-lbl">✅ Đạt</div>'
                f'<div class="metric-num c-green">{pass_count}</div>'
                f'<div class="metric-lbl">điểm</div></div>',
                unsafe_allow_html=True)
    m3.markdown(f'<div class="metric-box"><div class="metric-lbl">❌ Không đạt</div>'
                f'<div class="metric-num c-red">{fail_count}</div>'
                f'<div class="metric-lbl">điểm</div></div>',
                unsafe_allow_html=True)
    m4.markdown(f'<div class="metric-box"><div class="metric-lbl">Xếp loại</div>'
                f'<div class="metric-num" style="color:{rating_color};font-size:1.6rem;line-height:1.2">'
                f'Loại {rating}<br>'
                f'<span style="font-size:0.75rem;font-weight:600">{pct}% đạt</span></div></div>',
                unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # Banner cảnh báo
    if crit_fails:
        fail_descs = ['[' + c + '] ' + next(x["desc"] for x in _active_items if x["code"] == c)
                      for c in crit_fails]
        st.markdown(
            '<div class="alert-critical">'
            '<div class="alert-title">🔴 VI PHẠM MỤC BẮT BUỘC (*) — BÁO BAN GIÁM HIỆU NGAY</div>'
            '<div class="alert-body">' + "<br>".join(f"• {d}" for d in fail_descs) + '</div>'
            '</div>',
            unsafe_allow_html=True,
        )
    elif alert_key == "MAJOR":
        st.markdown(
            '<div class="alert-major"><div class="alert-title">🟠 Nhà cung cấp CHƯA ĐẠT — '
            'Yêu cầu khắc phục trước bữa ăn tiếp theo</div></div>',
            unsafe_allow_html=True,
        )
    elif alert_key == "MINOR":
        st.markdown(
            '<div class="alert-minor"><div class="alert-title">🟡 Cần cải thiện — '
            'Thông báo nhà cung cấp trong 24 giờ</div></div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div class="alert-ok"><div class="alert-title">✅ Nhà cung cấp ĐẠT chuẩn — '
            'Lưu hồ sơ và tiếp tục theo dõi</div></div>',
            unsafe_allow_html=True,
        )

    # ── AI Phân tích rủi ro tổng hợp ─────────────────────────────────────────
    st.markdown('<div class="sec-hdr">🤖 AI Phân tích rủi ro tổng hợp (Tuỳ chọn — chạy trước khi tạo báo cáo)</div>',
                unsafe_allow_html=True)
    st.markdown(
        '<div class="sf-card" style="border-left:3px solid #7C3AED;margin-bottom:10px">'
        '<div class="sf-card-title">Mục đích &amp; Cách đánh giá của AI</div>'
        '<div class="sf-card-body">'
        '<b>AI xem xét:</b> (1) Danh sách mục Không Đạt + ghi chú mô tả lỗi của người kiểm tra, '
        '(2) Nhận xét hình ảnh từ Claude Vision (nếu đã chạy từng mục). '
        '<br><b>AI đánh giá:</b> Mức độ rủi ro ngộ độc thực phẩm theo khung HACCP — '
        'Critical Control Point nào bị vi phạm, nguyên nhân có thể, mức độ nghiêm trọng. '
        '<br><b>AI đề xuất:</b> Biện pháp khắc phục cụ thể, có dẫn chiếu điều khoản Luật ATTP Việt Nam '
        '(NĐ 15/2018, QCVN 8-1:2011, TTLT 13/2016). '
        '<br><b>Kết quả AI được đính kèm vào báo cáo Word.</b>'
        '</div></div>',
        unsafe_allow_html=True,
    )

    _ai_disabled = not api_key or fail_count == 0
    _ai_hint = ("Chỉ cần thiết khi có mục Không Đạt" if fail_count == 0
                else ("Cần nhập API key" if not api_key else f"Phân tích {fail_count} mục không đạt"))
    if st.button(f"🤖 Chạy phân tích AI ({_ai_hint})",
                 disabled=_ai_disabled, use_container_width=False):
        _fails = []
        for _it in _active_items:
            _c = _it["code"]
            if st.session_state.sup_r.get(_c) == "❌ Không Đạt":
                _note_txt = st.session_state.sup_notes.get(_c, "").strip()
                _vision_txt = st.session_state.sup_vision.get(_c, "")
                _entry = f"[{_c}] {_it['desc']}"
                if _note_txt:
                    _entry += f" — Ghi chú: {_note_txt}"
                if _vision_txt:
                    _entry += f" — Vision AI nhận xét: {_vision_txt}"
                _fails.append(_entry)

        _prompt = (
            f"Đây là kết quả kiểm tra nhà cung cấp suất ăn học đường '{sup_name}' "
            f"tại trường '{sup_school}' ngày {sup_date.strftime('%d/%m/%Y')}. "
            f"Xếp loại: {rating} ({pct}% đạt — {pass_count}/{_n_active} điểm). "
            f"Các vi phạm phát hiện:\n" + "\n".join(f"  {x}" for x in _fails) + "\n\n"
            "Hãy viết phân tích ngắn gọn (khoảng 200 từ) bằng tiếng Việt bao gồm:\n"
            "1. Đánh giá mức độ rủi ro ngộ độc thực phẩm theo HACCP (CCP nào bị vi phạm)\n"
            "2. Nguyên nhân có thể và hậu quả tiềm ẩn\n"
            "3. Biện pháp khắc phục ưu tiên, dẫn chiếu điều khoản pháp luật cụ thể\n"
            "4. Khuyến nghị hành động tiếp theo (tạm dừng / tiếp tục / tăng tần suất giám sát)"
        )
        try:
            _client = anthropic.Anthropic(api_key=api_key)
            with st.spinner("AI đang phân tích toàn bộ kết quả kiểm tra..."):
                _amsg = _client.messages.create(
                    model=MODEL, max_tokens=800,
                    messages=[{"role": "user", "content": _prompt}],
                )
            _ai_text = _amsg.content[0].text if _amsg.content else ""
            st.session_state["sup_ai_analysis"] = _ai_text
            st.markdown(
                '<div class="sf-card" style="border-left:3px solid #7C3AED">'
                '<div class="sf-card-title">🤖 Kết quả phân tích AI tổng hợp</div>'
                f'<div class="sf-card-body">{_ai_text}</div>'
                '</div>',
                unsafe_allow_html=True,
            )
        except Exception as _ae:
            st.error(f"Lỗi AI: {_ae}")

    if "sup_ai_analysis" in st.session_state and st.session_state.sup_ai_analysis:
        st.caption("✅ Kết quả AI đã được lưu — sẽ được đính kèm vào báo cáo Word khi tạo.")

    # ── Validate ─────────────────────────────────────────────────────────────
    # Mục chưa chọn (None = chưa chấm)
    _unselected = [item["code"] for item in _active_items
                   if st.session_state.sup_r.get(item["code"]) is None]
    # Mục Không Đạt thiếu bằng chứng
    _missing_evid = [
        item["code"] for item in _active_items
        if st.session_state.sup_r.get(item["code"]) == "❌ Không Đạt"
        and item["code"] not in st.session_state.sup_imgs
        and len(st.session_state.sup_notes.get(item["code"], "").strip()) < 10
    ]
    can_submit = (
        bool(sup_school.strip()) and
        bool(sup_name.strip()) and
        bool(sup_inspector.strip()) and
        len(_unselected) == 0 and
        len(_missing_evid) == 0
    )

    # ── Tạo báo cáo Word + lưu DB ─────────────────────────────────────────────
    st.markdown('<div class="sec-hdr">📄 Tạo báo cáo & Lưu hồ sơ</div>', unsafe_allow_html=True)

    if not can_submit:
        _miss = []
        if not sup_school.strip():    _miss.append("Tên trường")
        if not sup_name.strip():      _miss.append("Tên nhà cung cấp")
        if not sup_inspector.strip(): _miss.append("Người kiểm tra")
        if _unselected:
            _miss.append(f"Chưa chấm {len(_unselected)} mục: {', '.join(_unselected)}")
        if _missing_evid:
            _miss.append(f"Mục Không Đạt chưa có ghi chú/ảnh: {', '.join(_missing_evid)}")
        st.warning("⚠️ Chưa đủ điều kiện tạo báo cáo: " + " · ".join(_miss))

    if st.button("📄 Tải báo cáo", type="primary",
                 disabled=not can_submit, use_container_width=True,
                 key="sup_submit_btn"):
        guard_key = f"sup_saved_{sup_school}_{sup_date}_{sup_inspector}"
        already_saved = st.session_state.get(guard_key, False)

        ai_narrative = st.session_state.get("sup_ai_analysis", f"Xếp loại {rating}")

        # Lưu DB (chỉ 1 lần)
        if not already_saved and db_ok():
            _res_dict = {}
            for _it in _active_items:
                _c = _it["code"]
                _v = st.session_state.sup_r.get(_c, "")
                _res_dict[_c] = "Đạt" if _v == "✅ Đạt" else "Không Đạt"
            _notes_dict = {c: st.session_state.sup_notes.get(c, "") for c in _res_dict}
            _sid = db_save_checklist(
                school=sup_school, date_str=str(sup_date),
                inspector=sup_inspector, menu=f"NCC: {sup_name}",
                level="—", results=_res_dict, notes=_notes_dict,
                alert_level=alert_key, pass_count=pass_count, fail_count=fail_count,
                ai_narrative=ai_narrative[:500],
                extra_results={"contract": sup_contract, "supplier": sup_name},
                check_type="nha_cung_cap",
            )
            if _sid:
                st.session_state[guard_key] = True
                st.success("✅ Đã lưu vào database!")
            else:
                st.warning("Không lưu được DB — kiểm tra kết nối Supabase.")
        elif already_saved:
            st.info("Phiên này đã lưu DB — chỉ xuất lại file Word.")
        else:
            st.warning("Database chưa kết nối — chỉ tạo file Word, không lưu DB.")

        # Tạo Word
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO

            doc = Document()
            _sty = doc.styles["Normal"]
            _sty.font.name = "Times New Roman"
            _sty.font.size = Pt(13)
            for _sec in doc.sections:
                _sec.top_margin    = Cm(2.5); _sec.bottom_margin = Cm(2.5)
                _sec.left_margin   = Cm(3.0); _sec.right_margin  = Cm(2.0)

            def _wr(para, text, bold=False, size=13, color=None):
                r = para.add_run(text)
                r.font.name = "Times New Roman"; r.font.size = Pt(size); r.bold = bold
                if color: r.font.color.rgb = color
                return r

            # Quốc hiệu
            for _txt in ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", "Độc lập – Tự do – Hạnh phúc"):
                _p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _wr(_p, _txt, bold=True)

            doc.add_paragraph("")
            _p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _wr(_p, "BIÊN BẢN KIỂM TRA NHÀ CUNG CẤP SUẤT ĂN HỌC ĐƯỜNG", bold=True, size=14)

            _p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _wr(_p, f"Ngày {sup_date.strftime('%d/%m/%Y')} · NCC: {sup_name} · Trường: {sup_school}")

            doc.add_paragraph("")
            for _lbl, _val in [
                ("Tên trường", sup_school), ("Nhà cung cấp", sup_name),
                ("Người kiểm tra", sup_inspector), ("Ngày kiểm tra", sup_date.strftime("%d/%m/%Y")),
                ("Số hợp đồng", sup_contract or "—"),
                ("Kết quả tổng hợp", f"Loại {rating} — {pct}% đạt ({pass_count}/{_n_active} điểm · {'Giao hàng' if _delivery_mode else 'Toàn diện'})"),
            ]:
                _p = doc.add_paragraph()
                _wr(_p, f"{_lbl}: ", bold=True)
                _wr(_p, _val)

            doc.add_paragraph("")
            _mode_txt = "KIỂM TRA GIAO HÀNG (10 ĐIỂM · S03–S12)" if _delivery_mode else "KIỂM TRA TOÀN DIỆN (12 ĐIỂM · S01–S12)"
            _p = doc.add_paragraph(); _wr(_p, f"KẾT QUẢ {_mode_txt}", bold=True)

            # Bảng 5 cột
            _tbl = doc.add_table(rows=1, cols=5)
            _tbl.style = "Table Grid"
            _hdrs = ["Mã", "Nội dung kiểm tra", "Kết quả", "Ghi chú", "Vision AI"]
            for _ci, _h in enumerate(_hdrs):
                _tbl.rows[0].cells[_ci].text = _h
                for _pp in _tbl.rows[0].cells[_ci].paragraphs:
                    for _rr in _pp.runs:
                        _rr.bold = True
                        _rr.font.name = "Times New Roman"; _rr.font.size = Pt(11)

            for _it in _active_items:
                _c = _it["code"]
                _row = _tbl.add_row().cells
                _row[0].text = f"{_c}{'(*)' if _it['critical'] else ''}"
                _row[1].text = _it["desc"]
                _rv = st.session_state.sup_r.get(_c, "")
                _row[2].text = "Đạt" if _rv == "✅ Đạt" else "Không Đạt"
                _row[3].text = st.session_state.sup_notes.get(_c, "")
                _row[4].text = st.session_state.sup_vision.get(_c, "")
                if _row[2].text == "Không Đạt":
                    for _pp in _row[2].paragraphs:
                        for _rr in _pp.runs:
                            _rr.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                for _cell in _row:
                    for _pp in _cell.paragraphs:
                        for _rr in _pp.runs:
                            _rr.font.name = "Times New Roman"; _rr.font.size = Pt(11)

            # Ảnh minh chứng trong Word (embed nếu có)
            _img_count = sum(1 for c in st.session_state.sup_imgs
                             if st.session_state.sup_r.get(c) == "❌ Không Đạt")
            if _img_count > 0:
                doc.add_paragraph("")
                _p = doc.add_paragraph(); _wr(_p, "ẢNH MINH CHỨNG CÁC MỤC KHÔNG ĐẠT", bold=True)
                for _it2 in _active_items:
                    _c2 = _it2["code"]
                    if (_c2 in st.session_state.sup_imgs and
                            st.session_state.sup_r.get(_c2) == "❌ Không Đạt"):
                        _p2 = doc.add_paragraph()
                        _wr(_p2, f"[{_c2}] {_it2['desc']}", bold=True, size=11)
                        try:
                            from io import BytesIO as _BIO
                            _ibuf = _BIO(st.session_state.sup_imgs[_c2]["bytes"])
                            doc.add_picture(_ibuf, width=Cm(10))
                        except Exception:
                            _p3 = doc.add_paragraph()
                            _wr(_p3, "(Không nhúng được ảnh — xem ảnh trong hồ sơ digital)", size=11)

            # Kết luận
            doc.add_paragraph("")
            _concl = {
                "A": f"Nhà cung cấp {sup_name} ĐẠT chuẩn (Loại A). Tiếp tục hợp đồng bình thường.",
                "B": f"Nhà cung cấp {sup_name} xếp Loại B. Thông báo khắc phục trong 24 giờ.",
                "C": f"Nhà cung cấp {sup_name} KHÔNG ĐẠT (Loại C). Báo Ban Giám Hiệu ngay. "
                     f"Xem xét tạm dừng hợp đồng.",
            }[rating]
            _p = doc.add_paragraph(); _wr(_p, "KẾT LUẬN: ", bold=True)
            _wr(_p, _concl)
            if crit_fails:
                _wr(_p, f"\n⚠️ Vi phạm mục bắt buộc: {', '.join(crit_fails)}",
                    color=RGBColor(0xDC, 0x26, 0x26))

            # AI analysis trong Word
            if ai_narrative and "Xếp loại" not in ai_narrative:
                doc.add_paragraph("")
                _p = doc.add_paragraph(); _wr(_p, "PHÂN TÍCH RỦI RO (AI):", bold=True)
                _p2 = doc.add_paragraph(); _wr(_p2, ai_narrative, size=12)

            # Chữ ký
            doc.add_paragraph("")
            _p = doc.add_paragraph(
                f"......., ngày {sup_date.strftime('%d')} tháng "
                f"{sup_date.strftime('%m')} năm {sup_date.strftime('%Y')}"
            )
            _p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for _rr in _p.runs: _rr.font.name = "Times New Roman"; _rr.font.size = Pt(13)

            _sig_tbl = doc.add_table(rows=1, cols=2)
            _s1, _s2 = _sig_tbl.rows[0].cells
            _s1.text = "NGƯỜI KIỂM TRA\n(Ký, ghi rõ họ tên)"
            _s2.text = "ĐẠI DIỆN NHÀ CUNG CẤP\n(Ký, ghi rõ họ tên)"
            for _cell in [_s1, _s2]:
                for _pp in _cell.paragraphs:
                    _pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for _rr in _pp.runs:
                        _rr.bold = True
                        _rr.font.name = "Times New Roman"; _rr.font.size = Pt(13)

            _buf = BytesIO()
            doc.save(_buf); _buf.seek(0)
            _fn = f"KiemTraNCC_{sup_name.replace(' ', '_')}_{sup_date.strftime('%Y%m%d')}.docx"
            st.download_button(
                "⬇️ Tải báo cáo (.docx)", data=_buf.getvalue(), file_name=_fn,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as _we:
            st.error(f"Lỗi tạo Word: {_we}")

    # ── Task#5: Hồ sơ & Chứng nhận Nhà Cung Cấp (BGH quản lý) ───────────────
    if (role == "Ban Giám Hiệu" or st.session_state.get("is_super")) and db_ok():
        st.markdown('<div class="sec-hdr">📋 Hồ sơ Nhà Cung Cấp — Chứng nhận & Hết hạn</div>',
                    unsafe_allow_html=True)
        _s5_school = st.session_state.get("user_school","") or sup_school
        _ncc_reg5  = db_get_ncc_registry(school=_s5_school)
        _today5    = now_vn().date()

        # Alert hết hạn
        _exp_warns = []
        for _n5 in _ncc_reg5:
            for _f5, _l5 in [("license_expiry","Giấy phép"), ("attp_expiry","Chứng nhận ATTP")]:
                _e5 = _n5.get(_f5)
                if _e5:
                    try:
                        import datetime as _dt5
                        _e5d = _dt5.date.fromisoformat(_e5)
                        _dl5 = (_e5d - _today5).days
                        if _dl5 <= 30:
                            _exp_warns.append(
                                f"{'🚨' if _dl5 <= 0 else '⚠️'} "
                                f"<b>{_n5['ncc_name']}</b> — {_l5} "
                                f"{'HẾT HẠN' if _dl5 <= 0 else f'còn {_dl5} ngày'} "
                                f"({_e5d.strftime('%d/%m/%Y')})"
                            )
                    except Exception:
                        pass
        if _exp_warns:
            st.markdown(
                '<div style="background:#FEF2F2;border:1.5px solid #FCA5A5;border-radius:8px;'
                'padding:10px 16px;margin-bottom:8px">'
                '<b style="color:#991B1B">🚨 Chứng nhận sắp/đã hết hạn:</b><br>'
                + "<br>".join(f'<span style="font-size:0.85rem;color:#991B1B">{w}</span>'
                               for w in _exp_warns)
                + '</div>', unsafe_allow_html=True,
            )
        if _ncc_reg5:
            import pandas as _pd5
            st.dataframe(_pd5.DataFrame([{
                "Tên NCC": n["ncc_name"], "Giấy phép": n.get("license_no","—"),
                "Hết hạn GP": (n.get("license_expiry","") or "—")[:10],
                "Hết hạn ATTP": (n.get("attp_expiry","") or "—")[:10],
                "SĐT": n.get("phone","—"),
            } for n in _ncc_reg5]), use_container_width=True, hide_index=True)
        with st.expander("➕ Thêm / Cập nhật NCC"):
            _n5c1, _n5c2 = st.columns(2)
            _n5nm  = _n5c1.text_input("Tên NCC", placeholder="Công ty TNHH...", key="n5_name")
            _n5lic = _n5c2.text_input("Số giấy phép", placeholder="01/GPCSSX-2024", key="n5_lic")
            _n5c3, _n5c4, _n5c5 = st.columns(3)
            _n5le  = _n5c3.date_input("Hết hạn Giấy phép", value=None, key="n5_le", format="DD/MM/YYYY")
            _n5ae  = _n5c4.date_input("Hết hạn ATTP cert", value=None, key="n5_ae", format="DD/MM/YYYY")
            _n5ph  = _n5c5.text_input("SĐT", placeholder="0901...", key="n5_ph")
            if st.button("💾 Lưu", key="n5_save", type="primary"):
                if _n5nm.strip() and _s5_school:
                    if db_save_ncc_registry(_s5_school, _n5nm.strip(), _n5lic.strip(),
                                             str(_n5le) if _n5le else "",
                                             str(_n5ae) if _n5ae else "", _n5ph.strip()):
                        st.success(f"✅ Đã lưu: {_n5nm}"); st.rerun()
                else:
                    st.warning("Điền tên NCC và có tài khoản trường.")
    elif db_ok():
        # BGS / Y Tế: chỉ xem
        _s5_school_view = st.session_state.get("user_school","")
        _ncc_view = db_get_ncc_registry(school=_s5_school_view)
        if _ncc_view:
            st.markdown('<div class="sec-hdr">📋 Danh sách Nhà Cung Cấp được phê duyệt</div>',
                        unsafe_allow_html=True)
            import pandas as _pd5v
            st.dataframe(_pd5v.DataFrame([{
                "Tên NCC": n["ncc_name"],
                "Giấy phép": n.get("license_no","—"),
                "Hết hạn GP": (n.get("license_expiry","") or "—")[:10],
                "Hết hạn ATTP": (n.get("attp_expiry","") or "—")[:10],
            } for n in _ncc_view]), use_container_width=True, hide_index=True)


def tab_ncc_bgh(school: str = ""):
    """Tab 🏭 Nhà Cung Cấp — Dashboard tổng hợp dành riêng cho Ban Giám Hiệu."""
    import pandas as _pd_ncc
    import plotly.graph_objects as _go_ncc
    import plotly.express as _px_ncc

    # ── Banner header ──────────────────────────────────────────────────────────
    st.markdown(
        '<div style="background:linear-gradient(135deg,#3B0764 0%,#7C3AED 60%,#4C1D95 100%);'
        'border-radius:12px;padding:14px 22px;margin-bottom:14px">'
        '<div style="color:white;font-size:1.05rem;font-weight:700">🏭 Quản lý Nhà Cung Cấp</div>'
        '<div style="color:#DDD6FE;font-size:0.8rem">'
        'Hồ sơ chứng nhận · Hiệu suất đánh giá · Cảnh báo hết hạn · Lịch kiểm tra'
        '</div></div>', unsafe_allow_html=True,
    )

    if not db_ok():
        st.warning("Cần kết nối database để xem thông tin nhà cung cấp.")
        return

    _today = now_vn().date()

    # ── Alert: cert expiry ─────────────────────────────────────────────────────
    _reg = db_get_ncc_registry(school=school)
    _exp_alerts = []
    for _n in _reg:
        for _f, _l in [("license_expiry","Giấy phép"), ("attp_expiry","Chứng nhận ATTP")]:
            _e = _n.get(_f)
            if _e:
                try:
                    import datetime as _dtn
                    _ed = _dtn.date.fromisoformat(_e)
                    _dl = (_ed - _today).days
                    if _dl <= 30:
                        _exp_alerts.append(
                            f"{'🚨' if _dl <= 0 else '⚠️'} <b>{_n['ncc_name']}</b> — "
                            f"{_l} {'HẾT HẠN' if _dl <= 0 else f'còn {_dl} ngày'} "
                            f"({_ed.strftime('%d/%m/%Y')})"
                        )
                except Exception: pass

    # Alert: monthly check
    _this_m = now_vn().strftime("%Y-%m")
    try:
        _ncc_month_ses = _get_sb().table("checklist_sessions").select("check_date")\
            .eq("check_type","nha_cung_cap").eq("school_name",school)\
            .gte("check_date",f"{_this_m}-01").execute().data or []
    except Exception:
        _ncc_month_ses = []
    if not _ncc_month_ses:
        _exp_alerts.append(f"⏰ Tháng {now_vn().month:02d}/{now_vn().year} chưa có đánh giá NCC toàn diện (12 điểm)")

    if _exp_alerts:
        st.markdown(
            '<div style="background:#FEF2F2;border:1.5px solid #FCA5A5;border-radius:8px;'
            'padding:12px 16px;margin-bottom:12px">'
            '<div style="font-weight:700;color:#991B1B;margin-bottom:6px">🚨 Cảnh báo NCC</div>'
            + "".join(f'<div style="font-size:0.85rem;color:#991B1B;margin:3px 0">{a}</div>'
                       for a in _exp_alerts)
            + '</div>', unsafe_allow_html=True,
        )

    # ── NCC Performance Dashboard ──────────────────────────────────────────────
    try:
        _ncc_ses = _get_sb().table("checklist_sessions")\
            .select("check_date,pass_count,total_items,alert_level,menu_today")\
            .eq("check_type","nha_cung_cap").eq("school_name",school)\
            .order("check_date",desc=True).limit(100).execute().data or []
    except Exception:
        _ncc_ses = []

    if _ncc_ses:
        _ncc_df = _pd_ncc.DataFrame([{
            "Ngày": s["check_date"],
            "Tỷ lệ": round(s["pass_count"]/max(s["total_items"],1)*100, 1),
            "NCC": (s.get("menu_today","") or "").replace("NCC:","").strip()[:30],
            "Loại": ("A" if s.get("pass_count",0) >= 10
                     else "B" if s.get("pass_count",0) >= 8 else "C"),
        } for s in _ncc_ses])

        # KPI tổng quan
        _ncc_total = len(_ncc_df)
        _ncc_a = (_ncc_df["Loại"]=="A").sum()
        _ncc_b = (_ncc_df["Loại"]=="B").sum()
        _ncc_c = (_ncc_df["Loại"]=="C").sum()
        _ncc_avg = _ncc_df["Tỷ lệ"].mean()

        k1,k2,k3,k4 = st.columns(4)
        for _kc, _kv, _kl, _kclr in [
            (k1, _ncc_total, "Tổng lần đánh giá", "c-blue"),
            (k2, _ncc_a, "✅ Loại A (≥83%)", "c-green"),
            (k3, _ncc_b, "🟡 Loại B (67–82%)", "c-orange"),
            (k4, _ncc_c, "🔴 Loại C (<67%)", "c-red" if _ncc_c>0 else "c-green"),
        ]:
            _kc.markdown(f'<div class="metric-box"><div class="metric-lbl">{_kl}</div>'
                          f'<div class="metric-num {_kclr}">{_kv}</div></div>',
                          unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        # Charts: Donut A/B/C + Trend line
        _ch1, _ch2 = st.columns(2)
        with _ch1:
            _fig_pie = _go_ncc.Figure(_go_ncc.Pie(
                labels=["Loại A","Loại B","Loại C"],
                values=[_ncc_a, _ncc_b, _ncc_c], hole=0.45,
                marker_colors=["#16A34A","#F59E0B","#DC2626"],
                textfont_size=13, textinfo="percent+label",
                hovertemplate="%{label}: %{value} lần<extra></extra>",
            ))
            _fig_pie.update_layout(
                plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                font=dict(family="Inter",size=12), margin=dict(l=10,r=10,t=30,b=10),
                height=260, title=dict(text="Phân bố xếp loại NCC", font=dict(size=13)),
                showlegend=False,
                annotations=[dict(text=f"<b>{_ncc_avg:.0f}%</b><br>TB", x=0.5, y=0.5,
                                   showarrow=False, font_size=14)],
            )
            st.plotly_chart(_fig_pie, use_container_width=True)

        with _ch2:
            try:
                _ncc_trend = _ncc_df.sort_values("Ngày").tail(20).copy()
                _ncc_trend["Ngày_fmt"] = _pd_ncc.to_datetime(_ncc_trend["Ngày"])\
                    .dt.strftime("%d/%m")
                _fig_trend = _go_ncc.Figure(_go_ncc.Scatter(
                    x=_ncc_trend["Ngày_fmt"], y=_ncc_trend["Tỷ lệ"],
                    mode="lines+markers",
                    line=dict(color="#7C3AED", width=2.5),
                    marker=dict(size=8, color=["#16A34A" if v>=83 else "#F59E0B" if v>=67 else "#DC2626"
                                               for v in _ncc_trend["Tỷ lệ"]]),
                    hovertemplate="Ngày %{x}<br>Tỷ lệ: %{y:.0f}%<extra></extra>",
                ))
                _fig_trend.add_hline(y=83, line_dash="dot", line_color="#16A34A",
                                      annotation_text=" Loại A", annotation_font_size=10)
                _fig_trend.add_hline(y=67, line_dash="dot", line_color="#F59E0B",
                                      annotation_text=" Loại B", annotation_font_size=10)
                _fig_trend.update_layout(
                    plot_bgcolor="white", paper_bgcolor="#F8FAFC",
                    font=dict(family="Inter",size=11), margin=dict(l=10,r=10,t=30,b=10),
                    height=260, title=dict(text="Xu hướng điểm NCC theo thời gian", font=dict(size=13)),
                    xaxis=dict(showgrid=False), yaxis=dict(range=[0,110], ticksuffix="%"),
                    showlegend=False,
                )
                st.plotly_chart(_fig_trend, use_container_width=True)
            except Exception: pass

        # Bảng lịch sử
        with st.expander(f"📋 Lịch sử {_ncc_total} lần đánh giá NCC"):
            _ncc_show = _ncc_df.copy()
            try:
                _ncc_show["Ngày"] = _pd_ncc.to_datetime(_ncc_show["Ngày"])\
                    .dt.strftime("%d/%m/%Y")
            except Exception: pass
            st.dataframe(_ncc_show, use_container_width=True, hide_index=True)
    else:
        st.markdown(
            '<div style="background:#F5F3FF;border:1px dashed #DDD6FE;border-radius:12px;'
            'padding:24px;text-align:center;margin:8px 0">'
            '<div style="font-size:2rem;margin-bottom:8px">🏭</div>'
            '<div style="font-size:0.95rem;font-weight:600;color:#5B21B6;margin-bottom:6px">'
            'Chưa có kết quả đánh giá NCC</div>'
            '<div style="font-size:0.82rem;color:#64748B">'
            '→ Vào tab 🏭 Nhà Cung Cấp (BGS/Y Tế) để thực hiện đánh giá đầu tiên'
            '</div></div>', unsafe_allow_html=True,
        )

    # ── NCC Registry: Hồ sơ & Chứng nhận ─────────────────────────────────────
    st.markdown('<div class="sec-hdr">📋 Hồ sơ & Chứng nhận</div>', unsafe_allow_html=True)
    if _reg:
        _reg_df = _pd_ncc.DataFrame([{
            "Tên NCC": n["ncc_name"], "Số GP": n.get("license_no","—"),
            "Hết hạn GP": (n.get("license_expiry","") or "—")[:10],
            "Hết hạn ATTP": (n.get("attp_expiry","") or "—")[:10],
            "SĐT": n.get("phone","—"),
        } for n in _reg])
        st.dataframe(_reg_df, use_container_width=True, hide_index=True)
    else:
        st.info("Chưa có hồ sơ NCC nào. Thêm bên dưới.")

    with st.expander("➕ Thêm / Cập nhật hồ sơ NCC"):
        _r1, _r2 = st.columns(2)
        _rn = _r1.text_input("Tên NCC", placeholder="Công ty TNHH Bếp Xanh", key="rncc_nm")
        _rl = _r2.text_input("Số giấy phép", placeholder="01/GPCSSX-2024", key="rncc_lic")
        _r3, _r4, _r5 = st.columns(3)
        _rle = _r3.date_input("Hết hạn Giấy phép", value=None, key="rncc_le", format="DD/MM/YYYY")
        _rae = _r4.date_input("Hết hạn ATTP cert", value=None, key="rncc_ae", format="DD/MM/YYYY")
        _rph = _r5.text_input("SĐT", placeholder="0901...", key="rncc_ph")
        if st.button("💾 Lưu hồ sơ NCC", key="rncc_save", type="primary"):
            if _rn.strip():
                if db_save_ncc_registry(school or st.session_state.get("user_school",""),
                                         _rn.strip(), _rl.strip(),
                                         str(_rle) if _rle else "", str(_rae) if _rae else "",
                                         _rph.strip()):
                    st.success(f"✅ Đã lưu hồ sơ: {_rn}"); st.rerun()
            else:
                st.warning("Điền tên nhà cung cấp.")


def show_onboarding_banner(school: str, profiles: list, sessions_count: int):
    """G5: Banner hướng dẫn bắt đầu cho trường mới — hiện đầu trang BGH."""
    if st.session_state.get("onboarding_dismissed"):
        return

    # Kiểm tra mức độ hoàn thành setup
    _has_yte  = any(p.get("role") == "y_te_hoc_duong"  for p in profiles)
    _has_bgs  = any(p.get("role") == "ban_giam_sat"    for p in profiles)
    _has_ph   = any(p.get("role") == "phu_huynh"       for p in profiles)
    _has_data = sessions_count > 0
    _steps = [_has_yte, _has_bgs, _has_ph, _has_data]
    _done  = sum(_steps)

    if _done == 4:
        # Tất cả bước hoàn thành — ẩn banner
        st.session_state.onboarding_dismissed = True
        return

    st.markdown(
        '<div style="background:linear-gradient(135deg,#0F2651 0%,#1D4ED8 100%);'
        'border-radius:14px;padding:20px 24px;margin-bottom:16px">'
        '<div style="display:flex;justify-content:space-between;align-items:flex-start;'
        'flex-wrap:wrap;gap:12px">'
        '<div>'
        '<div style="color:white;font-size:1.05rem;font-weight:700;margin-bottom:4px">'
        f'🎉 Chào mừng {school} đến với SchoolFood AI!</div>'
        '<div style="color:#BFDBFE;font-size:0.82rem">'
        f'Đã hoàn thành {_done}/4 bước thiết lập — hoàn tất để kích hoạt đầy đủ tính năng</div>'
        '</div>'
        f'<div style="background:rgba(255,255,255,0.15);border-radius:20px;padding:4px 14px;'
        f'color:white;font-size:0.82rem;font-weight:700">{_done}/4 bước</div>'
        '</div>'
        '</div>',
        unsafe_allow_html=True,
    )

    # Progress bar
    _pct = int(_done / 4 * 100)
    st.markdown(
        f'<div style="background:#E2E8F0;border-radius:4px;height:6px;margin:-8px 0 12px 0">'
        f'<div style="background:#16A34A;height:6px;border-radius:4px;width:{_pct}%;'
        f'transition:width 0.5s"></div></div>',
        unsafe_allow_html=True,
    )

    # 4 bước setup
    _step_defs = [
        ("🏥 Tạo tài khoản Y Tế Học Đường",
         "Người kiểm thực 3 bước hàng ngày — vai trò quan trọng nhất",
         _has_yte),
        ("👥 Tạo tài khoản Ban Giám Sát (Đại Diện PHHS)",
         "Người thực hiện checklist 20 điểm — 2 lần/tuần",
         _has_bgs),
        ("👨‍👩‍👧 Tạo tài khoản Phụ Huynh (tuỳ chọn)",
         "Phụ huynh xem kết quả và gửi phản hồi",
         _has_ph),
        ("✅ Thực hiện kiểm tra bữa ăn đầu tiên",
         "Sau khi có đủ tài khoản, mời Y Tế / BGS đăng nhập và bắt đầu kiểm tra",
         _has_data),
    ]

    _c1, _c2 = st.columns(2)
    for i, (title, desc, done) in enumerate(_step_defs):
        _col = _c1 if i % 2 == 0 else _c2
        _icon = "✅" if done else "⬜"
        _bg   = "#F0FDF4" if done else "#F8FAFC"
        _bd   = "#86EFAC" if done else "#E2E8F0"
        _col.markdown(
            f'<div style="background:{_bg};border:1px solid {_bd};border-radius:8px;'
            f'padding:10px 14px;margin:4px 0">'
            f'<div style="font-size:0.83rem;font-weight:600;color:#1E293B">'
            f'{_icon} {title}</div>'
            f'<div style="font-size:0.75rem;color:#64748B;margin-top:2px">{desc}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    _ob1, _ob2 = st.columns([3, 1])
    _ob1.caption("💡 Bước 1–3: Vào tab **👤 Quản lý tài khoản** → Thêm người dùng mới")
    if _ob2.button("✕ Ẩn hướng dẫn", key="dismiss_onboarding"):
        st.session_state.onboarding_dismissed = True
        st.rerun()
    st.markdown('<div style="height:8px"></div>', unsafe_allow_html=True)


def show_login_page():
    """Trang đăng nhập — hiển thị trước khi vào app chính."""
    # Logo + tiêu đề
    st.markdown(
        '<div style="text-align:center;padding:50px 20px 32px">'
        '<div style="font-size:3.2rem">🍱</div>'
        '<div style="font-size:2rem;font-weight:800;color:#1B3B6F;margin-top:10px">SchoolFood AI</div>'
        '<div style="font-size:0.95rem;color:#64748B;margin-top:6px">'
        'Nền tảng giám sát An toàn Thực phẩm bữa ăn học đường</div>'
        '</div>',
        unsafe_allow_html=True,
    )

    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        st.markdown(
            '<div style="background:white;border:1px solid #E2E8F0;border-radius:16px;'
            'padding:32px 28px;box-shadow:0 4px 24px rgba(0,0,0,0.09)">'
            '<div style="font-size:1.1rem;font-weight:700;color:#1E293B;'
            'text-align:center;margin-bottom:24px">Đăng nhập tài khoản</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        # Lấy email đã lưu từ lần đăng nhập trước (trong cùng phiên trình duyệt)
        _saved_email = st.session_state.get("last_login_email", "")

        with st.form("login_form", clear_on_submit=False):
            _email    = st.text_input("📧 Email", value=_saved_email,
                                       placeholder="ten@truong.edu.vn")
            _password = st.text_input("🔒 Mật khẩu", type="password", placeholder="••••••••")
            _submit   = st.form_submit_button("Đăng nhập →", type="primary",
                                               use_container_width=True)

        if _submit:
            if not _email or not _password:
                st.error("Vui lòng nhập đầy đủ email và mật khẩu.")
            elif not db_ok():
                st.error("❌ Database chưa kết nối. Liên hệ quản trị viên thiết lập Supabase.")
            else:
                try:
                    with st.spinner("Đang xác thực..."):
                        _user = db_auth_login(_email, _password)
                    _profile = db_get_profile(_user["id"])
                    if not _profile:
                        st.error("❌ Tài khoản chưa được cấu hình. Liên hệ Ban Giám Hiệu nhà trường.")
                    elif not _profile.get("is_active", True):
                        st.error("❌ Tài khoản đã bị tạm khóa. Liên hệ Ban Giám Hiệu.")
                    else:
                        try:
                            _get_sb().table("user_profiles").update({
                                "last_login": now_vn().isoformat()
                            }).eq("id", _user["id"]).execute()
                        except Exception:
                            pass
                        # Ghi nhớ email cho lần sau (trong cùng phiên trình duyệt)
                        st.session_state.last_login_email = _email.strip().lower()
                        st.session_state.auth_user    = _user
                        st.session_state.user_profile = _profile
                        st.rerun()
                except Exception as _le:
                    st.error(f"❌ {_le}")

        st.caption("Chưa có tài khoản? Liên hệ Ban Giám Hiệu nhà trường để được cấp.")

        # Quên mật khẩu
        with st.expander("🔑 Quên mật khẩu?"):
            _fe = st.text_input("Email đã đăng ký", key="reset_email")
            if st.button("Gửi link đặt lại mật khẩu", key="reset_btn"):
                if _fe and db_ok():
                    if db_auth_reset_password(_fe):
                        st.success("✅ Đã gửi email — kiểm tra hộp thư (kể cả spam).")
                    else:
                        st.error("Không gửi được. Kiểm tra kết nối DB.")
                else:
                    st.warning("Nhập email và đảm bảo database đã kết nối.")

    # Demo mode
    st.markdown('<div style="text-align:center;margin-top:20px">', unsafe_allow_html=True)
    if st.button("🔓 Dùng chế độ Demo (không cần đăng nhập)", use_container_width=False):
        st.session_state.auth_user    = {"id": "demo", "email": "demo@demo.vn", "demo": True}
        st.session_state.user_profile = {
            "full_name": "Demo — Chưa đăng nhập", "role": "ban_giam_sat",
            "school_name": "", "is_active": True,
        }
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown(
        '<div style="text-align:center;font-size:0.75rem;color:#94A3B8;margin-top:32px">'
        '⚖️ NĐ 15/2018 · TTLT 13/2016 · QĐ 3958/QĐ-BYT 2025 · SchoolFood AI v2.0'
        '</div>',
        unsafe_allow_html=True,
    )


def tab_user_management(school: str = ""):
    """Ban Giám Hiệu: Quản lý tài khoản người dùng."""
    st.markdown(
        '<div class="sf-card">'
        '<div class="sf-card-title">👤 Quản lý tài khoản người dùng</div>'
        '<div class="sf-card-body">'
        'Tạo tài khoản, phân vai trò và quản lý người dùng trong hệ thống SchoolFood AI</div>'
        '</div>',
        unsafe_allow_html=True,
    )

    if not db_ok():
        st.warning("Database chưa kết nối — không thể quản lý tài khoản.")
        return

    # is_super từ session_state để biết có lock trường không
    _caller_is_super = st.session_state.get("is_super", False)

    # ── Danh sách người dùng (ẩn super admin) ────────────────────────────────
    st.markdown('<div class="sec-hdr">📋 Danh sách tài khoản trong trường</div>',
                unsafe_allow_html=True)
    _all_p  = db_get_all_profiles(school=school if school else "")
    # Ẩn super admin khỏi danh sách BGH thấy
    _profiles = [p for p in _all_p if not p.get("is_super_admin", False)]

    if _profiles:
        import pandas as _pd_um
        _df_um = _pd_um.DataFrame([{
            "Email": p.get("email", ""),
            "Họ và tên": p.get("full_name", ""),
            "Vai trò": ROLE_VN.get(p.get("role", ""), p.get("role", "")),
            "Trường": p.get("school_name", ""),
            "Trạng thái": "✅ Hoạt động" if p.get("is_active") else "🔒 Tạm khóa",
            "Đăng nhập cuối": (
                (__import__("datetime").datetime.fromisoformat(
                    (p.get("last_login") or "")[:19].replace("T"," ")
                ) + __import__("datetime").timedelta(hours=7)
                ).strftime("%d/%m/%Y %H:%M")
                if p.get("last_login") and len(p.get("last_login","")) >= 16 else "—"
            ),
        } for p in _profiles])
        st.dataframe(_df_um, use_container_width=True, hide_index=True)

        # Bật/tắt tài khoản
        st.markdown('<div class="sec-hdr">🔧 Cập nhật trạng thái</div>', unsafe_allow_html=True)
        _emails    = [p.get("email", "") for p in _profiles]
        _sel_email = st.selectbox("Chọn tài khoản", _emails, key="um_sel_email")
        _sel_p     = next((p for p in _profiles if p.get("email") == _sel_email), None)
        if _sel_p:
            _cur_active = _sel_p.get("is_active", True)
            _col1, _ = st.columns(2)
            if _cur_active:
                if _col1.button("🔒 Tạm khóa tài khoản này", key="um_deact"):
                    if db_toggle_profile(_sel_p["id"], False):
                        st.success("✅ Đã tạm khóa!"); st.rerun()
            else:
                if _col1.button("✅ Kích hoạt lại", key="um_act"):
                    if db_toggle_profile(_sel_p["id"], True):
                        st.success("✅ Đã kích hoạt!"); st.rerun()
    else:
        st.info("Chưa có tài khoản nào trong trường này. Hãy thêm bên dưới.")

    # ── Thêm người dùng mới ───────────────────────────────────────────────────
    st.markdown('<div class="sec-hdr">➕ Thêm người dùng cho trường</div>',
                unsafe_allow_html=True)
    st.caption(
        "Tạo tài khoản cho giáo viên / phụ huynh trong trường. "
        "Người dùng nhận mật khẩu từ bạn và đổi sau lần đầu đăng nhập."
    )
    with st.form("add_user_form", clear_on_submit=True):
        _nu_c1, _nu_c2 = st.columns(2)
        _nu_email = _nu_c1.text_input("📧 Email người dùng", placeholder="ten@truong.edu.vn")
        _nu_name  = _nu_c2.text_input("👤 Họ và tên đầy đủ", placeholder="Nguyễn Văn A")
        _nu_c3, _nu_c4 = st.columns(2)

        # Vai trò: BGH được tạo BGH khác (tối đa 2/trường); super admin không giới hạn
        _bgh_count = sum(1 for p in _profiles if p.get("role") == "ban_giam_hieu")
        _bgh_full  = (not _caller_is_super) and (_bgh_count >= 2)
        if _caller_is_super:
            _nu_roles_allowed = list(ROLE_VN.values())
        else:
            # BGH được tạo tất cả vai trò kể cả BGH (nhưng tối đa 2 BGH/trường)
            _nu_roles_allowed = [v for k, v in ROLE_VN.items()
                                 if not (k == "ban_giam_hieu" and _bgh_full)]
        _nu_role = _nu_c3.selectbox("Vai trò", _nu_roles_allowed)
        if _bgh_full and not _caller_is_super:
            st.caption("⚠️ Trường này đã có 2 tài khoản Ban Giám Hiệu — tối đa cho phép.")

        # Trường: Super Admin tự nhập, BGH bị khóa theo trường mình
        if _caller_is_super:
            _nu_school = _nu_c4.text_input("🏫 Tên trường", value=school or "",
                                            placeholder="THCS Nguyễn Du")
        else:
            _nu_school = _nu_c4.text_input("🏫 Tên trường", value=school,
                                            disabled=True, placeholder=school)

        # Cấp học mặc định — chỉ hiện khi tạo tài khoản Y Tế (ảnh hưởng kiem_thuc)
        _nu_default_level = "Tiểu Học (6–11 tuổi)"
        if _nu_role == "Y Tế Học Đường":
            _nu_default_level = st.selectbox(
                "📚 Cấp học phụ trách (mặc định khi đăng nhập)",
                ["Tiểu Học (6–11 tuổi)", "THCS (12–15 tuổi)", "THPT (16–18 tuổi)"],
                help="Dùng cho trường đa cấp — Y Tế sẽ thấy tiêu chuẩn dinh dưỡng đúng cấp khi đăng nhập"
            )

        _nu_pw = st.text_input("🔒 Mật khẩu tạm thời (≥ 6 ký tự)",
                                type="password", placeholder="Người dùng sẽ đổi sau lần đầu")
        _nu_submit = st.form_submit_button("➕ Tạo tài khoản", type="primary",
                                            use_container_width=True)

    if _nu_submit:
        import re as _re_email
        _email_clean = _nu_email.strip().lower() if _nu_email else ""
        _email_valid = bool(_re_email.match(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$',
                                             _email_clean)) if _email_clean else False
        if not all([_nu_email, _nu_name, _nu_pw, _nu_school]):
            st.warning("⚠️ Vui lòng điền đầy đủ: email, họ tên, trường, mật khẩu.")
        elif not _email_valid:
            st.error(f"❌ Email không đúng định dạng: **{_email_clean}**\n\nVD đúng: `ten@truong.edu.vn` · Không được có khoảng trắng hoặc ký tự đặc biệt.")
        elif len(_nu_pw) < 6:
            st.warning("Mật khẩu tạm thời cần ≥ 6 ký tự.")
        else:
            try:
                with st.spinner("Đang tạo tài khoản..."):
                    _uid = db_auth_signup(_email_clean, _nu_pw)
                _ok = db_save_profile(_uid, _email_clean, _nu_name,
                                       ROLE_KEY.get(_nu_role, "phu_huynh"), _nu_school,
                                       default_level=_nu_default_level)
                if _ok:
                    st.success(
                        f"✅ Đã tạo tài khoản **{_nu_name}** ({_nu_email}) — "
                        f"Vai trò: {_nu_role} · Trường: {_nu_school}"
                    )
                    st.info("💡 Người dùng đăng nhập được ngay — không cần xác nhận email.")
                    st.rerun()
                else:
                    st.warning("Tài khoản auth đã tạo nhưng không lưu được profile — kiểm tra DB.")
            except Exception as _ue:
                st.error(f"❌ Lỗi: {_ue}")


def tab_about():
    st.markdown("""<div class="sf-card">
        <div class="sf-card-title">Về SchoolFood AI</div>
        <div class="sf-card-body">Nền tảng giám sát ATTP bữa ăn học đường — giúp mỗi bên thực hiện đúng vai trò, đúng thời điểm, có bằng chứng rõ ràng.</div>
    </div>""", unsafe_allow_html=True)

    for s in SCHEDULE:
        st.markdown(f"""<div class="sf-card" style="padding:14px 18px;border-left:3px solid {s['color']}">
            <div class="sf-card-title">{s['role']}</div>
            <div class="sf-card-body">{s['freq']} · {s['what']}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown('<div class="sec-hdr">Tài liệu pháp luật đã tải</div>', unsafe_allow_html=True)
    pdfs = sorted(LEGAL_DIR.glob("*.pdf"))
    for p in pdfs:
        st.success(f"✅ {p.name} ({p.stat().st_size // 1024} KB)")
    if not pdfs:
        st.warning("Chưa có file PDF trong 07_Legal_Regulations/")
    try:
        import pypdf  # noqa: F401
        st.success("✅ pypdf đã cài — AI đọc được nội dung văn bản pháp luật thực tế")
    except ImportError:
        st.warning("Chưa cài pypdf. Chạy: pip install pypdf")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(page_title="SchoolFood AI", page_icon="🍱",
                       layout="wide", initial_sidebar_state="collapsed")
    inject_css()

    # ── G7: Dark mode CSS ────────────────────────────────────────────────────
    if st.session_state.get("dark_mode"):
        st.markdown("""<style>
        .stApp, [data-testid="stAppViewContainer"] {
            background-color: #0F172A !important;
        }
        .stApp > header { background-color: #0F172A !important; }
        .sf-card, .metric-box, .schedule-card {
            background: #1E293B !important;
            border-color: #334155 !important;
            color: #E2E8F0 !important;
        }
        .sf-card-title, .sec-hdr, .metric-lbl { color: #CBD5E1 !important; }
        .sf-card-body, .alert-body            { color: #94A3B8 !important; }
        .metric-num                            { color: #F1F5F9 !important; }
        .stMarkdown, .stText, p, span, div, label { color: #CBD5E1; }
        .stTextInput input, .stTextArea textarea, .stSelectbox select,
        [data-testid="stTextInput"] input,
        [data-testid="stTextArea"] textarea    {
            background: #1E293B !important;
            color: #E2E8F0 !important;
            border-color: #334155 !important;
        }
        .stTabs [data-baseweb="tab-list"]      { background: #1E293B !important; }
        .stTabs [data-baseweb="tab"]           { color: #94A3B8 !important; }
        .stTabs [aria-selected="true"]         { color: #38BDF8 !important; }
        .main .block-container                 { background: #0F172A !important; }
        [data-testid="stSidebar"]              { background: #1E293B !important; }
        .stButton > button                     {
            background: #1E293B !important;
            color: #E2E8F0 !important;
            border-color: #334155 !important;
        }
        .stButton > button[kind="primary"]     {
            background: #2563EB !important;
            color: white !important;
        }
        .stDataFrame, .stTable                 {
            background: #1E293B !important;
            color: #E2E8F0 !important;
        }
        [data-baseweb="select"] > div          {
            background: #1E293B !important;
            color: #E2E8F0 !important;
            border-color: #334155 !important;
        }
        .alert-ok      { background: #052e16 !important; border-color: #166534 !important; }
        .alert-minor   { background: #1c1a00 !important; border-color: #854d0e !important; }
        .alert-major   { background: #1c0a00 !important; border-color: #7c2d12 !important; }
        .alert-critical{ background: #1a0000 !important; border-color: #991b1b !important; }
        </style>""", unsafe_allow_html=True)

    # Header chính — không dùng HTML comments để tránh Streamlit hiển thị raw text
    _circles = (
        '<div style="position:absolute;top:-40px;right:-40px;width:220px;height:220px;'
        'border-radius:50%;background:rgba(255,255,255,0.05);pointer-events:none"></div>'
        '<div style="position:absolute;bottom:-60px;right:80px;width:160px;height:160px;'
        'border-radius:50%;background:rgba(255,255,255,0.04);pointer-events:none"></div>'
    )
    _badges = (
        '<span style="background:rgba(255,255,255,0.15);color:white;padding:4px 12px;'
        'border-radius:20px;font-size:0.75rem;font-weight:600;border:1px solid rgba(255,255,255,0.2)">'
        '⚖️ NĐ 15/2018</span>'
        '<span style="background:rgba(255,255,255,0.15);color:white;padding:4px 12px;'
        'border-radius:20px;font-size:0.75rem;font-weight:600;border:1px solid rgba(255,255,255,0.2)">'
        '🤖 AI Vision</span>'
        '<span style="background:rgba(34,197,94,0.25);color:#86EFAC;padding:4px 12px;'
        'border-radius:20px;font-size:0.75rem;font-weight:600;border:1px solid rgba(134,239,172,0.3)">'
        '🟢 Live</span>'
    )
    _stat = lambda val, lbl, clr: (
        f'<div style="flex:1;min-width:80px;text-align:center;padding:0 8px">'
        f'<div style="color:{clr};font-size:1.8rem;font-weight:800;line-height:1.1">{val}</div>'
        f'<div style="color:#BFDBFE;font-size:0.82rem;margin-top:5px;font-weight:500">{lbl}</div></div>'
    )
    _sep = '<div style="width:1px;background:rgba(255,255,255,0.15);margin:0 2px"></div>'
    _stats = (
        _stat("20", "Điểm kiểm tra", "white") + _sep +
        _stat("7",  "Mục bắt buộc",  "#FCA5A5") + _sep +
        _stat("4",  "Cấp cảnh báo",  "#6EE7B7") + _sep +
        _stat("6",  "Văn bản pháp lý","#FDE68A") + _sep +
        _stat("🤖", "AI phân tích",   "#C4B5FD")
    )
    # ── WOW Header — AI-powered, đột phá ──────────────────────────────────────
    st.markdown(
        # Outer container: multi-layer gradient (navy → blue → AI purple)
        '<div style="background:linear-gradient(135deg,#0C1445 0%,#1E3A8A 35%,#1D4ED8 65%,#5B21B6 100%);'
        'border-radius:20px;padding:28px 32px 22px 32px;margin-bottom:16px;'
        'position:relative;overflow:hidden;'
        'box-shadow:0 20px 60px rgba(29,78,216,0.4),0 4px 20px rgba(91,33,182,0.3)">'

        # Background decorative elements
        '<div style="position:absolute;top:-60px;right:-60px;width:280px;height:280px;'
        'border-radius:50%;background:radial-gradient(circle,rgba(124,58,237,0.2),transparent 70%);'
        'pointer-events:none"></div>'
        '<div style="position:absolute;bottom:-80px;left:10%;width:200px;height:200px;'
        'border-radius:50%;background:radial-gradient(circle,rgba(59,130,246,0.15),transparent 70%);'
        'pointer-events:none"></div>'
        '<div style="position:absolute;top:50%;right:25%;width:140px;height:140px;'
        'border-radius:50%;background:radial-gradient(circle,rgba(16,185,129,0.1),transparent 70%);'
        'pointer-events:none"></div>'

        # Top row: logo + title + badges
        '<div style="display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:12px">'
        '<div style="display:flex;align-items:center;gap:16px">'
        # Emoji in glowing circle
        '<div style="background:rgba(255,255,255,0.12);backdrop-filter:blur(10px);'
        'border:1px solid rgba(255,255,255,0.2);border-radius:16px;'
        'padding:12px 14px;font-size:2.8rem;line-height:1;'
        'box-shadow:0 0 20px rgba(124,58,237,0.4)">🍱</div>'
        '<div>'
        '<div style="color:white;font-size:2.2rem;font-weight:900;letter-spacing:-1px;line-height:1;'
        'text-shadow:0 2px 20px rgba(255,255,255,0.3)">'
        'SchoolFood <span style="background:linear-gradient(90deg,#60A5FA,#A78BFA);'
        '-webkit-background-clip:text;-webkit-text-fill-color:transparent;'
        'background-clip:text">AI</span></div>'
        '<div style="color:rgba(255,255,255,0.65);font-size:0.8rem;margin-top:5px;'
        'display:flex;align-items:center;gap:8px">'
        '<span style="background:rgba(16,185,129,0.3);border:1px solid rgba(16,185,129,0.5);'
        'border-radius:20px;padding:2px 10px;font-weight:700;color:#6EE7B7;font-size:0.72rem">'
        '● LIVE</span>'
        '<span style="opacity:0.7">Powered by Claude AI · NĐ 15/2018 · TTLT 13/2016</span>'
        '</div></div></div>'
        # Right: badge pills
        f'<div style="display:flex;gap:6px;flex-wrap:wrap;align-items:flex-start;padding-top:6px">{_badges}</div>'
        '</div>'

        # Description
        '<p style="color:rgba(219,234,254,0.9);font-size:0.98rem;margin:16px 0 14px 0;'
        'line-height:1.7;font-weight:400;max-width:700px">'
        'Nền tảng giám sát An toàn Thực phẩm học đường — '
        '<b style="color:white">AI phân tích ảnh</b>, '
        '<b style="color:#A78BFA">câu hỏi chống gian lận</b>, '
        '<b style="color:#6EE7B7">alert real-time</b>, '
        '<b style="color:#FCD34D">báo cáo chuẩn hành chính</b>'
        '</p>'

        # Stats bar
        f'<div style="display:flex;gap:0;border-top:1px solid rgba(255,255,255,0.12);'
        f'padding-top:14px;flex-wrap:wrap">{_stats}</div>'
        '</div>',
        unsafe_allow_html=True,
    )

    # ── G3: Kiểm tra xác thực ────────────────────────────────────────────────
    # Nếu DB kết nối và chưa đăng nhập → hiện trang login
    # Nếu DB không có (demo/local) → dùng chế độ selectbox cũ
    _use_auth = db_ok()
    _auth_user    = st.session_state.get("auth_user")
    _user_profile = st.session_state.get("user_profile")

    if _use_auth and not _auth_user:
        show_login_page()
        return  # Dừng main(), hiện login page

    # ── Đọc API key ──────────────────────────────────────────────────────────
    import os
    api_key = (
        (st.secrets.get("ANTHROPIC_API_KEY", "") if hasattr(st, "secrets") else "")
        or os.environ.get("ANTHROPIC_API_KEY", "")
    )

    # ── Thanh thông tin người dùng / điều khiển ──────────────────────────────
    if _use_auth and _user_profile:
        # ── Chế độ đã đăng nhập ──────────────────────────────────────────────
        _pf           = _user_profile
        _is_super     = _pf.get("is_super_admin", False)
        _role_key     = _pf.get("role", "phu_huynh")
        _school_pf    = _pf.get("school_name", "")
        _is_demo      = _auth_user.get("demo", False)

        # ── Thanh thông tin user — thiết kế mới 1 hàng ──────────────────────
        _role_display = ROLE_VN.get(_role_key, "Phụ Huynh")
        _clr     = ROLE_CLR_MAP.get(_role_display, "#64748B")
        _dm_now  = st.session_state.get("dark_mode", False)
        _bg_bar  = "#1E293B" if _dm_now else "white"
        _brd_bar = "#334155" if _dm_now else "#E2E8F0"
        _name_c  = "#F1F5F9" if _dm_now else "#1E293B"
        _email_c = "#94A3B8"
        _dm_on   = _dm_now

        # Badge info HTML
        _badge = (
            f'<span style="background:#7C3AED;color:white;border-radius:6px;'
            f'padding:3px 12px;font-size:0.75rem;font-weight:800;letter-spacing:0.03em">⚡ ADMIN</span>'
            if _is_super else
            f'<span style="background:{_clr};color:white;border-radius:6px;'
            f'padding:3px 12px;font-size:0.75rem;font-weight:700">{_role_display}</span>'
        )
        _school_html = (
            f'<span style="background:#EFF6FF;color:#2563EB;border-radius:5px;'
            f'padding:2px 8px;font-size:0.72rem;font-weight:600">🏫 {_school_pf}</span>'
            if _school_pf and not _is_super else ''
        )
        _demo_html = (
            '<span style="background:#FEF9C3;color:#92400E;border-radius:5px;'
            'padding:2px 8px;font-size:0.7rem;font-weight:600">🔓 Demo</span>'
            if _is_demo else ''
        )
        _api_html = (
            '<span style="background:#DCFCE7;color:#166534;border-radius:5px;'
            'padding:2px 8px;font-size:0.72rem;font-weight:600">✅ AI</span>'
            if api_key else
            '<span style="background:#F1F5F9;color:#64748B;border-radius:5px;'
            'padding:2px 8px;font-size:0.72rem">AI off</span>'
        )

        st.markdown(
            f'<div style="background:{_bg_bar};border:1px solid {_brd_bar};border-radius:12px;'
            f'padding:8px 16px;margin-bottom:8px;box-shadow:0 1px 4px rgba(0,0,0,0.06)">'
            f'<div style="display:flex;align-items:center;justify-content:space-between;'
            f'flex-wrap:wrap;gap:8px">'
            # Trái: badge + tên + email + trường
            f'<div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap">'
            f'{_badge}'
            f'<span style="font-size:0.9rem;color:{_name_c};font-weight:700">'
            f'{_pf.get("full_name","")}</span>'
            f'<span style="font-size:0.77rem;color:{_email_c}">'
            f'{_auth_user.get("email","")}</span>'
            f'{_school_html}{_demo_html}'
            f'</div>'
            # Phải: API status badge
            f'<div style="display:flex;align-items:center;gap:6px">'
            f'{_api_html}'
            f'</div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

        # Hàng điều khiển: role switcher (admin) + API key input + dark mode + logout
        if _is_super:
            _ctrl = st.columns([3, 2, 0.6, 0.8])
            with _ctrl[0]:
                role = st.selectbox(
                    "🔧 Xem vai trò",
                    ["Ban Giám Hiệu", "Ban Giám Sát (Đại Diện PHHS)",
                     "Y Tế Học Đường", "Phụ Huynh"],
                    key="admin_role_switch", label_visibility="collapsed",
                )
            with _ctrl[1]:
                if not api_key:
                    _mk = st.text_input("Key", type="password",
                                        placeholder="sk-ant-... (tuỳ chọn)",
                                        label_visibility="collapsed")
                    if _mk: api_key = _mk
            if _ctrl[2].button("🌙" if not _dm_on else "☀️", use_container_width=True,
                                help="Chuyển Dark/Light mode"):
                st.session_state.dark_mode = not _dm_on; st.rerun()
            if _ctrl[3].button("🚪 Đăng xuất", use_container_width=True):
                for _k in ("auth_user","user_profile","admin_role_switch"):
                    st.session_state.pop(_k, None)
                st.rerun()
        else:
            role = _role_display
            _ctrl = st.columns([3, 0.6, 0.8])
            with _ctrl[0]:
                if not api_key:
                    _mk = st.text_input("Key", type="password",
                                        placeholder="sk-ant-... (Claude API Key, tuỳ chọn)",
                                        label_visibility="collapsed")
                    if _mk: api_key = _mk
            if _ctrl[1].button("🌙" if not _dm_on else "☀️", use_container_width=True,
                                help="Chuyển Dark/Light mode"):
                st.session_state.dark_mode = not _dm_on; st.rerun()
            if _ctrl[2].button("🚪 Đăng xuất", use_container_width=True):
                for _k in ("auth_user","user_profile"):
                    st.session_state.pop(_k, None)
                st.rerun()

        # Đổi mật khẩu: đặt ở cuối tab Hướng dẫn (xem tab_guide function)

        # Level — tất cả vai trò đều lấy từ profile (Y Tế chọn trong tab_kiem_thuc)
        _default_lvl = _pf.get("default_level") or "Tiểu Học (6–11 tuổi)"
        level = _default_lvl
        loc = _school_pf or "TP.HCM"
        # Admin không bị lock trường — school chỉ lock cho các vai trò thường
        st.session_state.user_school  = "" if _is_super else _school_pf
        st.session_state.user_role    = role
        st.session_state.is_bgh       = (role == "Ban Giám Hiệu")
        st.session_state.is_super     = _is_super

    else:
        # ── Chế độ demo/local: giữ selectbox cũ ─────────────────────────────
        st.markdown(
            '<div style="background:#FEF9C3;border:1px solid #FDE68A;border-radius:10px;'
            'padding:8px 16px;margin-bottom:8px;font-size:0.8rem;color:#92400E">'
            '🔓 <b>Chế độ Demo</b> — Database chưa kết nối, dùng selectbox để chọn vai trò</div>',
            unsafe_allow_html=True,
        )
        ctrl1, ctrl2, ctrl3, ctrl4 = st.columns([2, 2, 1.5, 2.5])
        with ctrl1:
            role = st.selectbox(
                "Vai trò",
                ["Phụ Huynh", "Ban Giám Sát (Đại Diện PHHS)",
                 "Y Tế Học Đường", "Ban Giám Hiệu"],
            )
        with ctrl2:
            level = st.selectbox(
                "Cấp trường",
                ["Tiểu Học (6–11 tuổi)", "THCS (12–15 tuổi)", "THPT (16–18 tuổi)"],
            )
        with ctrl3:
            loc = st.text_input("Tỉnh/TP", value="TP.HCM")
        with ctrl4:
            if api_key:
                st.text_input("AI", value="✅ AI đã kết nối", disabled=True,
                              label_visibility="visible")
            else:
                _mk2 = st.text_input("Claude API Key", type="password",
                                     placeholder="sk-ant-...",
                                     help="Không có key? Checklist vẫn dùng được đầy đủ")
                if _mk2:
                    api_key = _mk2
        _school_pf = ""
        st.session_state.user_school  = ""
        st.session_state.user_role    = role
        st.session_state.is_bgh       = (role == "Ban Giám Hiệu")

    # ── Mô tả vai trò (dùng chung) ───────────────────────────────────────────
    DESCS = {
        "Phụ Huynh":                    "Xem thực đơn, kết quả kiểm tra và gửi phản hồi",
        "Ban Giám Sát (Đại Diện PHHS)": "Kiểm tra bếp ăn 2 lần/tuần, tạo báo cáo chính thức theo luật",
        "Y Tế Học Đường":               "Ghi kiểm thực 3 bước hàng ngày, xác nhận mẫu lưu thức ăn",
        "Ban Giám Hiệu":                "Xem tổng quan tình hình ATTP, duyệt báo cáo và quản lý nhà cung cấp",
    }
    ROLE_CLR = ROLE_CLR_MAP

    # Tính nhắc nhở
    _t_info   = _REMINDER_TIMES.get(role)
    _now      = now_vn()
    _reminder_txt = ""
    if _t_info and _now.weekday() < 5:
        _it     = _t_info["hour"] * 60 + _t_info["min"]
        _ct     = _now.hour * 60 + _now.minute
        _ml     = _it - _ct
        _is_day = _now.weekday() in _t_info["days"]
        if _is_day and 0 <= _ml <= 15:
            _reminder_txt = (
                f'&nbsp;&nbsp;<span style="background:#FEF9C3;color:#92400E;font-weight:700;'
                f'padding:2px 10px;border-radius:12px;font-size:0.78rem">'
                f'⏰ NHẮC: Còn {_ml} phút đến giờ kiểm tra!</span>'
            )
        elif _is_day and _ml > 0:
            _rt  = _it - 15
            _reminder_txt = (
                f'&nbsp;&nbsp;<span style="color:#64748B;font-size:0.78rem">'
                f'⏰ Nhắc nhở lúc {_rt//60:02d}:{_rt%60:02d} — Kiểm tra {_t_info["hour"]:02d}:{_t_info["min"]:02d}</span>'
            )
        elif _is_day and _ml <= 0:
            _reminder_txt = (
                f'&nbsp;&nbsp;<span style="color:#94A3B8;font-size:0.78rem">'
                f'Đã qua giờ kiểm tra hôm nay</span>'
            )
        elif not _is_day:
            _next_d = next((d for d in sorted(_t_info["days"]) if d > _now.weekday()), min(_t_info["days"]))
            _day_vn = ["Thứ 2","Thứ 3","Thứ 4","Thứ 5","Thứ 6"][_next_d]
            _reminder_txt = (
                f'&nbsp;&nbsp;<span style="color:#64748B;font-size:0.78rem">'
                f'📅 Lịch tiếp theo: {_day_vn} {_t_info["hour"]:02d}:{_t_info["min"]:02d}</span>'
            )

    st.markdown(
        f'<div style="display:flex;justify-content:space-between;align-items:center;'
        f'font-size:0.82rem;color:#475569;margin-bottom:6px;flex-wrap:wrap;gap:4px">'
        f'<span><b style="color:{ROLE_CLR.get(role,"#64748B")}">{role}:</b> '
        f'{DESCS.get(role,"")}{_reminder_txt}</span>'
        f'<span style="color:#94A3B8">🔴 Khẩn cấp: <b>115</b> &nbsp;·&nbsp; Cục ATTP: <b>1800 6838</b></span>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Hiển thị banner nhắc nhở nổi bật nếu trong 15 phút
    show_reminder_banner(role)

    # ── G5: Onboarding cho BGH mới ───────────────────────────────────────────
    # Onboarding check — chỉ chạy khi cần, không thêm DB call
    if role == "Ban Giám Hiệu" and _use_auth and _school_pf:
        if not st.session_state.get("onboarding_dismissed"):
            _ob_profiles = db_get_all_profiles(school=_school_pf)
            _ob_profiles = [p for p in _ob_profiles if not p.get("is_super_admin", False)]
            _ob_sessions = 0  # Sẽ check trong banner nếu cần
            show_onboarding_banner(_school_pf, _ob_profiles, _ob_sessions)

    # Lịch sử — gắn cờ đỏ nếu có CRITICAL gần đây
    # Dùng cache: tránh gọi DB thêm lần nữa chỉ để check critical flag
    _hist_label = "📊 Lịch sử"  # Flag đỏ sẽ hiện trong alert banner

    # ── Tabs theo từng vai trò — mỗi vai trò thấy đúng chức năng mình cần ─────
    if role == "Phụ Huynh":
        # Phụ Huynh: xem kết quả đơn giản + phản hồi, không cần dashboard kỹ thuật
        _tabs = st.tabs([
            "💬 Hỏi đáp AI",
            "🍱 Góc Phụ Huynh",
            "🚨 Khẩn cấp",
            "📖 Hướng dẫn",
        ])
        with _tabs[0]: tab_chat(api_key, role, level, loc)
        with _tabs[1]: tab_parent_view(api_key)
        with _tabs[2]: tab_emergency(api_key)
        with _tabs[3]: tab_guide()

    elif role == "Y Tế Học Đường":
        # Y Tế: kiểm thực 3 bước + kiểm tra NCC (Bước 1 yêu cầu check nguyên liệu đầu vào)
        # TTLT 13/2016 Điều 9 khoản a: Y tế kiểm tra nguồn gốc nguyên liệu → cần tab NCC
        _tabs = st.tabs([
            "💬 Hỏi đáp AI",
            "🏥 Kiểm thực 3 bước",
            "🏭 Nhà Cung Cấp",
            _hist_label,
            "🚨 Khẩn cấp",
            "📖 Hướng dẫn",
        ])
        with _tabs[0]: tab_chat(api_key, role, level, loc)
        with _tabs[1]: tab_kiem_thuc(api_key, level)
        with _tabs[2]: tab_supplier(api_key, role=role)
        with _tabs[3]: tab_history(role=role, school_filter="" if _is_super else _school_pf)
        with _tabs[4]: tab_emergency(api_key)
        with _tabs[5]: tab_guide()

    elif role == "Ban Giám Sát (Đại Diện PHHS)":
        # BGS: checklist + NCC + lịch sử + lịch chuẩn mực + hướng dẫn
        _tabs = st.tabs([
            "💬 Hỏi đáp AI",
            "✅ Checklist kiểm tra",
            "🏭 Nhà Cung Cấp",
            _hist_label,
            "📅 Lịch & Chuẩn mực",
            "🚨 Khẩn cấp",
            "📖 Hướng dẫn",
        ])
        with _tabs[0]: tab_chat(api_key, role, level, loc)
        with _tabs[1]: tab_checklist(api_key)
        with _tabs[2]: tab_supplier(api_key, role=role)
        with _tabs[3]: tab_history(role=role, school_filter="" if _is_super else _school_pf)
        with _tabs[4]: tab_schedule()
        with _tabs[5]: tab_emergency(api_key)
        with _tabs[6]: tab_guide()

    else:  # Ban Giám Hiệu
        # BGH: dashboard + NCC tab riêng + lịch + khẩn cấp + hướng dẫn + quản lý TK
        _bgh_tabs = ["💬 Hỏi đáp AI", _hist_label, "🏭 Nhà Cung Cấp",
                     "📅 Lịch & Chuẩn mực", "🚨 Khẩn cấp", "📖 Hướng dẫn"]
        if _use_auth:
            _bgh_tabs.append("👤 Quản lý tài khoản")
        _tabs = st.tabs(_bgh_tabs)
        with _tabs[0]: tab_chat(api_key, role, level, loc)
        with _tabs[1]: tab_history(role=role, school_filter="" if _is_super else _school_pf)
        with _tabs[2]: tab_ncc_bgh(school="" if _is_super else _school_pf)
        with _tabs[3]: tab_schedule()
        with _tabs[4]: tab_emergency(api_key)
        with _tabs[5]: tab_guide()
        if _use_auth and len(_tabs) > 6:
            with _tabs[6]: tab_user_management(school=_school_pf)


if __name__ == "__main__":
    main()
